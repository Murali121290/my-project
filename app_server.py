from flask import (
    Flask, render_template, render_template_string, request, redirect, flash,
    send_file, url_for, jsonify, session, send_from_directory, g
)
import os
import io
import shutil
import zipfile
from dotenv import load_dotenv
import utils.url_validator
import ReferenceConversion
import utils.track_changes
from ReferenceConversion import _looks_like_inline_citation

load_dotenv()

ALTTEXT_URL = os.environ.get("ALTTEXT_URL", "http://alttext_app:5000")
ALTTEXT_SERVICE_TOKEN = os.environ.get("ALTTEXT_SERVICE_TOKEN", "")

import warnings
# Suppress sqlite3 datetime adapter deprecation warning (Python 3.12+)
warnings.filterwarnings('ignore', message='.*default datetime adapter is deprecated.*', category=DeprecationWarning)

import traceback
import uuid
import json
import re
import sqlite3
import socket
import time
import threading
import concurrent.futures
import requests as _requests_lib
try:
    import psycopg2
    from psycopg2 import pool, extras
except ImportError:
    psycopg2 = None

# Windows/Linux Compatibility
# Word automation removed for Linux deployment
HAS_WIN32COM = False
pythoncom = None
win32 = None

from pathlib import Path
from threading import Lock
from datetime import datetime, timedelta, timezone
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.middleware.proxy_fix import ProxyFix
from flask_wtf.csrf import CSRFProtect
from functools import wraps
from waitress import serve
from contextlib import contextmanager
from queue import Queue, Empty
import logging
from logging.handlers import RotatingFileHandler
from collections import defaultdict
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table

from highlighter.core_highlighter_docx import process_docx
import ReferencesStructing
# APA validation now lives in validation_core (CitationProcessor)
from validation_core import CitationProcessor, ValidationReport
import ReferenceConversion
import utils.track_changes
from ReferenceConversion import _looks_like_inline_citation

# Compatibility shims — keep these names alive in case any code references them
def validate_document_multi_style(file_path, style=None):
    """Shim: run CitationProcessor and return ValidationReport."""
    p = CitationProcessor(file_path)
    return p.run()

def generate_apa_report(report_or_results, filename=""):
    """Shim: accept either a ValidationReport or the old dict."""
    if isinstance(report_or_results, ValidationReport):
        return f"Document: {filename}\n\n{report_or_results.summary()}"
    return str(report_or_results)

def apply_citation_formatting(file_path, results):
    """Shim: no-op — CitationProcessor applies formatting in-place."""
    return 0

def insert_comments_in_document(file_path, results, *args, **kwargs):
    """Shim: no-op — CitationProcessor inserts comments in-place."""
    from docx import Document
    return Document(file_path), 0
import tempfile
from io import BytesIO
from extractor import extract_from_file, write_permission_log

from extractor_ai import extract_from_file_ai
import bias_scanner

def _now_utc():
    return datetime.now(timezone.utc)
# -----------------------
# Configuration
# -----------------------
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, "S4C-Processed-Documents")
COMMON_MACRO_FOLDER = os.path.join(BASE_DIR, "S4c-Macros")
DEFAULT_MACRO_NAME = 'CE_Tool.dotm'
REPORT_FOLDER = "reports"
DATABASE = os.path.join(BASE_DIR, "reference_validator.db")
LOG_FILE = os.path.join(BASE_DIR, 'user_activity.log')

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(COMMON_MACRO_FOLDER, exist_ok=True)
os.makedirs(REPORT_FOLDER, exist_ok=True)
os.makedirs("logs", exist_ok=True)

ALLOWED_EXTENSIONS = {'.doc', '.docx'}
WORD_START_RETRIES = 3
WORD_LOCK = Lock()
TOKEN_TTL = timedelta(hours=1)

# -----------------------
# Route-Specific Macro Configuration
# -----------------------
ROUTE_MACROS = {
    'language': {
        'name': 'Language Editing',
        'description': 'Language editing and grammar correction tools',
        'icon': 'edit',
        'macros': [
            "LanguageEdit.GrammarCheck_WithErrorHandling",
            "LanguageEdit.SpellCheck_Advanced",
            "LanguageEdit.StyleConsistency_Check",
            "LanguageEdit.ReadabilityAnalysis",
            "LanguageEdit.TerminologyValidation"
        ]
    },
    'technical': {
        'name': 'Technical Editing',
        'description': 'Technical document formatting and validation tools',
        'icon': 'cog',
        'macros': [
            "Referencevalidation.ValidateBWNumCite_WithErrorHandling",
            "ReferenceRenumber.Reorderbasedonseq",
            "Copyduplicate.duplicate4",
            "citationupdateonly.citationupdate",
            "techinal.technicalhighlight"
        ]
    },
    'macro_processing': {
        'name': 'Reference Processing',
        'description': 'Reference validation and citation tools',
        'icon': 'bookmark',
        'macros': [
            "Referencevalidation.ValidateBWNumCite_WithErrorHandling",
            "ReferenceRenumber.Reorderbasedonseq",
            "Copyduplicate.duplicate4",
            "citationupdateonly.citationupdate",
            "Prediting.Preditinghighlight",
            "msrpre.GenerateDashboardReport",
        ]
    },
    'ppd': {
        'name': 'PPD Processing',
        'description': 'PPD final processing tools (from PPD_Final.py)',
        'icon': 'magic',
        'macros': [
            "PPD_HTML.GenerateDocument",
            "PPD_HTML.Generate_HTML_WORDReport",
        ]
    }
}

ROUTE_MACROS['credit_extractor'] = {
    'name': 'Credit / Permission Log',
    'description': 'Caption & credit line extraction for permissions',
    'icon': 'file-text',
    'macros': []
}

ROUTE_MACROS['word_to_xml'] = {
    'name': 'Word to XML Converter',
    'description': 'Convert Word documents to XML format',
    'icon': 'file-code',
    'macros': []
}

ROUTE_MACROS['validation'] = {
    'name': 'Reference Validation',
    'description': 'Automated reference structuring and validation',
    'icon': 'check-circle',
    'macros': []
}

ROUTE_MACROS['bias_scan'] = {
    'name': 'Bias Scanner',
    'description': 'Scan documents for bias terms and generate reports',
    'icon': 'search',
    'macros': []
}

# Flask app
app = Flask(__name__)

# Apply ProxyFix for Nginx (handles HTTPS, X-Forwarded-Proto, etc.)
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1, x_prefix=1)

# Fix for CSRF token missing in multi-worker environment
# Ensure secret key is consistent across workers by storing it in a file if not in env
secret_key_path = os.path.join(BASE_DIR, '.flask_secret_key')
if os.environ.get('SECRET_KEY'):
    app.secret_key = os.environ.get('SECRET_KEY')
elif os.path.exists(secret_key_path):
    with open(secret_key_path, 'rb') as f:
        app.secret_key = f.read()
else:
    # Generate and save a new key so it persists across restarts and workers
    generated_key = os.urandom(24)
    try:
        with open(secret_key_path, 'wb') as f:
            f.write(generated_key)
        app.secret_key = generated_key
    except IOError:
        # Fallback if cannot write to file
        app.secret_key = 'fallback-secret-key-change-this-in-prod'

csrf = CSRFProtect(app)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['COMMON_MACRO_FOLDER'] = COMMON_MACRO_FOLDER
app.config['REPORT_FOLDER'] = REPORT_FOLDER
app.config['DATABASE'] = DATABASE
app.config['MAX_CONTENT_LENGTH'] = 200 * 1024 * 1024  # 200 MB

# Session Hardening
app.config['SESSION_COOKIE_NAME'] = 's4c_session'
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=1)
app.config['SESSION_REFRESH_EACH_REQUEST'] = True

# Token-based download map
download_tokens = {}

ROUTE_PERMISSIONS = {
    'language': ['COPYEDIT', 'ADMIN'],
    'technical': ['COPYEDIT', 'ADMIN'],
    'macro_processing': ['COPYEDIT', 'ADMIN'],
    'ppd': ['PERMISSIONS','COPYEDIT', 'PPD', 'PM', 'ADMIN'],
    'credit_extractor': ['PERMISSIONS', 'PM', 'PPD','ADMIN'],
    'word_to_xml': ['ADMIN'],
    'bias_scan': ['COPYEDIT', 'ADMIN'],
    'alttext': ['COPYEDIT', 'ADMIN', 'PERMISSIONS', 'PM', 'PPD'],
}

def get_user_role():
    return session.get('role') or (g.user.get('role') if g.user else None)

def has_role(*roles):
    role = get_user_role()
    return role is not None and role.upper() in [r.upper() for r in roles]

def role_required(allowed_roles):
    def decorator(f):
        @wraps(f)
        def wrapped(*args, **kwargs):
            if 'user_id' not in session:
                flash("Please log in to continue.")
                return redirect(url_for('login'))

            if not has_role(*allowed_roles) and not session.get('is_admin'):
                flash("You don't have permission to access this page.", "error")
                return redirect(url_for('dashboard'))

            return f(*args, **kwargs)
        return wrapped
    return decorator

def process_credit_extractor_job(job_id, temp_dir, file_paths, original_filenames, user_id, username, extraction_method="manual", api_key=None):
    with app.app_context():
        # Helper to update progress file
        def update_progress(updates):
            try:
                p_path = os.path.join(temp_dir, "progress.json")
                current = {}
                if os.path.exists(p_path):
                    with open(p_path, "r") as f:
                        current = json.load(f)
                current.update(updates)
                with open(p_path, "w") as f:
                    json.dump(current, f)
            except Exception as ex:
                print(f"Progress update failed: {ex}")

        # Initialize progress file
        update_progress({
            "total": len(file_paths),
            "current": 0,
            "status": "Starting",
            "folder": temp_dir
        })

        all_results = []
        
        try:
            for idx, path in enumerate(file_paths, start=1):
                filename = original_filenames[idx-1]
                update_progress({
                    "current": idx,
                    "status": f"Processing {filename}"
                })

                if extraction_method == "ai":
                    all_results.extend(extract_from_file_ai(path, api_key))
                else:
                    all_results.extend(extract_from_file(path))

            if not all_results:
                update_progress({"status": "No captions found"})
                return

            output_xlsx = os.path.join(temp_dir, "permission_log.xlsx")
            write_permission_log(all_results, output_xlsx)

            # Always return the Excel file directly
            final_path = output_xlsx
            processed_files = ["permission_log.xlsx"]

            # Register download token (Optional if relying solely on file system, but kept for consistency)
            token = uuid.uuid4().hex
            download_tokens[token] = {
                "path": temp_dir,
                "expires": _now_utc() + TOKEN_TTL,
                "user": username,
                "route_type": "credit_extractor"
            }

            # DB logging
            with db_pool.get_connection() as db:
                db.execute(
                    '''INSERT INTO macro_processing
                       (user_id, token, original_filenames, processed_filenames, selected_tasks, route_type)
                       VALUES (?, ?, ?, ?, ?, ?)''',
                    (
                        user_id,
                        token,
                        json.dumps(original_filenames),
                        json.dumps(processed_files),
                        json.dumps({"route_type": "credit_extractor"}),
                        "credit_extractor"
                    )
                )
                db.commit()

            update_progress({
                "status": "Completed",
                "download_token": token,
                "zip_path": final_path # Used by download_zip
            })

        except Exception as e:
            update_progress({"status": f"Failed: {e}"})

@app.route("/credit-extractor", methods=["GET", "POST"])
@csrf.exempt
@role_required(ROUTE_PERMISSIONS.get('credit_extractor', ['ADMIN']))
def credit_extractor():
    if request.method == "POST":
        files = request.files.getlist("files")

        if not files or all(f.filename == "" for f in files):
            return jsonify({"error": "No files selected"}), 400

        extraction_method = request.form.get("extraction_method", "manual")
        api_key = request.form.get("api_key", "")

        # Use token as job_id for consistency
        token = uuid.uuid4().hex
        job_id = token
        
        # Save files synchronously before threading
        temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], token)
        os.makedirs(temp_dir, exist_ok=True)

        saved_paths = []
        original_filenames = []
        
        try:
            for f in files:
                if f.filename:
                    safe_name = secure_filename(f.filename) or f"document_{len(saved_paths)}"
                    path = os.path.join(temp_dir, safe_name)
                    f.save(path)
                    saved_paths.append(path)
                    original_filenames.append(f.filename)
        except Exception as e:
            try:
                shutil.rmtree(temp_dir)
            except:
                pass
            return jsonify({"error": f"File save failed: {e}"}), 500

        # In-memory update for same-worker immediate feedback (optional)
        app.config.setdefault("PROGRESS_DATA", {})
        app.config["PROGRESS_DATA"][job_id] = {
            "total": len(saved_paths),
            "current": 0,
            "status": "Starting"
        }

        batch_queue.submit(
            job_id=job_id,
            route_type='credit_extractor',
            user_id=session['user_id'],
            username=session['username'],
            target_fn=process_credit_extractor_job,
            fn_args=(job_id, temp_dir, saved_paths, original_filenames,
                     session['user_id'], session['username'], extraction_method, api_key),
            payload_dict={
                'temp_dir': temp_dir,
                'saved_paths': saved_paths,
                'original_filenames': original_filenames,
                'extraction_method': extraction_method,
                'api_key': api_key,
            }
        )

        return jsonify({"job_id": job_id})

    system_key_configured = bool(os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY") or app.config.get("GEMINI_API_KEY"))
    return render_template("upload_credit.html", system_key_configured=system_key_configured)

def process_bias_scan_job(job_id, temp_dir, file_paths, original_filenames, user_id, username):
    """
    Process bias scanning job in background.
    Scans documents for bias terms and creates ZIP with highlighted docs and Excel report.
    """
    with app.app_context():
        # Helper to update progress file
        def update_progress(updates):
            try:
                p_path = os.path.join(temp_dir, "progress.json")
                current = {}
                if os.path.exists(p_path):
                    with open(p_path, "r") as f:
                        current = json.load(f)
                current.update(updates)
                with open(p_path, "w") as f:
                    json.dump(current, f)
            except Exception as ex:
                print(f"Progress update failed: {ex}")

        # Initialize progress file
        update_progress({
            "total": len(file_paths),
            "current": 0,
            "status": "Starting bias scan",
            "folder": temp_dir
        })

        try:
            # Load bias terms
            bias_terms_path = os.path.join(BASE_DIR, "bias_terms.csv")
            term_category_map, categories = bias_scanner.load_bias_terms(bias_terms_path)
            
            if not term_category_map:
                update_progress({"status": "Failed: bias_terms.csv not found or empty"})
                return

            # Create output directories
            word_out_dir = os.path.join(temp_dir, "word")
            os.makedirs(word_out_dir, exist_ok=True)
            
            all_report_rows = []
            
            # Process each file
            for idx, path in enumerate(file_paths, start=1):
                filename = original_filenames[idx-1]
                update_progress({
                    "current": idx,
                    "status": f"Scanning {filename}"
                })

                # Scan document
                highlighted_path, report_rows = bias_scanner.scan_docx(
                    path, 
                    term_category_map, 
                    word_out_dir
                )
                all_report_rows.extend(report_rows)

            # Generate Excel report
            excel_path = os.path.join(temp_dir, "bias_report.xlsx")
            bias_scanner.write_excel(all_report_rows, excel_path)

            # Create ZIP file
            zip_path = os.path.join(temp_dir, "bias_scan_output.zip")
            bias_scanner.create_zip(word_out_dir, excel_path, zip_path)

            # Clean up temporary PDF files
            bias_scanner.cleanup_pdf_files()

            # Register download token
            token = uuid.uuid4().hex
            download_tokens[token] = {
                "path": temp_dir,
                "expires": _now_utc() + TOKEN_TTL,
                "user": username,
                "route_type": "bias_scan"
            }

            # DB logging
            with db_pool.get_connection() as db:
                db.execute(
                    '''INSERT INTO macro_processing
                       (user_id, token, original_filenames, processed_filenames, selected_tasks, route_type)
                       VALUES (?, ?, ?, ?, ?, ?)''',
                    (
                        user_id,
                        token,
                        json.dumps(original_filenames),
                        json.dumps(["bias_scan_output.zip"]),
                        json.dumps({"route_type": "bias_scan", "terms_found": len(all_report_rows)}),
                        "bias_scan"
                    )
                )
                db.commit()

            update_progress({
                "status": "Completed",
                "download_token": token,
                "zip_path": zip_path,
                "terms_found": len(all_report_rows)
            })

        except Exception as e:
            update_progress({"status": f"Failed: {e}"})
            traceback.print_exc()

@app.route("/bias-scan", methods=["GET", "POST"])
@csrf.exempt
@role_required(ROUTE_PERMISSIONS.get('bias_scan', ['ADMIN']))
def bias_scan():
    """Bias scanner route handler"""
    if request.method == "POST":
        files = request.files.getlist("files")

        if not files or all(f.filename == "" for f in files):
            return jsonify({"error": "No files selected"}), 400

        # Validate file types
        for f in files:
            if f.filename and not allowed_file(f.filename):
                return jsonify({"error": f"Invalid file type: {f.filename}. Only .doc and .docx files are allowed."}), 400

        # Use token as job_id for consistency
        token = uuid.uuid4().hex
        job_id = token 
        
        # Save files synchronously before threading
        temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], token)
        os.makedirs(temp_dir, exist_ok=True)

        saved_paths = []
        original_filenames = []
        
        try:
            for f in files:
                if f.filename:
                    safe_name = secure_filename(f.filename) or f"document_{len(saved_paths)}.docx"
                    path = os.path.join(temp_dir, safe_name)
                    f.save(path)
                    saved_paths.append(path)
                    original_filenames.append(f.filename)
        except Exception as e:
            try:
                shutil.rmtree(temp_dir)
            except:
                pass
            return jsonify({"error": f"File save failed: {e}"}), 500

        # In-memory update for same-worker immediate feedback (optional)
        app.config.setdefault("PROGRESS_DATA", {})
        app.config["PROGRESS_DATA"][job_id] = {
            "total": len(saved_paths),
            "current": 0,
            "status": "Starting"
        }

        batch_queue.submit(
            job_id=job_id,
            route_type='bias_scan',
            user_id=session['user_id'],
            username=session['username'],
            target_fn=process_bias_scan_job,
            fn_args=(job_id, temp_dir, saved_paths, original_filenames,
                     session['user_id'], session['username']),
            payload_dict={
                'temp_dir': temp_dir,
                'saved_paths': saved_paths,
                'original_filenames': original_filenames,
            }
        )

        return jsonify({"job_id": job_id})

    return render_template("upload_bias.html")


def process_word_to_xml_job(job_id, temp_dir, file_paths, original_filenames, user_id, username):
    """
    Process Word to XML conversion job using Perl script.
    Creates ZIP file containing original Word files and generated XML files.
    """
    with app.app_context():
        import subprocess
        
        # Helper to update progress file
        def update_progress(updates):
            try:
                p_path = os.path.join(temp_dir, "progress.json")
                current = {}
                if os.path.exists(p_path):
                    with open(p_path, "r") as f:
                        current = json.load(f)
                current.update(updates)
                with open(p_path, "w") as f:
                    json.dump(current, f)
            except Exception as ex:
                print(f"Progress update failed: {ex}")

        # Initialize progress file
        update_progress({
            "total": len(file_paths),
            "current": 0,
            "status": "Starting conversion",
            "folder": temp_dir
        })

        try:
            # Path to Word to XML tools
            wordtoxml_dir = os.path.join(BASE_DIR, "wordtoxml")
            perl_script = os.path.join(wordtoxml_dir, "Word2XML_Books.pl")
            
            # Check if Perl script exists
            if not os.path.exists(perl_script):
                update_progress({"status": f"Failed: Perl script not found at {perl_script}"})
                return

            # Update progress
            update_progress({
                "current": 0,
                "status": f"Converting {len(file_paths)} file(s) to XML"
            })

            # Execute Perl script with temp_dir as argument
            # The Perl script expects a directory path and processes all .docx files in it
            try:
                result = subprocess.run(
                    ["perl", perl_script, temp_dir],
                    cwd=wordtoxml_dir,
                    capture_output=True,
                    text=True,
                    timeout=300  # 5 minute timeout
                )
                
                if result.returncode != 0:
                    error_msg = result.stderr or "Unknown error during conversion"
                    update_progress({"status": f"Failed: {error_msg}"})
                    return
                    
            except subprocess.TimeoutExpired:
                update_progress({"status": "Failed: Conversion timeout (5 minutes)"})
                return
            except Exception as e:
                update_progress({"status": f"Failed: {str(e)}"})
                return

            # Check if XML files were created in html subdirectory
            html_dir = os.path.join(temp_dir, "html")
            if not os.path.exists(html_dir):
                update_progress({"status": "Failed: No XML output generated"})
                return

            # Find generated XML files
            xml_files = [f for f in os.listdir(html_dir) if f.endswith('.xml')]
            
            if not xml_files:
                update_progress({"status": "Failed: No XML files found in output"})
                return

            # Save execution log - REMOVED per user request ("no need merge logs")
            # The perl script generates its own .log files in the html folder
            
            update_progress({
                "current": len(file_paths),
                "status": "Creating ZIP file"
            })

            # Create ZIP file containing ONLY XML and Log files
            zip_path = os.path.join(temp_dir, "word_to_xml_output.zip")
            with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
                # Add generated XML and LOG files from html directory
                if os.path.exists(html_dir):
                    for filename in os.listdir(html_dir):
                        if filename.lower().endswith(('.xml', '.log')):
                            file_path = os.path.join(html_dir, filename)
                            z.write(file_path, arcname=filename) # Flattened structure

            # Register download token
            token = uuid.uuid4().hex
            download_tokens[token] = {
                "path": temp_dir,
                "expires": _now_utc() + TOKEN_TTL,
                "user": username,
                "route_type": "word_to_xml"
            }

            # DB logging
            with db_pool.get_connection() as db:
                db.execute(
                    '''INSERT INTO macro_processing
                       (user_id, token, original_filenames, processed_filenames, selected_tasks, route_type)
                       VALUES (?, ?, ?, ?, ?, ?)''',
                    (
                        user_id,
                        token,
                        json.dumps(original_filenames),
                        json.dumps(["word_to_xml_output.zip"]),
                        json.dumps({"route_type": "word_to_xml", "xml_files": xml_files}),
                        "word_to_xml"
                    )
                )
                db.commit()

            update_progress({
                "status": "Completed",
                "download_token": token,
                "zip_path": zip_path,
                "xml_files_count": len(xml_files)
            })

        except Exception as e:
            update_progress({"status": f"Failed: {str(e)}"})
            traceback.print_exc()

@app.route("/word-to-xml", methods=["GET", "POST"])
@csrf.exempt
@role_required(ROUTE_PERMISSIONS.get('word_to_xml', ['ADMIN']))
def word_to_xml():
    """Word to XML conversion route handler"""
    if request.method == "POST":
        files = request.files.getlist("files")

        if not files or all(f.filename == "" for f in files):
            return jsonify({"error": "No files selected"}), 400

        # Validate file types
        for f in files:
            if f.filename and not allowed_file(f.filename):
                return jsonify({"error": f"Invalid file type: {f.filename}. Only .doc and .docx files are allowed."}), 400

        # Use token as job_id for consistency
        token = uuid.uuid4().hex
        job_id = token 
        
        # Save files synchronously before threading
        temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], token)
        os.makedirs(temp_dir, exist_ok=True)

        saved_paths = []
        original_filenames = []
        
        try:
            for f in files:
                if f.filename:
                    safe_name = secure_filename(f.filename) or f"document_{len(saved_paths)}.docx"
                    path = os.path.join(temp_dir, safe_name)
                    f.save(path)
                    saved_paths.append(path)
                    original_filenames.append(f.filename)
        except Exception as e:
            try:
                shutil.rmtree(temp_dir)
            except:
                pass
            return jsonify({"error": f"File save failed: {e}"}), 500

        # In-memory update for same-worker immediate feedback (optional)
        app.config.setdefault("PROGRESS_DATA", {})
        app.config["PROGRESS_DATA"][job_id] = {
            "total": len(saved_paths),
            "current": 0,
            "status": "Starting"
        }

        batch_queue.submit(
            job_id=job_id,
            route_type='word_to_xml',
            user_id=session['user_id'],
            username=session['username'],
            target_fn=process_word_to_xml_job,
            fn_args=(job_id, temp_dir, saved_paths, original_filenames,
                     session['user_id'], session['username']),
            payload_dict={
                'temp_dir': temp_dir,
                'saved_paths': saved_paths,
                'original_filenames': original_filenames,
            }
        )

        return jsonify({"job_id": job_id})

    return render_template("upload_word_to_xml.html")


@app.route("/progress/<job_id>", methods=["GET"])

def get_progress(job_id):
    """
    Get progress status for a background job.
    Reads from progress.json file in the job's temp directory.
    """
    try:
        # Always prefer file-based progress (updated by background thread)
        temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], job_id)
        progress_file = os.path.join(temp_dir, "progress.json")

        if os.path.exists(progress_file):
            with open(progress_file, "r") as f:
                progress_data = json.load(f)
            return jsonify(progress_data)

        # Fall back to in-memory (only useful before file is written)
        if job_id in app.config.get("PROGRESS_DATA", {}):
            return jsonify(app.config["PROGRESS_DATA"][job_id])

        return jsonify({"status": "Not found", "total": 0, "current": 0}), 404
    except Exception as e:
        return jsonify({"status": f"Error: {e}", "total": 0, "current": 0}), 500


# -----------------------
# Database Connection Pool (SQLite + Postgres)
# -----------------------
class PostgresWrapper:
    def __init__(self, conn, pool_ref):
        self.conn = conn
        self.pool_ref = pool_ref
        self.cursor = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        self.is_postgres = True

    def execute(self, query, params=None):
        # Translate SQLite ? placeholders to Postgres %s
        # This is a basic regex replace; for complex queries with literal ? it might be risky, 
        # but for this app's simple queries it's sufficient.
        pg_query = query.replace('?', '%s')
        
        try:
            if params:
                self.cursor.execute(pg_query, params)
            else:
                self.cursor.execute(pg_query)
        except Exception:
            self.conn.rollback()
            raise
        return self.cursor

    def fetchone(self):
        return self.cursor.fetchone()

    def fetchall(self):
        return self.cursor.fetchall()

    def commit(self):
        self.conn.commit()

    def rollback(self):
        self.conn.rollback()

    def close(self):
        self.cursor.close()
        # Return connection to pool
        if self.pool_ref:
            self.pool_ref.putconn(self.conn)

class SQLiteWrapper:
    """Wrapper for SQLite connections to provide consistent interface with PostgreSQL"""
    def __init__(self, conn, pool_ref):
        self.conn = conn
        self.pool_ref = pool_ref
        self.is_postgres = False

    def execute(self, query, params=None):
        try:
            if params:
                return self.conn.execute(query, params)
            else:
                return self.conn.execute(query)
        except Exception:
            self.conn.rollback()
            raise

    def fetchone(self):
        # SQLite cursor is returned by execute
        pass

    def fetchall(self):
        # SQLite cursor is returned by execute
        pass

    def commit(self):
        self.conn.commit()

    def rollback(self):
        self.conn.rollback()

    def close(self):
        # Return connection to pool
        if self.pool_ref:
            try:
                self.pool_ref.put(self.conn, block=False)
            except:
                self.conn.close()
        else:
            self.conn.close()

class DatabasePool:
    def __init__(self, database_path, pool_size=5):
        self.database_path = database_path
        self.lock = threading.Lock()
        
        self.db_host = os.environ.get('DB_HOST')
        self.db_name = os.environ.get('DB_NAME')
        self.db_user = os.environ.get('DB_USER')
        self.db_pass = os.environ.get('DB_PASSWORD')
        self.db_port = os.environ.get('DB_PORT', '5432')
        
        self.use_postgres = (self.db_host is not None and psycopg2 is not None)
        
        if self.use_postgres:
            # Initialize Postgres Pool
            self.pg_pool = psycopg2.pool.ThreadedConnectionPool(
                minconn=1,
                maxconn=pool_size,
                host=self.db_host,
                database=self.db_name,
                user=self.db_user,
                password=self.db_pass,
                port=self.db_port
            )
            print(f"✅ Connected to PostgreSQL at {self.db_host}")
        else:
            # Initialize SQLite Pool
            self.pool = Queue(maxsize=pool_size)
            for _ in range(pool_size):
                conn = sqlite3.connect(database_path, check_same_thread=False)
                conn.row_factory = sqlite3.Row
                conn.execute("PRAGMA journal_mode=WAL")
                conn.execute("PRAGMA synchronous=NORMAL")
                # cache_size pragma
                self.pool.put(conn)
            print(f"✅ using SQLite at {database_path}")

    @contextmanager
    def get_connection(self):
        if self.use_postgres:
            conn = self.pg_pool.getconn()
            wrapper = PostgresWrapper(conn, self.pg_pool)
            try:
                yield wrapper
            except Exception:
                wrapper.rollback()
                raise
            finally:
                # Wrapper.close() puts it back, but let's ensure safety
                # Ideally the user calls wrapper.close()? No, the context manager should handle it.
                # Actually, our wrapper.close() puts it back. 
                # But typical context manager usage in this app is `with get_connection() as db: ...`
                # So we should close/put back here.
                wrapper.close()
        else:
            conn = None
            wrapper = None
            try:
                conn = self.pool.get(timeout=5)
                wrapper = SQLiteWrapper(conn, self.pool)
                yield wrapper
            except Empty:
                conn = sqlite3.connect(self.database_path, check_same_thread=False)
                conn.row_factory = sqlite3.Row
                wrapper = SQLiteWrapper(conn, None)
                yield wrapper
            except Exception:
                if wrapper:
                    wrapper.rollback()
                raise
            finally:
                if wrapper:
                    wrapper.close()


db_pool = DatabasePool(DATABASE)


# -----------------------
# Batch Queue Manager
# -----------------------
_ROUTE_JOB_FUNCTIONS = {
    'credit_extractor': 'process_credit_extractor_job',
    'bias_scan':        'process_bias_scan_job',
    'word_to_xml':      'process_word_to_xml_job',
    'ppd':              'process_ppd_job',
    'validation':       'process_validation_job',
    'technical':        'process_technical_job',
}


class BatchQueueManager:
    MAX_WORKERS = 4

    def __init__(self, db_pool_ref, app_ref):
        self._db = db_pool_ref
        self._app = app_ref
        self._executor = concurrent.futures.ThreadPoolExecutor(
            max_workers=self.MAX_WORKERS, thread_name_prefix="batchworker"
        )
        self._active_futures = {}   # job_id → Future (None while slot is being reserved)
        self._lock = Lock()

    def submit(self, job_id, route_type, user_id, username,
               target_fn, fn_args, payload_dict, priority=0):
        """Enqueue a job. Runs immediately if a worker slot is free, else persists as pending."""
        try:
            with self._db.get_connection() as db:
                db.execute(
                    "INSERT INTO job_queue "
                    "(job_id, route_type, user_id, username, status, priority, payload) "
                    "VALUES (?, ?, ?, ?, 'pending', ?, ?)",
                    (job_id, route_type, user_id, username, priority, json.dumps(payload_dict))
                )
                db.commit()
        except Exception as e:
            print(f"BatchQueue submit DB error for {job_id}: {e}")
        self._try_dispatch(job_id, target_fn, fn_args)

    def _try_dispatch(self, job_id, target_fn, fn_args):
        """Dispatch a job to the executor if a worker slot is available."""
        with self._lock:
            if len(self._active_futures) >= self.MAX_WORKERS:
                return
            self._active_futures[job_id] = None  # Reserve slot immediately

        self._mark_running(job_id)
        future = self._executor.submit(target_fn, *fn_args)
        future.add_done_callback(lambda f, jid=job_id: self._on_complete(jid, f))
        with self._lock:
            self._active_futures[job_id] = future

    def _on_complete(self, job_id, future):
        """Called by executor when a job finishes (success, failure, or cancel)."""
        with self._lock:
            self._active_futures.pop(job_id, None)

        if future.cancelled():
            self._mark_done(job_id, 'cancelled')
        elif future.exception():
            self._mark_done(job_id, 'failed', str(future.exception()))
        else:
            self._mark_done(job_id, 'completed')

        self._drain_pending()

    def _drain_pending(self):
        """Promote pending DB jobs into the executor when slots are free."""
        with self._lock:
            slots_free = self.MAX_WORKERS - len(self._active_futures)
        if slots_free <= 0:
            return

        try:
            with self._db.get_connection() as db:
                rows = db.execute(
                    "SELECT job_id, route_type, user_id, username, payload "
                    "FROM job_queue WHERE status='pending' "
                    "ORDER BY priority DESC, queued_at ASC LIMIT ?",
                    (slots_free,)
                ).fetchall()
        except Exception as e:
            print(f"BatchQueue _drain_pending DB error: {e}")
            return

        for row in rows:
            try:
                jid      = row['job_id']
                rtype    = row['route_type']
                uid      = row['user_id']
                uname    = row['username']
                payload  = json.loads(row['payload'])

                fn_name  = _ROUTE_JOB_FUNCTIONS.get(rtype)
                if not fn_name:
                    continue
                target_fn = globals().get(fn_name)
                if not target_fn:
                    continue
                fn_args = self._reconstruct_fn_args(rtype, jid, uid, uname, payload)

                with self._lock:
                    if len(self._active_futures) >= self.MAX_WORKERS:
                        break
                    self._active_futures[jid] = None

                self._mark_running(jid)
                future = self._executor.submit(target_fn, *fn_args)
                future.add_done_callback(lambda f, j=jid: self._on_complete(j, f))
                with self._lock:
                    self._active_futures[jid] = future

            except Exception as e:
                print(f"BatchQueue drain error for job {jid}: {e}")

    def _reconstruct_fn_args(self, route_type, job_id, user_id, username, payload):
        """Reconstruct fn_args tuple from stored payload for crash recovery."""
        p = payload
        if route_type == 'credit_extractor':
            return (job_id, p['temp_dir'], p['saved_paths'], p['original_filenames'],
                    user_id, username, p.get('extraction_method', 'manual'), p.get('api_key', ''))
        elif route_type == 'bias_scan':
            return (job_id, p['temp_dir'], p['saved_paths'], p['original_filenames'], user_id, username)
        elif route_type == 'word_to_xml':
            return (job_id, p['temp_dir'], p['saved_paths'], p['original_filenames'], user_id, username)
        elif route_type == 'ppd':
            return (job_id, p['unique_folder'], p['saved'], p['combined_dashboard'],
                    p['book_title'], p['safe_title'], username, user_id)
        elif route_type == 'validation':
            return (job_id, p['processing_dir'], p['saved_paths'], p['original_filenames'],
                    p['options'], user_id, username)
        elif route_type == 'technical':
            return (job_id, p['unique_folder'], p['saved_paths'], p['original_filenames'],
                    p['run_te'], user_id, username)
        else:
            raise ValueError(f"Unknown route_type for reconstruction: {route_type}")

    def cancel(self, job_id, requesting_user_id):
        """Cancel a pending or running job. Returns (success, message)."""
        try:
            with self._db.get_connection() as db:
                row = db.execute(
                    "SELECT user_id, status FROM job_queue WHERE job_id=?", (job_id,)
                ).fetchone()
        except Exception as e:
            return False, f"DB error: {e}"

        if not row:
            return False, "Job not found"

        if int(row['user_id']) != int(requesting_user_id):
            return False, "Forbidden"

        status = row['status']
        if status == 'pending':
            try:
                with self._db.get_connection() as db:
                    db.execute(
                        "UPDATE job_queue SET status='cancelled', completed_at=? WHERE job_id=?",
                        (_now_utc(), job_id)
                    )
                    db.commit()
            except Exception as e:
                return False, f"DB error: {e}"
            return True, "Cancelled"

        if status == 'running':
            with self._lock:
                future = self._active_futures.get(job_id)
            if future:
                future.cancel()
            try:
                with self._db.get_connection() as db:
                    db.execute(
                        "UPDATE job_queue SET status='cancelled', completed_at=? WHERE job_id=?",
                        (_now_utc(), job_id)
                    )
                    db.commit()
            except Exception as e:
                return False, f"DB error: {e}"
            return True, "Cancel requested (job may already be running)"

        return False, f"Cannot cancel job in status '{status}'"

    def list_jobs(self, user_id, is_admin=False, limit=100):
        """List jobs for a user (or all jobs if admin)."""
        try:
            with self._db.get_connection() as db:
                if is_admin:
                    rows = db.execute(
                        "SELECT job_id, route_type, username, status, priority, payload, "
                        "queued_at, started_at, completed_at, downloaded_at, error_msg "
                        "FROM job_queue ORDER BY queued_at DESC LIMIT ?",
                        (limit,)
                    ).fetchall()
                else:
                    rows = db.execute(
                        "SELECT job_id, route_type, username, status, priority, payload, "
                        "queued_at, started_at, completed_at, downloaded_at, error_msg "
                        "FROM job_queue WHERE user_id=? ORDER BY queued_at DESC LIMIT ?",
                        (user_id, limit)
                    ).fetchall()
            result = []
            for row in rows:
                d = dict(row)
                try:
                    p = json.loads(d.get('payload', '{}'))
                    d['files'] = p.get('original_filenames', p.get('saved_paths', []))
                    if not d['files'] and 'saved' in p:
                        d['files'] = [os.path.basename(f) for f in p.get('saved', [])]
                except Exception:
                    d['files'] = []
                d.pop('payload', None)
                for k in ['queued_at', 'started_at', 'completed_at', 'downloaded_at']:
                    if d.get(k) and not isinstance(d[k], str):
                        d[k] = str(d[k])
                result.append(d)
            return result
        except Exception as e:
            print(f"BatchQueue list_jobs error: {e}")
            return []

    def _mark_running(self, job_id):
        try:
            with self._db.get_connection() as db:
                db.execute(
                    "UPDATE job_queue SET status='running', started_at=? WHERE job_id=?",
                    (_now_utc(), job_id)
                )
                db.commit()
        except Exception as e:
            print(f"BatchQueue _mark_running error for {job_id}: {e}")

    def _mark_done(self, job_id, status, error_msg=None):
        try:
            with self._db.get_connection() as db:
                db.execute(
                    "UPDATE job_queue SET status=?, completed_at=?, error_msg=? WHERE job_id=?",
                    (status, _now_utc(), error_msg, job_id)
                )
                db.commit()
        except Exception as e:
            print(f"BatchQueue _mark_done error for {job_id}: {e}")


batch_queue = BatchQueueManager(db_pool, app)


# -----------------------
# Word Processor (Disabled for Linux)
# -----------------------
class OptimizedDocumentProcessor:
    def __init__(self):
        pass

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        pass

    def process_documents_batch(self, file_paths, selected_tasks, route_type):
        return ["Word automation (macros) is not supported in this Linux environment."]


# -----------------------
# Utility Functions
# -----------------------
def get_ip_address():
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        ip_address = s.getsockname()[0]
        s.close()
        return ip_address
    except Exception:
        return "127.0.0.1"


def log_activity(username, action, details=""):
    with open(LOG_FILE, "a", encoding="utf-8") as f:
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        f.write(f"{timestamp} - {username} - {action} - {details}\n")


def log_errors(error_list):
    with open(LOG_FILE, "a", encoding="utf-8") as log_file:
        for err in error_list:
            log_file.write(f"{datetime.now().isoformat()} - ERROR - {err}\n")


def allowed_file(filename):
    return any(filename.lower().endswith(ext) for ext in ALLOWED_EXTENSIONS)


def setup_logging():
    if not app.debug:
        file_handler = RotatingFileHandler('logs/s4c.log', maxBytes=10240000, backupCount=10)
        file_handler.setFormatter(logging.Formatter(
            '%(asctime)s %(levelname)s: %(message)s [in %(pathname)s:%(lineno)d]'
        ))
        file_handler.setLevel(logging.INFO)
        app.logger.addHandler(file_handler)
        app.logger.setLevel(logging.INFO)


def cleanup_expired_tokens():
    current_time = _now_utc()
    expired_tokens = []

    for token, data in list(download_tokens.items()):
        expires = data.get("expires")
        if not expires:
            expired_tokens.append(token)
            continue

        expires = _ensure_utc(expires)

        if current_time > expires:
            expired_tokens.append(token)

    for token in expired_tokens:
        try:
            info = download_tokens.get(token)
            if not info:
                continue

            path = info.get("path")
            if path and os.path.exists(path):
                shutil.rmtree(path, ignore_errors=True)

            log_activity(
                info.get("user", "system"),
                f"TOKEN_EXPIRED_{info.get('route_type', 'UNKNOWN').upper()}",
                token[:8]
            )

            download_tokens.pop(token, None)

        except Exception as e:
            log_errors([f"Token cleanup failed ({token}): {e}"])



def kill_word_processes():
    pass


def save_uploaded_file(file, folder):
    try:
        filename = secure_filename(file.filename)
        file_path = os.path.join(folder, filename)

        with open(file_path, 'wb') as f:
            file.save(f)

        return file_path, None
    except Exception as e:
        return None, str(e)


# -----------------------
# Template Filters
# -----------------------
@app.template_filter('from_json')
def from_json_filter(value):
    try:
        return json.loads(value)
    except (ValueError, TypeError):
        return value


@app.template_filter('format_date')
def format_date_filter(value):
    try:
        if isinstance(value, str):
            dt = datetime.strptime(value, '%Y-%m-%d %H:%M:%S')
        else:
            dt = value
        return dt.strftime('%b %d, %Y %I:%M %p')
    except (ValueError, AttributeError):
        return value


# -----------------------
# Database Functions
# -----------------------
def get_db():
    return db_pool.get_connection()


def init_db():
    with app.app_context():
        with db_pool.get_connection() as db:
            is_postgres = getattr(db, 'is_postgres', False)
            pk_type = "SERIAL PRIMARY KEY" if is_postgres else "INTEGER PRIMARY KEY AUTOINCREMENT"
            
            # Helper to create table safely
            def create_table_safe(query):
                try:
                    db.execute(query)
                    db.commit()
                except Exception as e:
                    # If it's a "relation already exists" or unique violation (race condition), ignore.
                    # Postgres error code 42P07 is duplicate_table, but we catch generic Exception here for simplicity with SQLite too
                    # "UniqueViolation" seen in logs was "duplicate key value violates unique constraint" on a system catalog
                    # which is odd for "CREATE TABLE", but suggests a race on pg_type.
                    # Just logging and moving on is safest for init logic.
                    if "already exists" in str(e) or "UniqueViolation" in str(e):
                        db.rollback()
                    else:
                        raise e

            # Create tables
            create_table_safe(f'''CREATE TABLE IF NOT EXISTS users (
                            id {pk_type},
                            username TEXT UNIQUE NOT NULL,
                            password TEXT NOT NULL,
                            email TEXT,
                            is_admin BOOLEAN DEFAULT FALSE,
                            role TEXT DEFAULT 'USER',
                            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')

            create_table_safe(f'''CREATE TABLE IF NOT EXISTS files (
                            id {pk_type},
                            user_id INTEGER NOT NULL,
                            original_filename TEXT NOT NULL,
                            stored_filename TEXT NOT NULL,
                            report_filename TEXT,
                            upload_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                            FOREIGN KEY (user_id) REFERENCES users(id))''')

            create_table_safe(f'''CREATE TABLE IF NOT EXISTS validation_results (
                            id {pk_type},
                            file_id INTEGER NOT NULL,
                            total_references INTEGER,
                            total_citations INTEGER,
                            missing_references TEXT,
                            unused_references TEXT,
                            sequence_issues TEXT,
                            FOREIGN KEY (file_id) REFERENCES files(id))''')

            create_table_safe(f'''CREATE TABLE IF NOT EXISTS macro_processing (
                            id {pk_type},
                            user_id INTEGER NOT NULL,
                            token TEXT UNIQUE NOT NULL,
                            original_filenames TEXT NOT NULL,
                            processed_filenames TEXT NOT NULL,
                            selected_tasks TEXT NOT NULL,
                            processing_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                            errors TEXT,
                            route_type TEXT DEFAULT 'general',
                            FOREIGN KEY (user_id) REFERENCES users(id))''')

            create_table_safe(f'''CREATE TABLE IF NOT EXISTS job_queue (
                            id           {pk_type},
                            job_id       TEXT UNIQUE NOT NULL,
                            route_type   TEXT NOT NULL,
                            user_id      INTEGER NOT NULL,
                            username     TEXT NOT NULL,
                            status       TEXT NOT NULL DEFAULT 'pending',
                            priority     INTEGER NOT NULL DEFAULT 0,
                            payload      TEXT NOT NULL,
                            queued_at    TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                            started_at   TIMESTAMP,
                            completed_at TIMESTAMP,
                            downloaded_at TIMESTAMP,
                            error_msg    TEXT,
                            FOREIGN KEY (user_id) REFERENCES users(id))''')

            # --- MIGRATION: Check and Add 'route_type' to macro_processing if missing ---
            try:
                # Check for column existence (SQLite/Postgres agnostic check)
                try:
                    # Try selecting the column. If it fails, it doesn't exist.
                    db.execute("SELECT route_type FROM macro_processing LIMIT 1")
                except Exception:
                     # Column missing, add it
                     print("Migrating: Adding 'route_type' column to macro_processing table...")
                     db.rollback() # Clear error state for Postgres
                     alter_query = "ALTER TABLE macro_processing ADD COLUMN route_type TEXT DEFAULT 'general'"
                     db.execute(alter_query)
                     db.commit()
                     print("Migration successful.")
            except Exception as e:
                print(f"Migration check failed (might already exist or other error): {e}")
                db.rollback()

            # Create indexes for performance
            try:
                db.execute("CREATE INDEX IF NOT EXISTS idx_files_user_id ON files(user_id)")
                db.execute("CREATE INDEX IF NOT EXISTS idx_files_upload_date ON files(upload_date)")
                db.execute("CREATE INDEX IF NOT EXISTS idx_macro_user_id ON macro_processing(user_id)")
                db.execute("CREATE INDEX IF NOT EXISTS idx_macro_route_type ON macro_processing(route_type)")
                db.execute("CREATE INDEX IF NOT EXISTS idx_jq_status ON job_queue(status)")
                db.execute("CREATE INDEX IF NOT EXISTS idx_jq_user_id ON job_queue(user_id)")
            except Exception as e:
                # ignore specific index errors or just log
                print(f"Index creation warning: {e}")

            # Create default admin safely
            try:
                # Check existence first
                admin_user = db.execute("SELECT * FROM users WHERE username=%s" if is_postgres else "SELECT * FROM users WHERE username=?", ('admin',)).fetchone()
                if not admin_user:
                    hashed_password = generate_password_hash("admin123", method='pbkdf2:sha256')
                    query = "INSERT INTO users (username,password,email,is_admin) VALUES (%s,%s,%s,%s)" if is_postgres else "INSERT INTO users (username,password,email,is_admin) VALUES (?,?,?,?)"
                    db.execute(query, ('admin', hashed_password, 'admin@example.com', True))
                    db.commit()
            except Exception as e:
                # If race condition causes unique violation on insert, ignore
                if "UniqueViolation" in str(e) or "UNIQUE constraint failed" in str(e):
                    db.rollback()
                else:
                    print(f"Error creating admin user: {e}")

def migrate_add_role_column():
    """Ensure the 'role' column exists for legacy DBs."""
    try:
        with db_pool.get_connection() as db:
            is_postgres = getattr(db, 'is_postgres', False)
            
            cols = []
            if is_postgres:
                # Postgres check
                cur = db.execute("SELECT column_name FROM information_schema.columns WHERE table_name = 'users'")
                cols = [r[0] for r in cur.fetchall()]
            else:
                # SQLite check
                cur = db.execute("PRAGMA table_info(users)")
                cols = [r["name"] for r in cur.fetchall()]

            if "role" not in cols:
                db.execute("ALTER TABLE users ADD COLUMN role TEXT DEFAULT 'USER'")
                db.commit()
                app.logger.info("Added 'role' column to users table")
    except Exception as e:
        log_errors([f"Migration error adding role column: {e}"])

@app.context_processor
def inject_current_role():
    return {'current_role': get_user_role()}
# -----------------------
# Enhanced Reference Validator
# -----------------------
# -----------------------
# Enhanced Reference Validator (Logic from Referencenumvalidation.py)
# -----------------------

def iter_document_paragraphs(doc):
    """
    Iterate through all paragraphs in the document body in order,
    including those inside tables.
    """
    body = doc._element.body
    for child in body:
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            table = Table(child, doc)
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p


def get_numbers(text):
    """
    Extract numbers from text like '1', '2-5', '1, 3, 5'.
    Handles ranges "1-5" -> [1, 2, 3, 4, 5].
    """
    nums = []
    # Matches: (start)-(end) OR (single)
    # Allows hyphen, en dash, em dash
    pattern = re.compile(r'(\d+)\s*[-–—]\s*(\d+)|(\d+)')
    
    for start, end, single in pattern.findall(text):
        if start and end:
            try:
                s, e = int(start), int(end)
                if s <= e:
                    nums.extend(range(s, e + 1))
            except ValueError:
                pass
        elif single:
            try:
                nums.append(int(single))
            except ValueError:
                pass
    return nums


def format_numbers(nums):
    """
    Format a list of numbers into a string like '1-3, 5'.
    Collapses ranges of 3 or more (e.g. 1,2,3 -> 1-3).
    """
    nums = sorted(set(nums))
    if not nums:
        return ""

    parts = []
    if not nums:
        return ""

    start = prev = nums[0]

    for n in nums[1:]:
        if n == prev + 1:
            prev = n
        else:
            length = prev - start + 1
            if length >= 3:
                parts.append(f"{start}-{prev}")
            elif length == 2:
                parts.append(f"{start},{prev}")
            else:
                parts.append(str(start))
            start = prev = n

    length = prev - start + 1
    if length >= 3:
        parts.append(f"{start}-{prev}")
    elif length == 2:
        parts.append(f"{start},{prev}")
    else:
        parts.append(str(start))

    return ", ".join(parts)


def is_citation_run(run):
    """
    Determine if a run is part of a citation.
    Checks for 'cite_bib' style OR superscript with number-like content.
    """
    if run.style and run.style.name == "cite_bib":
        return True
    if run.font.superscript:
        text = run.text.strip()
        if not text:
            return False
        # Must look like numbers/ranges/separators
        if re.match(r'^[\d,\-–—\s]+$', text):
            return True
    return False


class ReferenceProcessor:
    def __init__(self, doc):
        self.doc = doc
        
    def get_references_in_bibliography(self):
        """
        Returns a Set of IDs found in the bibliography sections (REF-N style).
        Also returns a list of objects for reordering later.
        """
        refs_found = set()
        ref_objects = [] # list of dicts: {'id': int, 'para': p, 'run': r}

        for para in self.doc.paragraphs:
            if para.style and para.style.name == "REF-N":
                found_id = None
                bib_run = None
                
                # Try finding styled run
                for run in para.runs:
                    if run.style and run.style.name == "bib_number":
                        nums = get_numbers(run.text)
                        if nums:
                            found_id = nums[0]
                            bib_run = run
                            break
                            
                # Fallback: Check start of text if no styled run
                if found_id is None:
                    match = re.match(r'^(\d+)', para.text.strip())
                    if match:
                        found_id = int(match.group(1))
                
                if found_id is not None:
                    refs_found.add(found_id)
                    ref_objects.append({
                        'id': found_id,
                        'para': para,
                        'run': bib_run
                    })
                    
        return refs_found, ref_objects

    def get_citations_in_text(self):
        """
        Scans document for citations.
        Returns:
            all_cited_ids: list of all IDs in order of appearance (with duplicates)
            appearance_order: list of unique IDs in order of first appearance
        """
        all_cited_ids = []
        appearance_order = []
        seen = set()
        
        # Regex for fallback pattern ^1-3^
        citation_pattern = re.compile(r'\^([\d,\-–—\s]+)\^')

        for para in iter_document_paragraphs(self.doc):
            # 1. Process runs
            current_group = []
            
            for run in para.runs:
                if is_citation_run(run):
                    current_group.append(run)
                else:
                    if current_group:
                        # Flush group
                        text = "".join(r.text for r in current_group)
                        nums = get_numbers(text)
                        all_cited_ids.extend(nums)
                        for n in nums:
                            if n not in seen:
                                seen.add(n)
                                appearance_order.append(n)
                        current_group = []
                    
                    # Check fallback pattern in non-citation run
                    matches = citation_pattern.findall(run.text)
                    for m in matches:
                        nums = get_numbers(m)
                        all_cited_ids.extend(nums)
                        for n in nums:
                            if n not in seen:
                                seen.add(n)
                                appearance_order.append(n)
            
            # Flush trailing group
            if current_group:
                text = "".join(r.text for r in current_group)
                nums = get_numbers(text)
                all_cited_ids.extend(nums)
                for n in nums:
                    if n not in seen:
                        seen.add(n)
                        appearance_order.append(n)
                        
        return all_cited_ids, appearance_order

    def find_duplicates(self, ref_objects):
        """
        Finds duplicate references using fuzzy matching (difflib).
        Returns a list of dicts: {'id': int, 'text': str, 'duplicate_of': int, 'score': float}
        """
        duplicates = []
        processed_refs = [] # list of (id, clean_text)
        
        # 1. Pre-process all candidates
        for obj in ref_objects:
            full_text = obj['para'].text.strip()
            # Remove leading numbering like "1. ", "[1] "
            clean_text = re.sub(r'^\[?\d+\]?[\.\s]*', '', full_text)
            processed_refs.append({'id': obj['id'], 'text': clean_text})
            
        # 2. Compare O(N^2)
        # We only check forward to avoid double reporting (A=B, B=A)
        # We assume the *earlier* ID is the "original" and later is "duplicate"
        n = len(processed_refs)
        for i in range(n):
            ref_a = processed_refs[i]
            
            # Skip if strict duplicate logic already caught it? 
            # No, let's just do fuzzy for all.
            
            for j in range(i + 1, n):
                ref_b = processed_refs[j]
                
                # Metric: similarity ratio
                # Quick check: length difference shouldn't be too huge
                len_a = len(ref_a['text'])
                len_b = len(ref_b['text'])
                if len_a == 0 or len_b == 0: 
                    continue
                    
                # Optimization: Length ratio check
                if min(len_a, len_b) / max(len_a, len_b) < 0.6:
                    continue
                    
                ratio = difflib.SequenceMatcher(None, ref_a['text'], ref_b['text']).ratio()
                
                # Threshold: 0.85 (85% similar)
                # The user's example is extremely similar, probably > 90%
                if ratio > 0.85:
                    duplicates.append({
                        'id': ref_b['id'], # The later one is the duplicate
                        'text': ref_b['text'][:100] + "...",
                        'duplicate_of': ref_a['id'],
                        'score': round(ratio * 100, 1)
                    })
                    
        return duplicates

    def get_validation_stats(self):
        bib_refs, ref_objects = self.get_references_in_bibliography()
        all_cited, _ = self.get_citations_in_text()
        
        unique_cited = set(all_cited)
        
        # Missing: Cited but not in Bib
        missing = sorted(unique_cited - bib_refs)
        
        # Unused: In Bib but not Cited
        unused = sorted(bib_refs - unique_cited)
        
        # Duplicates
        duplicates = self.find_duplicates(ref_objects)
        
        # Sequence Issues
        sequence_issues = []
        seen_in_seq = []
        previous_max = 0
        
        for n in all_cited:
            if n not in seen_in_seq:
                if n < previous_max:
                     pass
                
                if n != len(seen_in_seq) + 1:
                     sequence_issues.append({
                         "position": len(seen_in_seq) + 1,
                         "current": n,
                         "expected": len(seen_in_seq) + 1
                     })
                
                seen_in_seq.append(n)
                previous_max = max(previous_max, n)
                
        return {
            "total_references": len(bib_refs),
            "total_citations": len(all_cited),
            "missing_references": missing,
            "unused_references": unused,
            "duplicate_references": duplicates,
            "sequence_issues": sequence_issues,
            "is_perfect": (not missing and not unused and not sequence_issues and not duplicates)
        }

    def renumber(self):
        """
        Renumber citations and reorder bibliography.
        Returns: mapping (Old -> New)
        """
        _, appearance_order = self.get_citations_in_text()
        
        # Create Mapping
        mapping = {} 
        new_id = 1
        for old_id in appearance_order:
            mapping[old_id] = new_id
            new_id += 1
            
        # 1. Update Citations in Text
        citation_pattern = re.compile(r'\^([\d,\-–—\s]+)\^')
        
        for para in iter_document_paragraphs(self.doc):
            current_group = []
            
            for run in para.runs:
                if is_citation_run(run):
                    current_group.append(run)
                else:
                    if current_group:
                        # Replace
                        text = "".join(r.text for r in current_group)
                        nums = get_numbers(text)
                        if nums:
                            new_nums = [mapping.get(n, n) for n in nums]
                            new_text = format_numbers(new_nums)
                            current_group[0].text = new_text
                            for r in current_group[1:]:
                                r.text = ""
                        current_group = []
                    
                    # Pattern replacement
                    def replace_func(m):
                         nums = get_numbers(m.group(1))
                         new_nums = [mapping.get(n, n) for n in nums]
                         return "^" + format_numbers(new_nums) + "^"
                    
                    new_run_text = citation_pattern.sub(replace_func, run.text)
                    if new_run_text != run.text:
                        run.text = new_run_text

            if current_group:
                text = "".join(r.text for r in current_group)
                nums = get_numbers(text)
                if nums:
                    new_nums = [mapping.get(n, n) for n in nums]
                    new_text = format_numbers(new_nums)
                    current_group[0].text = new_text
                    for r in current_group[1:]:
                        r.text = ""

        # 2. Reorder Bibliography
        _, ref_objects = self.get_references_in_bibliography()
        
        # Sort objects into Cited and Uncited
        cited_refs = []
        uncited_refs = []
        
        for obj in ref_objects:
            if obj['id'] in mapping:
                obj['new_id'] = mapping[obj['id']]
                cited_refs.append(obj)
            else:
                uncited_refs.append(obj)
        
        if not ref_objects:
            return mapping

        # Find anchor (min index)
        body = self.doc._element.body
        
        indices = []
        for obj in ref_objects:
            try:
                idx = body.index(obj['para']._element)
                indices.append(idx)
            except ValueError:
                pass 
        
        if not indices:
            return mapping
            
        anchor = min(indices)
        
        # Remove all
        for obj in ref_objects:
             p = obj['para']._element
             if p.getparent() == body:
                 body.remove(p)
                 
        # Insert Cited (Sorted)
        cited_refs.sort(key=lambda x: x['new_id'])
        
        insert_idx = anchor
        for obj in cited_refs:
            # Update ID text
            if obj['run']:
                obj['run'].text = str(obj['new_id'])
            
            body.insert(insert_idx, obj['para']._element)
            insert_idx += 1
            
        # Insert Uncited (Appended after cited)
        for obj in uncited_refs:
            body.insert(insert_idx, obj['para']._element)
            insert_idx += 1
            
        return mapping


def process_document(file_path):
    doc = Document(file_path)
    processor = ReferenceProcessor(doc)
    
    # Check BEFORE
    before_stats = processor.get_validation_stats()
    
    # DECISION:
    # 1. If Unused References exist -> ABORT renumbering.
    if before_stats["unused_references"]:
        return doc, before_stats, before_stats, {}, "Failed: Document validation failed due to unused references."

    # 2. If Perfect -> No need.
    if before_stats["is_perfect"]:
        return doc, before_stats, before_stats, {}, "Validation completed."
        
    # 3. If Missing Refs -> Can't safely renumber usually
    if before_stats["missing_references"]:
         return doc, before_stats, before_stats, {}, "Failed: Missing references detected."

    # DO RENUMBER
    mapping = processor.renumber()
    
    # Check AFTER (Validate result)
    after_stats = processor.get_validation_stats()
    
    # Determine status message
    changes_made = False
    if mapping:
        for k, v in mapping.items():
            if k != v:
                changes_made = True
                break

    if before_stats["duplicate_references"]:
        count = len(before_stats['duplicate_references'])
        prefix = "Renumbering" if changes_made else "Validation"
        status_msg = f"{prefix} completed with {count} duplicate{'s' if count > 1 else ''}."
    elif changes_made:
        status_msg = "Renumbering completed successfully."
    else:
        status_msg = "Validation completed." # Fallback if no changes and no duplicates but not 'perfect' initially (e.g. sequence issues resolved to identity?)

    return doc, before_stats, after_stats, mapping, status_msg



# -----------------------
# Authentication (update load_logged_in_user)
# -----------------------
@app.before_request
def load_logged_in_user():
    user_id = session.get('user_id')
    if user_id is None:
        g.user = None
    else:
        with db_pool.get_connection() as db:
            user = db.execute("SELECT * FROM users WHERE id = ?", (user_id,)).fetchone()
            g.user = dict(user) if user else None
            if g.user:
                session['role'] = g.user.get('role', 'USER')



@app.before_request
def require_login():
    if request.endpoint in (
        'login', 'logout', 'static',
        'download_report',  # ✅ add this
        'register', 'reset_database',  # we'll secure this below
        'macro_download'
    ):
        return None
    if not session.get('user_id'):
        flash("Please log in to continue.")
        return redirect(url_for('login'))


def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('is_admin'):
            flash("Admin privileges required", "error")
            return redirect(url_for('dashboard'))
        return f(*args, **kwargs)

    return decorated_function

# -----------------------
# HTML to Excel (remove images)
# -----------------------
import pandas as pd
from bs4 import BeautifulSoup
import os
from pathlib import Path
from datetime import datetime
import chardet

import chardet  # at top of file with other imports

# -----------------------
# HTML to Excel (remove images)
# -----------------------
def html_to_excel_no_images(html_path, output_dir):
    """
    Converts an HTML file to an .xls file by removing <img> tags and writing
    the resulting HTML to a .xls file so Excel can open it.
    Returns the output file path or None on failure.
    """
    try:
        # read raw bytes and detect encoding
        with open(html_path, "rb") as f:
            raw_data = f.read()

        encoding = None
        try:
            detected = chardet.detect(raw_data)
            encoding = detected.get("encoding") or "utf-8"
        except Exception:
            encoding = "utf-8"

        try:
            html_content = raw_data.decode(encoding, errors="ignore")
        except Exception:
            html_content = raw_data.decode("utf-8", errors="ignore")

        # Remove <img> tags (handles attributes and self-closing)
        html_no_images = re.sub(r"<img\b[^>]*>", "", html_content, flags=re.IGNORECASE)

        # Also remove inline base64 images in style attributes (background-image:url(data:...))
        html_no_images = re.sub(r'url\(\s*data:[^)]+\)', 'url()', html_no_images, flags=re.IGNORECASE)

        # Build a safe output filename
        base = Path(html_path).stem
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_file = os.path.join(output_dir, f"{base}_{timestamp}.xls")

        with open(output_file, "w", encoding="utf-8") as f:
            f.write(html_no_images)

        return output_file
    except Exception as e:
        log_errors([f"HTML to Excel conversion failed for {html_path}: {e}"])
        return None


# -----------------------
# Real openpyxl Excel Report Builder
# -----------------------
def build_excel_report(chapters_data: list, output_path: str) -> bool:
    """
    Build a real .xlsx file from chapters_data.
    One sheet — all chapters stacked vertically, colour-coded sections.
    Returns True on success, False on failure.
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font as XFont, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        from bs4 import BeautifulSoup

        NUM_COLS = 7

        # ── Style helpers ────────────────────────────────────────
        def _fill(hex_c):
            return PatternFill("solid", fgColor=hex_c)

        def _font(bold=False, color="000000", size=9, italic=False):
            return XFont(bold=bold, color=color, size=size, italic=italic, name="Segoe UI")

        def _border():
            s = Side(style="thin", color="DDDDDD")
            return Border(left=s, right=s, top=s, bottom=s)

        def _centre():
            return Alignment(horizontal="center", vertical="center", wrap_text=True)

        def _left():
            return Alignment(horizontal="left", vertical="center", wrap_text=True)

        def _write(ws, row, col, value, bold=False, fg=None, fc="000000",
                   sz=9, align=None, italic=False, bdr=True):
            c = ws.cell(row=row, column=col, value=value)
            c.font = _font(bold=bold, color=fc, size=sz, italic=italic)
            c.alignment = align or _left()
            if fg:
                c.fill = _fill(fg)
            if bdr:
                c.border = _border()
            return c

        def _section(ws, row, label, bg="1E2235", cols=NUM_COLS):
            ws.row_dimensions[row].height = 20
            _write(ws, row, 1, label, bold=True, fg=bg, fc="FFFFFF", sz=10,
                   align=_centre(), bdr=True)
            ws.merge_cells(start_row=row, start_column=1,
                           end_row=row, end_column=cols)

        def _thead(ws, row, labels, bg="4361EE", cols=NUM_COLS):
            ws.row_dimensions[row].height = 18
            for i, lbl in enumerate(labels[:cols], 1):
                _write(ws, row, i, lbl, bold=True, fg=bg, fc="FFFFFF",
                       sz=9, align=_centre())
            for i in range(len(labels) + 1, cols + 1):
                _write(ws, row, i, "", fg=bg)

        def _drow(ws, row, values, bg="FFFFFF", cols=NUM_COLS):
            ws.row_dimensions[row].height = 16
            for i, v in enumerate(values[:cols], 1):
                _write(ws, row, i, str(v) if v is not None else "", fg=bg, sz=9)
            for i in range(len(values) + 1, cols + 1):
                _write(ws, row, i, "", fg=bg)

        def _blank(ws, row, h=6):
            ws.row_dimensions[row].height = h

        def _autofit_columns(ws, min_width=10, max_width=80):
            """Resize each column to fit its widest cell value."""
            from openpyxl.cell.cell import MergedCell
            col_max: dict = {}
            for row in ws.iter_rows():
                for cell in row:
                    if isinstance(cell, MergedCell) or cell.value is None:
                        continue
                    col = get_column_letter(cell.column)
                    # Account for newlines inside a cell
                    width = max(len(line) for line in str(cell.value).splitlines()) if cell.value else 0
                    col_max[col] = max(col_max.get(col, min_width), width)
            for col, width in col_max.items():
                ws.column_dimensions[col].width = min(width + 2, max_width)

        def _parse_table(html_str):
            """Extract list-of-lists from first <table> in html_str."""
            if not html_str:
                return []
            soup = BeautifulSoup(html_str, "html.parser")
            rows = []
            for tr in soup.find_all("tr"):
                cells = [td.get_text(" ", strip=True) for td in tr.find_all(["td", "th"])]
                if any(cells):
                    rows.append(cells)
            return rows

        # ── Workbook setup ────────────────────────────────────────
        wb = Workbook()
        wb.remove(wb.active)  # Remove default sheet

        # Helper to safely sum numeric values
        def safe_sum(key):
            total = 0
            for ch in chapters_data:
                val = ch.get(key, 0)
                try:
                    if isinstance(val, str):
                        val = val.replace(',', '')
                    total += float(val)
                except (ValueError, TypeError):
                    pass
            return round(total)

        # ── 1. Optional: Main Summary Sheet (if multiple chapters) ─
        if len(chapters_data) > 1:
            ws_summary = wb.create_sheet(title="Overall Summary")
            sum_cols = 6
            col_widths_sum = [15, 45, 15, 15, 18, 18]
            for i, w in enumerate(col_widths_sum, 1):
                ws_summary.column_dimensions[get_column_letter(i)].width = w
            
            rs = 1
            ws_summary.row_dimensions[rs].height = 30
            _write(ws_summary, rs, 1, "S4Carlisle  —  Combined Manuscript Summary",
                   bold=True, fg="1E2235", fc="FFFFFF", sz=13, align=_centre(), bdr=False)
            ws_summary.merge_cells(start_row=rs, start_column=1, end_row=rs, end_column=sum_cols)
            rs += 1; _blank(ws_summary, rs); rs += 1

            # Overall Metrics
            _section(ws_summary, rs, "  OVERALL METRICS", "2C3E7A", cols=sum_cols); rs += 1
            metrics = [
                ("Total Chapters Processed", len(chapters_data)),
                ("Total Word Count", safe_sum("total_words")),
                ("Total CE MS Pages", safe_sum("ce_pages")),
            ]
            for lbl, val in metrics:
                ws_summary.row_dimensions[rs].height = 16
                _write(ws_summary, rs, 1, lbl, bold=True, fg="EEF2FF", sz=9)
                _write(ws_summary, rs, 2, str(val), fg="FFFFFF", sz=9)
                ws_summary.merge_cells(start_row=rs, start_column=2, end_row=rs, end_column=sum_cols)
                rs += 1
            _blank(ws_summary, rs); rs += 1

            # Chapter Breakdown Table
            _section(ws_summary, rs, "  CHAPTER BREAKDOWN", "1A5276", cols=sum_cols); rs += 1
            headers = ["Chapter No.", "File Name", "Total Words", "CE Pages", "Missing Cits", "Fmt Issues"]
            _thead(ws_summary, rs, headers, bg="1A5276", cols=sum_cols); rs += 1
            
            for i, ch in enumerate(chapters_data):
                row_vals = [
                    ch.get("chapter_number", "—"),
                    ch.get("doc_name", "—"),
                    ch.get("total_words", "—"),
                    ch.get("ce_pages", "—"),
                    ch.get("missing_citations", "—"),
                    ch.get("fmt_issues", "—")
                ]
                bg = "F0F2F8" if i % 2 == 0 else "FFFFFF"
                _drow(ws_summary, rs, row_vals, bg=bg, cols=sum_cols); rs += 1
            
            ws_summary.freeze_panes = "A2"
            _autofit_columns(ws_summary)

        # ── 2. Loop over chapters for Individual Sheets ────────────
        COLORS = {
            "chapter":    "1E2235",
            "info":       "2C3E7A",
            "wc":         "1A5276",
            "summary":    "1E2235",
            "citations":  "2C3E7A",
            "special":    "6C3483",
            "fmt":        "1A5276",
            "comments":   "1B6B43",
            "unnumbered": "784212",
        }

        import re
        from pathlib import Path

        for idx, ch in enumerate(chapters_data):
            # Generate a safe Excel sheet name
            doc_name = str(ch.get("doc_name", f"Chapter_{idx+1}"))
            safe_title = re.sub(r'[\\*?:/\[\]]', '', Path(doc_name).stem).strip()
            if not safe_title: safe_title = f"Chapter_{idx+1}"
            if len(safe_title) > 31: safe_title = safe_title[:15] + ".." + safe_title[-14:]
            
            orig_title = safe_title
            counter = 1
            while safe_title in wb.sheetnames:
                suffix = f"_{counter}"
                safe_title = orig_title[:31 - len(suffix)] + suffix
                counter += 1

            ws = wb.create_sheet(title=safe_title)

            col_widths = [30, 16, 16, 16, 16, 18, 44]
            for i, w in enumerate(col_widths, 1):
                ws.column_dimensions[get_column_letter(i)].width = w

            r = 1

            # ── Top banner ────────────────────────────────────────────
            ws.row_dimensions[r].height = 30
            _write(ws, r, 1, "S4Carlisle  —  Manuscript Analysis Dashboard",
                   bold=True, fg="1E2235", fc="FFFFFF", sz=13, align=_centre(), bdr=False)
            ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=NUM_COLS)
            r += 1; _blank(ws, r); r += 1

            # ── Chapter divider banner ────────────────────────────
            ch_label = f"{ch.get('chapter_number','—')}  —  {ch.get('chapter_title','—')}"
            ws.row_dimensions[r].height = 26
            _write(ws, r, 1, ch_label, bold=True, fg="4361EE", fc="FFFFFF",
                   sz=12, align=_centre(), bdr=False)
            ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=NUM_COLS)
            r += 1

            # ── 1. Chapter Info ───────────────────────────────────
            _section(ws, r, "  CHAPTER INFORMATION", COLORS["info"]); r += 1
            meta = [
                ("Chapter No.",   ch.get("chapter_number", "—")),
                ("Chapter Title", ch.get("chapter_title",  "—")),
                ("Author(s)",     ch.get("authors",        "—")),
                ("File",          ch.get("doc_name",       "—")),
                ("Date",          ch.get("date",           "—")),
                ("Analyst",       ch.get("analyst",        "—")),
            ]
            for lbl, val in meta:
                ws.row_dimensions[r].height = 16
                _write(ws, r, 1, lbl, bold=True, fg="EEF2FF", sz=9)
                _write(ws, r, 2, val, fg="FFFFFF", sz=9)
                ws.merge_cells(start_row=r, start_column=2,
                               end_row=r, end_column=NUM_COLS)
                r += 1
            _blank(ws, r); r += 1

            # ── 2. Word Count & Pages ─────────────────────────────
            _section(ws, r, "  WORD COUNT & PAGES", COLORS["wc"]); r += 1
            _thead(ws, r, ["Metric", "Value", "", "", "", "", ""],
                   bg=COLORS["wc"]); r += 1
            wc_rows = [
                ("Total Word Count",     ch.get("total_words", "—")),
                ("Total MS Pages",       ch.get("total_pages", "—")),
                ("CE MS Pages",          ch.get("ce_pages",    "—")),
                ("Main Text Page Count", ch.get("pages",       "—")),
            ]
            for i, (lbl, val) in enumerate(wc_rows):
                bg = "F0F2F8" if i % 2 == 0 else "FFFFFF"
                ws.row_dimensions[r].height = 16
                _write(ws, r, 1, lbl, fg=bg, sz=9)
                _write(ws, r, 2, str(val), bold=True, fg="FFFFFF", sz=10)
                for c in range(3, NUM_COLS + 1):
                    _write(ws, r, c, "", fg="FFFFFF")
                r += 1
            _blank(ws, r); r += 1

            # ── 3. Analysis Summary ───────────────────────────────
            _section(ws, r, "  ANALYSIS SUMMARY", COLORS["summary"]); r += 1
            sum_rows = _parse_table(ch.get("detailed_summary", ""))
            if sum_rows:
                _thead(ws, r, sum_rows[0][:NUM_COLS]); r += 1
                for i, row_vals in enumerate(sum_rows[1:]):
                    bg = "F0F2F8" if i % 2 == 0 else "FFFFFF"
                    _drow(ws, r, row_vals, bg=bg); r += 1
            else:
                _drow(ws, r, ["No summary data available"], bg="FFFFFF"); r += 1
            _blank(ws, r); r += 1

            # ── 4. Citations / Captions Detail ────────────────────
            _section(ws, r, "  CITATIONS & CAPTIONS DETAIL", COLORS["citations"]); r += 1
            cit_rows = _parse_table(ch.get("msr_content", ""))
            if cit_rows:
                _thead(ws, r, cit_rows[0][:NUM_COLS], bg=COLORS["citations"]); r += 1
                for i, row_vals in enumerate(cit_rows[1:]):
                    text = " ".join(str(v) for v in row_vals).lower()
                    if "missing" in text and "citation" in text:
                        bg = "FFF8E1"
                    elif "missing" in text and "caption" in text:
                        bg = "FDE8F0"
                    else:
                        bg = "F0F2F8" if i % 2 == 0 else "FFFFFF"
                    _drow(ws, r, row_vals, bg=bg); r += 1
            else:
                _drow(ws, r, ["No citation data available"], bg="FFFFFF"); r += 1
            _blank(ws, r); r += 1

            # ── 5. Special Characters ─────────────────────────────
            _section(ws, r, "  SPECIAL CHARACTERS", COLORS["special"]); r += 1
            spec_rows = _parse_table(ch.get("spec_content", ""))
            if spec_rows:
                _thead(ws, r, spec_rows[0][:NUM_COLS], bg=COLORS["special"]); r += 1
                for i, row_vals in enumerate(spec_rows[1:]):
                    _drow(ws, r, row_vals,
                          bg="F0F2F8" if i % 2 == 0 else "FFFFFF"); r += 1
            else:
                _drow(ws, r, ["No special characters found"], bg="FFFFFF"); r += 1
            _blank(ws, r); r += 1

            # ── 6. Formatting Issues ──────────────────────────────
            _section(ws, r, "  FORMATTING ISSUES", COLORS["fmt"]); r += 1
            fmt_rows = _parse_table(ch.get("fmt_content", ""))
            if fmt_rows:
                _thead(ws, r, fmt_rows[0][:NUM_COLS], bg=COLORS["fmt"]); r += 1
                for i, row_vals in enumerate(fmt_rows[1:]):
                    _drow(ws, r, row_vals,
                          bg="FFF3E0" if i % 2 == 0 else "FFFFFF"); r += 1
            else:
                _drow(ws, r, ["No formatting issues found"], bg="FFFFFF"); r += 1
            _blank(ws, r); r += 1

            # ── 7. Comments ───────────────────────────────────────
            _section(ws, r, "  COMMENTS", COLORS["comments"]); r += 1
            com_rows = _parse_table(ch.get("comment_content", ""))
            if com_rows:
                _thead(ws, r, com_rows[0][:NUM_COLS], bg=COLORS["comments"]); r += 1
                for i, row_vals in enumerate(com_rows[1:]):
                    _drow(ws, r, row_vals,
                          bg="F0F2F8" if i % 2 == 0 else "FFFFFF"); r += 1
            else:
                _drow(ws, r, ["No comments found"], bg="FFFFFF"); r += 1
            _blank(ws, r); r += 1

            # ── 8. Unnumbered Elements ────────────────────────────
            _section(ws, r, "  UNNUMBERED ELEMENTS", COLORS["unnumbered"]); r += 1
            unn_rows = _parse_table(ch.get("unnumbered_content", ""))
            if unn_rows:
                _thead(ws, r, unn_rows[0][:NUM_COLS], bg=COLORS["unnumbered"]); r += 1
                for i, row_vals in enumerate(unn_rows[1:]):
                    _drow(ws, r, row_vals,
                          bg="FFF3E0" if i % 2 == 0 else "FFFFFF"); r += 1
            else:
                _drow(ws, r, ["No unnumbered elements found"], bg="FFFFFF"); r += 1

            # ── Chapter separator gap ─────────────────────────────
            r += 1; _blank(ws, r, h=12); r += 1

            # ── Footer ────────────────────────────────────────────────
            ws.row_dimensions[r].height = 16
            _write(ws, r, 1, "Generated by S4Carlisle Manuscript Analysis Tool",
                   italic=True, fg="1E2235", fc="FFFFFF", sz=8,
                   align=_centre(), bdr=False)
            ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=NUM_COLS)

            ws.freeze_panes = "A2"
            _autofit_columns(ws)
            ws.print_area = f"A1:{get_column_letter(NUM_COLS)}{r}"
            ws.page_setup.orientation = "landscape"
            ws.page_setup.fitToWidth = 1
            ws.page_setup.fitToHeight = 0

        if not wb.sheetnames:
            wb.create_sheet(title="Dashboard")

        wb.save(output_path)
        return True

    except Exception as e:
        log_errors([f"build_excel_report failed: {e}"])
        return False


# -----------------------
# Generic Route Handler
# -----------------------
def _process_macro_request(route_type):
    # Stub function since automation is removed
    flash("This feature relies on Microsoft Word automation and is not available.")
    return redirect(url_for(route_type))

# -----------------------
# Routes
# -----------------------
@app.route('/', strict_slashes=False)
def index():
    if 'user_id' in session:
        return redirect(url_for('dashboard'))
    else:
        return redirect(url_for('login'))


@app.route('/login', methods=['GET', 'POST'], strict_slashes=False)
def login():
    if 'user_id' in session:
        return redirect(url_for('dashboard'))

    if request.method == "POST":
        username = request.form['username']
        password = request.form['password']

        try:
            with db_pool.get_connection() as db:
                # Explicitly handle Postgres vs SQLite syntax if wrapper's replace isn't trusted,
                # though wrapper should handle it. Use ? for consistency with wrapper logic.
                user = db.execute("SELECT id, username, password, is_admin FROM users WHERE username=?",
                                  (username,)).fetchone()

                if user:
                    stored_hash = user['password']
                    # Legacy hash compatibility check - unlikely needed for new setups but kept just in case
                    if stored_hash.startswith('$'):
                        stored_hash = stored_hash[1:]

                    if check_password_hash(stored_hash, password):
                        session['user_id'] = user['id']
                        session['username'] = user['username']
                        session['is_admin'] = bool(user['is_admin'])
                        # Handle Postgres RealDictRow vs SQLite Row access if needed, coverage:
                        # user['role'] might be missing from SELECT above? 
                        # Wait, login doesn't select role! 
                        # We should probably select role to cache it in session? 
                        # Existing code didn't, but `get_user_role` checks session['role'] OR g.user.
                        # Let's add role to session for performance.
                        
                        log_activity(username, "LOGIN")
                        return redirect(url_for('dashboard'))
                    else:
                        print(f"Login failed for {username}: Password mismatch")
                else:
                    print(f"Login failed for {username}: User not found")

        except Exception as e:
            print(f"Login error: {e}")
            log_errors([f"Login Exception: {e}"])

        flash("Invalid username or password", "error")

    return render_template('login.html')


@app.route("/register", methods=["GET", "POST"], strict_slashes=False)
def register():
    if request.method == "POST":
        username = request.form['username']
        password = request.form['password']
        email = request.form.get('email', '')

        with db_pool.get_connection() as db:
            try:
                hashed = generate_password_hash(password, method='pbkdf2:sha256')
                db.execute("INSERT INTO users (username,password,email) VALUES (?,?,?)",
                           (username, hashed, email))
                db.commit()
                flash("Registration successful", "success")
                return redirect(url_for('login'))
            except sqlite3.IntegrityError:
                db.rollback()
                flash("Username/email already exists", "error")

    return render_template("register.html")


@app.route('/logout', strict_slashes=False)
def logout():
    user = session.get('username')
    if user:
        log_activity(user, "LOGOUT")
    session.clear()
    flash("Logged out successfully.")
    return redirect(url_for('login'))

def handle_macro_route(route_type, template_name):
    if 'user_id' not in session:
        flash("Please log in to continue.")
        return redirect(url_for('login'))

    # Linux Compatibility - Disable Macro POSTs
    if request.method == 'POST':
        flash("This feature relies on Microsoft Word automation and is not available on Linux servers.", "error")
        # Can also return here if we don't want to attempt processing even if code stub exists
        # return render_template(template_name, ...)
        
    # We still allow rendering the page so users can see the UI, but actions fail nicely.
    # However, since we removed the logic, calling _process_macro_request (which we'll keep as stub) is safe-ish.
    if request.method == 'POST':
         # In a real scenario we might just block it.
         pass
         
    # Stub config for rendering (names are fine, execution is disabled)
    download_token = request.args.get('download_token')
    route_config = ROUTE_MACROS.get(route_type, {})

    return render_template(template_name,
                           download_token=download_token,
                           route_config=route_config,
                           macro_names=route_config.get('macros', []))
# -----------------------
# Routes (patched with role_required)
# -----------------------
@app.route('/language', methods=['GET', 'POST'], strict_slashes=False)
@role_required(ROUTE_PERMISSIONS.get('language', ['ADMIN']))
def language():
    return handle_macro_route('language', 'language_edit.html')

@app.route('/macro_processing', methods=['GET', 'POST'])
@role_required(ROUTE_PERMISSIONS.get('macro_processing', ['ADMIN']))
def macro_processing():
    return handle_macro_route('macro_processing', 'macro_processing.html')

from jinja2 import Template

def process_book_indexer_job(job_id, temp_dir, saved_paths, api_key, model_name, output_filename, user_id, username):
    """
    Background worker for generating a book index using Gemini.
    Zips the resulting DOCX so it works with the existing batch queue downloader.
    """
    with app.app_context():
        import pdfplumber
        from docx import Document
        from docx.shared import Inches
        from google import genai
        from google.genai import types
        import book_indexer_core
        import zipfile
        
        def update_progress(status, pct):
            try:
                p_path = os.path.join(temp_dir, "progress.json")
                current = {}
                if os.path.exists(p_path):
                    with open(p_path, "r") as f:
                        current = json.load(f)
                current["status"] = status
                current["progress"] = pct
                with open(p_path, "w") as f:
                    json.dump(current, f)
                if job_id in app.config.get("PROGRESS_DATA", {}):
                    app.config["PROGRESS_DATA"][job_id]["status"] = status
            except Exception as e:
                logging.error(f"Error updating progress: {e}")

        try:
            update_progress("Extracting text...", 10)
            pages_text = []
            global_page = 1
            for path in saved_paths:
                with pdfplumber.open(path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text and text.strip():
                            pages_text.append(f"[PAGE {global_page}]\n{text.strip()}")
                        global_page += 1
            
            CHUNK_SIZE = 30
            chunks = []
            for i in range(0, len(pages_text), CHUNK_SIZE):
                chunks.append("\n\n".join(pages_text[i:i + CHUNK_SIZE]))

            client = genai.Client(api_key=api_key)
            merged_index = {}
            api_warnings = []
            
            update_progress("Querying Gemini AI...", 30)
            for chunk_idx, chunk_text in enumerate(chunks):
                pct = 30 + int((chunk_idx / max(1, len(chunks))) * 50)
                update_progress(f"Processing PDF part {chunk_idx + 1}/{len(chunks)}", pct)
                
                prompt = book_indexer_core.PROMPT_TEMPLATE.format(text=chunk_text)
                response = client.models.generate_content(
                    model=model_name,
                    contents=prompt,
                    config=types.GenerateContentConfig(
                        temperature=0.15,
                        max_output_tokens=65536
                    )
                )

                partial_text = ""
                try:
                    if response.candidates:
                        candidate = response.candidates[0]
                        finish = getattr(candidate, 'finish_reason', None)
                        finish_str = str(finish)
                        if finish_str in ('1', 'STOP', 'FinishReason.STOP',
                                          '2', 'MAX_TOKENS', 'FinishReason.MAX_TOKENS'):
                            partial_text = response.text or ""
                            if finish_str in ('2', 'MAX_TOKENS', 'FinishReason.MAX_TOKENS'):
                                api_warnings.append(f"PDF part {chunk_idx + 1}: output truncated")
                        else:
                            api_warnings.append(f"PDF part {chunk_idx + 1}: skipped finish_reason={finish}")
                    else:
                        api_warnings.append(f"PDF part {chunk_idx + 1}: blocked")
                except Exception as resp_err:
                    api_warnings.append(f"PDF part {chunk_idx + 1}: error {resp_err}")

                if partial_text:
                    cleaned_text = book_indexer_core.clean_llm_response(partial_text)
                    book_indexer_core.parse_partial_index(cleaned_text, merged_index)

            if not merged_index:
                raise Exception("No index entries generated. " + "; ".join(api_warnings))

            update_progress("Generating DOCX...", 85)
            doc = Document()
            doc.add_heading('Index', 0)
            
            for term_key in sorted(merged_index.keys()):
                entry = merged_index[term_key]
                display = entry['display']
                pages_str = book_indexer_core.format_pages(entry['pages'])
                
                see = entry.get('see')
                see_also = entry.get('see_also', [])

                if see:
                    line = f"{display}. See {see}"
                else:
                    line = f"{display}, {pages_str}" if pages_str else display
                    if see_also:
                        line += f". See also {', '.join(see_also)}"
                doc.add_paragraph(line)

                for sub_key in sorted(entry['sub'].keys()):
                    sub_entry = entry['sub'][sub_key]
                    sub_display = sub_entry['display']
                    sub_pages_str = book_indexer_core.format_pages(sub_entry['pages'])
                    sub_line = f"{sub_display}, {sub_pages_str}" if sub_pages_str else sub_display
                    p = doc.add_paragraph(sub_line)
                    p.paragraph_format.left_indent = Inches(0.25)
            
            docx_path = os.path.join(temp_dir, output_filename)
            doc.save(docx_path)
            
            update_progress("Zipping output...", 95)
            zip_filename = "Book_Indexer_Result.zip"
            zip_path = os.path.join(temp_dir, zip_filename)
            
            with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
                zf.write(docx_path, arcname=output_filename)

            update_progress("Completed", 100)
            
            try:
                p_path = os.path.join(temp_dir, "progress.json")
                if os.path.exists(p_path):
                    with open(p_path, "r") as f:
                        final_p = json.load(f)
                    final_p["status"] = "Completed"
                    final_p["zip_path"] = zip_path
                    with open(p_path, "w") as f:
                        json.dump(final_p, f)
            except:
                pass

            return True

        except Exception as e:
            error_msg = f"Failed: {str(e)}"
            logging.error(error_msg)
            update_progress(error_msg, 0)
            try:
                p_path = os.path.join(temp_dir, "progress.json")
                if os.path.exists(p_path):
                    with open(p_path, "r") as f:
                        final_p = json.load(f)
                    final_p["status"] = error_msg
                    final_p["error"] = error_msg
                    with open(p_path, "w") as f:
                        json.dump(final_p, f)
            except:
                pass
            return False

def process_ppd_job(job_id, unique_folder, saved, combined_dashboard,
                    book_title, safe_title, username, user_id):
    with app.app_context():
        from word_analyzer_docx import (
            CitationAnalyzer,
            BoxTagLinker,
            build_element_mapping_html,
            extract_with_docx,
            remove_tags_keep_formatting_docx,
            generate_formatting_html,
            generate_multilingual_html,
            build_comments_html,
            build_export_highlight_html,
            build_unnumbered_tab_html,
            build_detailed_summary_table,
            build_combined_dashboard_html,
            build_new_combined_dashboard_html,
            extract_chapter_metadata,
            count_references_and_body_wc,
            count_total_words_via_txt,
            count_unnumbered_elements,
            extract_unnumbered_image_markups,
            build_unnumbered_image_markups_html,
            build_text_page_map,
            _page_from_map,
            get_xml_comments,
        )
        from docx import Document as _DocxDocument

        # Helper to update progress file
        def update_progress(updates):
            try:
                p_path = os.path.join(unique_folder, "progress.json")
                current = {}
                if os.path.exists(p_path):
                    with open(p_path, "r") as f:
                        current = json.load(f)
                current.update(updates)
                with open(p_path, "w") as f:
                    json.dump(current, f)
            except Exception as ex:
                print(f"Progress update failed: {ex}")

        per_chapter_files = []   # individual dashboard + individual Excel
        combined_files    = []   # combined dashboard + combined Excel
        chapters_data = []

        def _chapter_sort_key(raw):
            """Return integer for sorting; chapters with no digit sort last."""
            m = re.search(r'\d+', raw or '')
            return int(m.group()) if m else 2**31

        def _chapter_num_display(raw):
            """Normalize any chapter_number variant to 'Chapter N'.
            e.g. 'Chapter 1', 'chapter 2', 'CHAPTER 3', 'Ch. 4', '5' → 'Chapter N'."""
            m = re.search(r'\d+', raw or '')
            return f"Chapter {m.group()}" if m else raw

        # --- START LINUX BATCH PDF CONVERSION OPTIMIZATION ---
        # Pre-convert all uploaded docs to PDF simultaneously to avoid LibreOffice cold-starts
        import subprocess
        import shutil
        lo_cmd = shutil.which("libreoffice") or shutil.which("soffice")
        if not lo_cmd and os.name == "nt" and os.path.exists(r"C:\Program Files\LibreOffice\program\soffice.exe"):
            lo_cmd = r"C:\Program Files\LibreOffice\program\soffice.exe"
        
        if lo_cmd and saved:
            update_progress({"status": "Batch processing..."})
            try:
                cmd = [lo_cmd, "--headless", "--convert-to", "pdf", "--outdir", unique_folder] + [os.path.abspath(p) for p in saved]
                subprocess.run(cmd, timeout=300, capture_output=True)
            except Exception as e:
                app.logger.warning(f"Batch PDF conversion failed: {e}")
        # --- END LINUX BATCH PDF CONVERSION OPTIMIZATION ---

        for i, path in enumerate(saved, 1):
            fname = os.path.basename(path)
            update_progress({
                "current": i,
                "status": f"Processing {fname}"
            })

            try:
                # --- Step 1: extract paragraphs (fallback page numbers) ---
                paras, comments, imgs, foot, end = extract_with_docx(path)

                # --- Step 1b: compute dtypes and unnumbered counts from original file ---
                # Must happen BEFORE tag removal so inline markers like <UNFIG...> are present
                analyzer = CitationAnalyzer()
                doc_data = [(t, p, c) for (t, p, c, _) in paras]
                dtypes = analyzer.analyze_document_citations(doc_data)
                unnumbered_counts = count_unnumbered_elements(path, dtypes, paras=paras)

                # Extract image markup placeholders BEFORE tags are stripped
                img_markup_items = extract_unnumbered_image_markups(path)
                unnumbered_counts["image_placeholders"] = len(img_markup_items)

                # --- Step 1c: word counts on ORIGINAL file (before tag removal) ---
                # Must run here so <KP>, </KP>, <AU> etc. are still present in the file.
                # Word counts these tag tokens as words; our count must match.
                _txt_wc = count_total_words_via_txt(path)
                ref_count, body_wc, _raw_total_wc, ref_style = count_references_and_body_wc(path)

                # --- Step 2: strip inline tags from .docx on disk ---
                remove_tags_keep_formatting_docx(path)

                # --- Step 3: build PDF page map ONCE (LibreOffice/docx2pdf) ---
                pdf_total_pages, text_page_map = build_text_page_map(path)

                # Remap paragraph page numbers with real PDF values
                paras = [
                    (t, _page_from_map(text_page_map, t, p), c, h)
                    for (t, p, c, h) in paras
                ]

                # Rebuild dtypes with PDF page numbers now that paras is remapped.
                # paras still holds original text (inline markers intact from pre-tag-removal),
                # so analyze_document_citations produces correct labels with PDF pages.
                doc_data = [(t, p, c) for (t, p, c, _) in paras]
                dtypes = analyzer.analyze_document_citations(doc_data)

                # Update callout pages using the remapped paragraphs
                from word_analyzer_docx import _CALLOUT_RE, _PAGE_REF_RE
                callout_pages = []
                callouts_count = 0
                for t, p, c, h in paras:
                    if _CALLOUT_RE.search(t) or _PAGE_REF_RE.search(t):
                        callouts_count += 1
                        if p not in callout_pages:
                            callout_pages.append(p)
                unnumbered_counts["callouts"] = callouts_count
                unnumbered_counts["callout_pages"] = sorted(callout_pages)

                # --- Step 4: load Document once for all analysis passes ---
                _doc = _DocxDocument(path)

                table_count = len(dtypes.get("Table", {}).get("Caption", {}))

                # Re-fetch comments now that we have PDF page map and remapped paras
                comments = get_xml_comments(_doc, text_page_map=text_page_map, paras=paras)

                # --- Step 5: generate all HTML sections (reuse _doc) ---
                fmt_html = generate_formatting_html(
                    path, used_word=False, text_page_map=text_page_map, doc=_doc, paras=paras
                )
                # multilingual loads its own copy because it highlights+saves the doc
                spec_html = generate_multilingual_html(path, text_page_map=text_page_map)
                com_html = build_comments_html(comments)

                # Chapter metadata (uses cleaned _doc for title/author extraction)
                chapter_number, chapter_title, authors = extract_chapter_metadata(path, doc=_doc)

                # Word counts already computed on original file at Step 1c (before tag removal).
                # TXT-based total is primary (includes <KP> tokens, matches Word); para-scan fallback.
                total_wc = _txt_wc if _txt_wc > 0 else _raw_total_wc

                box_linker = BoxTagLinker(chapter_number=chapter_number)
                box_linker.scan(doc_data)
                box_linker.validate()

                summary_html, stats = build_detailed_summary_table(
                    dtypes, imgs, table_count, foot, end,
                    fmt_html, spec_html, com_html,
                    ref_count=ref_count,
                    ref_style=ref_style,
                    unnumbered_counts=unnumbered_counts,
                    chapter_number=chapter_number,
                    box_linker=box_linker,
                )
                msr_html = build_element_mapping_html(dtypes, "Figure",     chapter_number)
                msr_html += build_element_mapping_html(dtypes, "Table",      chapter_number)
                msr_html += build_element_mapping_html(dtypes, "Exhibit",    chapter_number)
                msr_html += build_element_mapping_html(dtypes, "Appendix",   chapter_number)
                msr_html += build_element_mapping_html(dtypes, "Case Study", chapter_number)
                msr_html += box_linker.build_html()
                exp_html = build_export_highlight_html(paras)
                unnumbered_html = build_unnumbered_tab_html(unnumbered_counts)
                if img_markup_items:
                    unnumbered_html += "<h3>Unnumbered Image Placeholders</h3>"
                    unnumbered_html += build_unnumbered_image_markups_html(img_markup_items)

                # Total words for fallback if something goes wrong
                para_words_only = sum(len(t.split()) for (t, _, _, _) in paras)

                # Total page count: PDF (LibreOffice) primary; LRPB max fallback; estimate last resort
                _lrpb_max_page = max((p for _, p, _, _ in paras), default=0) if paras else 0
                actual_total_pages = (pdf_total_pages    if pdf_total_pages > 0
                                      else _lrpb_max_page if _lrpb_max_page > 0
                                      else (len(paras) // 40) + 1)

                # Body Words & Body Pages
                final_body_wc = body_wc if body_wc > 0 else para_words_only
                body_pages = round(final_body_wc / 250) if final_body_wc > 0 else 1

                # CE Pages (Total document words / 250)
                final_total_wc = total_wc if total_wc > 0 else para_words_only
                ce_pages_val = round(final_total_wc / 250) if final_total_wc > 0 else 1

                # Render individual file dashboard using Jinja template
                _xl_link = f"{Path(path).stem}_Analysis.xlsx"
                single_chapter_data = [{
                    "doc_name":          fname,
                    "book_title":        book_title or "—",
                    "chapter_number":    chapter_number or "—",
                    "chapter_title":     chapter_title or "—",
                    "authors":           authors or "—",
                    "total_pages":       actual_total_pages,
                    "pages":             body_pages,
                    "words":             final_body_wc,
                    "total_words":       final_total_wc,
                    "ce_pages":          ce_pages_val,
                    "date":              _now_utc().strftime("%d-%m-%Y"),
                    "analyst":           username,
                    "detailed_summary":  summary_html,
                    "msr_content":       msr_html,
                    "fmt_content":       fmt_html,
                    "spec_content":      spec_html,
                    "comment_content":   com_html,
                    "export_highlight":  exp_html,
                    "unnumbered_content": unnumbered_html,
                    "missing_citations": stats.get("missing_citations", 0),
                    "missing_captions":  stats.get("missing_captions", 0),
                    "fmt_issues":        stats.get("fmt_issues", 0),
                    "fig_missing_cit":   stats.get("fig_missing_cit", 0),
                    "fig_missing_cap":   stats.get("fig_missing_cap", 0),
                    "tab_missing_cit":   stats.get("tab_missing_cit", 0),
                    "tab_missing_cap":   stats.get("tab_missing_cap", 0),
                    "box_missing_cit":   stats.get("box_missing_cit", 0),
                    "box_missing_cap":   stats.get("box_missing_cap", 0),
                    "total_citations":   stats.get("total_citations", 0),
                    "total_captions":    stats.get("total_captions", 0),
                    "fig_total_cit":     stats.get("fig_total_cit", 0),
                    "fig_total_cap":     stats.get("fig_total_cap", 0),
                    "tab_total_cit":     stats.get("tab_total_cit", 0),
                    "tab_total_cap":     stats.get("tab_total_cap", 0),
                    "box_total_cit":     stats.get("box_total_cit", 0),
                    "box_total_cap":     stats.get("box_total_cap", 0),
                    "excel_link":        _xl_link,
                }]
                html = build_combined_dashboard_html(
                    single_chapter_data,
                    css="", js="",
                    logo_b64=get_base64_logo()
                )

                out_html = os.path.join(unique_folder, Path(path).stem + "_Dashboard.html")
                with open(out_html, "w", encoding="utf-8") as f:
                    f.write(html)

                per_chapter_files.append(out_html)

                # Generate individual Excel report for this chapter alone
                single_xl_path = os.path.join(unique_folder, _xl_link)
                if build_excel_report(single_chapter_data, single_xl_path):
                    per_chapter_files.append(single_xl_path)
                else:
                    app.logger.error(f"Single Excel report failed for {fname}")

                # Collect chapter data for optional combined dashboard
                chapters_data.append({
                    "doc_name":          fname,
                    "chapter_number":    chapter_number or "—",
                    "chapter_number_display": _chapter_num_display(chapter_number or ""),
                    "chapter_title":     chapter_title or "—",
                    "authors":           authors or "—",
                    "total_pages":       actual_total_pages,
                    "pages":             body_pages,
                    "words":             final_body_wc,
                    "total_words":       final_total_wc,
                    "ce_pages":          ce_pages_val,
                    "date":              _now_utc().strftime("%d-%m-%Y"),
                    "analyst":           username,
                    "detailed_summary":  summary_html,
                    "msr_content":       msr_html,
                    "fmt_content":       fmt_html,
                    "spec_content":      spec_html,
                    "comment_content":   com_html,
                    "export_highlight":  exp_html,
                    "unnumbered_content": unnumbered_html,
                    "missing_citations": stats.get("missing_citations", 0),
                    "missing_captions":  stats.get("missing_captions", 0),
                    "fmt_issues":        stats.get("fmt_issues", 0),
                    "fig_missing_cit":   stats.get("fig_missing_cit", 0),
                    "fig_missing_cap":   stats.get("fig_missing_cap", 0),
                    "tab_missing_cit":   stats.get("tab_missing_cit", 0),
                    "tab_missing_cap":   stats.get("tab_missing_cap", 0),
                    "box_missing_cit":   stats.get("box_missing_cit", 0),
                    "box_missing_cap":   stats.get("box_missing_cap", 0),
                    "total_citations":   stats.get("total_citations", 0),
                    "total_captions":    stats.get("total_captions", 0),
                    "fig_total_cit":     stats.get("fig_total_cit", 0),
                    "fig_total_cap":     stats.get("fig_total_cap", 0),
                    "tab_total_cit":     stats.get("tab_total_cit", 0),
                    "tab_total_cap":     stats.get("tab_total_cap", 0),
                    "box_total_cit":     stats.get("box_total_cit", 0),
                    "box_total_cap":     stats.get("box_total_cap", 0),
                    "excel_link":        _xl_link,
                })

                # Excel is generated once after all chapters (see below)

            except Exception as e:
                app.logger.error(f"Failed processing {fname}: {e}")
                update_progress({"status": f"Failed: {e}"})
                break

        # Sort chapters into numerical order before combined dashboard
        chapters_data.sort(key=lambda ch: _chapter_sort_key(ch.get("chapter_number", "")))

        # Generate combined dashboard if requested and 2+ chapters processed
        if combined_dashboard and len(chapters_data) > 1:
            combined_html = build_new_combined_dashboard_html(
                chapters_data,
                logo_b64=get_base64_logo(),
                book_title=book_title,
            )
            combined_path = os.path.join(unique_folder, f"{safe_title}_Manuscript Analysis Report.html")
            with open(combined_path, "w", encoding="utf-8") as f:
                f.write(combined_html)
            combined_files.append(combined_path)

        # Generate combined Excel report (all chapters, single sheet)
        if chapters_data:
            xl_name = f"{safe_title}_Manuscript Analysis Report.xlsx"
            xl_path = os.path.join(unique_folder, xl_name)
            if build_excel_report(chapters_data, xl_path):
                combined_files.append(xl_path)
            else:
                log_errors(["Combined Excel report generation failed"])

        # Create ZIP — only include the relevant output set
        zip_path = os.path.join(unique_folder, f"{safe_title}_Manuscript_Analysis_Report.zip")
        if combined_dashboard and len(chapters_data) > 1:
            files_to_zip = combined_files   # Combined_Dashboard.html + Manuscript_Analysis.xlsx only
        else:
            files_to_zip = per_chapter_files  # {stem}_Dashboard.html + {stem}_Analysis.xlsx only
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
            for f in files_to_zip:
                if os.path.exists(f):
                    z.write(f, arcname=os.path.basename(f))

        update_progress({
            "status": "Completed",
            "current": len(saved),
            "zip_path": zip_path
        })

        # Job is already logged in job_queue table, no additional logging needed



@app.route("/ppd", methods=["GET", "POST"])
@csrf.exempt
@role_required(ROUTE_PERMISSIONS.get('ppd', ['ADMIN']))
def ppd():
    if request.method == "GET":
        return render_template("ppd.html")

    # -----------------------
    # Upload Handling
    # -----------------------
    uploaded = request.files.getlist("docfiles")
    if not uploaded:
        return jsonify({"error": "No files uploaded"}), 400

    # Unique job token folder
    token = uuid.uuid4().hex
    unique_folder = os.path.join(app.config['UPLOAD_FOLDER'], token)
    os.makedirs(unique_folder, exist_ok=True)

    # Register token immediately so it's downloadable from history
    download_tokens[token] = {
        'path': unique_folder,
        'expires': datetime.now() + TOKEN_TTL,
        'user': session.get('username'),
        'route_type': 'ppd'
    }

    saved = []
    for f in uploaded:
        fn = os.path.basename(f.filename)

        if not fn.lower().endswith((".doc", ".docx")):
            continue

        # sanitize filename for Windows
        fn = re.sub(r'[<>:"/\\|?*]', "_", fn)

        save_path = os.path.join(unique_folder, fn)
        f.save(save_path)
        saved.append(save_path)

    if not saved:
        return jsonify({"error": "No valid .doc/.docx files uploaded"}), 400

    # Combined dashboard flag — only relevant when 2+ files uploaded
    combined_dashboard = request.form.get('combined_dashboard', 'false').lower() == 'true'

    # Book title — required
    book_title = request.form.get('book_title', '').strip()
    if not book_title:
        return jsonify({"error": "Book Title is required"}), 400
    safe_title = re.sub(r'[<>:"/\\|?*\s]+', '_', book_title).strip('_')

    # Capture username before thread starts
    username = session.get("username") or "Analyst"

    # Job ID = Token (for multi-worker support via file system)
    job_id = token
    
    # Initialize progress file
    progress_file = os.path.join(unique_folder, "progress.json")
    initial_progress = {
        "total": len(saved),
        "current": 0,
        "status": "Starting",
        "folder": unique_folder
    }
    try:
        with open(progress_file, "w") as f:
            json.dump(initial_progress, f)
    except Exception as e:
        app.logger.error(f"Failed to create progress file: {e}")

    # -----------------------
    # Background job (batch queue)
    # -----------------------
    current_user_id = session.get('user_id')
    batch_queue.submit(
        job_id=job_id,
        route_type='ppd',
        user_id=current_user_id,
        username=username,
        target_fn=process_ppd_job,
        fn_args=(job_id, unique_folder, saved, combined_dashboard,
                 book_title, safe_title, username, current_user_id),
        payload_dict={
            'unique_folder': unique_folder,
            'saved': saved,
            'combined_dashboard': combined_dashboard,
            'book_title': book_title,
            'safe_title': safe_title,
        }
    )
    return jsonify({"job_id": job_id})

@app.route("/progress/<job_id>")
def progress(job_id):
    # Try reading from file system first (multi-worker support)
    try:
        # Check file system using job_id as token
        token_path = os.path.join(app.config['UPLOAD_FOLDER'], job_id)
        progress_path = os.path.join(token_path, "progress.json")
        
        if os.path.exists(progress_path):
            with open(progress_path, "r") as f:
                return jsonify(json.load(f))
                
        # Fallback: check if job_id is in config (legacy or very early state)
        if job_id in app.config.get("PROGRESS_DATA", {}):
            return jsonify(app.config["PROGRESS_DATA"][job_id])
            
    except Exception:
        pass
        
    return jsonify({})


@app.route("/download_zip/<job_id>")
def download_zip(job_id):
    zip_path = None
    folder_path = None
    
    # 1. Check in-memory (legacy)
    data = app.config.get("PROGRESS_DATA", {}).get(job_id)
    if data and "zip_path" in data:
        zip_path = data["zip_path"]
        folder_path = data.get("folder")
    
    # 2. Check file system (multi-worker)
    if not zip_path:
        # job_id is the token
        possible_folder = os.path.join(app.config['UPLOAD_FOLDER'], job_id)
        possible_progress = os.path.join(possible_folder, "progress.json")
        
        if os.path.exists(possible_progress):
            with open(possible_progress, "r") as f:
                file_data = json.load(f)
                if file_data.get("status") == "Completed":
                    if "zip_path" in file_data:
                        zip_path = file_data["zip_path"]
                    else:
                        # Fallback check for standard zip name
                        cand = os.path.join(possible_folder, "Reference_Process.zip")
                        if os.path.exists(cand):
                            zip_path = cand
                        # Legacy fallback
                        elif "zip_path" not in file_data:
                            # try checking for old name just in case? 
                            pass
                            
                    if zip_path:
                        folder_path = possible_folder

    if not zip_path:
        return "Not ready", 404

    if not os.path.exists(zip_path):
        return "ZIP not found", 404

    # Read ZIP into memory
    try:
        with open(zip_path, "rb") as f:
            zip_bytes = f.read()
    except Exception:
        return "Failed reading zip", 500

    # ----- MARK AS DOWNLOADED IN BATCH QUEUE -----
    try:
        with db_pool.get_connection() as db:
            db.execute(
                "UPDATE job_queue SET status='downloaded', downloaded_at=? WHERE job_id=?",
                (_now_utc(), job_id)
            )
            db.commit()
    except Exception:
        pass  # non-fatal — download still proceeds
    # -----------------------------------------------

    # ----- AUTO CLEANUP -----
    try:
        if folder_path and os.path.exists(folder_path):
            shutil.rmtree(folder_path, ignore_errors=True)
        if job_id in app.config.get("PROGRESS_DATA", {}):
            del app.config["PROGRESS_DATA"][job_id]
    except Exception as e:
        app.logger.error(f"Cleanup error for job {job_id}: {e}")
    # -------------------------

    # Return ZIP to client
    download_filename = os.path.basename(zip_path) if zip_path else "download.zip"
    
    return send_file(
        io.BytesIO(zip_bytes),
        mimetype="application/zip",
        as_attachment=True,
        download_name=download_filename
    )


# -----------------------
# Batch Queue Routes
# -----------------------
@app.route("/batch-queue")
def batch_queue_view():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    is_admin = bool(session.get('is_admin'))
    jobs = batch_queue.list_jobs(session['user_id'], is_admin)
    return render_template("batch_queue.html", jobs=jobs, is_admin=is_admin)


@app.route("/batch-queue/api")
def batch_queue_api():
    if 'user_id' not in session:
        return jsonify({"error": "Not logged in"}), 401
    jobs = batch_queue.list_jobs(session['user_id'], bool(session.get('is_admin')))
    return jsonify({"jobs": jobs})


@app.route("/batch-queue/cancel/<job_id>", methods=["POST"])
@csrf.exempt
def batch_queue_cancel(job_id):
    if 'user_id' not in session:
        return jsonify({"error": "Not logged in"}), 401
    ok, msg = batch_queue.cancel(job_id, session['user_id'])
    return jsonify({"ok": ok, "message": msg}), (200 if ok else 400)


# -----------------------
# Book Indexer Route
# -----------------------
@app.route('/book-indexer', methods=['GET'])
@role_required(ROUTE_PERMISSIONS.get('book_indexer', ['ADMIN', 'USER']))
def book_indexer_ui():
    return render_template('book_indexer.html')

@app.route('/book-indexer/api/extract', methods=['POST'])
@csrf.exempt
@role_required(ROUTE_PERMISSIONS.get('book_indexer', ['ADMIN', 'USER']))
def book_indexer_extract():
    api_key = request.form.get('api_key') or os.getenv('GEMINI_API_KEY')
    model_name = request.form.get('model', 'gemini-2.5-pro')
    pdf_files = request.files.getlist('pdf_files')

    if not api_key:
        return jsonify({"error": "Missing Gemini API key in request and .env"}), 400

    if not pdf_files or len(pdf_files) == 0 or pdf_files[0].filename == '':
        return jsonify({"error": "No PDF files provided"}), 400

    token = uuid.uuid4().hex
    job_id = token 
    
    temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], token)
    os.makedirs(temp_dir, exist_ok=True)

    saved_paths = []
    
    try:
        for f in pdf_files:
            if f.filename:
                safe_name = secure_filename(f.filename) or f"document_{len(saved_paths)}.pdf"
                path = os.path.join(temp_dir, safe_name)
                f.save(path)
                saved_paths.append(path)
    except Exception as e:
        try:
            shutil.rmtree(temp_dir)
        except:
            pass
        return jsonify({"error": f"File save failed: {e}"}), 500

    app.config.setdefault("PROGRESS_DATA", {})
    app.config["PROGRESS_DATA"][job_id] = {
        "total": len(saved_paths),
        "current": 0,
        "status": "Starting"
    }

    if len(pdf_files) == 1:
        base_name = os.path.splitext(secure_filename(pdf_files[0].filename))[0]
    else:
        base_name = "Combined"
    output_filename = f"{base_name}_index.docx"

    batch_queue.submit(
        job_id=job_id,
        route_type='book_indexer',
        user_id=session['user_id'],
        username=session.get('username', 'unknown'),
        target_fn=process_book_indexer_job,
        fn_args=(job_id, temp_dir, saved_paths, api_key, model_name, output_filename,
                 session['user_id'], session.get('username', 'unknown')),
        payload_dict={
            'temp_dir': temp_dir,
            'saved_paths': saved_paths,
            'model': model_name
        }
    )

    return jsonify({"job_id": job_id})

# -----------------------
# File Validation Route
# -----------------------
@app.route("/progress/<job_id>")
def check_progress(job_id):
    """
    Generic progress check route.
    Checks in-memory config first, then disk-based progress.json.
    """
    # 1. Check in-memory first (fastest)
    if "PROGRESS_DATA" in app.config and job_id in app.config["PROGRESS_DATA"]:
        return jsonify(app.config["PROGRESS_DATA"][job_id])
        
    # 2. Check disk (persistence/multi-worker)
    # We search in UPLOAD_FOLDER/{job_id}/progress.json
    try:
        progress_path = os.path.join(app.config['UPLOAD_FOLDER'], job_id, "progress.json")
        if os.path.exists(progress_path):
            with open(progress_path, "r") as f:
                data = json.load(f)
            return jsonify(data)
    except Exception as e:
        pass
        
    return jsonify({"status": "Unknown", "current": 0, "total": 0})

def process_validation_job(job_id, processing_dir, file_paths, original_filenames, options, user_id, username):
    """
    Background worker for validation.
    """
    with app.app_context():
        # Helper to update progress
        def update_progress(updates):
            # 1. Update In-Memory
            if "PROGRESS_DATA" not in app.config:
                app.config["PROGRESS_DATA"] = {}
            
            # Init if missing
            if job_id not in app.config["PROGRESS_DATA"]:
                app.config["PROGRESS_DATA"][job_id] = {"current": 0, "total": len(file_paths), "status": "Starting..."}
            
            app.config["PROGRESS_DATA"][job_id].update(updates)
            
            # 2. Update Disk (progress.json)
            try:
                p_path = os.path.join(processing_dir, "progress.json")
                current_data = {}
                if os.path.exists(p_path):
                    with open(p_path, "r") as f:
                        current_data = json.load(f)
                current_data.update(updates)
                with open(p_path, "w") as f:
                    json.dump(current_data, f)
            except Exception:
                pass
                
        update_progress({"status": "Initializing...", "total": len(file_paths)})
        
        processed_file_paths = []
        
        run_structuring = options.get('run_structuring', False)
        run_validation = options.get('run_validation', False)
        run_name_year = options.get('run_name_year', False)
        run_gemini = options.get('run_gemini', False)
        target_style = options.get('target_style', 'Auto')
        
        is_report_only = options.get('is_report_only', False)
        
        try:
            for idx, filepath in enumerate(file_paths):
                filename = original_filenames[idx]
                base_name = os.path.splitext(filename)[0]
                
                # Create per-file output folder
                file_output_dir = os.path.join(processing_dir, f"{base_name}_Results")
                os.makedirs(file_output_dir, exist_ok=True)
                
                update_progress({"status": f"Processing {filename}...", "current": idx})
                
                # Consolidated Log Buffer
                log_buffer = []
                log_buffer.append(f"PROCESS LOG FOR: {filename}")
                log_buffer.append(f"DATE: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                log_buffer.append("="*60)
                
                # Work on a copy in the temp dir initially
                current_filepath = filepath
                
                before = {'status': 'Skipped'}
                after = {'status': 'Skipped'}
                mapping = {}
                structuring_log_content = ""
                apa_log_content = ""
                val_msg = "Not validated"
                
                # -------------------------------------------------
                # 1. Structuring
                # -------------------------------------------------
                if run_structuring:
                    update_progress({"status": f"Structuring {filename}..."})
                    log_buffer.append("\n--- STRUCTURING ---")
                    try:
                        # process_docx_file takes input and output dir
                        # It generates _fixed.docx and _fix_log.txt in output_dir
                        struct_res = process_docx_file(Path(current_filepath), Path(file_output_dir), target_style=target_style)
                        
                        if struct_res.get('log_file') and struct_res.get('log_file').exists():
                            with open(struct_res.get('log_file'), 'r', encoding='utf-8') as lf:
                                structuring_log_content = lf.read()
                                log_buffer.append(structuring_log_content)
                            # We absorb the log content into our main log, 
                            # but we can also keep the individual file if user wants it.
                            # For neatness, maybe just keep our main log? 
                            # User asked for _log.txt, let's keep our consolidated one.
                            # struct_res['log_file'].unlink() # Remove component log to avoid clutter?
                            pass
                        else:
                            log_buffer.append("No structuring log generated.")
                            
                        if struct_res.get('output_docx') and struct_res.get('output_docx').exists():
                            current_filepath = str(struct_res.get('output_docx'))
                            log_buffer.append("Structuring successful.")
                        else:
                            log_buffer.append("Structuring failed to produce output.")
                            
                    except Exception as e:
                        log_errors([f"Structuring error {filename}: {e}"])
                        log_buffer.append(f"Structuring error: {e}")

                # -------------------------------------------------
                # 1b. Conversion (Separate Step)
                # -------------------------------------------------
                if run_gemini:
                    update_progress({"status": f"Converting {filename}..."})
                    log_buffer.append("\n--- CONVERSION ---")
                    try:
                        from ReferenceConversion import process_conversion
                        # source_style="Auto" auto-detects per reference; target_style is the desired output format
                        # "Auto" means keep each reference in its detected style (validate, don't convert)
                        # Only convert when user explicitly chose "AMA" or "APA" as target
                        conv_target = target_style if target_style in ("AMA", "APA") else "Auto"
                        conv_res = process_conversion(Path(current_filepath), Path(file_output_dir), source_style="Auto", target_style=conv_target)
                        
                        if conv_res.get('log_file') and conv_res.get('log_file').exists():
                            with open(conv_res.get('log_file'), 'r', encoding='utf-8') as lf:
                                log_buffer.append(lf.read())
                        
                        if conv_res.get('output_docx') and conv_res.get('output_docx').exists():
                            current_filepath = str(conv_res.get('output_docx'))
                            log_buffer.append("Conversion successful.")
                        else:
                            log_buffer.append("Conversion failed to produce output.")
                            
                    except Exception as e:
                        log_errors([f"Conversion error {filename}: {e}"])
                        log_buffer.append(f"Conversion error: {e}")



                # -------------------------------------------------
                # 2. Validation (Check References)
                # -------------------------------------------------
                if run_validation:
                    update_progress({"status": f"Validating {filename}..."})
                    log_buffer.append("\n--- NUMERICAL VALIDATION ---")
                    try:
                        doc, before, after, mapping, val_msg = process_document(current_filepath)
                        log_buffer.append(f"Result: {val_msg}")
                        log_buffer.append(f"Before Stats: {before}")
                        
                        has_citations = bool(mapping)
                        is_perfect = before.get('is_perfect', False)
                        
                        # If validation changes things or we just need to pass the doc forward
                        if has_citations or (not is_perfect):
                             temp_val_path = os.path.join(file_output_dir, f"temp_val_{uuid.uuid4().hex}.docx")
                             doc.save(temp_val_path)
                             current_filepath = temp_val_path
                    except Exception as e:
                        log_errors([f"Validation error {filename}: {e}"])
                        log_buffer.append(f"Error during validation: {e}")

                # -------------------------------------------------
                # 3. Name & Year
                # -------------------------------------------------
                if run_name_year:
                    update_progress({"status": f"Name & Year Check {filename}..."})
                    log_buffer.append("\n--- NAME & YEAR VALIDATION ---")
                    try:
                        from validation_core import CitationProcessor
                        
                        temp_ny_path = os.path.join(file_output_dir, f"temp_ny_{uuid.uuid4().hex}.docx")
                        
                        processor = CitationProcessor(current_filepath)
                        report = processor.run()
                        processor.save(temp_ny_path)
                        
                        comment_count = len(report.issues)
                        formatted_count = report.stats.get('matched', 0) + report.stats.get('format_fixed', 0)
                        
                        log_buffer.append(f"Comments inserted: {comment_count}")
                        log_buffer.append(f"Formatting applied: {formatted_count}")
                        
                        report_text = f"Document: {filename}\n\n{report.summary()}"
                        apa_log_content = str(report_text)
                        
                        if comment_count > 0 or formatted_count > 0:
                            current_filepath = temp_ny_path
                        else:
                            if os.path.exists(temp_ny_path):
                                os.remove(temp_ny_path)
                            
                        if not run_validation:
                             before['total_references'] = report.total_refs
                    except Exception as e:
                        log_errors([f"Name/Year error {filename}: {e}"])
                        log_buffer.append(f"Error during Name/Year check: {e}")

                # -------------------------------------------------
                # 4. Finalize & Save Results
                # -------------------------------------------------
                
                # A. Log File
                final_log_name = f"{base_name}_log.txt"
                final_log_path = os.path.join(file_output_dir, final_log_name)
                with open(final_log_path, "w", encoding="utf-8") as f:
                    f.write("\n".join(log_buffer))
                    f.write("\n\n--- Report Details ---\n")
                    f.write(apa_log_content)
                
                # B. HTML Report
                try:
                    res = {
                        'filename': filename,
                        'status_msg': val_msg,
                        'error': None,
                        'before': before,
                        'after': after,
                        'mapping': mapping,
                        'structuring_log': structuring_log_content,
                        'apa_log': apa_log_content
                    }
                    
                    with app.test_request_context():
                        g.user = None 
                        html_report = render_template(
                            'result_content.html', 
                            results_list=[res], 
                            offline_mode=True, 
                            now=datetime.now,
                            token=job_id
                        )
                    
                    report_filename = f"{base_name}_result.html"
                    report_path = os.path.join(file_output_dir, report_filename)
                    with open(report_path, "w", encoding="utf-8") as rf:
                        rf.write(html_report)
                except Exception as e:
                    log_errors([f"HTML Report generation failed for {filename}: {e}"])

                # C. Final Document (Conditionally)
                if not is_report_only:
                    final_doc_name = f"{base_name}_Processed.docx"
                    final_doc_path = os.path.join(file_output_dir, final_doc_name)
                    
                    # Ensure we are copying the latest 'current_filepath'
                    if os.path.abspath(current_filepath) != os.path.abspath(final_doc_path):
                        shutil.copy2(current_filepath, final_doc_path)
                
                # Cleanup temp intermediates in file_output_dir if any? 
                # (We won't traverse and delete temp_* files to be safe, but OS cleans temp or we can rely on containing folder deletion)
                
                # DB Logging (Simplified)
                try:
                    with db_pool.get_connection() as db:
                        cursor = db.execute(
                            'INSERT INTO files (user_id, original_filename, stored_filename, report_filename) VALUES (?, ?, ?, ?)',
                            (user_id, filename, filename, final_log_name)
                        )
                        db.execute(
                            'INSERT INTO validation_results (file_id, total_references, total_citations, missing_references, unused_references, sequence_issues) VALUES (?, ?, ?, ?, ?, ?)',
                            (cursor.lastrowid, 
                             before.get('total_references', 0),
                             before.get('total_citations', 0),
                             "", "", "")
                        )
                        db.commit()
                except Exception as ex:
                    print(f"DB Log Error: {ex}")
            
            # End of loop
            
            update_progress({"status": "Finalizing..."})
            
            # Create ZIP
            zip_name = "Reference_Process.zip"
            zip_path = os.path.join(processing_dir, zip_name)
            
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as z:
                # Walk through processing_dir
                for root, dirs, files in os.walk(processing_dir):
                    for file in files:
                        if file == "progress.json" or file == zip_name or file.endswith(".docx"): # skip input docs in root, loop handled moved ones
                            # Wait, input docs are in processing_dir root. Result folders are subdirs.
                            # We want to ZIP the subdirs.
                            pass
                        
                        file_abs_path = os.path.join(root, file)
                        rel_path = os.path.relpath(file_abs_path, processing_dir)
                        
                        # Include: 
                        # 1. Anything inside a "_Results" folder
                        # 2. Skip input files in root (checked by not having path separator?)
                        
                        if "_Results" in rel_path:
                            # Cleanup temp files inside results?
                            if file.startswith("temp_") or file.endswith("_fix_log.txt") or file.endswith("_fixed.docx"): # Structuring artifacts we might not want if we have _Processed
                                # Logic check: _fixed.docx is the result of structuring. 
                                # If we ran structuring but not report only, it became _Processed.docx? 
                                # No, current_filepath became _fixed.docx. Then it was copied to _Processed.docx.
                                # So _fixed.docx is redundant if _Processed exists.
                                pass 
                            
                            z.write(file_abs_path, arcname=rel_path)
            
            # Register Download
            download_tokens[job_id] = {
                "path": processing_dir,
                "expires": _now_utc() + TOKEN_TTL,
                "user": username,
                "route_type": "validation",
                "zip_path": zip_path
            }
            
            update_progress({
                "status": "Completed", 
                "download_token": job_id,
                "current": len(file_paths),
                "zip_path": zip_path
            })
            
        except Exception as e:
            update_progress({"status": f"Failed: {str(e)}"})
            log_errors([f"Job {job_id} failed: {e}"])

@app.route("/validate", methods=["GET", "POST"], strict_slashes=False)
def validate_file():
    if 'user_id' not in session:
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({"success": False, "message": "Please log in to continue"})
        return redirect(url_for('login'))

    if request.method == "POST":
        uploaded_files = request.files.getlist('files')
        if not uploaded_files or not uploaded_files[0].filename:
             uploaded_files = request.files.getlist('file')

        if not uploaded_files or not uploaded_files[0].filename:
            return jsonify({"success": False, "message": "No files selected"})

        token = uuid.uuid4().hex
        processing_dir = os.path.join(app.config['UPLOAD_FOLDER'], token)
        os.makedirs(processing_dir, exist_ok=True)
        
        saved_paths = []
        original_filenames = []
        
        try:
            file_errors = []
            
            for file in uploaded_files:
                filename = secure_filename(file.filename)
                if not allowed_file(filename):
                    continue
                filepath = os.path.join(processing_dir, filename)
                file.save(filepath)
                
                # --- NEW VALIDATION LOGIC START ---
                if filename.lower().endswith('.docx'):
                    try:
                        from docx import Document
                        
                        doc = Document(filepath)
                        has_ref_style = False
                        has_ref_open = False
                        has_ref_close = False
                        
                        for para in doc.paragraphs:
                            if para.style and para.style.name and para.style.name.upper() in ['REF-N', 'REF-U']:
                                has_ref_style = True
                            if '<ref-open>' in para.text:
                                has_ref_open = True
                            if '<ref-close>' in para.text:
                                has_ref_close = True
                                
                            if has_ref_style and has_ref_open and has_ref_close:
                                break
                                
                        missing = []
                        if not has_ref_style:
                            missing.append("style 'REF-N' or 'REF-U'")
                        if not has_ref_open:
                            missing.append("'<ref-open>' tag")
                        if not has_ref_close:
                            missing.append("'<ref-close>' tag")
                            
                        if missing:
                            file_errors.append(f"{filename} (Missing: {', '.join(missing)})")
                            
                    except Exception as e:
                        file_errors.append(f"{filename} (Error reading file: {str(e)})")
                # --- NEW VALIDATION LOGIC END ---
                
                saved_paths.append(filepath)
                original_filenames.append(file.filename)
                
            # If any files failed validation, abort the entire batch
            if file_errors:
                import shutil
                shutil.rmtree(processing_dir, ignore_errors=True)
                error_msg = "Please fix the following files and try again: | " + " | ".join(file_errors)
                return jsonify({"success": False, "message": error_msg})
                
            # Collect Options
            options = {
                'is_report_only': str(request.form.get('report_only')).lower() in ['true', 'on', '1'],
                'run_validation': str(request.form.get('run_validation')).lower() in ['true', 'on', '1'],
                'run_structuring': str(request.form.get('run_structuring')).lower() in ['true', 'on', '1'],
                'run_name_year': str(request.form.get('run_name_year_validation')).lower() in ['true', 'on', '1'],
                'run_gemini': str(request.form.get('run_gemini')).lower() in ['true', 'on', '1'],
                'target_style': request.form.get('target_style', 'Auto')
            }
            
            # Use report only flag to force validation logic if not explicitly checked but needed
            if options['is_report_only']:
                options['run_validation'] = True

            # Start Background Job (batch queue)
            batch_queue.submit(
                job_id=token,
                route_type='validation',
                user_id=session['user_id'],
                username=session.get('username', 'unknown'),
                target_fn=process_validation_job,
                fn_args=(token, processing_dir, saved_paths, original_filenames,
                         options, session['user_id'], session.get('username', 'unknown')),
                payload_dict={
                    'processing_dir': processing_dir,
                    'saved_paths': saved_paths,
                    'original_filenames': original_filenames,
                    'options': options,
                }
            )

            return jsonify({"success": True, "job_id": token, "message": "Processing started"})

        except Exception as e:
            return jsonify({"success": False, "message": f"Start failed: {str(e)}"})

    # GET Request - Render Page
    return render_template("upload.html")


import base64

def get_base64_logo():
    logo_path = os.path.join(app.static_folder, "images", "S4c.png")
    try:
        with open(logo_path, "rb") as f:
            return base64.b64encode(f.read()).decode("utf-8")
    except Exception as e:
        app.logger.warning(f"Logo not found or failed to load: {e}")
        return ""

# -----------------------
# Dashboard
# -----------------------
@app.route("/dashboard", strict_slashes=False)
def dashboard():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    with db_pool.get_connection() as db:
        if session.get('is_admin'):
            recent_files = db.execute('''SELECT f.*, u.username 
                                       FROM files f 
                                       JOIN users u ON f.user_id = u.id 
                                       ORDER BY f.upload_date DESC LIMIT 5''').fetchall()

            recent_macro = db.execute('''SELECT m.*, u.username 
                                       FROM macro_processing m
                                       JOIN users u ON m.user_id = u.id 
                                       ORDER BY m.processing_date DESC LIMIT 5''').fetchall()

            # Route-specific stats
            route_stats = {}
            for route_type in ROUTE_MACROS.keys():
                count = db.execute("SELECT COUNT(*) FROM macro_processing WHERE route_type = ?",
                                   (route_type,)).fetchone()[0]
                route_stats[route_type] = count

            admin_stats = {
                'total_users': db.execute("SELECT COUNT(*) FROM users").fetchone()[0],
                'total_files': db.execute("SELECT COUNT(*) FROM files").fetchone()[0],
                'total_validations': db.execute("SELECT COUNT(*) FROM validation_results").fetchone()[0],
                'total_macro': db.execute("SELECT COUNT(*) FROM macro_processing").fetchone()[0],
                'route_stats': route_stats
            }
        else:
            recent_files = db.execute('''SELECT * FROM files 
                                       WHERE user_id=? 
                                       ORDER BY upload_date DESC LIMIT 5''',
                                      (session['user_id'],)).fetchall()

            recent_macro = db.execute('''SELECT * FROM macro_processing 
                                       WHERE user_id=? 
                                       ORDER BY processing_date DESC LIMIT 5''',
                                      (session['user_id'],)).fetchall()

            # User-specific route stats
            route_stats = {}
            for route_type in ROUTE_MACROS.keys():
                count = db.execute("SELECT COUNT(*) FROM macro_processing WHERE user_id = ? AND route_type = ?",
                                   (session['user_id'], route_type)).fetchone()[0]
                route_stats[route_type] = count

            admin_stats = {'route_stats': route_stats}

    return render_template("dashboard.html",
                           recent_files=recent_files,
                           recent_macro=recent_macro,
                           admin_stats=admin_stats,
                           route_macros=ROUTE_MACROS)


# -----------------------
# Download Route
# -----------------------
@app.route('/macro-download', strict_slashes=False)
def macro_download():
    token = request.args.get('token')
    if not token:
        flash("Invalid download request.")
        return redirect(url_for('dashboard'))

    token_data = download_tokens.get(token)
    
    # Check database if not in memory (multi-worker or restart scenario)
    if not token_data:
        try:
            with db_pool.get_connection() as db:
                result = db.execute(
                    "SELECT * FROM macro_processing WHERE token = ?",
                    (token,)
                ).fetchone()
                
                if result:
                    # Reconstruct token_data from database record
                    token_data = {
                        'path': os.path.join(app.config['UPLOAD_FOLDER'], token),
                        'route_type': result['route_type'] if result['route_type'] else 'general',
                        'user_id': result['user_id']
                    }
                    # Check if processing is complete
                    if not result['errors'] or result['errors'] == '':
                        # Assume completed if no errors recorded
                        pass
                    else:
                        # Has errors, might not be ready
                        token_data = None
        except Exception as e:
            app.logger.error(f"Error checking database for token {token}: {e}")
    
    # Fallback to file system
    if not token_data:
        possible_folder = os.path.join(app.config['UPLOAD_FOLDER'], token)
        if os.path.exists(possible_folder):
            token_data = {
                'path': possible_folder,
                'route_type': 'recovered' 
            }

    if not token_data:
        flash("Invalid or expired download token.")
        return redirect(url_for('dashboard'))

    # Skip expiry check for database-recovered tokens (they don't have 'expires' field)
    if 'expires' in token_data and is_token_expired(token_data):
        cleanup_token_data(token)
        flash("Download token has expired.")
        return redirect(url_for('dashboard'))

    user_folder = token_data['path']
    route_type = token_data.get('route_type', 'general')

    if not os.path.exists(user_folder):
        flash("No files found for this download token.")
        return redirect(url_for('dashboard'))

    try:
        # Special handling for credit_extractor: Direct Excel download
        if route_type == 'credit_extractor':
            excel_path = os.path.join(user_folder, 'permission_log.xlsx')
            if os.path.exists(excel_path):
                with open(excel_path, 'rb') as f:
                    memory_file = io.BytesIO(f.read())
                
                try:
                    shutil.rmtree(user_folder)
                    if token in download_tokens: del download_tokens[token]
                except Exception as e:
                    log_errors([f"Cleanup error: {str(e)}"])
                    
                return send_file(memory_file, 
                               mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                               as_attachment=True, 
                               download_name=f"Permission_Log_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")

        memory_file = io.BytesIO()
        with zipfile.ZipFile(memory_file, 'w', zipfile.ZIP_DEFLATED, compresslevel=6) as zipf:
            for root, _, files in os.walk(user_folder):
                for file in files:
                    # Filter for validation route: only _Processed.docx, _log.txt and _report.html
                    if route_type == 'validation':
                        lower_name = file.lower()
                        if not (lower_name.endswith('_processed.docx') or
                                lower_name.endswith('_log.txt') or
                                lower_name.endswith('_report.html')):
                            continue

                    # Filter for credit_extractor: only the output Excel/zip, not source docs or progress.json
                    elif route_type == 'credit_extractor':
                        lower_name = file.lower()
                        if not (lower_name == 'permission_log.xlsx' or lower_name == 'permission_logs.zip'):
                            continue

                    file_path = os.path.join(root, file)
                    if os.path.getsize(file_path) < 50 * 1024 * 1024:
                        arcname = os.path.relpath(file_path, user_folder)
                        zipf.write(file_path, arcname)

        memory_file.seek(0)

        try:
            shutil.rmtree(user_folder)
            del download_tokens[token]
        except Exception as e:
            log_errors([f"Cleanup error: {str(e)}"])

        route_name = ROUTE_MACROS.get(route_type, {}).get('name', 'Processed')
        zip_filename = f"{route_name.replace(' ', '_')}_Documents_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"

        return send_file(memory_file,
                         mimetype='application/zip',
                         as_attachment=True,
                         download_name=zip_filename)

    except Exception as e:
        flash(f"Download failed: {str(e)}")
        log_errors([f"Download error for token {token}: {str(e)}"])
        return redirect(url_for('dashboard'))


# -----------------------
# File History
# -----------------------
@app.route('/history', strict_slashes=False)
def file_history():
    if not g.user:
        flash("Please log in to view file history", "error")
        return redirect(url_for('login'))

    page = int(request.args.get('page', 1))
    per_page = 10
    offset = (page - 1) * per_page
    route_filter = request.args.get('route', 'all')

    with get_db() as conn:
        # Filter logic
        filter_condition = ""
        params = []

        if route_filter != "all":
            if route_filter == "validation":
                filter_condition = "WHERE type = 'validation'"
            else:
                filter_condition = "WHERE type = 'macro' AND route_type = ?"
                params.append(route_filter)

        # Admin vs User-specific
        if session.get("is_admin"):
            user_condition = ""
        else:
            user_condition = "AND user_id = ?" if filter_condition else "WHERE user_id = ?"
            params.append(g.user["id"])

        # Unified query
        query = f"""
            SELECT * FROM (
                SELECT f.id,
                       f.original_filename AS original_filename,
                       f.upload_date AS date,
                       f.report_filename,
                       v.total_references,
                       v.total_citations,
                       u.username,
                       'validation' AS type,
                       '' AS route_type,
                       '' AS token,
                       '' AS selected_tasks,
                       '' AS original_filenames,
                       f.user_id
                FROM files f
                LEFT JOIN validation_results v ON f.id = v.file_id
                JOIN users u ON f.user_id = u.id

                UNION ALL

                SELECT m.id,
                       '' AS original_filename,
                       m.processing_date AS date,
                       '' AS report_filename,
                       0 AS total_references,
                       0 AS total_citations,
                       u.username,
                       'macro' AS type,
                       m.route_type AS route_type, 
                       m.token,
                       m.selected_tasks,
                       m.original_filenames,
                       m.user_id
                FROM macro_processing m
                JOIN users u ON m.user_id = u.id
            ) combined
            {filter_condition}
            {user_condition}
            ORDER BY date DESC
            LIMIT ? OFFSET ?
        """

        params.extend([per_page, offset])
        cursor = conn.execute(query, params)
        history = cursor.fetchall()

        # Count total records for pagination
        count_query = f"""
            SELECT COUNT(*) FROM (
                SELECT f.id, f.user_id, 'validation' AS type, '' AS route_type
                FROM files f
                UNION ALL
                SELECT m.id, m.user_id, 'macro' AS type, m.route_type AS route_type
                FROM macro_processing m
            ) combined
            {filter_condition}
            {user_condition}
        """
        cursor = conn.execute(count_query, params[:-2])  # exclude LIMIT/OFFSET
        total_records = cursor.fetchone()[0]

        # Get overall stats
        stats_user_condition = ""
        stats_params = []
        if not session.get("is_admin"):
            stats_user_condition = "WHERE user_id = ?"
            stats_params = [g.user["id"]]
            
        stats_query = f"""
            SELECT type, route_type, COUNT(*) as count FROM (
                SELECT id, user_id, 'validation' AS type, 'validation' AS route_type FROM files
                UNION ALL
                SELECT id, user_id, 'macro' AS type, route_type FROM macro_processing
            ) combined
            {stats_user_condition}
            GROUP BY type, route_type
        """
        cursor = conn.execute(stats_query, stats_params)
        raw_stats = cursor.fetchall()
        
        overview_stats = {
            'total': 0,
            'validation': 0,
            'macro': 0,
            'by_route': {}
        }
        for row in raw_stats:
            c = row['count']
            overview_stats['total'] += c
            if row['type'] == 'validation':
                overview_stats['validation'] += c
                overview_stats['by_route']['validation'] = c
            else:
                overview_stats['macro'] += c
                overview_stats['by_route'][row['route_type']] = c

    total_pages = (total_records + per_page - 1) // per_page

    return render_template(
        "file_history.html",
        history=history,
        page=page,
        total_pages=total_pages,
        route_filter=route_filter,
        route_macros=ROUTE_MACROS,
        overview_stats=overview_stats
    )
# -----------------------
# Admin Routes
# -----------------------
@app.route("/admin", strict_slashes=False)
@admin_required
def admin_dashboard():
    with db_pool.get_connection() as db:
        route_stats = {}
        for route_type in ROUTE_MACROS.keys():
            count = db.execute(
                "SELECT COUNT(*) FROM macro_processing WHERE route_type = ?",
                (route_type,)
            ).fetchone()[0]
            route_stats[route_type] = count

        # totals
        total_users = db.execute("SELECT COUNT(*) FROM users").fetchone()[0]
        total_files = db.execute("SELECT COUNT(*) FROM files").fetchone()[0]
        total_validations = db.execute("SELECT COUNT(*) FROM validation_results").fetchone()[0]
        total_macro = db.execute("SELECT COUNT(*) FROM macro_processing").fetchone()[0]

        # roles (defensive: handle missing column)
        try:
            role_counts = db.execute(
                "SELECT role, COUNT(*) as count FROM users GROUP BY role"
            ).fetchall()
            role_stats = {
                (r["role"] if r["role"] else "USER"): r["count"] for r in role_counts
            }
        except sqlite3.OperationalError as e:
            log_errors([f"Role stats query failed: {e}"])
            role_stats = {}

        admin_stats = {
            'total_users': total_users,
            'total_files': total_files,
            'total_validations': total_validations,
            'total_macro': total_macro,
            'route_stats': route_stats
        }

    return render_template(
        "admin_dashboard.html",
        admin_stats=admin_stats,
        route_macros=ROUTE_MACROS,
        role_stats=role_stats   # ✅ now passed to template
    )


@app.route("/admin/user/<int:user_id>/change-role", methods=["POST"], strict_slashes=False)
@admin_required
def admin_change_role(user_id):
    new_role = request.form.get('role', '').upper()
    if not new_role:
        flash("No role provided", "error")
        return redirect(url_for('admin_users'))

    if user_id == session.get('user_id'):
        flash("Cannot change your own role", "error")
        return redirect(url_for('admin_users'))

    with db_pool.get_connection() as db:
        user = db.execute("SELECT * FROM users WHERE id=?", (user_id,)).fetchone()
        if not user:
            flash("User not found", "error")
            return redirect(url_for('admin_users'))

        db.execute("UPDATE users SET role=? WHERE id=?", (new_role, user_id))
        db.commit()
        flash("User role updated", "success")
        log_activity(session['username'], 'CHANGE_ROLE', f"user:{user['username']} -> {new_role}")
    return redirect(url_for('admin_users'))
# -----------------------
# Admin User Management
# -----------------------
@app.route("/admin/users", strict_slashes=False)
@admin_required
def admin_users():
    with db_pool.get_connection() as db:
        users = db.execute(
            'SELECT id, username, email, is_admin, role, created_at FROM users ORDER BY created_at DESC').fetchall()
    return render_template("admin_users.html", users=users)


@app.route("/admin/create-user", methods=["GET", "POST"], strict_slashes=False)
@admin_required
def admin_create_user():
    if request.method == "POST":
        username = request.form['username']
        password = request.form['password']
        email = request.form.get('email', '')
        is_admin = 'is_admin' in request.form
        role = request.form.get('role', 'USER').upper()

        with db_pool.get_connection() as db:
            try:
                hashed = generate_password_hash(password, method='pbkdf2:sha256')
                db.execute("INSERT INTO users (username,password,email,is_admin,role) VALUES (?,?,?,?,?)",
                           (username, hashed, email, is_admin, role))
                db.commit()
                flash("User created successfully", "success")
                return redirect(url_for('admin_users'))
            except sqlite3.IntegrityError:
                db.rollback()
                flash("Username/email exists", "error")

    return render_template("admin_create_user.html")


@app.route('/admin/change_password/<int:user_id>', methods=['GET', 'POST'], strict_slashes=False)
@admin_required
def admin_change_password(user_id):
    with db_pool.get_connection() as db:
        user = db.execute("SELECT * FROM users WHERE id = ?", (user_id,)).fetchone()

        if not user:
            flash("User not found.", "error")
            return redirect(url_for('admin_users'))

        if request.method == 'POST':
            new_password = request.form['new_password']
            hashed = generate_password_hash(new_password)
            db.execute("UPDATE users SET password = ? WHERE id = ?", (hashed, user_id))
            db.commit()
            flash(f"Password updated for {user['username']}.", "success")
            return redirect(url_for('admin_users'))

    return render_template("admin_change_password.html", user=user)


@app.route("/admin/user/<int:user_id>/toggle-admin", methods=["POST"], strict_slashes=False)
@admin_required
def admin_toggle_admin(user_id):
    if user_id == session.get('user_id'):
        flash("Cannot change your own admin status", "error")
        return redirect(url_for('admin_users'))

    with db_pool.get_connection() as db:
        user = db.execute("SELECT is_admin FROM users WHERE id=?", (user_id,)).fetchone()
        if not user:
            flash("User not found", "error")
            return redirect(url_for('admin_users'))

        new_status = not bool(user['is_admin'])
        db.execute("UPDATE users SET is_admin=? WHERE id=?", (new_status, user_id))
        db.commit()
        status_text = "granted" if new_status else "revoked"
        flash(f"Admin privileges {status_text}", "success")

    return redirect(url_for('admin_users'))


@app.route("/admin/user/<int:user_id>/delete", methods=["POST"], strict_slashes=False)
@admin_required
def admin_delete_user(user_id):
    # Prevent admins from deleting themselves
    if user_id == session.get('user_id'):
        flash("Cannot delete your own account", "error")
        return redirect(url_for('admin_users'))

    try:
        with db_pool.get_connection() as db:
            # Check macro history
            macro_count = db.execute(
                "SELECT COUNT(*) FROM macro_processing WHERE user_id=?",
                (user_id,)
            ).fetchone()[0]

            if macro_count > 0:
                flash("Cannot delete user with macro history", "error")
                return redirect(url_for('admin_users'))

            # Check files
            user_files = db.execute(
                "SELECT COUNT(*) FROM files WHERE user_id=?",
                (user_id,)
            ).fetchone()[0]

            if user_files > 0:
                flash("Cannot delete user with files", "error")
                return redirect(url_for('admin_users'))

            # At this point it's safe to delete user
            # Optionally remove any related rows (safety) - will cascade if you used FK cascade, but we'll be explicit
            try:
                db.execute("DELETE FROM validation_results WHERE file_id IN (SELECT id FROM files WHERE user_id=?)", (user_id,))
            except Exception:
                # ignore if validation_results references don't exist
                pass

            try:
                db.execute("DELETE FROM files WHERE user_id=?", (user_id,))
            except Exception:
                # ignore if no files
                pass

            db.execute("DELETE FROM macro_processing WHERE user_id=?", (user_id,))  # should be zero if earlier check passed
            db.execute("DELETE FROM users WHERE id=?", (user_id,))
            db.commit()

            flash("User deleted successfully", "success")
            log_activity(session.get('username', 'system'), "DELETE_USER", f"user_id:{user_id}")
            return redirect(url_for('admin_users'))

    except Exception as e:
        log_errors([f"Error deleting user {user_id}: {e}", traceback.format_exc()])
        flash("An error occurred while deleting the user", "error")
        return redirect(url_for('admin_users'))



@app.route("/admin/files")
@admin_required
def admin_files():
    page = request.args.get('page', 1, type=int)
    per_page = 10
    offset = (page - 1) * per_page

    with db_pool.get_connection() as db:
        files = db.execute('''SELECT f.*, u.username, v.total_references, v.total_citations
                           FROM files f
                           JOIN users u ON f.user_id = u.id
                           LEFT JOIN validation_results v ON f.id = v.file_id
                           ORDER BY f.upload_date DESC LIMIT ? OFFSET ?''',
                           (per_page, offset)).fetchall()

        total_count = db.execute("SELECT COUNT(*) FROM files").fetchone()[0]
        total_pages = (total_count + per_page - 1) // per_page

    return render_template("admin_files.html", files=files, page=page, total_pages=total_pages)


@app.route("/admin/file/<int:file_id>/delete", methods=["POST"])
@admin_required
def admin_delete_file(file_id):
    with db_pool.get_connection() as db:
        file = db.execute("SELECT * FROM files WHERE id=?", (file_id,)).fetchone()
        if not file:
            flash("File not found", "error")
            return redirect(url_for('admin_files'))

        # Delete the file from storage
        try:
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], file['stored_filename'])
            if os.path.exists(file_path):
                os.remove(file_path)

            # Delete report file if exists
            if file['report_filename']:
                report_path = os.path.join(REPORT_FOLDER, file['report_filename'])
                if os.path.exists(report_path):
                    os.remove(report_path)
        except Exception as e:
            flash(f"Error deleting file: {str(e)}", "error")
            return redirect(url_for('admin_files'))

        # Delete from database
        db.execute("DELETE FROM validation_results WHERE file_id=?", (file_id,))
        db.execute("DELETE FROM files WHERE id=?", (file_id,))
        db.commit()

        flash("File deleted successfully", "success")
        return redirect(url_for('admin_files'))


@app.route("/admin/stats")
@admin_required
def admin_stats():
    with db_pool.get_connection() as db:
        # Get recent files
        # Get recent files (Combined Validation + Macros)
        recent_files = db.execute('''
            SELECT * FROM (
                SELECT f.id,
                       f.original_filename AS original_filename,
                       f.upload_date AS date,
                       u.username,
                       'validation' AS type,
                       '' AS route_type,
                       '' AS original_filenames
                FROM files f
                JOIN users u ON f.user_id = u.id

                UNION ALL

                SELECT m.id,
                       '' AS original_filename,
                       m.processing_date AS date,
                       u.username,
                       'macro' AS type,
                       m.route_type AS route_type,
                       m.original_filenames
                FROM macro_processing m
                JOIN users u ON m.user_id = u.id
            ) combined
            ORDER BY date DESC LIMIT 20
        ''').fetchall()

        # Month-wise User Activity (Last 6 Months)
        from datetime import datetime
        now = datetime.now()
        month_headers = []
        # Generate last 6 months list [(year, month), ...]
        current_y, current_m = now.year, now.month
        for i in range(6):
            y, m = current_y, current_m - i
            while m <= 0:
                m += 12
                y -= 1
            month_headers.append(f"{y}-{m:02d}")
        
        # Prepare placeholders for SQL IN clause
        placeholders = ','.join(['?'] * len(month_headers))
        
        query = f'''
            SELECT u.username, strftime('%Y-%m', activity_date) as month, COUNT(*) as count
            FROM (
                SELECT user_id, upload_date as activity_date FROM files
                UNION ALL
                SELECT user_id, processing_date as activity_date FROM macro_processing
            ) a
            JOIN users u ON a.user_id = u.id
            WHERE strftime('%Y-%m', activity_date) IN ({placeholders})
            GROUP BY u.username, month
        '''
        
        raw_stats = db.execute(query, month_headers).fetchall()
        
        # Organize data: {username: {'total': 0, 'months': {'2023-01': 0, ...}}}
        user_map = {}
        
        # Initialize users first to ensure we catch those who have 0 activity in this period but exist? 
        # Or just show active ones? The previous query showed ALL users with total count.
        # Let's get ALL users totals first to keep consistency with previous view.
        
        all_users = db.execute('''
            SELECT u.username, 
                   ((SELECT COUNT(*) FROM files f WHERE f.user_id = u.id) + 
                    (SELECT COUNT(*) FROM macro_processing m WHERE m.user_id = u.id)) as total_count
            FROM users u
            ORDER BY total_count DESC
        ''').fetchall()
        
        for u in all_users:
            user_map[u['username']] = {
                'username': u['username'],
                'total': u['total_count'],
                'months': {m: 0 for m in month_headers}
            }
            
        # Fill in monthly data
        for row in raw_stats:
            uname = row['username']
            month = row['month']
            count = row['count']
            if uname in user_map and month in user_map[uname]['months']:
                user_map[uname]['months'][month] = count

        # Convert to list sorted by total count desc
        users_data = sorted(user_map.values(), key=lambda x: x['total'], reverse=True)

        # Get total counts
        total_users = db.execute("SELECT COUNT(*) FROM users").fetchone()[0]
        total_files = db.execute("SELECT COUNT(*) FROM files").fetchone()[0]
        total_validations = db.execute("SELECT COUNT(*) FROM validation_results").fetchone()[0]
        total_macro = db.execute("SELECT COUNT(*) FROM macro_processing").fetchone()[0]

        # Role stats
        role_counts = db.execute("SELECT role, COUNT(*) as count FROM users GROUP BY role").fetchall()
        role_stats = {r["role"]: r["count"] for r in role_counts}

    return render_template(
        "admin_stats.html",
        recent_files=recent_files,
        users_data=users_data,
        month_headers=month_headers,
        admin_stats={
            'total_users': total_users,
            'total_files': total_files,
            'total_validations': total_validations,
            'total_macro': total_macro
        },
        role_stats=role_stats,
        route_macros=ROUTE_MACROS
    )


@app.route('/doi_finder')
def doi_finder():
    """DOI Correction and Metadata Finder"""
    if 'user_id' not in session:
        flash("Please log in to continue.")
        return redirect(url_for('login'))

    return render_template('doi_finder.html')


@app.route('/api/log-action', methods=['POST'])
def log_action_api():
    """API to log client-side actions like DOI searches"""
    if 'user_id' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
        
    try:
        data = request.get_json()
        action_type = data.get('action_type', 'unknown')
        details = data.get('details', {})
        
        user_id = session.get('user_id')
        token = uuid.uuid4().hex  # Generate a dummy token for the existing schema
        
        with db_pool.get_connection() as db:
            db.execute('''INSERT INTO macro_processing 
                          (user_id, token, original_filenames, processed_filenames, selected_tasks, route_type)
                          VALUES (?, ?, ?, ?, ?, ?)''',
                       (user_id, 
                        token, 
                        json.dumps([details.get('query', 'single_lookup')]), # Store query/filename here
                        json.dumps([]), 
                        json.dumps(details), 
                        action_type)) # Use route_type to store the action (e.g., 'doi_finder')
            db.commit()
            
        return jsonify({'status': 'logged'})
    except Exception as e:
        app.logger.error(f"Failed to log action: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/smart-format', methods=['POST'])
def smart_format_api():
    """
    Heuristic formatting for non-journal references (Book, Web, Thesis)
    when DOI/PubMed lookup fails.
    """
    try:
        data = request.get_json()
        raw_ref = data.get('reference', '').strip()
        style = data.get('style', 'apa').lower() # 'apa' or 'ama'
        
        if not raw_ref:
             return jsonify({'error': 'No reference text provided'}), 400
             
        # Detect Style mode for parsing
        # (The parsing function expects raw string)
        if style == 'ama':
            parsed = parse_ama_reference_raw(raw_ref)
            target_style = 'REF-N'
        else:
            parsed = parse_apa_reference_raw(raw_ref)
            target_style = 'REF-U'
            
        # Refine type based on content if not set
        if 'manual_type' not in parsed:
             # simple heuristics
             lower_ref = raw_ref.lower()
             if 'dissertation' in lower_ref or 'thesis' in lower_ref:
                 parsed['manual_type'] = 'thesis'
             elif 'http' in lower_ref or 'www.' in lower_ref:
                 parsed['manual_type'] = 'web'
             elif 'isbn' in lower_ref:
                 parsed['manual_type'] = 'book'
             else:
                 # Default to book for fallback formatting if it looks like Author. Title. Pub.
                 parsed['manual_type'] = 'book'

        # Generate segments
        segments = generate_fallback_citation(parsed, raw_ref, style_mode=target_style)
        
        # Convert segments to HTML string
        # Segments are tuples (text, style_name)
        # We want to return a clean HTML string.
        # Check usage in doi_finder.html -> it expects text.
        # But for APA/AMA we might want some formatting (italics)?
        # generate_fallback_citation returns list of (text, style).
        # We need to render this to HTML.
        
        html_out = ""
        for text, s_name in segments:
             if s_name:
                 # Map docx styles to HTML
                 if s_name in ('bib_journal', 'bib_book', 'bib_title', 'bib_confproceedings'):
                     # Italics
                     html_out += f"<i>{text}</i>"
                 elif s_name == 'bib_volume':
                     if style == 'ama':
                         # AMA volume is bold? No, usually italics or standard depending.
                         # Let's stick to standard or minimal italics.
                         # Actually checking standard AMA: Journal Title. Year;Volume(Issue):Pages.
                         # ReferencesStructing logic handles punctuation.
                         html_out += text
                     else:
                         # APA Volume is italic
                         html_out += f"<i>{text}</i>"
                 else:
                     html_out += text
             else:
                 html_out += text
                 
        return jsonify({
            'formatted': html_out,
            'type': parsed.get('manual_type', 'unknown')
        })

    except Exception as e:
        app.logger.error(f"Smart Format Error: {e}")
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route("/admin/macro-stats")
@admin_required
def admin_macro_stats():
    with db_pool.get_connection() as db:
        macro_records = db.execute('''SELECT selected_tasks, processing_date, errors, route_type
                                    FROM macro_processing 
                                    ORDER BY processing_date DESC''').fetchall()

    route_stats = {}
    error_stats = {}
    daily_stats = {}

    for record in macro_records:
        route_type = record['route_type'] or 'unknown'

        # Count by route
        if route_type not in route_stats:
            route_stats[route_type] = 0
        route_stats[route_type] += 1

        # Count errors by route
        if record['errors']:
            try:
                error_count = len(json.loads(record['errors']))
                if route_type not in error_stats:
                    error_stats[route_type] = 0
                error_stats[route_type] += error_count
            except:
                pass

        # Daily stats
        date = record['processing_date'][:10]
        if date not in daily_stats:
            daily_stats[date] = {}
        if route_type not in daily_stats[date]:
            daily_stats[date][route_type] = 0
        daily_stats[date][route_type] += 1

    return render_template("admin_macro_stats.html",
                           route_stats=route_stats,
                           error_stats=error_stats,
                           daily_stats=daily_stats,
                           route_macros=ROUTE_MACROS)


@app.route("/admin/macro-history")
@admin_required
def admin_macro_history():
    page = request.args.get('page', 1, type=int)
    per_page = 10
    offset = (page - 1) * per_page

    with db_pool.get_connection() as db:
        macro_history = db.execute('''SELECT m.*, u.username
                                    FROM macro_processing m
                                    JOIN users u ON m.user_id = u.id
                                    ORDER BY m.processing_date DESC LIMIT ? OFFSET ?''',
                                   (per_page, offset)).fetchall()

        total_count = db.execute("SELECT COUNT(*) FROM macro_processing").fetchone()[0]
        total_pages = (total_count + per_page - 1) // per_page

    return render_template("admin_macro_history.html",
                           macro_history=macro_history,
                           page=page,
                           total_pages=total_pages,
                           macro_names=ROUTE_MACROS.get('macro_processing', {}).get('macros', []))


# -----------------------
# Report Routes
# -----------------------
@app.route("/report/<filename>")
def download_report(filename):
    if 'user_id' not in session:
        return redirect(url_for('login'))

    safe_filename = secure_filename(filename)
    if safe_filename != filename:
        flash("Invalid filename", "error")
        return redirect(url_for('dashboard'))

    with db_pool.get_connection() as db:
        if session.get('is_admin'):
            file_exists = db.execute('SELECT 1 FROM files WHERE report_filename=?',
                                     (filename,)).fetchone()
        else:
            file_exists = db.execute('SELECT 1 FROM files WHERE report_filename=? AND user_id=?',
                                     (filename, session['user_id'])).fetchone()

        if not file_exists:
            flash("No permission to access this report", "error")
            return redirect(url_for('dashboard'))

    report_path = os.path.join(REPORT_FOLDER, filename)
    if not os.path.exists(report_path):
        flash("Report file not found", "error")
        return redirect(url_for('dashboard'))

    try:
        return send_from_directory(REPORT_FOLDER, filename, as_attachment=True, download_name=f"report_{filename}")
    except FileNotFoundError:
        flash("Report file could not be downloaded", "error")
        return redirect(url_for('dashboard'))


# -----------------------
# Reset Routes
# -----------------------
@app.route('/macro-reset', methods=['POST'])
def macro_reset_application():
    try:
        if os.path.exists(app.config['UPLOAD_FOLDER']):
            shutil.rmtree(app.config['UPLOAD_FOLDER'])
        os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
        download_tokens.clear()
        return jsonify({"success": True, "message": "All files are cleared"})
    except Exception as e:
        return jsonify({"success": False, "message": str(e)})


@app.route("/reset-db", methods=["POST"])
@admin_required
def reset_database():
    if os.path.exists(DATABASE):
        os.remove(DATABASE)
    init_db()
    return "Database reset successfully! New admin created: username='admin', password='admin123'"




# -----------------------
# Background Tasks
# -----------------------
def start_background_cleanup():
    def cleanup_worker():
        while True:
            try:
                cleanup_expired_tokens()
                # Audit stale futures (belt-and-suspenders)
                with batch_queue._lock:
                    stale = [jid for jid, f in batch_queue._active_futures.items()
                             if f is not None and f.done()]
                    for jid in stale:
                        del batch_queue._active_futures[jid]
                time.sleep(300)  # Run every 5 minutes
            except Exception as e:
                log_errors([f"Background cleanup error: {str(e)}"])

    cleanup_thread = threading.Thread(target=cleanup_worker, daemon=True)
    cleanup_thread.start()


# -----------------------
# Error Handlers
# -----------------------
from werkzeug.exceptions import NotFound, RequestEntityTooLarge

@app.errorhandler(413)
@app.errorhandler(RequestEntityTooLarge)
def handle_file_too_large(error):
    limit_mb = app.config.get('MAX_CONTENT_LENGTH', 0) // (1024 * 1024)
    if request.is_json or request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return jsonify({
            'error': f'File too large. Please upload a file smaller than {limit_mb} MB.',
            'code': 413
        }), 413
    return (
        f'<h3>File Too Large</h3>'
        f'<p>The uploaded file exceeds the {limit_mb}&nbsp;MB limit. '
        f'Please use a smaller file and try again.</p>'
        f'<a href="javascript:history.back()">&#8592; Go Back</a>'
    ), 413

@app.errorhandler(Exception)
def handle_unexpected_error(error):
    if isinstance(error, NotFound):
        return "Not Found", 404

    app.logger.error(f'Unexpected error: {error}')
    if app.debug:
        return str(error), 500
    return 'An unexpected error occurred', 500


# -----------------------
# Application Initialization
# -----------------------
def validate_route_configuration():
    errors = []

    for route_type, config in ROUTE_MACROS.items():
        # Credit Extractor doesn't use macros, so skip empty check for it
        if route_type != 'credit_extractor' and not config.get('macros'):
            # Just a warning now, not an error
            print(f"Warning: Route '{route_type}' has no macros defined")

        if not config.get('name'):
            errors.append(f"Route '{route_type}' has no name defined")

    # On Linux/Docker, we don't need the macro template file check
    if HAS_WIN32COM:
        macro_path = os.path.join(COMMON_MACRO_FOLDER, DEFAULT_MACRO_NAME)
        if not os.path.exists(macro_path):
            errors.append(f"Macro template file not found: {macro_path}")

    if errors:
        for error in errors:
            log_errors([f"Configuration error: {error}"])
        return False

    return True


def initialize_optimized_app():
    if not validate_route_configuration():
        print("Warning: Route configuration validation failed")

    # Initialize DB
    init_db()

    # 🔹 Ensure schema upgrades (e.g. add 'role' column if missing)
    try:
        migrate_add_role_column()
    except Exception as e:
        log_errors([f"Migration failed during startup: {e}"])

    setup_logging()
    start_background_cleanup()

    # Recover any jobs that were 'pending' or 'running' when server last died
    try:
        with db_pool.get_connection() as db:
            db.execute("UPDATE job_queue SET status='pending' WHERE status='running'")
            db.commit()
        batch_queue._drain_pending()
    except Exception as e:
        log_errors([f"Batch queue startup recovery failed: {e}"])

    # populate PPD macros into route configuration on startup (safe guard in case module missing)
    try:
        if hasattr(ppd, 'macro_names') and isinstance(ppd.macro_names, (list, tuple)):
            ROUTE_MACROS['ppd']['macros'] = ppd.macro_names
    except Exception as e:
        log_errors([f"Failed to load PPD macro names: {e}"])

    app.logger.info("Application initialized with route-specific macro processing")

    return app


from datetime import datetime, timezone

# Make sure _now_utc() and is_token_expired() are defined as above

def process_technical_job(job_id, unique_folder, saved_paths, original_filenames,
                          run_te, user_id, username):
    """Background worker for technical editing."""
    with app.app_context():
        def update_progress(updates):
            try:
                p_path = os.path.join(unique_folder, "progress.json")
                current = {}
                if os.path.exists(p_path):
                    with open(p_path, "r") as f:
                        current = json.load(f)
                current.update(updates)
                with open(p_path, "w") as f:
                    json.dump(current, f)
            except Exception as ex:
                print(f"Progress update failed: {ex}")

        update_progress({"total": len(saved_paths), "current": 0, "status": "Starting"})
        try:
            processed_files = []
            for idx, input_path in enumerate(saved_paths, 1):
                filename = original_filenames[idx - 1]
                update_progress({"current": idx, "status": f"Processing {filename}"})
                output_path = input_path
                if run_te:
                    print(f"[TECH] Processing Technical QA: {filename}")
                    process_docx(input_path, output_path, skip_validation=True)
                else:
                    shutil.copy(input_path, output_path)
                processed_files.append(filename)

            # Create ZIP
            zip_path = os.path.join(unique_folder, "Technical_Documents.zip")
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as z:
                for fname in processed_files:
                    fpath = os.path.join(unique_folder, fname)
                    if os.path.exists(fpath):
                        z.write(fpath, arcname=fname)

            # Save manifest
            with open(os.path.join(unique_folder, "manifest.txt"), "w") as mf:
                mf.write("\n".join(processed_files))

            # DB log
            try:
                with db_pool.get_connection() as db:
                    db.execute(
                        '''INSERT INTO macro_processing
                           (user_id, token, original_filenames, processed_filenames, selected_tasks, route_type)
                           VALUES (?, ?, ?, ?, ?, ?)''',
                        (user_id, job_id, json.dumps(original_filenames),
                         json.dumps(processed_files),
                         json.dumps({'route_type': 'technical', 'run_technical_editing': run_te,
                                     'task_indices': ['4'] if run_te else []}),
                         'technical')
                    )
                    db.commit()
            except Exception as e:
                print(f"DB Logging Error (Technical): {e}")

            update_progress({"status": "Completed", "zip_path": zip_path,
                             "download_token": job_id})
        except Exception as e:
            update_progress({"status": f"Failed: {str(e)}"})
            log_errors([f"Job {job_id} failed: {e}"])

@app.route('/technical', methods=['GET', 'POST'])
@role_required(ROUTE_PERMISSIONS.get('technical', ['ADMIN']))
def technical():

    if request.method == 'POST':
        uploaded_files = request.files.getlist("word_files[]")

        run_te = request.form.get("run_technical_editing") == "1"

        if not uploaded_files:
            return jsonify({"error": "No files uploaded"}), 400

        # Create ONE unique folder per job
        token = uuid.uuid4().hex
        unique_folder = os.path.join(UPLOAD_FOLDER, token)
        os.makedirs(unique_folder, exist_ok=True)

        saved_paths = []
        original_filenames = []

        for f in uploaded_files:
            filename = secure_filename(f.filename)
            if not filename:
                continue
            input_path = os.path.join(unique_folder, filename)
            f.save(input_path)
            saved_paths.append(input_path)
            original_filenames.append(f.filename)

        if not saved_paths:
            shutil.rmtree(unique_folder, ignore_errors=True)
            return jsonify({"error": "No valid files uploaded"}), 400

        # Initialize progress
        with open(os.path.join(unique_folder, "progress.json"), "w") as pf:
            json.dump({"total": len(saved_paths), "current": 0, "status": "Starting"}, pf)

        batch_queue.submit(
            job_id=token,
            route_type='technical',
            user_id=session['user_id'],
            username=session.get('username', 'unknown'),
            target_fn=process_technical_job,
            fn_args=(token, unique_folder, saved_paths, original_filenames,
                     run_te, session['user_id'], session.get('username', 'unknown')),
            payload_dict={
                'unique_folder': unique_folder,
                'saved_paths': saved_paths,
                'original_filenames': original_filenames,
                'run_te': run_te,
            }
        )

        return jsonify({"job_id": token})

    return render_template("technical_edit.html")



# =========================
# UTC-SAFE DATETIME HELPERS
# =========================

def _now_utc():
    """Return timezone-aware UTC datetime."""
    return datetime.now(timezone.utc)


def _ensure_utc(dt):
    """Convert naive datetime to UTC-aware."""
    if dt is None:
        return None
    if dt.tzinfo is None:
        return dt.replace(tzinfo=timezone.utc)
    return dt


def is_token_expired(token_info):
    """Safe expiration check for tokens."""
    expires = token_info.get("expires")
    if not expires:
        return True

    expires = _ensure_utc(expires)
    return _now_utc() > expires


# -----------------------
# Global Cleanup Helper for ALL Routes
# -----------------------
def cleanup_token_data(token):
    """
    Safely remove:
      - the temp folder
      - the zip file
      - the token entry from download_tokens
    Works for Technical, Macro, Language, and any other route using tokens.
    """
    try:
        token_info = download_tokens.get(token)
        if not token_info:
            return

        folder = token_info.get("path")

        # Remove processed folder
        if folder and os.path.isdir(folder):
            shutil.rmtree(folder)

        # Remove ZIP
        zip_path = folder + ".zip" if folder else None
        if zip_path and os.path.exists(zip_path):
            os.remove(zip_path)

        # Remove token entry
        download_tokens.pop(token, None)

    except Exception as e:
        log_errors([f"CLEANUP ERROR (token={token}): {str(e)}"])


@app.route('/favicon.ico')
def favicon():
    return send_from_directory(os.path.join(app.root_path, 'static', 'images'),
                               'S4c.png', mimetype='image/vnd.microsoft.icon')

@app.route('/technical/download/<token>')
def technical_download(token):
    # Try in-memory first
    info = download_tokens.get(token)
    
    # Fallback to file system for multi-worker support
    if not info:
        possible_folder = os.path.join(app.config['UPLOAD_FOLDER'], token)
        if os.path.exists(possible_folder):
            info = {
                "path": possible_folder,
                "expires": _now_utc() + TOKEN_TTL  # Renew/Assume valid if folder exists
            }

    if not info:
        flash("Invalid or expired token.")
        return redirect(url_for('dashboard'))

    # Check expiry only if we have explicit expiry info (memory)
    # If recovered from disk, we trust folder existence for now (or could check creation time)
    if 'expires' in info and is_token_expired(info):
        cleanup_token_data(token)
        flash("Token expired.")
        return redirect(url_for('dashboard'))

    folder = info["path"]
    zip_path = folder + ".zip"

    try:
        if not os.path.exists(zip_path):
            shutil.make_archive(folder, 'zip', folder)

        mem = io.BytesIO()
        with open(zip_path, 'rb') as f:
            mem.write(f.read())
        mem.seek(0)

    except Exception as e:
        log_errors([f"ZIP read/create failure: {e}"])
        flash("Download failed.")
        return redirect(url_for('dashboard'))

    # Optional: cleanup after download if one-time use
    # cleanup_token_data(token) 

    return send_file(
        mem,
        as_attachment=True,
        mimetype="application/zip",
        download_name=f"Technical_Documents_{_now_utc().strftime('%Y%m%d_%H%M%S')}.zip"
    )

# -----------------------
# Alttext Integration
# -----------------------

def _alttext_headers():
    return {
        "X-Service-Token": ALTTEXT_SERVICE_TOKEN,
        "X-Forwarded-For": request.remote_addr,
    }


@app.route("/alttext-api/<path:subpath>", methods=["GET", "POST", "PUT", "DELETE", "PATCH"])
@csrf.exempt
def alttext_proxy(subpath):
    if "user_id" not in session:
        return jsonify({"error": "Unauthorized"}), 401
    target = f"{ALTTEXT_URL}/{subpath}"
    if request.query_string:
        target += "?" + request.query_string.decode()
    headers = _alttext_headers()
    if request.content_type:
        headers["Content-Type"] = request.content_type
    try:
        resp = _requests_lib.request(
            method=request.method,
            url=target,
            headers=headers,
            data=request.get_data(),
            timeout=3600,
            stream=True,
        )
    except _requests_lib.exceptions.ConnectionError:
        return jsonify({"error": "Alttext service unavailable"}), 503
    ct = resp.headers.get("Content-Type", "application/json")
    skip = {"content-encoding", "transfer-encoding", "connection", "content-length"}
    fwd_headers = {k: v for k, v in resp.headers.items() if k.lower() not in skip}
    if any(x in ct for x in ("pdf", "octet-stream", "xlsx", "zip", "spreadsheet")):
        def _stream():
            for chunk in resp.iter_content(65536):
                if chunk:
                    yield chunk
        return app.response_class(_stream(), status=resp.status_code,
                                   headers=fwd_headers, content_type=ct)
    return (resp.content, resp.status_code, fwd_headers)


@app.route("/alttext/dashboard")
@role_required(ROUTE_PERMISSIONS.get('alttext', ['ADMIN']))
def alttext_dashboard():
    return render_template("alttext/dashboard.html", title="Alt Text Dashboard")


@app.route("/alttext/upload")
@role_required(ROUTE_PERMISSIONS.get('alttext', ['ADMIN']))
def alttext_upload():
    return render_template("alttext/upload.html", title="Process Documents")


@app.route("/alttext/batches")
@role_required(ROUTE_PERMISSIONS.get('alttext', ['ADMIN']))
def alttext_batches():
    return render_template("alttext/batches.html", title="Batch History")


@app.route("/alttext/batch/<int:batch_id>")
@role_required(ROUTE_PERMISSIONS.get('alttext', ['ADMIN']))
def alttext_batch_details(batch_id):
    return render_template("alttext/batch_details.html", batch_id=batch_id)


@app.route("/alttext/markup")
@role_required(ROUTE_PERMISSIONS.get('alttext', ['ADMIN']))
def alttext_markup():
    return render_template("alttext/mark.html", title="Markup Tool", no_padding=True)


@app.route("/alttext/review/<int:job_id>")
@role_required(ROUTE_PERMISSIONS.get('alttext', ['ADMIN']))
def alttext_review(job_id):
    try:
        r = _requests_lib.get(
            f"{ALTTEXT_URL}/api/job/{job_id}/data",
            headers=_alttext_headers(),
            timeout=10,
        )
        pdf_filename = r.json().get("pdf_filename", "") if r.ok else ""
    except Exception:
        pdf_filename = ""
    return render_template("alttext/review.html", job_id=job_id,
                           pdf_filename=pdf_filename, no_padding=True)


@app.route("/alttext/downloads")
@role_required(ROUTE_PERMISSIONS.get('alttext', ['ADMIN']))
def alttext_downloads():
    try:
        r = _requests_lib.get(
            f"{ALTTEXT_URL}/api/output-files",
            headers=_alttext_headers(),
            timeout=10,
        )
        files = r.json().get("files", []) if r.ok else []
    except Exception:
        files = []
    return render_template("alttext/download.html", title="Generated Outputs", files=files)


# -----------------------
# Main Execution
# -----------------------
from waitress import serve

# 🔹 create app globally so waitress-serve can see it
app = initialize_optimized_app()

if __name__ == '__main__':
    print("=== S4C APPLICATION STARTUP ===")
    host_ip = get_ip_address()
    print(f"Your IP address: {host_ip}")

    port = 8081

    print(f"\nAccess URLs:")
    print(f"Local: http://localhost:{port}")
    print(f"Network: http://{host_ip}:{port}")
    print("=================================\n")

    # run with waitress directly if launched via python
    serve(app, host="0.0.0.0", port=port, threads=4)