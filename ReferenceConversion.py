import os
import json
import logging
import re
from typing import Optional, Dict, List, Tuple
from pathlib import Path
from dataclasses import dataclass

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import copy

# ─────────────────────────────────────────────
# OPTIONAL spaCy NER INTEGRATION (lazy-loaded)
# ─────────────────────────────────────────────
_spacy_nlp = None
_SPACY_AVAILABLE = False

def _get_spacy_nlp():
    """Lazy-load spaCy model for NER. Returns None if unavailable."""
    global _spacy_nlp, _SPACY_AVAILABLE
    if _spacy_nlp is not None:
        return _spacy_nlp
    try:
        import spacy
        _spacy_nlp = spacy.load("en_core_web_sm")
        _SPACY_AVAILABLE = True
        logger.info("spaCy en_core_web_sm loaded for NER support in sentence casing.")
    except (ImportError, OSError) as e:
        _spacy_nlp = False  # sentinel: tried and failed
        logger.debug(f"spaCy NER unavailable ({e}); using dictionary-only approach.")
    return _spacy_nlp

from gemini_ref_converter import convert_reference, CitationStyle, BIB_FIELDS, CONVERSION_MAP, DEFAULT_MODEL

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s: %(message)s")
logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────
# JOURNAL ABBREVIATION EXPANSION (HYBRID)
# ─────────────────────────────────────────────
_JOURNAL_ABBREV_CACHE = None

def _load_journal_abbreviations() -> Dict[str, str]:
    """Load journal abbreviations from local JSON file."""
    global _JOURNAL_ABBREV_CACHE
    if _JOURNAL_ABBREV_CACHE is not None:
        return _JOURNAL_ABBREV_CACHE

    try:
        abbrev_file = Path(__file__).parent / "journal_abbreviations.json"
        if abbrev_file.exists():
            with open(abbrev_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
                _JOURNAL_ABBREV_CACHE = data.get("journal_abbreviations", {})
                logger.debug(f"Loaded {len(_JOURNAL_ABBREV_CACHE)} journal abbreviations")
                return _JOURNAL_ABBREV_CACHE
    except Exception as e:
        logger.debug(f"Failed to load journal abbreviations: {e}")

    _JOURNAL_ABBREV_CACHE = {}
    return _JOURNAL_ABBREV_CACHE

def _lookup_journal_local(abbrev: str) -> Optional[str]:
    """Look up journal abbreviation in local table (fast, cost-free)."""
    if not abbrev:
        return None
    abbr_dict = _load_journal_abbreviations()
    # Try exact lowercase match first
    abbrev_lower = abbrev.lower()
    result = abbr_dict.get(abbrev_lower)
    if result:
        return result
    # Try with trailing period removed if not found
    if abbrev_lower.endswith('.'):
        result = abbr_dict.get(abbrev_lower[:-1])
        if result:
            return result
    # Try with trailing period added if not found
    result = abbr_dict.get(abbrev_lower + '.')
    return result if result else None

def _expand_journal_via_gemini(abbrev: str) -> Optional[str]:
    """Ask Gemini to expand unknown journal abbreviation."""
    if not abbrev:
        return None
    try:
        from gemini_ref_converter import _build_prompt, _call_gemini_api
        prompt = f"""What is the full name of this academic journal abbreviation?
Abbreviation: {abbrev}

Return ONLY the full journal name, nothing else. If unknown, return the abbreviation unchanged."""

        response = _call_gemini_api(prompt, max_tokens=100)
        if response and response.strip():
            expanded = response.strip().strip('"\'')
            if expanded and expanded.lower() != abbrev.lower():
                logger.info(f"Gemini expanded '{abbrev}' -> '{expanded}'")
                return expanded
    except Exception as e:
        logger.debug(f"Gemini expansion failed for '{abbrev}': {e}")

    return None

def _expand_journal_hybrid(abbrev: str) -> Optional[str]:
    """Hybrid expansion: try local lookup first, then Gemini."""
    if not abbrev:
        return None

    # Try 1: Local lookup table (fast, free)
    expanded = _lookup_journal_local(abbrev)
    if expanded:
        logger.info(f"[Journal Abbrev] Local lookup: '{abbrev}' -> '{expanded}'")
        return expanded

    # Try 2: Gemini expansion (for unknown abbreviations)
    expanded = _expand_journal_via_gemini(abbrev)
    if expanded:
        return expanded

    return None

SKIP_SAME_STYLE = os.environ.get("REFERENCE_CONVERTER_SKIP_SAME_STYLE", "1").strip().lower() not in ("0", "false", "no")
MAX_CONVERSION_WORKERS = int(os.environ.get("REFERENCE_CONVERTER_MAX_WORKERS", "2"))

BIB_NUMBER_PATTERNS = [
    re.compile(r'^\s*\[(\d+)\]\.\s+'),
    re.compile(r'^\s*\((\d+)\)\.\s+'),
    re.compile(r'^\s*\[(\d+)\]\s+'),
    re.compile(r'^\s*\((\d+)\)\s+'),
    re.compile(r'^\s*(\d+)\.\s+'),
    re.compile(r'^\s*(\d+)\s+'),
]


def extract_bib_number(text):
    for pat in BIB_NUMBER_PATTERNS:
        m = pat.match(text)
        if m:
            return int(m.group(1)), pat.pattern
    return None, None


def detect_source_style(raw_text: str) -> CitationStyle:
    stripped = raw_text.strip()
    if re.match(r'^\[?\d+\]?\.?\s+', stripped):
        return CitationStyle.AMA
    if re.search(r'\.\s+\d{4};', stripped):
        return CitationStyle.AMA
    if re.search(r'\bpp?\.\s*\d+[-–]\d+', stripped):
        return CitationStyle.APA
    if re.search(r'\(\d{4}[a-z]?\)', stripped):
        return CitationStyle.APA
    if re.search(r'https://doi\.org/', stripped):
        return CitationStyle.APA
    if re.search(r',\s+\d+\(\d+\),\s+\d+', stripped):
        return CitationStyle.APA
    return CitationStyle.APA


def detect_ref_type_from_metadata(metadata: Dict) -> str:
    return metadata.get("bib_reftype") or "unknown"


def _blank_metadata() -> Dict:
    meta = {field: "" for field in BIB_FIELDS}
    meta["bib_reftype"] = "unknown"
    return meta


def _lookup_year(item: Dict) -> str:
    for date_key in ("published-print", "published-online", "issued", "created"):
        dp = (item.get(date_key) or {}).get("date-parts")
        if dp and dp[0] and dp[0][0]:
            return _format_ama_date_parts(dp[0])
    if item.get("year"):
        return str(item["year"])
    return ""


def _lookup_ref_type(item: Dict) -> str:
    raw_type = (item.get("type") or "").lower()
    if raw_type in ("journal-article", "journal"):
        return "journal"
    if raw_type in ("book",):
        return "book"
    if raw_type in ("book-chapter", "chapter"):
        return "book_chapter"
    if raw_type in ("proceedings-article", "conference-paper"):
        return "conference"
    if raw_type in ("dissertation", "thesis"):
        return "thesis"
    if raw_type in ("report",):
        return "report"
    if raw_type in ("website", "webpage", "posted-content"):
        return "website"
    if item.get("container-title") and (item.get("volume") or item.get("page")):
        return "journal"
    if item.get("publisher"):
        return "book"
    if item.get("URL"):
        return "website"
    return "unknown"


def _expand_page_range(start: str, end: str) -> Tuple[str, str]:
    start = str(start or "").strip()
    end = str(end or "").strip()
    if not start or not end:
        return start, end
    start_match = re.match(r'^([A-Za-z]*)(\d+)([A-Za-z]*)$', start)
    end_match = re.match(r'^([A-Za-z]*)(\d+)([A-Za-z]*)$', end)
    if not start_match or not end_match:
        return start, end

    s_prefix, s_digits, s_suffix = start_match.groups()
    e_prefix, e_digits, e_suffix = end_match.groups()
    if e_prefix or e_suffix or s_prefix != e_prefix or s_suffix != e_suffix:
        return start, end
    if len(e_digits) >= len(s_digits):
        return start, end

    expanded = s_digits[:len(s_digits) - len(e_digits)] + e_digits
    if int(expanded) < int(s_digits):
        carry_len = len(s_digits) - len(e_digits)
        if carry_len > 0:
            expanded = str(int(s_digits[:carry_len]) + 1) + e_digits
    return start, f"{e_prefix}{expanded}{e_suffix}"


def _page_parts(page: str) -> Tuple[str, str]:
    if not page:
        return "", ""
    page = str(page).strip()
    match = re.split(r"[-–—]", page, maxsplit=1)
    if len(match) == 2:
        return _expand_page_range(match[0].strip(), match[1].strip())
    return page, ""


_AMA_MONTHS = (
    "", "Jan", "Feb", "Mar", "Apr", "May", "Jun",
    "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"
)


def _format_ama_date_parts(parts: List) -> str:
    if not parts:
        return ""
    year = str(parts[0]) if len(parts) >= 1 and parts[0] else ""
    month = int(parts[1]) if len(parts) >= 2 and str(parts[1]).isdigit() else None
    day = int(parts[2]) if len(parts) >= 3 and str(parts[2]).isdigit() else None
    if not year:
        return ""
    if month and 1 <= month <= 12:
        month_str = _AMA_MONTHS[month]
        if day:
            return f"{year};{month_str} {day}"
        return f"{year};{month_str}"
    return year


def _metadata_from_lookup(item: Dict, target_style: CitationStyle) -> Dict:
    meta = _blank_metadata()
    ref_type = _lookup_ref_type(item)
    title_val = item.get("title") or ""
    title = title_val[0] if isinstance(title_val, list) and title_val else str(title_val or "")
    container_val = item.get("container-title") or ""
    container = container_val[0] if isinstance(container_val, list) and container_val else str(container_val or "")
    short_container_val = item.get("short-container-title") or ""
    short_container = short_container_val[0] if isinstance(short_container_val, list) and short_container_val else ""
    page = item.get("page") or ""
    fpage, lpage = _page_parts(page)
    authors = item.get("author") or []
    surnames = []
    fnames = []
    for author in authors:
        family = (author.get("family") or "").strip()
        given = (author.get("given") or "").strip()
        literal = (author.get("literal") or "").strip()
        if family or given:
            surnames.append(family or given)
            fnames.append(given if family else "")
        elif literal:
            surnames.append(literal)
            fnames.append("")
    meta["bib_reftype"] = ref_type
    meta["bib_title"] = title
    meta["bib_article"] = title
    meta["bib_book"] = title if ref_type in ("book", "edited_book") else ""
    meta["bib_journal"] = short_container if target_style == CitationStyle.AMA and short_container else container
    meta["bib_surname"] = "|".join(surnames)
    meta["bib_fname"] = "|".join(fnames)
    meta["bib_year"] = _lookup_year(item)
    meta["bib_volume"] = str(item.get("volume") or "")
    meta["bib_issue"] = str(item.get("issue") or "")
    meta["bib_fpage"] = fpage
    meta["bib_lpage"] = lpage
    meta["bib_doi"] = str(item.get("DOI") or "").replace("https://doi.org/", "").replace("doi:", "").strip()
    meta["bib_url"] = str(item.get("URL") or "").strip()
    meta["bib_publisher"] = str(item.get("publisher") or "")
    meta["bib_institution"] = str(item.get("institution") or item.get("publisher") or "")
    meta["bib_school"] = str(item.get("school") or "")
    meta["bib_conference"] = str(item.get("event") or item.get("container-title") or "")
    meta["bib_reportnum"] = str(item.get("report-number") or "")
    return meta


# ─────────────────────────────────────────────
# TITLE CASE / SENTENCE CASE HELPERS
# ─────────────────────────────────────────────

_PROPER_NOUNS = re.compile(
    r'\b('
    r'COVID(?:-19)?|HIV(?:/AIDS)?|AIDS|SARS(?:-CoV(?:-2)?)?|MERS|'
    r'COPD|PCR|DNA|RNA|mRNA|CT|MRI|ICU|ECG|EKG|'
    r'USA|UK|UN|EU|US|WHO|FDA|CDC|NIH|NHS|'
    r'English|French|German|Spanish|Italian|Chinese|Japanese|'
    r'Korean|Arabic|Russian|Portuguese|Dutch|Swedish|'
    r'American|European|Asian|African|Australian|Canadian|'
    r'British|Indian|Brazilian|Mexican|'
    r'Cochrane|PubMed|CrossRef|MEDLINE'
    r')\b',
    re.IGNORECASE,
)

_DOTTED_ABBREV_RE = re.compile(r'[A-Z](?:\.[A-Z])+\.?')

_ABBREV_BEFORE_PERIOD = re.compile(
    r'\b(?:e\.g|i\.e|et\s+al|Dr|Mr|Mrs|Ms|Prof|Rev|St|Fig|Tab|Vol|No|'
    r'pp|Suppl|Proc|Natl|Acad|Sci|approx|ca|cf|vs|ed|eds|'
    r'Jan|Feb|Mar|Apr|Jun|Jul|Aug|Sep|Oct|Nov|Dec|Ltd|Inc|Corp)\.',
    re.IGNORECASE,
)

# ─────────────────────────────────────────────
# PROPER NOUN DICTIONARY (user-extensible)
# ─────────────────────────────────────────────
_PROPER_NOUN_DICT: Optional[Dict] = None
_PROPER_NOUN_RE: Optional[re.Pattern] = None
_PROPER_NOUN_LOWERCASE: Optional[frozenset] = None
_PROPER_NOUN_CANONICAL: Optional[Dict[str, str]] = None

def _load_proper_noun_dict() -> Dict:
    """Load proper_nouns.json once and cache. Returns empty dict on failure."""
    global _PROPER_NOUN_DICT, _PROPER_NOUN_RE, _PROPER_NOUN_LOWERCASE, _PROPER_NOUN_CANONICAL
    if _PROPER_NOUN_DICT is not None:
        return _PROPER_NOUN_DICT

    json_path = Path(__file__).parent / "proper_nouns.json"
    try:
        with open(json_path, encoding="utf-8") as fh:
            data = json.load(fh)
    except (FileNotFoundError, json.JSONDecodeError) as e:
        logger.debug(f"proper_nouns.json not found or invalid ({e}); using built-in list only.")
        data = {}

    _PROPER_NOUN_DICT = data

    # Build combined regex: acronyms (exact case) + proper_nouns (title case)
    all_terms = (
        data.get("acronyms", []) +
        data.get("proper_nouns", []) +
        data.get("institution_names", []) +
        data.get("latin_phrases", []) +
        data.get("geographic_terms", [])
    )
    if all_terms:
        # Escape each term; sort longest-first to prevent partial matches
        escaped = sorted(
            (re.escape(t) for t in all_terms),
            key=len, reverse=True
        )
        _PROPER_NOUN_RE = re.compile(r'\b(?:' + '|'.join(escaped) + r')\b', re.IGNORECASE)
    else:
        _PROPER_NOUN_RE = None

    _PROPER_NOUN_LOWERCASE = frozenset(
        t.lower() for t in data.get("drug_names", [])
    )
    _PROPER_NOUN_CANONICAL = {t.lower(): t for t in all_terms}
    return _PROPER_NOUN_DICT

_SC_SMALL_WORDS = frozenset({
    "a", "an", "the", "and", "but", "or", "for", "nor", "on", "at",
    "to", "by", "in", "of", "up", "as", "is", "it", "its", "via",
    "per", "vs", "et", "with", "from", "into", "onto", "than",
})


def _to_sentence_case(text: str) -> str:
    """
    Sentence case conversion that preserves proper nouns and acronyms.

    - ALL CAPS (>70% uppercase): aggressive lowercase, then restore:
        1. Any fully-uppercase word >=3 chars in the original (FIX 3: e.g.
           SURUSS, MELAS, BJOG, MMWR — these are acronyms, not proper names)
        2. Mixed-case acronym patterns (e.g. DOHaD)
        3. Dotted abbreviations (U.S., U.K.)
        4. Known proper nouns whitelist + dictionary + NER
    - Title Case: only lowercase known small words; preserve all other
      capitalisation — naturally keeps proper nouns without a whitelist.
    """
    if not text:
        return text

    alpha_chars = [c for c in text if c.isalpha()]
    if not alpha_chars:
        return text

    _load_proper_noun_dict()  # idempotent — only loads once

    upper_ratio = sum(1 for c in alpha_chars if c.isupper()) / len(alpha_chars)
    is_all_caps = upper_ratio > 0.7

    if is_all_caps:
        # For all-caps input, only match truly mixed-case acronyms (lowercase→uppercase),
        # e.g. mRNA, DOHaD. Pure all-caps words are handled separately below.
        acronym_spans = []
        acronym_pattern = r'\b[a-z][A-Z][A-Za-z0-9]*\b'
        for m in re.finditer(acronym_pattern, text):
            acronym_spans.append((m.start(), m.end(), m.group()))
        for m in _DOTTED_ABBREV_RE.finditer(text):
            acronym_spans.append((m.start(), m.end(), m.group()))
        proper_spans = [(m.start(), m.end(), m.group()) for m in _PROPER_NOUNS.finditer(text)]

        # Collect short all-caps words (2–5 chars) that are likely true acronyms
        # e.g. HIV, DNA, PCR, BJOG, MELAS. Words longer than 5 chars are almost
        # always regular English words (MECHANISMS, ADVANCED, RESISTANCE…) that
        # should be lowercased. Common short words are excluded via the list below.
        _COMMON_TITLE_WORDS = frozenset({
            'the','and','for','are','but','not','you','all','can','had','her','was',
            'one','our','out','day','get','has','him','his','how','man','new','now',
            'old','see','two','way','who','boy','did','its','let','put','say','she',
            'too','use','study','review','case','report','data','analysis','methods',
            'results','discussion','conclusion','background','introduction','patients',
            'treatment','clinical','medical','health','disease','syndrome','approach',
            'risk','factor','effect','role','impact','outcomes','management','effects',
            'prevalence','incidence','mortality','morbidity','diagnosis','therapy',
            'surgery','screening','testing','monitoring','evaluation','assessment',
            'perspective','characteristics','features','findings','population','adults',
            'children','women','men','care','based','related','associated','induced',
            # Common 3–5 letter words that must NOT be restored as all-caps
            'age','aim','arm','cut','due','end','flu','gap','gut','hit','key','lab',
            'law','low','map','mix','non','per','pre','raw','red','run','sex','sub',
            'top','via','war','pain','cell','bone','skin','lung','gene','dose','drug',
            'form','flow','test','term','part','area','rate','site','type','mode',
            'base','oral','post','mild','high','deep','main','full','left','real',
            'open','free','need','make','move','stop','work','lack','help','fall',
            'hold','dead','cold','warm','flat','thin','vast','rare','safe','pure',
            'rich','weak','poor','dual','anti','semi','mono','poly','both','each',
            'such','well','long','late','fast','hard','slow','soft','wide','small',
            'acute','early','blood','brain','liver','heart','nasal','nerve','serum',
            'stage','phase','group','trial','level','index','scale','ratio','score',
            'value','local','focal','joint','novel','rapid','first','large','fluid',
            'under','above','after','below','inter','trans','extra','super','cross',
            'given','found','issue','field','since','other','these','those','their',
            'among','which','while','young','adult','human','model','mouse','total',
            'fatal','viral','toxic','solid','prior','proof','daily','brief','broad',
            'clear','close','colon','death','eight','three','seven','inner','outer',
            'whole','usual','valid','vital','range','route','urine','spleen','class',
        })
        allcaps_word_spans = [
            (m.start(), m.end(), m.group())
            for m in re.finditer(r'\b([A-Z]{2,5})\b', text)
            if m.group().lower() not in _COMMON_TITLE_WORDS
        ]

        result = text[0].upper() + text[1:].lower() if len(text) > 1 else text.upper()

        # Restore order: all-caps acronyms first, then finer patterns
        for start, end, word in allcaps_word_spans:
            result = result[:start] + word + result[end:]
        for start, end, word in acronym_spans:
            result = result[:start] + word + result[end:]
        for start, end, word in proper_spans:
            result = result[:start] + word + result[end:]

        # Extended proper noun dictionary from proper_nouns.json
        if _PROPER_NOUN_RE:
            for m in _PROPER_NOUN_RE.finditer(text):
                word = m.group()
                # Preserve case exactly as written in the dictionary entry
                result = result[:m.start()] + word + result[m.end():]

        # NER-based proper noun restoration (optional, only when spaCy available)
        nlp = _get_spacy_nlp()
        if nlp and isinstance(nlp, object) and nlp is not False and len(text) > 20:
            try:
                doc_ner = nlp(text)
                _NER_CAPITALIZE_LABELS = {"ORG", "GPE", "NORP", "FAC", "PRODUCT"}
                for ent in doc_ner.ents:
                    if ent.label_ in _NER_CAPITALIZE_LABELS:
                        ent_lower = ent.text.lower()
                        if ent_lower not in (_PROPER_NOUN_LOWERCASE or frozenset()):
                            # Re-apply original casing from the original text span
                            result = result[:ent.start_char] + ent.text + result[ent.end_char:]
            except Exception as ner_exc:
                logger.debug(f"NER processing failed ({ner_exc}); continuing with dictionary-only approach.")
    else:
        words = text.split(" ")
        result_words = []
        for i, word in enumerate(words):
            _is_mixed = len(word) > 1 and any(c.islower() for c in word[1:]) and any(c.isupper() for c in word[1:])
            if i == 0:
                if _is_mixed:
                    # Mixed-case brand name (StatPearls, UpToDate): ensure first letter
                    # is capitalised but preserve internal caps
                    result_words.append(word[0].upper() + word[1:])
                else:
                    result_words.append(word[0].upper() + word[1:].lower() if word else word)
            elif word.lower() in _SC_SMALL_WORDS:
                result_words.append(word.lower())
            elif word.isupper() and len(word) >= 2:
                # All-caps acronyms (HIV, DNA, BJOG): preserve
                result_words.append(word)
            elif _is_mixed:
                # Mixed-case brand names (StatPearls, UpToDate, mRNA): preserve
                result_words.append(word)
            else:
                result_words.append(word.lower())
        result = " ".join(result_words)

    # Capitalize after `:`, `.`, `?`, `!`, `;`, `—` with abbreviation disambiguation for `.`
    _abbrev_dots = (
        {m.end() - 1 for m in _ABBREV_BEFORE_PERIOD.finditer(result)}
        | {m.end() - 1 for m in _DOTTED_ABBREV_RE.finditer(result)}
    )

    def _cap_replacer(m):
        if m.group(1) == '.' and m.start(1) in _abbrev_dots:
            return m.group(0)
        return m.group(1) + m.group(2) + m.group(3).upper()

    result = re.sub(r'([:;—.?!])(\s+)([a-z])', _cap_replacer, result)
    return result


def _to_ama_sentence_case(text: str) -> str:
    """AMA titles keep sentence case but capitalize the first word after a colon."""
    text = _to_sentence_case(text)
    if not text:
        return text
    _load_proper_noun_dict()

    def _ama_colon_cap(match):
        return match.group(1) + match.group(2) + match.group(3).upper()

    text = re.sub(r'(:)(\s+)([a-z])', _ama_colon_cap, text)
    if _PROPER_NOUN_RE and _PROPER_NOUN_CANONICAL:
        def _restore_case(match):
            return _PROPER_NOUN_CANONICAL.get(match.group(0).lower(), match.group(0))
        text = _PROPER_NOUN_RE.sub(_restore_case, text)
    return text


def _split_italic_spans(
    text: str,
    italic_words: frozenset,
    style: Optional[str],
) -> List[Tuple[str, Optional[str]]]:
    """
    Split text into (chunk, style) segments where words matching italic_words
    get style 'bib_inline_italic' and the rest keep `style`.

    Args:
        text: The title or segment text to split
        italic_words: Frozenset of lowercase word stems to mark as italic
        style: The default style for non-italic segments (e.g., "bib_article")

    Returns:
        List of (chunk, style) tuples for appending to segments
    """
    if not italic_words:
        return [(text, style)]
    segs = []
    last = 0
    for m in re.finditer(r'\S+', text):
        bare = m.group().lower().strip('.,;:!?"\'()-')
        if bare and bare in italic_words:
            if m.start() > last:
                segs.append((text[last:m.start()], style))
            segs.append((m.group(), "bib_inline_italic"))
            last = m.end()
    if last < len(text):
        segs.append((text[last:], style))
    return segs if segs else [(text, style)]


_TITLE_CASE_SMALL = frozenset({
    "a","an","the","and","but","or","for","nor","on","at",
    "to","by","in","of","up","as","is","it","its","via","per","vs","et",
})


def _to_title_case(text: str) -> str:
    if not text:
        return text
    words = text.split()
    result = []
    for i, w in enumerate(words):
        if len(w) >= 2 and w.isupper():
            result.append(w)
        elif i == 0 or w.lower() not in _TITLE_CASE_SMALL:
            # Preserve mixed-case words like mRNA, CrossRef, DOHaD
            if len(w) > 1 and any(c.islower() for c in w[1:]) and any(c.isupper() for c in w[1:]):
                result.append(w)
            else:
                result.append(w[0].upper() + w[1:].lower() if len(w) > 1 else w.upper())
        else:
            result.append(w.lower())
    return " ".join(result)


# ─────────────────────────────────────────────
# PUBLISHER SUFFIX STRIPPER
# ─────────────────────────────────────────────

_PUB_SUFFIX_RE = re.compile(
    r',?\s+(?:Co\.|Ltd\.?|Limited|Inc\.?|LLC|L\.L\.C\.|Corp\.?|'
    r'GmbH|S\.A\.|Pvt\.?|Pty\.?|(?:Pty|Pvt)\.?\s+Ltd\.?)\s*$',
    re.IGNORECASE,
)


def _strip_publisher_suffixes(pub: str) -> str:
    if not pub:
        return pub
    cleaned = _PUB_SUFFIX_RE.sub("", pub).strip().rstrip(",").strip()
    return cleaned or pub


def _normalise_quotes(text: str) -> str:
    return (
        text
        .replace('\u2018', "'").replace('\u2019', "'")
        .replace('\u201c', '"').replace('\u201d', '"')
    )


def _to_curly_quotes(text: str) -> str:
    """Convert straight ASCII double-quotes to typographic curly quotes.
    Skips content inside URLs (avoids mangling query strings).
    """
    if '"' not in text:
        return text
    result = []
    for i, ch in enumerate(text):
        if ch == '"':
            prev = text[i - 1] if i > 0 else ' '
            result.append('\u201c' if (prev in ' \t\n\r([{' or i == 0) else '\u201d')
        else:
            result.append(ch)
    return ''.join(result)


def _normalise_double_periods(text: str) -> str:
    if not text:
        return text
    return re.sub(r'(?<!\.)\.\.(?!\.)', '.', text)


# ─────────────────────────────────────────────
# REF-TYPE HEURISTIC CORRECTOR
# ─────────────────────────────────────────────

def _fix_ref_type(meta: Dict, raw_text: str) -> Dict:
    rt = (meta.get("bib_reftype") or "").lower()
    fixed = dict(meta)
    if rt == "book" and fixed.get("bib_journal") and fixed.get("bib_volume"):
        fixed["bib_reftype"] = "journal"
        logger.info("  [TypeFix] 'book' → 'journal'  (has journal+volume)")
    elif rt == "book" and re.search(r'\d{4}\s*;\s*\d+[\s(:]', raw_text):
        fixed["bib_reftype"] = "journal"
        logger.info("  [TypeFix] 'book' → 'journal'  (year;volume pattern)")
    elif rt == "book" and re.search(r',\s*\*?\d+\*?\s*\(\d+\)\s*,\s*\d+', raw_text):
        fixed["bib_reftype"] = "journal"
        logger.info("  [TypeFix] 'book' → 'journal'  (APA volume(issue),page)")
    rt2 = (fixed.get("bib_reftype") or "").lower()
    if (rt2 in ("book", "journal") and fixed.get("bib_chaptertitle") and
            (fixed.get("bib_ed_surname") or re.search(r'\bIn[:\s]', raw_text))):
        fixed["bib_reftype"] = "book_chapter"
        logger.info("  [TypeFix] → 'book_chapter'")
    rt3 = (fixed.get("bib_reftype") or "").lower()
    if (rt3 == "book" and fixed.get("bib_ed_surname") and
            not fixed.get("bib_surname") and not fixed.get("bib_chaptertitle")):
        fixed["bib_reftype"] = "edited_book"
        logger.info("  [TypeFix] → 'edited_book'  (editors, no authors, no chapter)")
    if (fixed.get("bib_reftype", "book") == "book" and
            re.search(r'\b(?:eds?)\.\s+\w', raw_text, re.IGNORECASE) and
            not fixed.get("bib_surname")):
        fixed["bib_reftype"] = "edited_book"
        logger.info("  [TypeFix] → 'edited_book'  (ed./eds. in raw text)")
    rt4 = (fixed.get("bib_reftype") or "").lower()
    if (rt4 == "book" and (fixed.get("bib_conference") or
            re.search(r'\b(?:presented\s+at|proceedings\s+of|annual\s+(?:meeting|conference))\b',
                      raw_text, re.IGNORECASE))):
        fixed["bib_reftype"] = "conference"
        logger.info("  [TypeFix] → 'conference'")
    rt5 = (fixed.get("bib_reftype") or "").lower()
    if rt5 == "website" and fixed.get("bib_editionno"):
        fixed["bib_reftype"] = "book"
        logger.info("  [TypeFix] 'website' → 'book'  (has edition number)")
    rt6 = (fixed.get("bib_reftype") or "").lower()
    _raw_pub = (fixed.get("bib_publisher") or fixed.get("bib_institution") or "").strip()
    if rt6 == "website" and _raw_pub and not _raw_pub.lower().startswith("http"):
        fixed["bib_reftype"] = "report"
        logger.info("  [TypeFix] 'website' → 'report'  (has publisher/institution)")
    return fixed


_INLINE_CITATION_RE = re.compile(r'^\s*\([^)]{1,80}\)\s*[.,]?\s*$')


def _looks_like_inline_citation(text: str) -> bool:
    if _INLINE_CITATION_RE.match(text):
        return True
    if re.match(r'^\s*[\[\(]?\d[\d,\s\-–]+[\]\)]?\s*\.?\s*$', text):
        return True
    return False


# ─────────────────────────────────────────────
# FORMATTING FROM METADATA  (fallback)
# ─────────────────────────────────────────────

def format_apa_from_metadata(meta: Dict) -> str:
    return format_apa_from_metadata_old(meta)


def format_apa_from_metadata_old(meta: Dict) -> str:
    ref_type = meta.get("bib_reftype", "journal")
    parts = []
    surnames = [s.strip() for s in (meta.get("bib_surname") or "").split("|") if s.strip()]
    fnames   = [f.strip() for f in (meta.get("bib_fname")   or "").split("|") if f.strip()]
    has_etal = False
    if surnames and surnames[-1].lower().replace(".", "").strip() == "et al":
        has_etal = True
        surnames.pop()
        if len(fnames) > len(surnames): fnames.pop()
    authors  = []
    for i, surname in enumerate(surnames):
        initial = fnames[i] if i < len(fnames) else ""
        initials_fmt = " ".join(f"{p[0]}." for p in initial.split() if p) if initial else ""
        authors.append(f"{surname}, {initials_fmt}".strip(", "))
    if authors:
        if len(authors) > 20:
            author_str = ", ".join(authors[:19]) + ", ... " + authors[-1]
        elif has_etal:
            author_str = ", ".join(authors) + ", ..."
        elif len(authors) > 1:
            author_str = ", ".join(authors[:-1]) + ", & " + authors[-1]
        else:
            author_str = authors[0]
        parts.append(author_str + ".")
    year = meta.get("bib_year", "n.d.")
    parts.append(f"({year}).")
    if ref_type == "journal":
        title   = meta.get("bib_article", "")
        journal = _to_title_case(meta.get("bib_journal", ""))
        volume  = meta.get("bib_volume", "")
        issue   = meta.get("bib_issue", "")
        fpage   = meta.get("bib_fpage", "")
        lpage   = meta.get("bib_lpage", "")
        doi     = meta.get("bib_doi", "")
        if title:   parts.append(f"{_to_sentence_case(title)}.")
        vol_issue = f"*{journal}*" if journal else ""
        if volume:  vol_issue += f", *{volume}*"
        if issue:   vol_issue += f"({issue})"
        pages = f"{fpage}–{lpage}" if fpage and lpage else fpage or lpage
        if pages:   vol_issue += f", {pages}"
        if vol_issue: parts.append(vol_issue + ".")
        if doi:     parts.append(f"https://doi.org/{doi}")
    elif ref_type in ("book", "edited_book"):
        book_title = meta.get("bib_book") or ""
        edition    = meta.get("bib_editionno", "")
        publisher  = _strip_publisher_suffixes(meta.get("bib_publisher", ""))
        doi        = meta.get("bib_doi", "")
        url        = meta.get("bib_url", "")
        title_str  = f"*{_to_sentence_case(book_title)}*" if book_title else ""
        if edition and _ordinal(edition) not in ("1st", "1", "first"):
            title_str += f" ({_ordinal(edition)} ed.)"
        if title_str: parts.append(title_str + ".")
        if publisher: parts.append(publisher + ".")
        if doi:       parts.append(f"https://doi.org/{doi}")
        elif url:     parts.append(url)
    elif ref_type == "book_chapter":
        chapter   = meta.get("bib_chaptertitle") or ""
        book      = meta.get("bib_book", "")
        edition   = meta.get("bib_editionno", "")
        volume    = meta.get("bib_volume", "")
        fpage     = meta.get("bib_fpage", "")
        lpage     = meta.get("bib_lpage", "")
        publisher = _strip_publisher_suffixes(meta.get("bib_publisher", ""))
        doi       = meta.get("bib_doi", "")
        ed_surnames = [s.strip() for s in (meta.get("bib_ed_surname") or "").split("|") if s.strip()]
        ed_fnames   = [f.strip() for f in (meta.get("bib_ed_fname")   or "").split("|") if f.strip()]
        if chapter: parts.append(f"{_to_sentence_case(chapter)}.")
        editors = []
        for i, s in enumerate(ed_surnames):
            ini = ed_fnames[i] if i < len(ed_fnames) else ""
            initials_fmt = _format_initials_apa(ini)
            editor_name = f"{initials_fmt} {s}".strip() if initials_fmt else s
            editors.append(editor_name)
        ed_label = "Ed." if len(editors) == 1 else "Eds."
        if len(editors) > 1:
            in_str = "In " + ", ".join(editors[:-1]) + f" & {editors[-1]} ({ed_label}), " if editors else "In "
        else:
            in_str = "In " + ", ".join(editors) + f" ({ed_label}), " if editors else "In "
        book_str = f"*{_to_sentence_case(book)}*" if book else ""
        inner = []
        if edition and _ordinal(edition) not in ("1st", "1", "first"):
            inner.append(_ordinal(edition) + " ed.")
        if volume:
            inner.append("Vol. " + volume)
        if fpage:
            inner.append("pp. " + fpage + ("–" + lpage if lpage else ""))
        paren = " (" + ", ".join(inner) + ")" if inner else ""
        parts.append(in_str + book_str + paren + ".")
        if publisher: parts.append(publisher + ".")
        if doi:       parts.append(f"https://doi.org/{doi}")
    elif ref_type == "thesis":
        title  = meta.get("bib_title", "")
        deg    = meta.get("bib_deg", "Doctoral dissertation")
        school = meta.get("bib_school", "")
        url    = meta.get("bib_url", "")
        if title:
            bracket = f" [{deg}, {school}]" if school else f" [{deg}]"
            parts.append(f"*{_to_sentence_case(title)}*{bracket}.")
        if url: parts.append(url)
    elif ref_type == "conference":
        title    = meta.get("bib_title", "")
        conf     = meta.get("bib_conference", "")
        confloc  = meta.get("bib_conflocation", "")
        confdate = meta.get("bib_confdate", "")
        doi      = meta.get("bib_doi", "")
        if title: parts.append(f"*{_to_sentence_case(title)}* [Conference session].")
        conf_str = conf
        if confdate: conf_str += f", {confdate}"
        if confloc:  conf_str += f", {confloc}"
        if conf_str: parts.append(conf_str + ".")
        if doi: parts.append(f"https://doi.org/{doi}")
    elif ref_type in ("website", "ereference"):
        title    = meta.get("bib_title", "")
        site     = meta.get("bib_journal") or meta.get("bib_book", "")
        accessed = _clean_accessed(meta.get("bib_accessed", ""))
        url      = meta.get("bib_url", "")
        if title:    parts.append(f"{_to_sentence_case(title)}.")
        if site and not _is_url(site):
            parts.append(f"*{_to_title_case(site)}*.")
        if url:      parts.append(url)
    elif ref_type == "report":
        title  = meta.get("bib_title", "")
        repnum = meta.get("bib_reportnum", "")
        inst   = _strip_publisher_suffixes(meta.get("bib_institution", ""))
        doi    = meta.get("bib_doi", "")
        url    = meta.get("bib_url", "")
        title_str = f"*{_to_sentence_case(title)}*" if title else ""
        if repnum: title_str += f" (Report No. {repnum})"
        if title_str: parts.append(title_str + ".")
        if inst:      parts.append(inst + ".")
        if doi:       parts.append(f"https://doi.org/{doi}")
        elif url:     parts.append(url)
    return " ".join(parts)


def format_ama_from_metadata(meta: Dict) -> str:
    return format_ama_from_metadata_old(meta)


def format_ama_from_metadata_old(meta: Dict) -> str:
    ref_type = meta.get("bib_reftype", "journal")
    parts = []
    surnames = [s.strip() for s in (meta.get("bib_surname") or "").split("|") if s.strip()]
    fnames   = [f.strip() for f in (meta.get("bib_fname")   or "").split("|") if f.strip()]
    has_etal = False
    if surnames and surnames[-1].lower().replace(".", "").strip() == "et al":
        has_etal = True
        surnames.pop()
        if len(fnames) > len(surnames): fnames.pop()
    authors  = []
    for i, surname in enumerate(surnames):
        initial = fnames[i] if i < len(fnames) else ""
        if initial and len(initial) <= 3 and " " not in initial and not any(c.islower() for c in initial):
            initials_fmt = initial.replace(".", "")
        else:
            initials_fmt = "".join(p[0] for p in initial.split() if p) if initial else ""
        authors.append(f"{surname} {initials_fmt}".strip())
    if authors:
        if len(authors) <= 6 and not has_etal:
            author_str = ", ".join(authors)
        else:
            author_str = ", ".join(authors[:6]) + ", et al."
        parts.append(author_str + ".")
    if ref_type == "journal":
        title   = meta.get("bib_title", "")
        journal = meta.get("bib_journal", "")
        year    = meta.get("bib_year", "")
        volume  = meta.get("bib_volume", "")
        issue   = meta.get("bib_issue", "")
        fpage   = meta.get("bib_fpage", "")
        lpage   = meta.get("bib_lpage", "")
        doi     = meta.get("bib_doi", "")
        if title:   parts.append(f"{_to_sentence_case(title)}.")
        vol_str = journal or ""
        if year:    vol_str += f". {year}"
        if volume:  vol_str += f";{volume}"
        if issue:   vol_str += f"({issue})"
        pages = f"{fpage}-{lpage}" if fpage and lpage else fpage or lpage
        if pages:   vol_str += f":{pages}"
        if vol_str: parts.append(vol_str + ".")
        if doi:     parts.append(f"doi:{doi}")
    elif ref_type in ("book", "edited_book"):
        book_title = meta.get("bib_book") or meta.get("bib_title", "")
        edition    = meta.get("bib_editionno", "")
        publisher  = _strip_publisher_suffixes(meta.get("bib_publisher", ""))
        year       = meta.get("bib_year", "")
        doi        = meta.get("bib_doi", "")
        url        = meta.get("bib_url", "")
        title_str  = _to_sentence_case(book_title) if book_title else ""
        if edition and _ordinal(edition) not in ("1st", "1", "first"):
            title_str += f". {_ordinal(edition)} ed."
        if title_str: parts.append(title_str + ".")
        if publisher and year:
            parts.append(publisher + ";")
            parts.append(year + ".")
        elif publisher:
            parts.append(publisher + ".")
        if doi:       parts.append(f"doi:{doi}")
        elif url:     parts.append(url)
    elif ref_type == "book_chapter":
        chapter   = meta.get("bib_chaptertitle") or ""
        book      = meta.get("bib_book", "")
        edition   = meta.get("bib_editionno", "")
        fpage     = meta.get("bib_fpage", "")
        lpage     = meta.get("bib_lpage", "")
        publisher = _strip_publisher_suffixes(meta.get("bib_publisher", ""))
        year      = meta.get("bib_year", "")
        doi       = meta.get("bib_doi", "")
        ed_surnames = [s.strip() for s in (meta.get("bib_ed_surname") or "").split("|") if s.strip()]
        ed_fnames   = [f.strip() for f in (meta.get("bib_ed_fname")   or "").split("|") if f.strip()]
        if chapter: parts.append(f"{_to_sentence_case(chapter)}.")
        editors = []
        for i, s in enumerate(ed_surnames):
            ini = ed_fnames[i] if i < len(ed_fnames) else ""
            initials_fmt = "".join(p[0] for p in ini.split() if p) if ini else ""
            editors.append(f"{s} {initials_fmt}".strip())
        ed_label = "ed." if len(editors) == 1 else "eds."
        in_str = "In: " + ", ".join(editors) + f", {ed_label}. " if editors else "In: "
        book_str = _to_sentence_case(book) if book else ""
        if edition and _ordinal(edition) not in ("1st", "1", "first"):
            book_str += f". {_ordinal(edition)} ed"
        parts.append(in_str + book_str + ".")
        if publisher: parts.append(publisher + ";")
        if year:      parts.append(year + ".")
        pages = f"{fpage}-{lpage}" if fpage and lpage else fpage or lpage
        if pages:     parts[-1] = parts[-1].rstrip(".") + f":{pages}."
        if doi:       parts.append(f"doi:{doi}")
    elif ref_type == "thesis":
        title  = meta.get("bib_title", "")
        deg    = meta.get("bib_deg", "doctoral dissertation")
        school = meta.get("bib_school", "")
        year   = meta.get("bib_year", "")
        url    = meta.get("bib_url", "")
        if title:  parts.append(f"{_to_sentence_case(title)} [{deg}].")
        if school: parts.append(school + ";")
        if year:   parts.append(year + ".")
        if url:    parts.append(url)
    elif ref_type == "conference":
        title    = meta.get("bib_title", "")
        conf     = meta.get("bib_conference", "")
        confloc  = meta.get("bib_conflocation", "")
        confdate = meta.get("bib_confdate", "")
        doi      = meta.get("bib_doi", "")
        if title: parts.append(f"{_to_sentence_case(title)}.")
        conf_str = f"Paper presented at: {conf}" if conf else ""
        if confdate: conf_str += f"; {confdate}"
        if confloc:  conf_str += f"; {confloc}"
        if conf_str: parts.append(conf_str + ".")
        if doi:      parts.append(f"doi:{doi}")
    elif ref_type == "website":
        title    = meta.get("bib_title", "")
        site     = meta.get("bib_journal") or meta.get("bib_book", "")
        year     = meta.get("bib_year", "")
        accessed = _clean_accessed(meta.get("bib_accessed", ""))
        url      = meta.get("bib_url", "")
        if title:    parts.append(f"{_to_sentence_case(title)}.")
        if site and not _is_url(site): parts.append(f"{site}.")
        if year:     parts.append(f"Published {year}.")
        if accessed: parts.append(f"Accessed {accessed}.")
        if url:      parts.append(url)
    elif ref_type == "ereference":
        title       = meta.get("bib_title", "")
        book        = meta.get("bib_book", "") or meta.get("bib_journal", "")
        publisher   = _strip_publisher_suffixes(meta.get("bib_publisher", "") or meta.get("bib_institution", ""))
        year        = meta.get("bib_year", "")
        accessed    = _clean_accessed(meta.get("bib_accessed", ""))
        url         = meta.get("bib_url", "")
        ed_surnames = [s.strip() for s in (meta.get("bib_ed_surname") or "").split("|") if s.strip()]
        ed_fnames   = [f.strip() for f in (meta.get("bib_ed_fname")   or "").split("|") if f.strip()]
        if title: parts.append(f"{_to_sentence_case(title)}.")
        in_str = "In:"
        if ed_surnames:
            editors = []
            for i, s in enumerate(ed_surnames):
                ini = ed_fnames[i] if i < len(ed_fnames) else ""
                initials_fmt = "".join(p[0] for p in ini.split() if p) if ini else ""
                editors.append(f"{s} {initials_fmt}".strip())
            ed_label = "ed." if len(editors) == 1 else "eds."
            in_str += " " + ", ".join(editors) + f", {ed_label}."
        parts.append(in_str)
        if book:      parts.append(f"{book}.")
        if publisher: parts.append(f"{publisher};")
        if year:      parts.append(f"{year}.")
        if accessed:  parts.append(f"Accessed {accessed}.")
        if url:       parts.append(url)
    elif ref_type == "report":
        title    = meta.get("bib_title", "")
        repnum   = meta.get("bib_reportnum", "")
        inst     = _strip_publisher_suffixes(meta.get("bib_institution", ""))
        year     = meta.get("bib_year", "")
        doi      = meta.get("bib_doi", "")
        url      = meta.get("bib_url", "")
        accessed = _clean_accessed(meta.get("bib_accessed", ""))
        if title:    parts.append(f"{_to_sentence_case(title)}.")
        if inst:     parts.append(inst + ";")
        if year:     parts.append(year + ".")
        if repnum:   parts.append(f"Report No. {repnum}.")
        if doi:      parts.append(f"doi:{doi}")
        elif url:
            if accessed: parts.append(f"Accessed {accessed}.")
            parts.append(url)
    return " ".join(parts)


def _ordinal(n: str) -> str:
    try:
        m = re.search(r'^(\d+)', str(n).strip())
        if m:
            n_int = int(m.group(1))
            suffix = {1:"st",2:"nd",3:"rd"}.get(
                n_int % 10 if n_int % 100 not in (11,12,13) else 0, "th")
            return f"{n_int}{suffix}"
    except Exception:
        pass
    clean = str(n).lower().replace("edition","").replace("ed.","").replace("ed","").strip()
    return clean

# ─────────────────────────────────────────────
# PARAGRAPH FORMATTING HELPERS
# ─────────────────────────────────────────────

def _clear_paragraph_text(para) -> None:
    p_elem = para._p
    for r in list(p_elem.findall(qn("w:r"))):
        p_elem.remove(r)
    for hyperlink in p_elem.findall(qn("w:hyperlink")):
        for r in hyperlink.findall(qn("w:r")):
            for t in r.findall(qn("w:t")):
                t.text = ""


def _ensure_style(doc, styles, style_name):
    if style_name and styles is not None:
        try:
            from docx.enum.style import WD_STYLE_TYPE
            if style_name not in styles:
                new_style = styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
            else:
                new_style = styles[style_name]

            # For hyperlink styles, always ensure color and underline are present
            if style_name in ("bib_doi", "bib_url"):
                rPr = new_style.element.get_or_add_rPr()
                # Add color if not present
                color = rPr.find(qn("w:color"))
                if color is None:
                    color = OxmlElement("w:color")
                    color.set(qn("w:val"), "0563C1")
                    color.set(qn("w:themeColor"), "hyperlink")
                    rPr.append(color)
                # Add underline if not present
                u = rPr.find(qn("w:u"))
                if u is None:
                    u = OxmlElement("w:u")
                    u.set(qn("w:val"), "single")
                    rPr.append(u)

            return styles[style_name]
        except Exception:
            return style_name
    return style_name


_ITALIC_STYLES = {
    "bib_journal",
    "bib_book",
    "bib_title",
    "bib_volume",
    "bib_inline_italic",  # for within-title italic spans (e.g., *in vitro*)
}


def _append_hyperlink(para, text: str, url: str, doc=None, style_name: Optional[str] = None, italic: bool = False):
    part = para.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    if italic:
        ital = OxmlElement("w:i")
        rpr.append(ital)
    if style_name:
        style_val = _ensure_style(doc, doc.styles if doc is not None else None, style_name)
        style_id = style_val.name if hasattr(style_val, "name") else style_name
        rstyle = OxmlElement("w:rStyle")
        rstyle.set(qn("w:val"), style_id)
        rpr.append(rstyle)
    new_run.append(rpr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    para._p.append(hyperlink)
    return hyperlink


def _append_segment(para, text: str, style_name: Optional[str], doc=None, styles=None, allow_hyperlink: bool = True):
    if not text:
        return
    italic = style_name in _ITALIC_STYLES
    target = _style_link_target(style_name, text) if allow_hyperlink else ""
    if target:
        try:
            _append_hyperlink(para, text, target, doc=doc, style_name=style_name, italic=italic)
            return
        except Exception:
            logger.warning(f"Hyperlink insertion failed for '{text[:40]}'")
    run = para.add_run(text)
    if style_name:
        style_val = _ensure_style(doc, styles, style_name)
        try:
            run.style = style_val
        except Exception:
            pass
    if italic:
        run.italic = True


def _write_styled_runs(para, segments: List[Tuple[str, Optional[str]]], doc=None, original_text: str = None, is_conversion: bool = False) -> None:
    if original_text is None:
        original_text = para.text
    _clear_paragraph_text(para)
    styles = doc.styles if doc is not None else None

    prefix_num = ""
    prefix_sep = ""
    remaining_text = original_text

    bib_num, pattern = extract_bib_number(original_text)
    if bib_num is not None:
        prefix_num = str(bib_num)
        match = None
        for pat in BIB_NUMBER_PATTERNS:
            match = pat.match(original_text)
            if match:
                break
        if match:
            full_match = match.group(0)
            prefix_sep = full_match[len(match.group(match.lastindex if match.lastindex else 1)):]
            remaining_text = original_text[len(full_match):]

    if prefix_num:
        run = para.add_run(prefix_num)
        style_val = _ensure_style(doc, styles, "bib_number")
        try:
            run.style = style_val
        except Exception:
            pass
    if prefix_sep:
        para.add_run(prefix_sep)

    original_text = remaining_text

    try:
        from utils.track_changes import add_tracked_deletion, add_tracked_text
        use_track_changes = True
    except ImportError:
        use_track_changes = False

    if not use_track_changes:
        for text, style_name in segments:
            _append_segment(para, text, style_name, doc=doc, styles=styles, allow_hyperlink=True)
        return

    import difflib

    new_full_text = ""
    style_map = []
    for text, style_name in segments:
        if not text: continue
        new_full_text += text
        style_map.extend([style_name] * len(text))

    matcher = difflib.SequenceMatcher(None, original_text, new_full_text)

    for opcode, i1, i2, j1, j2 in matcher.get_opcodes():
        if opcode == 'equal':
            segment_text   = new_full_text[j1:j2]
            segment_styles = style_map[j1:j2]
            chunk_start = 0
            for k in range(len(segment_text) + 1):
                is_end        = (k == len(segment_text))
                style_changed = (k > 0 and k < len(segment_text) and segment_styles[k] != segment_styles[k-1])
                if is_end or style_changed:
                    chunk = segment_text[chunk_start:k]
                    if chunk:
                        style = segment_styles[chunk_start]
                        _append_segment(para, chunk, style, doc=doc, styles=styles, allow_hyperlink=True)
                    chunk_start = k
        elif opcode == 'delete':
            deleted_chunk = original_text[i1:i2]
            add_tracked_deletion(para, deleted_chunk, doc=doc, author="S4C Reference Converter")
        elif opcode in ('insert', 'replace'):
            if opcode == 'replace':
                deleted_chunk = original_text[i1:i2]
                add_tracked_deletion(para, deleted_chunk, doc=doc, author="S4C Reference Converter")
            segment_text   = new_full_text[j1:j2]
            segment_styles = style_map[j1:j2]
            chunk_start = 0
            for k in range(len(segment_text) + 1):
                is_end        = (k == len(segment_text))
                style_changed = (k > 0 and k < len(segment_text) and segment_styles[k] != segment_styles[k-1])
                if is_end or style_changed:
                    chunk = segment_text[chunk_start:k]
                    if chunk:
                        style = segment_styles[chunk_start]
                        if style:
                            _ensure_style(doc, styles, style)
                        try:
                            add_tracked_text(para, chunk, style=style, author="S4C Reference Converter", doc=doc)
                        except Exception:
                            _append_segment(para, chunk, style, doc=doc, styles=styles, allow_hyperlink=False)
                    chunk_start = k


def _write_cgrn_runs(para, text: str, doc=None, original_text: str = None) -> None:
    if original_text is None:
        original_text = para.text
    _clear_paragraph_text(para)

    import re as _re
    pattern = _re.compile(r'\*\*(.+?)\*\*|\*(.+?)\*')
    segments: List[Tuple[str, bool]] = []
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            segments.append((text[last:m.start()], False))
        if m.group(1) is not None:
            segments.append((m.group(1), False))
        else:
            segments.append((m.group(2), True))
        last = m.end()
    if last < len(text):
        segments.append((text[last:], False))

    try:
        from utils.track_changes import add_tracked_deletion, add_tracked_text
        use_tc = True
    except ImportError:
        use_tc = False

    if use_tc:
        new_plain = "".join(t for t, _ in segments)
        import difflib
        matcher = difflib.SequenceMatcher(None, original_text, new_plain)
        italic_map = []
        for t, ital in segments:
            italic_map.extend([ital] * len(t))
        for opcode, i1, i2, j1, j2 in matcher.get_opcodes():
            if opcode == 'equal':
                chunk = new_plain[j1:j2]
                if chunk:
                    run = para.add_run(chunk)
                    if any(italic_map[j1:j2]):
                        run.italic = True
            elif opcode == 'delete':
                add_tracked_deletion(para, original_text[i1:i2], doc=doc, author="S4C Reference Converter")
            elif opcode in ('insert', 'replace'):
                if opcode == 'replace':
                    add_tracked_deletion(para, original_text[i1:i2], doc=doc, author="S4C Reference Converter")
                chunk = new_plain[j1:j2]
                if chunk:
                    try:
                        run_obj = add_tracked_text(para, chunk, style=None, author="S4C Reference Converter", doc=doc)
                    except Exception:
                        run_obj = para.add_run(chunk)
                    if hasattr(run_obj, 'italic') and any(italic_map[j1:j2]):
                        run_obj.italic = True
    else:
        for seg_text, is_italic in segments:
            if not seg_text:
                continue
            run = para.add_run(seg_text)
            if is_italic:
                run.italic = True


def _set_paragraph_text(para, text: str, doc=None, original_text: str = None, is_conversion: bool = False) -> None:
    if original_text is None:
        original_text = para.text
    _clear_paragraph_text(para)
    styles = doc.styles if doc is not None else None

    prefix_text = ""
    remaining_text = original_text

    bib_num, pattern = extract_bib_number(original_text)
    if bib_num is not None:
        for pat in BIB_NUMBER_PATTERNS:
            match = pat.match(original_text)
            if match:
                prefix_text = match.group(0)
                remaining_text = original_text[len(prefix_text):]
                break

    if prefix_text:
        run = para.add_run(prefix_text)
        style_val = _ensure_style(doc, styles, "bib_number")
        try:
            run.style = style_val
        except Exception:
            pass

    original_text = remaining_text

    try:
        from utils.track_changes import add_tracked_deletion, add_tracked_text
        import difflib
        matcher = difflib.SequenceMatcher(None, original_text, text)
        for opcode, i1, i2, j1, j2 in matcher.get_opcodes():
            if opcode == 'equal':
                _append_segment(para, text[j1:j2], None, doc=doc, styles=styles, allow_hyperlink=False)
            elif opcode == 'delete':
                add_tracked_deletion(para, original_text[i1:i2], author="S4C Reference Converter", doc=doc)
            elif opcode in ('insert', 'replace'):
                if opcode == 'replace':
                    add_tracked_deletion(para, original_text[i1:i2], author="S4C Reference Converter", doc=doc)
                add_tracked_text(para, text[j1:j2], author="S4C Reference Converter", doc=doc)
    except ImportError:
        last = 0
        pattern = re.compile(r'(https?://\S+|doi:\s*10\.\d{4,9}/[-._;()/:A-Za-z0-9]+|https://doi\.org/10\.\d{4,9}/[-._;()/:A-Za-z0-9]+)', re.IGNORECASE)
        for m in pattern.finditer(text):
            if m.start() > last:
                _append_segment(para, text[last:m.start()], None, doc=doc, styles=styles, allow_hyperlink=False)
            token = m.group(1).rstrip(".,;)")
            token_style = "bib_doi" if (("doi" in token.lower() and "http" not in token.lower()) or "doi.org/" in token.lower()) else "bib_url"
            visible = token
            if token_style == "bib_doi" and token.lower().startswith("doi:"):
                visible = "doi:" + _normalize_doi_value(token)
            _append_segment(para, visible, token_style, doc=doc, styles=styles, allow_hyperlink=True)
            trailing = m.group(1)[len(m.group(1).rstrip(".,;)")):]
            if trailing:
                _append_segment(para, trailing, None, doc=doc, styles=styles, allow_hyperlink=False)
            last = m.end()
        if last < len(text):
            _append_segment(para, text[last:], None, doc=doc, styles=styles, allow_hyperlink=False)


# ─────────────────────────────────────────────
# DB JOURNAL NAME QUALIFIER STRIPPER
# ─────────────────────────────────────────────

_DB_QUALIFIER_PATTERN = re.compile(
    r'\s+\([A-Z][\w\s]+,\s+[A-Z][\w\s]+(:\s*\d{4})?\)'
)


def _strip_db_journal_qualifiers(raw_source: str, metadata: dict, final_text: str) -> tuple:
    journal = (metadata.get("bib_journal") or "").strip()
    if not journal:
        return metadata, final_text
    m = _DB_QUALIFIER_PATTERN.search(journal)
    if m and m.group(0).strip() not in raw_source:
        clean_journal = journal[:m.start()].strip()
        logger.info(f"  [JournalFix] Stripped DB qualifier: '{journal}' → '{clean_journal}'")
        metadata = dict(metadata)
        metadata["bib_journal"] = clean_journal
        bad_suffix = m.group(0)
        final_text = final_text.replace(journal, clean_journal)
        escaped = re.escape(bad_suffix.strip())
        final_text = re.sub(r'\s*' + escaped, '', final_text)
    return metadata, final_text


# ─────────────────────────────────────────────
# GEMINI OUTPUT PARSER
# ─────────────────────────────────────────────

def _fix_ama_author_format(text: str) -> str:
    """Fix AMA author format: 'Surname, Initial' → 'Surname Initial' (no comma)."""
    # Match pattern: Surname, FM at the start or after period/semicolon/colon/comma
    # This fixes Gemini output that incorrectly adds commas between surname and initials
    def fix_author(match):
        prefix = match.group(1) or ""
        surname = match.group(2)
        initials = match.group(3)
        return f"{prefix}{surname} {initials}"

    # Pattern: (start|period|semicolon|colon|comma) + space? + Surname, INITIALS
    # Matches all positions where authors might appear (including after other authors)
    # Handles: lowercase particles (van, de), all-caps surnames, mixed-case, apostrophes, and initials with periods
    text = re.sub(
        r'(^|[\.\;\:,][\s]?)((?:[a-z]{1,4} )?[A-Z][a-zA-Z\']+(?:\s+[a-z]{1,4})*),\s+([A-Z]{1,3}(?:\.[A-Z])*\.?)',
        fix_author,
        text,
        flags=re.MULTILINE
    )
    return text


def _truncate_ama_authors(text: str) -> str:
    """Truncate AMA authors to first 3 if 7+ authors are present.

    AMA style rule: Show first 3 + et al. for 7+ authors.
    Example: "Smith JA, Jones BC, Brown CD, Williams DE, Green EF, Black GH, White HI, et al."
    becomes: "Smith JA, Jones BC, Brown CD, et al."
    """
    # Pattern: Surname (with optional particles) followed by 1-3 capital letters (initials)
    # Handles: P, JA, ABC, with or without periods between them
    author_pattern = r'([A-Z][a-z]*(?:\s+[a-z]{1,4})?\s+[A-Z][A-Z]?[A-Z]?\.?)'

    # Count comma-separated authors at the start
    author_list = []
    pos = 0
    while pos < len(text):
        match = re.match(author_pattern, text[pos:])
        if match:
            author_name = match.group(1).strip()
            if author_name.lower() != 'et al':
                author_list.append(author_name)
            pos += len(match.group(1))
            # Skip comma and whitespace
            if text[pos:pos+1] == ',':
                pos += 1
                while text[pos:pos+1] == ' ':
                    pos += 1
            else:
                break
        else:
            break

    if len(author_list) > 6:
        # 7+ authors: keep first 3, add et al.
        first_three = ', '.join(author_list[:3])
        # Find where the author list ended and insert truncated version
        # Count how many characters were consumed
        author_section = ', '.join(author_list) + ', ' if author_list else ''
        author_section_len = re.match(r'^((?:[A-Z][a-z]*(?:\s+[a-z]+)?\s+[A-Z]+\.?(?:\s*,\s*)?)+)', text)
        if author_section_len:
            pos = len(author_section_len.group(1))
            remainder = text[pos:].lstrip(', ')
            # If remainder already starts with 'et al', don't add it again
            if remainder.lower().startswith('et al'):
                text = first_three + ', ' + remainder
            else:
                text = first_three + ', et al. ' + remainder

    return text


def _truncate_apa_authors(text: str) -> str:
    """Truncate APA authors to first 19 + ... + last if 21+ authors are present.

    APA style rule: Show first 19 ... last for 21+ authors.
    Example: "Smith, A. A., Jones, B. B., ..., Williams, Z. Z."
    """
    # Pattern for APA format: Lastname, Initial(s).
    # Handles: "Smith, A.", "Jones, A. B.", "Brown, A. B. C."
    author_pattern = r'([A-Z][a-z]*(?:\s+[a-z]{1,4})?\s*,\s*[A-Z]\.(?:\s+[A-Z]\.)*)'

    # Count comma-separated authors at the start
    author_list = []
    pos = 0
    while pos < len(text):
        match = re.match(author_pattern, text[pos:])
        if match:
            author_name = match.group(1).strip()
            author_list.append(author_name)
            pos += len(match.group(1))
            # Skip comma, ampersand, and whitespace
            ampersand_comma = text[pos:pos+3] if pos + 3 <= len(text) else ""
            if ampersand_comma.startswith(', &'):
                pos += 3
                while text[pos:pos+1] == ' ':
                    pos += 1
            elif text[pos:pos+1] == ',':
                pos += 1
                while text[pos:pos+1] == ' ':
                    pos += 1
            else:
                break
        else:
            break

    if len(author_list) > 20:
        # 21+ authors: keep first 19, add ellipsis, add last author
        first_nineteen = ', '.join(author_list[:19])
        last_author = author_list[-1]

        # Find where the author list ended
        author_section_match = re.match(
            r'^((?:[A-Z][a-z]*(?:\s+[a-z]+)?\s*,\s*[A-Z]\.(?:\s+[A-Z]\.)*(?:\s*,\s*)?)+)',
            text
        )
        if author_section_match:
            pos = len(author_section_match.group(1))
            remainder = text[pos:].lstrip(', &').lstrip()
            # Build new author section with ellipsis
            text = first_nineteen + ', ... ' + last_author + '. ' + remainder

    return text


def _fix_duplicate_title_website(text: str, title: str) -> str:
    """Remove duplicate title from website references.
    Example: 'Org. Title. Title. Website.' → 'Org. Title. Website.'
    """
    if not title or len(title) < 3:
        return text

    # For website references, look for pattern: title appears twice in sequence
    # Pattern: ". Title. Title." or ". Title. Title. Site" or "Title. Title."
    title_escaped = re.escape(title.strip().rstrip("."))

    # Match variations: "Title. Title." or "Title. Title " (with various spacing)
    patterns = [
        (f'({title_escaped})\\.\\s+\\1\\.', r'\1.'),  # "Title. Title."
        (f'({title_escaped})\\.\\s+\\1\\s', r'\1. '),  # "Title. Title "
    ]

    for pattern, replacement in patterns:
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)

    return text


def _fix_double_periods(text: str) -> str:
    """Fix double periods (..) that sometimes appear in output."""
    text = re.sub(r'\.{2,}', '.', text)
    return text


def _remove_reference_numbers(text: str) -> str:
    """Remove reference numbers from start of text.
    Example: '4. Physical activity...' → 'Physical activity...'
    This handles cases where numbering was accidentally included in the reference.
    """
    # Match: number(s), period, space at the start of the text
    text = re.sub(r'^\d+\.\s+', '', text)
    return text


def _parse_gemini_output_to_segments(text: str) -> List[Tuple[str, Optional[str]]]:
    raw_segs: List[Tuple[str, Optional[str]]] = []
    pattern = re.compile(r'\*\*(.+?)\*\*|\*(.+?)\*')
    last = 0
    for m in pattern.finditer(text):
        start, end = m.start(), m.end()
        if start > last:
            raw_segs.append((text[last:start], None))
        if m.group(1) is not None:
            raw_segs.append((m.group(1), "bib_bold"))
        else:
            raw_segs.append((m.group(2), "bib_journal"))
        last = end
    if last < len(text):
        raw_segs.append((text[last:], None))

    PAGE_RANGE = re.compile(
        r'([A-Za-z]?\d+[A-Za-z0-9]*)\s*[\u2013\u2014-]\s*([A-Za-z]?\d+[A-Za-z0-9]*)'
    )
    segs: List[Tuple[str, Optional[str]]] = []
    for seg_text, seg_style in raw_segs:
        if seg_style is not None or not seg_text:
            segs.append((seg_text, seg_style))
            continue
        last_pos = 0
        for pm in PAGE_RANGE.finditer(seg_text):
            before = seg_text[last_pos:pm.start()]
            if before:
                segs.append((before, None))
            raw_fpage = pm.group(1)
            raw_lpage = pm.group(2)
            dash_start = pm.start() + len(raw_fpage)
            dash_end   = pm.end() - len(raw_lpage)
            dash = seg_text[dash_start:dash_end].strip() or '\u2013'
            fpage, lpage = _expand_page_range(raw_fpage, raw_lpage)
            segs.append((fpage, "bib_fpage"))
            segs.append((dash, None))
            segs.append((lpage, "bib_lpage"))
            last_pos = pm.end()
        remainder = seg_text[last_pos:]
        if remainder:
            segs.append((remainder, None))
    return segs


# ─────────────────────────────────────────────
# SEGMENT BUILDERS
# ─────────────────────────────────────────────

def _is_organization(name: str) -> bool:
    if not name: return False
    keywords = {
        "committee", "group", "task force", "section", "association",
        "society", "department", "national", "center", "institute",
        "world health", "collaborative", "network", "council",
        "board", "organization", "agency", "university", "college"
    }
    lower_name = name.lower()
    return any(kw in lower_name for kw in keywords) or len(name.split()) > 3


def _split_pipe(value: Optional[str]) -> List[str]:
    if not value:
        return []
    return [v.strip() for v in str(value).split("|") if v.strip()]


def _format_initials_ama(initial: str) -> str:
    if not initial:
        return ""
    clean = str(initial).strip()
    if not clean:
        return ""
    if re.fullmatch(r"[A-Z](?:\.?[A-Z]){0,5}\.?", clean):
        return "".join(c.upper() for c in clean if c.isalpha())

    tokens = [
        tok for tok in re.split(r"[\s,.\-]+", clean)
        if tok and tok.lower() not in _NAME_SUFFIXES
    ]
    if tokens:
        return "".join(tok[0].upper() for tok in tokens if tok[0].isalpha())
    return "".join(c.upper() for c in clean if c.isalpha())


_NAME_SUFFIXES = frozenset({"jr","sr","ii","iii","iv","2nd","3rd","4th"})


def _format_initials_apa(initial: str) -> str:
    if not initial:
        return ""
    comma_parts  = [p.strip() for p in initial.split(",")]
    suffix_parts: list = []
    name_section = comma_parts[0]
    for part in comma_parts[1:]:
        cleaned = part.rstrip(".").lower()
        if cleaned in _NAME_SUFFIXES:
            suffix_parts.append(cleaned.capitalize() + "." if cleaned in {"jr","sr"} else part.strip())
        else:
            name_section += " " + part
    if any(c.islower() for c in name_section):
        clean     = re.sub(r"[^a-zA-Z\s]", " ", name_section)
        formatted = " ".join(w[0].upper() + "." for w in clean.split() if w)
    else:
        letters   = [c.upper() for c in name_section if c.isalpha()]
        formatted = " ".join(c + "." for c in letters)
    if suffix_parts:
        return formatted + ", " + " ".join(suffix_parts)
    return formatted


_REQUIRED_FIELDS_BY_TYPE: Dict[str, List[str]] = {
    "journal":      ["bib_title", "bib_journal", "bib_year"],
    "book":         ["bib_year"],
    "edited_book":  ["bib_year"],
    "book_chapter": ["bib_book"],
    "website":      ["bib_title", "bib_url"],
    "ereference":   ["bib_title"],
    "thesis":       ["bib_title", "bib_school"],
    "report":       ["bib_title"],
    "conference":   ["bib_title"],
}


@dataclass
class ValidationFinding:
    code: str
    message: str
    severity: str = "warning"
    fixable: bool = False


_DOI_CORE_RE = re.compile(r'10\.\d{4,9}/[-._;()/:A-Za-z0-9]+', re.IGNORECASE)
_URL_RE = re.compile(r'^https?://\S+$', re.IGNORECASE)


def _normalize_doi_value(value: str) -> str:
    raw = str(value or "").strip()
    if not raw:
        return ""
    raw = re.sub(r'^(?:https?://(?:dx\.)?doi\.org/|doi:\s*)', '', raw, flags=re.IGNORECASE).strip()
    m = _DOI_CORE_RE.search(raw)
    if not m:
        return ""
    return m.group(0).rstrip(".,;)")


def _normalize_url_value(value: str) -> str:
    raw = str(value or "").strip()
    if not raw:
        return ""
    raw = re.sub(r'^(?:available\s+from|retrieved\s+from|from)\s*:?\s*', '', raw, flags=re.IGNORECASE).strip()
    raw = _clean_url(raw)
    return raw if _URL_RE.match(raw) else ""


def _style_link_target(style_name: Optional[str], text: str) -> str:
    if style_name == "bib_doi":
        doi = _normalize_doi_value(text)
        return f"https://doi.org/{doi}" if doi else ""
    if style_name == "bib_url":
        return _normalize_url_value(text)
    return ""


def _normalize_reference_metadata(metadata: Optional[Dict]) -> Dict:
    meta = dict(metadata or {})
    meta["bib_doi"] = _normalize_doi_value(meta.get("bib_doi", ""))
    meta["bib_url"] = _normalize_url_value(meta.get("bib_url", ""))
    if meta.get("bib_fpage") and meta.get("bib_lpage"):
        meta["bib_fpage"], meta["bib_lpage"] = _expand_page_range(meta["bib_fpage"], meta["bib_lpage"])
    return meta


def _add_finding(findings: List[ValidationFinding], code: str, message: str, severity: str = "warning", fixable: bool = False) -> None:
    findings.append(ValidationFinding(code=code, message=message, severity=severity, fixable=fixable))


def _validate_converted_reference(
    metadata: Optional[Dict],
    final_text: str,
    target_style: str,
    ref_type: str,
) -> List[ValidationFinding]:
    findings: List[ValidationFinding] = []
    if not final_text or len(final_text.strip()) < 15:
        _add_finding(findings, "empty_output", "conversion produced no output", "error")
        return findings
    if not metadata:
        _add_finding(findings, "metadata_missing", "metadata extraction failed", "error")
        return findings
    if ref_type in ("unknown", ""):
        _add_finding(findings, "unknown_type", "reference type could not be determined", "warning")
    for field in _REQUIRED_FIELDS_BY_TYPE.get(ref_type, []):
        if not metadata.get(field):
            _add_finding(findings, f"missing_{field}", f"{field} missing", "warning")
    has_author = bool(
        metadata.get("bib_surname") or
        metadata.get("bib_organization") or
        metadata.get("bib_ed_surname")
    )
    if not has_author and ref_type not in ("website", "ereference"):
        _add_finding(findings, "author_missing", "author missing", "warning")

    text_without_urls = re.sub(r'https?://\S+', '', final_text)
    if re.search(r'(?<!\.)\.\.(?!\.)', text_without_urls):
        _add_finding(findings, "double_period", "double period in output", "warning", True)

    doi = _normalize_doi_value(metadata.get("bib_doi", ""))
    url = _normalize_url_value(metadata.get("bib_url", ""))
    if metadata.get("bib_doi") and not doi:
        _add_finding(findings, "doi_invalid", "DOI is malformed", "warning", True)
    if metadata.get("bib_url") and not url:
        _add_finding(findings, "url_invalid", "URL is malformed", "warning", True)

    if target_style == "APA":
        if ref_type not in ("website", "ereference", "thesis") and not re.search(r'\(\d{4}', final_text):
            _add_finding(findings, "apa_year_format", "year format incorrect", "warning", True)
        if doi and f"https://doi.org/{doi}" not in final_text:
            _add_finding(findings, "apa_doi_format", "APA DOI should use https://doi.org/ form", "warning", True)
        if ref_type in ("website", "ereference") and url and url not in final_text:
            _add_finding(findings, "apa_url_missing", "URL missing from APA reference", "warning", True)
    elif target_style == "AMA":
        if ref_type == "journal" and not re.search(r'\d{4}(?:;[A-Za-z]{3}(?: \d{1,2})?)?;', final_text):
            _add_finding(findings, "ama_year_volume", "year;volume format incorrect", "warning", True)
        if doi and f"doi:{doi}".lower() not in final_text.lower():
            _add_finding(findings, "ama_doi_format", "AMA DOI should use doi: form", "warning", True)
        if ref_type in ("website", "ereference", "report") and url and url not in final_text:
            _add_finding(findings, "ama_url_missing", "URL missing from AMA reference", "warning", True)

    return findings


def _rebuild_reference_from_metadata(metadata: Dict, target_style: str, gemini_text: str = "", italic_words: frozenset = frozenset()) -> Tuple[str, List[Tuple[str, Optional[str]]]]:
    if target_style == "AMA":
        segs = build_segments_ama(metadata, gemini_text)
    elif target_style == "APA":
        segs = build_segments_apa(metadata, gemini_text, italic_words=italic_words)
    else:
        return "", []
    return "".join(t for t, _ in segs), segs


def _auto_correct_reference(
    metadata: Optional[Dict],
    final_text: str,
    target_style: str,
    ref_type: str,
    gemini_text: str = "",
    italic_words: frozenset = frozenset(),
) -> Tuple[Dict, str, List[ValidationFinding], List[Tuple[str, Optional[str]]]]:
    meta = _normalize_reference_metadata(metadata)
    findings = _validate_converted_reference(meta, final_text, target_style, ref_type)
    fixable_codes = {f.code for f in findings if f.fixable}
    segs: List[Tuple[str, Optional[str]]] = []

    if target_style in ("APA", "AMA") and fixable_codes:
        rebuilt_text, rebuilt_segs = _rebuild_reference_from_metadata(meta, target_style, gemini_text, italic_words)
        if rebuilt_text:
            final_text = rebuilt_text
            segs = rebuilt_segs

    final_text = _normalise_quotes(_normalise_double_periods(final_text))
    findings = _validate_converted_reference(meta, final_text, target_style, ref_type)
    return meta, final_text, findings, segs


def _add_review_comment(doc, para, comment_text: str) -> None:
    try:
        from docx.text.run import Run as _Run
        runs = para.runs
        if not runs:
            return
        # Anchor to the year run (first run whose text contains a 4-digit year).
        # Falls back to runs[0] only when no year is found in any run.
        year_run = next(
            (r for r in runs if re.search(r'\d{4}', r.text)),
            runs[0]
        )
        doc.add_comment(
            runs=year_run,
            text=comment_text,
            author="S4C Reference Converter",
            initials="S4C",
        )
    except Exception as exc:
        logger.warning(f"Could not add review comment: {exc}")


def build_segments_ama(meta: Dict, gemini_text: str = "") -> List[Tuple[str, Optional[str]]]:
    segs: List[Tuple[str, Optional[str]]] = []
    ref_type = (meta.get("bib_reftype") or "journal").lower()

    if ref_type == "book" and meta.get("bib_chaptertitle") and meta.get("bib_fpage"):
        ref_type = "book_chapter"
    if ref_type == "book" and meta.get("bib_ed_surname") and not meta.get("bib_surname"):
        ref_type = "edited_book"

    surnames    = _split_pipe(meta.get("bib_surname"))
    fnames      = _split_pipe(meta.get("bib_fname"))
    has_etal = False
    if surnames and surnames[-1].lower().replace(".", "").strip() == "et al":
        has_etal = True
        surnames.pop()
        if len(fnames) > len(surnames): fnames.pop()
    n_auth      = len(surnames)
    ed_surnames = _split_pipe(meta.get("bib_ed_surname") or meta.get("bib-ed-surname"))
    ed_fnames   = _split_pipe(meta.get("bib_ed_fname")   or meta.get("bib-ed-fname"))

    if n_auth == 0:
        org = (meta.get("bib_organization") or
               meta.get("bib_institution") or
               (meta.get("bib_surname") if not meta.get("bib_fname") else "") or "")
        if not org:
            if ref_type in ("book", "edited_book", "report"):
                org = (meta.get("bib_publisher") or "").strip()
        if org:
            segs.append((org.rstrip("."), "bib_organization"))
            segs.append((".", None))
        elif ed_surnames and ref_type != "book_chapter":
            for i, es in enumerate(ed_surnames):
                if i > 0: segs.append((", ", None))
                segs.append((es, "bib_ed-surname"))
                ei = ed_fnames[i] if i < len(ed_fnames) else ""
                ei_str = _format_initials_ama(ei)
                if ei_str:
                    segs.append((" ", None))
                    segs.append((ei_str, "bib_ed-fname"))
            ed_label = "ed." if len(ed_surnames) == 1 else "eds."
            segs.append((f", {ed_label}", None))
    else:
        subset = surnames if n_auth <= 6 else surnames[:3]
        for i, surname in enumerate(subset):
            if i > 0: segs.append((", ", None))
            initial = fnames[i] if i < len(fnames) else ""
            initials_str = _format_initials_ama(initial)
            if not initials_str and _is_organization(surname):
                disp_name = surname[0].upper() + surname[1:] if i == 0 and surname else surname
                segs.append((disp_name, "bib_organization"))
            else:
                disp_name = surname[0].upper() + surname[1:] if i == 0 and surname else surname
                segs.append((disp_name, "bib_surname"))
                if initials_str:
                    segs.append((" ", None))
                    segs.append((initials_str, "bib_fname"))
        # Only append et al. if we have more than 6 authors
        # The has_etal flag from source doesn't matter if database enrichment gave us <=6 authors
        if n_auth > 6:
            segs.append((", ", None))
            segs.append(("et al.", "bib_etal"))
    if segs:
        last_text = segs[-1][0].rstrip()
        if last_text.endswith("."):
            segs.append((" ", None))
        else:
            segs.append((". ", None))

    chapter_title = meta.get("bib_chaptertitle") or ""
    main_title    = meta.get("bib_title") or ""
    book_title    = meta.get("bib_book") or ""

    def _ama_title(text: str) -> str:
        return _to_ama_sentence_case(text.rstrip("."))

    if ref_type == "book_chapter" and chapter_title:
        clean_title = _ama_title(chapter_title)
        segs.append((clean_title, "bib_chaptertitle"))
        segs.append((" " if re.search(r'[?!]$', clean_title) else ". ", None))
    elif main_title:
        # AMA book/edited_book titles use Title Case; other types use Sentence Case
        if ref_type in ("book", "edited_book"):
            clean_title = _to_title_case(main_title.rstrip("."))
        else:
            clean_title = _ama_title(main_title)
        segs.append((clean_title, "bib_article" if ref_type == "journal" else "bib_title"))
        segs.append((" " if re.search(r'[?!]$', clean_title) else ". ", None))

    if ref_type == "book_chapter":
        segs.append(("In: ", None))
        if ed_surnames:
            for i, es in enumerate(ed_surnames):
                if i > 0: segs.append((", ", None))
                ei = ed_fnames[i] if i < len(ed_fnames) else ""
                ei_str = _format_initials_ama(ei)
                if not ei_str and _is_organization(es):
                    segs.append((es, "bib_organization"))
                else:
                    segs.append((es, "bib_ed-surname"))
                    if ei_str:
                        segs.append((" ", None))
                        segs.append((ei_str, "bib_ed-fname"))
            ed_label = ", ed." if len(ed_surnames) == 1 else ", eds."
            segs.append((ed_label + " ", None))
        display_book = book_title or main_title or ""
        if display_book:
            # AMA book titles use Title Case
            segs.append((_to_title_case(display_book), "bib_book"))
            segs.append((". ", None))

    if ref_type == "journal":
        journal = meta.get("bib_journal") or ""
        year    = meta.get("bib_year") or ""
        volume  = meta.get("bib_volume") or ""
        issue   = meta.get("bib_issue") or ""
        fpage   = meta.get("bib_fpage") or ""
        lpage   = meta.get("bib_lpage") or ""
        if journal:
            segs.append((journal, "bib_journal"))
            segs.append((".", None))
        if year:
            segs.append((" ", None))
            segs.append((year, "bib_year"))
        if volume:
            if year and not year.rstrip().endswith(";"):
                segs.append((";", None))
            segs.append((volume, "bib_volume"))
        if issue:
            segs.append(("(", None))
            segs.append((issue, "bib_issue"))
            segs.append((")", None))
        if fpage:
            segs.append((":", None))
            segs.append((fpage, "bib_fpage"))
            if lpage:
                segs.append(("-", None))
                segs.append((lpage, "bib_lpage"))
        elif not volume and not issue and "Published online" in gemini_text:
            segs.append((". Published online", None))
        segs.append((".", None))

    elif ref_type in ("book", "edited_book", "book_chapter"):
        edition   = meta.get("bib_editionno") or ""
        publisher = _strip_publisher_suffixes(meta.get("bib_publisher") or "")
        year      = meta.get("bib_year") or ""
        if ref_type != "book_chapter" and book_title:
            segs.append((_to_sentence_case(book_title), "bib_book"))
            segs.append((". ", None))
        if edition and _ordinal(edition) not in ("1st", "1", "first"):
            segs.append((_ordinal(edition) + " ed. ", "bib_editionno"))
        org_used_as_author = (n_auth == 0 and publisher and
                              not meta.get("bib_organization") and
                              not meta.get("bib_institution") and
                              not (meta.get("bib_surname") if not meta.get("bib_fname") else ""))
        if publisher and not org_used_as_author:
            segs.append((publisher, "bib_publisher"))
            segs.append(("; ", None))
        if year:
            segs.append((year, "bib_year"))
        if ref_type == "book_chapter":
            fpage = meta.get("bib_fpage") or ""
            lpage = meta.get("bib_lpage") or ""
            if fpage:
                segs.append((":", None))
                segs.append((fpage, "bib_fpage"))
                if lpage:
                    segs.append(("-", None))
                    segs.append((lpage, "bib_lpage"))
        segs.append((".", None))

    elif ref_type == "conference":
        conf     = meta.get("bib_conference") or ""
        confloc  = meta.get("bib_conflocation") or ""
        confdate = meta.get("bib_confdate") or meta.get("bib_year") or ""
        if title := meta.get("bib_title") or "":
            clean_title = _ama_title(title)
            segs.append((clean_title, "bib_confpaper"))
            segs.append((" " if re.search(r'[?!]$', clean_title) else ". ", None))
        segs.append(("Paper presented at: ", None))
        if conf:
            segs.append((conf, "bib_conference"))
        if confdate:
            segs.append(("; ", None))
            segs.append((confdate, "bib_confdate"))
        if confloc:
            segs.append(("; ", None))
            segs.append((confloc, "bib_conflocation"))
        segs.append((".", None))

    elif ref_type == "thesis":
        title  = meta.get("bib_title") or ""
        deg    = meta.get("bib_deg") or "doctoral dissertation"
        school = meta.get("bib_school") or ""
        year   = meta.get("bib_year") or ""
        url    = meta.get("bib_url") or ""
        if title:
            segs.append((_ama_title(title), "bib_title"))
        bracket = f" [{deg}]."
        segs.append((bracket, None))
        if school:
            segs.append((" " + school + ";", None))
        if year:
            segs.append((" " + year + ".", None))
        if url:
            segs.append((" ", None))
            segs.append((url, "bib_url"))
        doi = (meta.get("bib_doi") or "").strip().lstrip("doi:").lstrip()
        if doi:
            segs.append((" doi:", "bib_doi"))
            segs.append((doi, "bib_doi"))
        return segs

    elif ref_type == "website":
        title    = meta.get("bib_title") or ""
        year     = meta.get("bib_year") or ""
        accessed = meta.get("bib_accessed") or ""
        url      = meta.get("bib_url") or ""
        site     = meta.get("bib_journal") or meta.get("bib_book") or ""
        # Suppress site name if it duplicates the sole org author already rendered above
        author_org = ""
        if n_auth == 1 and not (fnames[0] if fnames else ""):
            author_org = surnames[0] if surnames else ""
        if title:
            clean_title = _ama_title(title)
            segs.append((clean_title, "bib_title"))
            segs.append((" " if re.search(r'[?!]$', clean_title) else ". ", None))
        if site and not _is_url(site) and site.strip().rstrip(".") != author_org.strip().rstrip("."):
            segs.append((site, "bib_journal"))
            segs.append((". ", None))
        if year:
            segs.append(("Published ", None))
            segs.append((year, "bib_year"))
            segs.append((". ", None))
        if accessed:
            segs.append(("Accessed ", None))
            segs.append((accessed, "bib_accessed"))
            segs.append((". ", None))
        if url:
            segs.append((url, "bib_url"))

    elif ref_type == "ereference":
        title    = meta.get("bib_title") or ""
        book     = meta.get("bib_book") or meta.get("bib_journal") or ""
        pub      = _strip_publisher_suffixes(meta.get("bib_publisher") or meta.get("bib_institution") or "")
        year     = meta.get("bib_year") or ""
        accessed = meta.get("bib_accessed") or ""
        url      = meta.get("bib_url") or ""
        if title:
            clean_title = _ama_title(title)
            segs.append((clean_title, "bib_title"))
            segs.append((" " if re.search(r'[?!]$', clean_title) else ". ", None))
        segs.append(("In: ", None))
        if ed_surnames:
            for i, es in enumerate(ed_surnames):
                if i > 0:
                    segs.append((", ", None))
                ei = ed_fnames[i] if i < len(ed_fnames) else ""
                ei_str = _format_initials_ama(ei)
                segs.append((es, "bib_ed-surname"))
                if ei_str:
                    segs.append((" ", None))
                    segs.append((ei_str, "bib_ed-fname"))
            ed_label = "ed." if len(ed_surnames) == 1 else "eds."
            segs.append((f", {ed_label} ", None))
        if book:
            segs.append((_to_sentence_case(book), "bib_book"))
            segs.append((". ", None))
        if pub:
            segs.append((pub, "bib_publisher"))
            segs.append(("; ", None))
        if year:
            segs.append((year, "bib_year"))
            segs.append((".", None))
        if accessed:
            segs.append((" Accessed ", None))
            segs.append((accessed, "bib_accessed"))
            segs.append((".", None))
        if url:
            segs.append((" ", None))
            segs.append((url, "bib_url"))

    elif ref_type == "report":
        repnum   = meta.get("bib_reportnum") or ""
        inst     = _strip_publisher_suffixes(meta.get("bib_institution") or "")
        year     = meta.get("bib_year") or ""
        url      = meta.get("bib_url") or ""
        accessed = meta.get("bib_accessed") or ""
        doi      = (meta.get("bib_doi") or "").strip().lstrip("doi:").lstrip()
        if inst:
            segs.append((inst, "bib_institution"))
            segs.append(("; ", None))
        if year:
            segs.append((year, "bib_year"))
            segs.append((".", None))
        if repnum:
            segs.append((" Report No. " + repnum + ".", None))
        if doi:
            segs.append((" doi:", "bib_doi"))
            segs.append((doi, "bib_doi"))
        elif url:
            if accessed:
                segs.append((" Accessed ", None))
                segs.append((accessed, "bib_accessed"))
                segs.append((".", None))
            segs.append((" ", None))
            segs.append((url, "bib_url"))
        return segs

    doi = (meta.get("bib_doi") or "").strip().lstrip("doi:").lstrip()
    if doi and ref_type not in ("website", "ereference", "thesis"):
        segs.append((" doi:", "bib_doi"))
        segs.append((doi, "bib_doi"))

    return segs


_MONTH_DAY_YEAR_RE = re.compile(r'^([A-Za-z]+)\s+(\d{1,2}),\s*(\d{4})$')
_MONTH_YEAR_RE     = re.compile(r'^([A-Za-z]+)\s+(\d{4})$')


def _normalize_apa_year(y: str) -> str:
    y = y.strip()
    m = _MONTH_DAY_YEAR_RE.match(y)
    if m:
        return f"{m.group(3)}, {m.group(1)} {m.group(2)}"
    m = _MONTH_YEAR_RE.match(y)
    if m:
        return f"{m.group(2)}, {m.group(1)}"
    return y


_URL_TRAILING_RE = re.compile(
    r',?\s+on\s+\w+\s+\d{1,2},\s*\d{4}\.?$'
    r'|,?\s+on\s+\w+\s+\d{4}\.?$',
    re.IGNORECASE,
)


def _clean_url(url: str) -> str:
    url = url.strip().rstrip(".")
    url = _URL_TRAILING_RE.sub("", url).strip()
    return url


_ACCESSED_PREFIX_RE = re.compile(
    r'^(?:retrieved\s+from|retrieved\s*,?\s*from|from)\s+',
    re.IGNORECASE,
)


def _clean_accessed(accessed: str) -> str:
    accessed = accessed.strip()
    accessed = _ACCESSED_PREFIX_RE.sub("", accessed).strip()
    accessed = re.sub(r'\s*https?://\S+.*$', '', accessed, flags=re.IGNORECASE).strip()
    return accessed


def _is_url(value: str) -> bool:
    return bool(re.match(r'https?://', value.strip(), re.IGNORECASE))



def build_segments_apa(meta: Dict, gemini_text: str = "", italic_words: frozenset = frozenset()) -> List[Tuple[str, Optional[str]]]:
    segs: List[Tuple[str, Optional[str]]] = []
    ref_type = (meta.get("bib_reftype") or "journal").lower()

    surnames = _split_pipe(meta.get("bib_surname"))
    fnames   = _split_pipe(meta.get("bib_fname"))
    # Remove any stray ellipsis entries Gemini may have stored
    clean_pairs = [(s, fnames[i] if i < len(fnames) else "")
                   for i, s in enumerate(surnames)
                   if s.strip() not in ("…", "...", "\u2026")]
    surnames = [p[0] for p in clean_pairs]
    fnames   = [p[1] for p in clean_pairs]

    # ── FIX 1: et al. detection — clean pop preserving fname alignment ────
    has_etal = False
    if surnames and surnames[-1].lower().replace(".", "").strip() == "et al":
        has_etal = True
        surnames.pop()
        # Pop matching fname only if lists are aligned
        if fnames and len(fnames) >= len(surnames) + 1:
            fnames.pop()

    n_auth = len(surnames)

    # APA rule: org/group authors must appear last
    if n_auth > 1:
        ind_idx = [i for i in range(n_auth)
                   if _format_initials_apa(fnames[i] if i < len(fnames) else "") or not _is_organization(surnames[i])]
        org_idx = [i for i in range(n_auth) if i not in ind_idx]
        if org_idx:
            order    = ind_idx + org_idx
            surnames = [surnames[i] for i in order]
            fnames   = [fnames[i] if i < len(fnames) else "" for i in order]

    is_edited_book_primary = False
    if ref_type == "edited_book" and n_auth == 0:
        surnames = _split_pipe(meta.get("bib_ed_surname") or meta.get("bib-ed-surname"))
        fnames   = _split_pipe(meta.get("bib_ed_fname")   or meta.get("bib-ed-fname"))
        n_auth   = len(surnames)
        is_edited_book_primary = True

    if n_auth == 0:
        org = (meta.get("bib_organization") or
               meta.get("bib_institution") or
               (meta.get("bib_surname") if not meta.get("bib_fname") else "") or "")
        if not org:
            if ref_type in ("book", "edited_book", "report"):
                org = (meta.get("bib_publisher") or "").strip()
            elif ref_type == "website":
                site_val = (meta.get("bib_journal") or meta.get("bib_book") or "").strip()
                if site_val and not _is_url(site_val):
                    org = site_val
        if org:
            segs.append((org.rstrip("."), "bib_organization"))
            segs.append((".", None))
    else:
        subset = surnames if n_auth <= 20 else surnames[:19]
        for i, surname in enumerate(subset):
            if i > 0:
                segs.append((", ", None))
                # ── FIX 1: Only emit "& " when we have the FULL author list.
                # has_etal=True means source was truncated — "et al." follows
                # instead of "&", so no ampersand before last named author.
                if i == len(subset) - 1 and n_auth <= 20 and not has_etal:
                    segs.append(("& ", None))
            initial      = fnames[i] if i < len(fnames) else ""
            initials_str = _format_initials_apa(initial)
            if not initials_str and _is_organization(surname):
                disp_name = surname[0].upper() + surname[1:] if i == 0 and surname else surname
                segs.append((disp_name, "bib_organization"))
            else:
                disp_name = surname[0].upper() + surname[1:] if i == 0 and surname else surname
                segs.append((disp_name, "bib_surname"))
                if initials_str:
                    segs.append((", ", None))
                    segs.append((initials_str, "bib_fname"))

        if n_auth > 20:
            segs.append((", … ", None))
            segs.append((surnames[-1], "bib_surname"))
            last_initial = fnames[-1] if len(fnames) >= n_auth else ""
            initials_str = _format_initials_apa(last_initial)
            if initials_str:
                segs.append((", ", None))
                segs.append((initials_str, "bib_fname"))
        # NOTE: In APA 7th edition, we NEVER show "et al." because:
        # - If n_auth <= 20: we show all authors (APA requirement)
        # - If n_auth > 20: we show first 19 + ... + last (explicit APA format)
        # The has_etal flag indicates SOURCE was truncated, but if database
        # enrichment expanded it to <= 20 authors, we show them all (no et al.).
        # If database has > 20, we use "..." format instead.

    if is_edited_book_primary and n_auth > 0:
        ed_label = " (Ed.)" if n_auth == 1 else " (Eds.)"
        segs.append((ed_label, None))
        segs.append((".", None))

    if n_auth > 0 and not is_edited_book_primary:
        last_seg_text = segs[-1][0] if segs else ""
        if not last_seg_text.endswith("."):
            segs.append((".", None))

    segs.append((" (", None))
    raw_year = meta.get("bib_year") or "n.d."
    if ref_type in ("website", "ereference") and re.match(r'^\d{4}[a-z]?$', raw_year) and gemini_text:
        # Try APA-formatted date first: (YYYY, Month Day)
        _dm = re.search(r'\(' + re.escape(raw_year[:4]) + r',\s*([A-Za-z]+\s+\d{1,2})\)', gemini_text)
        if _dm:
            raw_year = raw_year[:4] + ", " + _dm.group(1)
        else:
            # Fall back to source date format: (Month Day, YYYY)
            _dm2 = re.search(r'\(([A-Za-z]+\s+\d{1,2}),?\s*' + re.escape(raw_year[:4]) + r'\)', gemini_text)
            if _dm2:
                raw_year = raw_year[:4] + ", " + _dm2.group(1)
    segs.append((_normalize_apa_year(raw_year), "bib_year"))
    segs.append(("). ", None))

    chapter_title = meta.get("bib_chaptertitle") or ""
    main_title    = meta.get("bib_title") or ""
    book_title    = meta.get("bib_book") or ""

    if ref_type == "thesis":
        title  = main_title or book_title or ""
        deg    = meta.get("bib_deg") or "Doctoral dissertation"
        school = meta.get("bib_school") or ""
        url    = meta.get("bib_url") or ""
        if title:
            segs.append((_to_sentence_case(title.rstrip(".")), "bib_title"))
        bracket = f" [{deg}"
        if school:
            bracket += f", {school}"
        bracket += "]."
        segs.append((bracket, None))
        if url:
            segs.append((" ", None))
            segs.append((url, "bib_url"))
        doi = (meta.get("bib_doi") or "").strip().lstrip("doi:").lstrip()
        if doi:
            segs.append((" https://doi.org/", "bib_doi"))
            segs.append((doi, "bib_doi"))
        return segs

    elif ref_type == "book_chapter" and (chapter_title or main_title):
        title_text  = chapter_title or main_title
        clean_title = _to_sentence_case(title_text.rstrip("."))
        segs.extend(_split_italic_spans(clean_title, italic_words, "bib_chaptertitle"))
        segs.append((" " if re.search(r'[?!]$', clean_title) else ". ", None))

    elif ref_type in ("book", "edited_book"):
        display_title = book_title or main_title or ""
        if display_title:
            clean_title = _to_sentence_case(display_title.rstrip("."))
            segs.append((clean_title, "bib_book"))
            _edi = meta.get("bib_editionno") or ""
            if not (_edi and _ordinal(_edi) not in ("1st", "1", "first")):
                segs.append((".", None))

    elif main_title:
        clean_title = _to_sentence_case(main_title.rstrip("."))
        style = "bib_article" if ref_type == "journal" else "bib_title"
        segs.extend(_split_italic_spans(clean_title, italic_words, style))
        segs.append((" " if re.search(r'[?!]$', clean_title) else ". ", None))

    if ref_type == "book_chapter":
        ed_surnames = _split_pipe(meta.get("bib_ed_surname") or meta.get("bib-ed-surname"))
        ed_fnames   = _split_pipe(meta.get("bib_ed_fname")   or meta.get("bib-ed-fname"))
        segs.append(("In ", None))
        if ed_surnames:
            for i, es in enumerate(ed_surnames):
                if i > 0:
                    segs.append((" & ", None) if i == len(ed_surnames) - 1 else (", ", None))
                ei           = ed_fnames[i] if i < len(ed_fnames) else ""
                initials_str = _format_initials_apa(ei)
                if initials_str:
                    segs.append((initials_str + " ", "bib_ed-fname"))
                if not initials_str and _is_organization(es):
                    segs.append((es, "bib_organization"))
                else:
                    segs.append((es, "bib_ed-surname"))
            ed_label = "(Ed.)," if len(ed_surnames) == 1 else "(Eds.),"
            segs.append((" " + ed_label + " ", None))
        display_book = book_title or main_title or ""
        if display_book:
            segs.append((_to_sentence_case(display_book.rstrip(".")), "bib_book"))
        edition   = meta.get("bib_editionno") or ""
        volume    = meta.get("bib_volume") or ""
        fpage     = meta.get("bib_fpage") or ""
        lpage     = meta.get("bib_lpage") or ""
        clean_ord = _ordinal(edition)
        inner_segs: list = []
        if edition and clean_ord not in ("1st", "1", "first"):
            inner_segs.append((clean_ord + " ed.", "bib_editionno"))
        if volume:
            if inner_segs: inner_segs.append((", ", None))
            inner_segs.append(("Vol. ", None))
            inner_segs.append((volume, "bib_volume"))
        if fpage:
            if inner_segs: inner_segs.append((", ", None))
            inner_segs.append(("pp. ", None))
            inner_segs.append((fpage, "bib_fpage"))
            if lpage:
                inner_segs.append(("–", None))
                inner_segs.append((lpage, "bib_lpage"))
        if inner_segs:
            segs.append((" (", None))
            segs.extend(inner_segs)
            segs.append((").", None))
        else:
            segs.append((".", None))
        publisher = _strip_publisher_suffixes(meta.get("bib_publisher") or "")
        if publisher:
            segs.append((" ", None))
            segs.append((publisher, "bib_publisher"))
            segs.append((".", None))

    elif ref_type == "journal":
        journal = _to_title_case(meta.get("bib_journal") or "")
        volume  = meta.get("bib_volume") or ""
        issue   = meta.get("bib_issue") or ""
        fpage   = meta.get("bib_fpage") or ""
        lpage   = meta.get("bib_lpage") or ""
        if journal:
            segs.append((journal, "bib_journal"))
        if volume:
            if journal:
                segs.append((", ", None))
            segs.append((volume, "bib_volume"))
        if issue:
            segs.append(("(", None))
            segs.append((issue, "bib_issue"))
            segs.append((")", None))
        if fpage:
            segs.append((", ", None))
            segs.append((fpage, "bib_fpage"))
            if lpage:
                segs.append(("–", None))
                segs.append((lpage, "bib_lpage"))
        elif not volume and not issue and "Advance online publication" in gemini_text:
            segs.append((". Advance online publication", None))
        segs.append((".", None))

    elif ref_type in ("book", "edited_book"):
        edition    = meta.get("bib_editionno") or ""
        _raw_pub   = meta.get("bib_publisher") or ""
        publisher  = _raw_pub if _raw_pub.strip() == "Author" else _strip_publisher_suffixes(_raw_pub)
        if edition and _ordinal(edition) not in ("1st", "1", "first"):
            segs.append((" (", None))
            segs.append((_ordinal(edition) + " ed.", "bib_editionno"))
            segs.append((").", None))
        org_used_as_author = (n_auth == 0 and publisher and
                              not meta.get("bib_organization") and
                              not meta.get("bib_institution") and
                              not (meta.get("bib_surname") if not meta.get("bib_fname") else ""))
        if publisher and not org_used_as_author:
            segs.append((" ", None))
            segs.append((publisher, "bib_publisher"))
            segs.append((".", None))

    elif ref_type == "website":
        site     = meta.get("bib_journal") or meta.get("bib_book") or ""
        accessed = _clean_accessed(meta.get("bib_accessed") or "")
        url      = _clean_url(meta.get("bib_url") or "")
        if _is_url(site):
            site = ""
        org_used_as_author = (n_auth == 0 and site and
                              not meta.get("bib_organization") and
                              not meta.get("bib_institution") and
                              not (meta.get("bib_surname") if not meta.get("bib_fname") else ""))
        if site and not org_used_as_author:
            segs.append((_to_title_case(site), "bib_journal"))
            segs.append((".", None))
        if accessed and url:
            segs.append((" Retrieved " + accessed + ", from ", None))
            segs.append((url, "bib_url"))
        elif url:
            segs.append((" ", None))
            segs.append((url, "bib_url"))

    elif ref_type == "ereference":
        ed_surnames = _split_pipe(meta.get("bib_ed_surname") or meta.get("bib-ed-surname"))
        ed_fnames   = _split_pipe(meta.get("bib_ed_fname")   or meta.get("bib-ed-fname"))
        ref_title   = meta.get("bib_book") or meta.get("bib_journal") or ""
        pub         = _strip_publisher_suffixes(meta.get("bib_publisher") or "")
        accessed    = _clean_accessed(meta.get("bib_accessed") or "")
        url         = _clean_url(meta.get("bib_url") or "")
        if _is_url(ref_title):
            ref_title = ""
        segs.append(("In ", None))
        if ed_surnames:
            for i, es in enumerate(ed_surnames):
                if i > 0:
                    segs.append((" & " if i == len(ed_surnames) - 1 else ", ", None))
                ei = ed_fnames[i] if i < len(ed_fnames) else ""
                initials_str = _format_initials_apa(ei)
                if initials_str:
                    segs.append((initials_str + " ", "bib_ed-fname"))
                segs.append((es, "bib_ed-surname"))
            ed_label = "(Ed.)," if len(ed_surnames) == 1 else "(Eds.),"
            segs.append((" " + ed_label + " ", None))
        if ref_title:
            segs.append((_to_title_case(ref_title), "bib_book"))
            segs.append((".", None))
        if pub:
            segs.append((" ", None))
            segs.append((pub, "bib_publisher"))
            segs.append((".", None))
        if accessed and url:
            segs.append((" Retrieved " + accessed + ", from ", None))
            segs.append((url, "bib_url"))
        elif url:
            segs.append((" ", None))
            segs.append((url, "bib_url"))

    elif ref_type == "conference":
        conf     = meta.get("bib_conference") or ""
        confloc  = meta.get("bib_conflocation") or ""
        confdate = meta.get("bib_confdate") or ""
        segs.append(("[Conference session]. ", None))
        if conf:
            segs.append((_to_title_case(conf), "bib_conference"))
        if confdate:
            segs.append((", " + confdate, None))
        if confloc:
            segs.append((", " + confloc, None))
        segs.append((".", None))

    elif ref_type == "report":
        repnum    = meta.get("bib_reportnum") or ""
        _raw_inst = meta.get("bib_institution") or meta.get("bib_publisher") or ""
        inst      = _raw_inst if _raw_inst.strip() == "Author" else _strip_publisher_suffixes(_raw_inst)
        if repnum:
            segs.append((" (Report No. " + repnum + ").", None))
        if inst:
            segs.append((" ", None))
            segs.append((inst, "bib_institution"))
            segs.append((".", None))

    doi = (meta.get("bib_doi") or "").strip().lstrip("doi:").lstrip()
    url = meta.get("bib_url") or ""
    if doi:
        segs.append((" https://doi.org/", "bib_doi"))
        segs.append((doi, "bib_doi"))
    elif url and ref_type not in ("website", "ereference", "thesis"):
        segs.append((" ", None))
        segs.append((url, "bib_url"))

    return segs


# ─────────────────────────────────────────────
# CONVERSION LOG ENTRY
# ─────────────────────────────────────────────

class ConversionLogEntry:
    def __init__(self, original: str, converted: str, ref_type: str,
                 source_style: str, target_style: str, notes: Optional[str] = None,
                 error: Optional[str] = None):
        self.original     = original
        self.converted    = converted
        self.ref_type     = ref_type
        self.source_style = source_style
        self.target_style = target_style
        self.notes        = notes
        self.error        = error

    def to_log_line(self) -> str:
        lines = [
            f"  TYPE:    {self.ref_type}",
            f"  FROM:    [{self.source_style}] {self.original}",
            f"  TO:      [{self.target_style}] {self.converted}",
        ]
        if self.notes:  lines.append(f"  NOTES:   {self.notes}")
        if self.error:  lines.append(f"  ERROR:   {self.error}")
        return "\n".join(lines)


# ─────────────────────────────────────────────
# MAIN PROCESSOR
# ─────────────────────────────────────────────

def process_conversion(
    input_docx: Path,
    output_dir: Optional[Path] = None,
    source_style: str = "Auto",
    target_style: str = "APA",
    model_name: str = DEFAULT_MODEL,
    prefer_gemini_output: bool = True,
) -> Dict[str, Path]:
    input_docx = Path(input_docx)
    if not input_docx.exists():
        raise FileNotFoundError(f"Input file not found: {input_docx}")

    target_style = target_style.strip().upper() if target_style.upper() != "AUTO" else "AUTO"
    if target_style not in ("AMA", "APA", "CGRN", "AUTO"):
        raise ValueError(f"target_style must be 'AMA', 'APA', 'CGRN', or 'AUTO', got: {target_style}")

    if output_dir is None:
        output_dir = input_docx.parent
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    if target_style == "APA":
        target_enum = CitationStyle.APA
    elif target_style == "CGRN":
        target_enum = CitationStyle.CGRN
    else:
        target_enum = CitationStyle.AMA

    stem             = input_docx.stem
    output_docx_path = output_dir / f"{stem}_Converted.docx"
    log_file_path    = output_dir / f"{stem}_conversion_log.txt"
    json_dump_path   = output_dir / f"{stem}_metadata_dump.json"

    doc = Document(input_docx)

    log_entries: List[ConversionLogEntry] = []
    json_records: List[Dict] = []
    log_header: List[str] = [
        f"Reference Conversion Log",
        f"Input:         {input_docx.name}",
        f"Source Style:  {source_style}",
        f"Target Style:  {target_style}",
        f"Model:         {model_name}",
        "=" * 60, ""
    ]

    total_count       = 0
    converted_count   = 0
    error_count       = 0
    ref_section_depth = 0

    from concurrent.futures import ThreadPoolExecutor, as_completed
    from ReferencesStructing import find_best_metadata_for_reference

    tasks = []

    for idx, para in enumerate(doc.paragraphs):
        raw_text = para.text.strip()
        if not raw_text:
            continue

        # Extract italic word spans from source paragraph
        italic_words: frozenset = frozenset()
        try:
            char_pos = 0
            italic_chars: set = set()
            for run in para.runs:
                run_len = len(run.text)
                if run.italic is True:
                    for i in range(char_pos, char_pos + run_len):
                        italic_chars.add(i)
                char_pos += run_len
            # Extract italic words: any word where all (or most) chars are italic
            if italic_chars:
                full_text = para.text  # un-stripped, position-aligned
                _italic_words = set()
                for m in re.finditer(r'\S+', full_text):
                    word_chars = set(range(m.start(), m.end()))
                    if word_chars & italic_chars:
                        _italic_words.add(m.group().lower().strip('.,;:!?"\''))
                italic_words = frozenset(_italic_words)
        except Exception:
            italic_words = frozenset()

        raw_lower = raw_text.lower()

        try:
            para_style_name_check = para.style.name or ''
        except Exception:
            para_style_name_check = ''

        is_cgrn_para = (para_style_name_check == '* ReferencesText')

        if "<ref-open>" in raw_lower or "<ref-close>" in raw_lower:
            open_count  = raw_lower.count("<ref-open>")
            close_count = raw_lower.count("<ref-close>")
            if open_count > 0:
                ref_section_depth += open_count
                logger.info(f"Entering reference section. Depth: {ref_section_depth}")
            if close_count > 0:
                ref_section_depth = max(0, ref_section_depth - close_count)
                logger.info(f"Exiting reference section. Depth: {ref_section_depth}")
            raw_text  = re.sub(r'(?i)<ref-open>\s*',  '', raw_text)
            raw_text  = re.sub(r'(?i)<ref-close>\s*', '', raw_text).strip()
            raw_lower = raw_text.lower()
            if not raw_text:
                continue

        if ref_section_depth == 0 and not is_cgrn_para:
            continue
        if len(raw_text) < 15:
            continue
        if _looks_like_inline_citation(raw_text):
            logger.debug(f"Skipping inline citation para: {raw_text[:60]}")
            continue

        total_count += 1
        try:
            para_style_name = para.style.name or ''
        except Exception:
            para_style_name = ''
        tasks.append({
            'doc_index':  idx,
            'para_obj':   para,
            'raw_text':   raw_text,
            'count':      total_count,
            'para_style': para_style_name,
            'italic_words': italic_words,
        })

    def process_task(task: dict):
        raw_text = task['raw_text']
        count    = task['count']
        logger.info(f"[{count}] Gemini API Call: {raw_text[:80]}...")

        if source_style.upper() == "AUTO":
            para_style = task.get('para_style', '')
            if para_style == 'REF-N':
                detected_source = CitationStyle.AMA
            elif para_style in ('REF-U', 'REF'):
                detected_source = CitationStyle.APA
            elif para_style == '* ReferencesText':
                detected_source = CitationStyle.CGRN
            else:
                detected_source = detect_source_style(raw_text)
        elif source_style.upper() == "CGRN":
            detected_source = CitationStyle.CGRN
        else:
            detected_source = CitationStyle.AMA if source_style.upper() == "AMA" else CitationStyle.APA

        task['detected_source'] = detected_source

        if target_style.upper() == "AUTO":
            t_enum = detected_source
            logger.info(f"  [{count}] Auto: strict formatting for {t_enum.value}")
        elif target_style.upper() == "CGRN":
            t_enum = CitationStyle.CGRN
        else:
            t_enum = CitationStyle.APA if target_style.upper() == "APA" else CitationStyle.AMA

        result = convert_reference(
            raw_text=raw_text,
            source_style=detected_source,
            target_style=t_enum,
            model_name=model_name,
            cr_item=None,
        )
        if not result:
            logger.error(
                f"  [{count}] Gemini conversion failed for: {raw_text[:100]}... "
                f"(Check logs above for Gemini error details — rate limit, timeout, API key, or service issue)"
            )

        cr_item = None
        try:
            temp_cr, source_db, score = find_best_metadata_for_reference(raw_text, detected_source.value)
            if temp_cr:
                is_journal = (
                    'pubmed' in source_db.lower()
                    or ('crossref' in source_db.lower() and temp_cr.get('type', '').lower() in ('journal-article', 'journal'))
                    or ('crossref' in source_db.lower() and not temp_cr.get('type') and temp_cr.get('container-title'))
                )
                # Detect source ref type from first-pass Gemini result so we can
                # reject a journal-article DB hit for a report/book reference.
                first_pass_type = (result.get("metadata", {}).get("bib_reftype") or "unknown").lower()
                source_is_journal = first_pass_type in ("journal", "unknown")
                if is_journal and not source_is_journal:
                    logger.info(
                        f"  [{count}] [DB Enrich] Rejected — DB is journal article but "
                        f"reference is '{first_pass_type}'. Skipping to avoid author contamination."
                    )
                elif is_journal and score >= 0.65:
                    cr_item = temp_cr
                    logger.info(f"  [{count}] [DB Enrich] Journal via {source_db} (Score: {score:.2f})")
                elif score >= 0.75:
                    cr_item = temp_cr
                    logger.info(f"  [{count}] [DB Enrich] General via {source_db} (Score: {score:.2f})")
                else:
                    logger.info(f"  [{count}] [DB Enrich] Ignored {source_db} (Score: {score:.2f}) — below threshold")
        except Exception as e:
            logger.warning(f"  [{count}] DB enrichment lookup failed: {e}")

        task['target_enum'] = t_enum
        task['result']      = result
        task['cr_item']     = cr_item
        task['skip']        = False
        return task

    if tasks:
        logger.info(f"Starting parallel conversions for {len(tasks)} references...")
        worker_count = min(MAX_CONVERSION_WORKERS, max(1, len(tasks)))
        logger.info(f"Using {worker_count} conversion worker(s).")
        with ThreadPoolExecutor(max_workers=worker_count) as executor:
            futures = [executor.submit(process_task, t) for t in tasks]
            for future in as_completed(futures):
                try:
                    future.result()
                except Exception as e:
                    logger.error(f"Error in parallel conversion task: {e}")

    for task in sorted(tasks, key=lambda x: x['doc_index']):
        count           = task['count']
        raw_text        = task['raw_text']
        para            = task['para_obj']
        result          = task['result']
        detected_source = task['detected_source']

        if task.get('skip'):
            logger.info(f"  [{count}] Skipping: kept original formatting.")
            continue

        if not result:
            error_count += 1
            entry = ConversionLogEntry(
                original=raw_text, converted="[FAILED]",
                ref_type="unknown", source_style=detected_source.value,
                target_style=target_style, error="No DB match and Gemini call skipped/failed",
            )
            log_entries.append(entry)
            logger.warning(f"  No usable conversion result for reference {count}")
            continue

        metadata   = result.get("metadata", {})
        ref_type   = detect_ref_type_from_metadata(metadata)
        gemini_out = result.get("formatted_output", "").strip()
        notes      = result.get("conversion_notes")

        # Heuristic ref-type correction
        metadata = _fix_ref_type(metadata, raw_text)
        ref_type = detect_ref_type_from_metadata(metadata)

        # ── FIX 2: Re-inject "et al" when source had it but Gemini dropped it.
        # This ensures has_etal=True fires in build_segments_apa/ama so the
        # truncation signal from the original source document is preserved.
        if "et al" in raw_text.lower() or "and others" in raw_text.lower():
            existing_surnames = [
                s.strip() for s in (metadata.get("bib_surname") or "").split("|")
                if s.strip()
            ]
            if existing_surnames and existing_surnames[-1].lower().replace(".", "").strip() != "et al":
                metadata = dict(metadata)
                metadata["bib_surname"] = "|".join(existing_surnames) + "|et al"
                # Maintain pipe-count alignment: append blank fname for "et al" sentinel
                fname_parts = (metadata.get("bib_fname") or "").split("|") if metadata.get("bib_fname") else []
                while len(fname_parts) < len(existing_surnames):
                    fname_parts.append("")
                fname_parts.append("")
                metadata["bib_fname"] = "|".join(fname_parts)
                logger.info(
                    f"  [{count}] [et al fix] Re-injected 'et al' into bib_surname "
                    f"({len(existing_surnames)} named authors retained)"
                )

        resolved_target = task['target_enum'].value

        cr_it = task.get('cr_item')
        is_journal = metadata.get("bib_reftype", "").lower() == "journal"
        pre_merge_journal = metadata.get("bib_journal", "")
        if cr_it:
            if cr_it.get("DOI") and (is_journal or not metadata.get("bib_doi")):
                db_doi = str(cr_it["DOI"]).replace("https://doi.org/","").replace("doi:","").strip()
                if db_doi:
                    metadata["bib_doi"] = db_doi

            if cr_it.get("title") and not metadata.get("bib_title"):
                raw_t = cr_it["title"]
                metadata["bib_title"] = raw_t[0] if isinstance(raw_t, list) else str(raw_t)

            if cr_it.get("URL") and not metadata.get("bib_url"):
                metadata["bib_url"] = str(cr_it["URL"]).strip()
            if cr_it.get("volume") and (is_journal or not metadata.get("bib_volume")):
                metadata["bib_volume"] = str(cr_it["volume"]).strip()
            if cr_it.get("issue") and (is_journal or not metadata.get("bib_issue")):
                metadata["bib_issue"] = str(cr_it["issue"]).strip()
            if cr_it.get("page") and (is_journal or not metadata.get("bib_fpage")):
                raw_page = str(cr_it["page"]).strip()
                fpage, lpage = _page_parts(raw_page)
                metadata["bib_fpage"] = fpage
                metadata["bib_lpage"] = lpage or metadata.get("bib_lpage", "")

            db_year = None
            for date_key in ("published-print", "published-online", "issued"):
                dp = cr_it.get(date_key, {}).get("date-parts")
                if dp and dp[0] and dp[0][0]:
                    db_year = _format_ama_date_parts(dp[0]) if resolved_target == "AMA" else str(dp[0][0])
                    break
            if not db_year and cr_it.get("year"):
                db_year = str(cr_it["year"])
            if db_year:
                if not metadata.get("bib_year"):
                    metadata["bib_year"] = db_year
                else:
                    src_year_match = re.search(r'(\d{4})', metadata.get("bib_year", ""))
                    db_year_match = re.search(r'(\d{4})', db_year)
                    src_yr = src_year_match.group(1) if src_year_match else ""
                    db_yr = db_year_match.group(1) if db_year_match else ""
                    if src_yr and db_yr and src_yr != db_yr:
                        if src_yr in raw_text:
                            logger.info(f"  [{count}] [DB Correction] Year kept {src_yr} — confirmed in source text")
                        else:
                            metadata["bib_year"] = db_year
                            logger.info(f"  [{count}] [DB Correction] Year → {db_year}")
                    elif resolved_target == "AMA" and len(db_year) > len(metadata.get("bib_year", "")):
                        metadata["bib_year"] = db_year
                        logger.info(f"  [{count}] [DB Correction] AMA date expanded → {db_year}")

            # Only apply DOI guard when DB lookup didn't provide a DOI
            if not cr_it.get("DOI"):
                _src_doi_m = re.search(r'10\.\d{4,9}/[-._;()/:A-Za-z0-9]+', raw_text, re.IGNORECASE)
                if _src_doi_m:
                    _src_doi = _src_doi_m.group(0).rstrip('.')
                    if _src_doi and _src_doi != metadata.get("bib_doi", ""):
                        metadata["bib_doi"] = _src_doi
                        logger.info(f"  [{count}] [DOI Guard] Restoring source DOI: {_src_doi}")

            if is_journal and resolved_target not in ("AMA",):
                # APA/CGRN: always use full container-title to expand abbreviations
                full = (cr_it.get("container-title") or [""])[0].strip()
                if full:
                    metadata["bib_journal"] = full
            elif not metadata.get("bib_journal"):
                if resolved_target == "AMA":
                    abbr = (cr_it.get("short-container-title") or [""])[0].strip()
                    full = (cr_it.get("container-title") or [""])[0].strip()
                    metadata["bib_journal"] = abbr or full
                else:
                    full = (cr_it.get("container-title") or [""])[0].strip()
                    metadata["bib_journal"] = full

        # ── HYBRID JOURNAL EXPANSION FALLBACK (when DB lookup failed) ────────
        # For journals without DB match (cr_it is None), try hybrid expansion:
        # 1. Local lookup table (fast, cost-free)
        # 2. Gemini expansion (for unknown abbreviations)
        if not cr_it and is_journal and metadata.get("bib_journal"):
            current_journal = metadata.get("bib_journal", "")
            # Check if it looks abbreviated (short, has dots, no spaces, etc.)
            # Better heuristic: abbreviated if has dots, or is very short (<3 chars) with no spaces
            word_count = len(current_journal.split())
            has_dots = '.' in current_journal
            is_single_word_caps = word_count == 1 and current_journal.isupper() and len(current_journal) <= 4
            # Marked as abbreviated if: has dots (likely abbreviation) or is a short all-caps acronym
            is_abbreviated = has_dots or is_single_word_caps
            if is_abbreviated:
                expanded = _expand_journal_hybrid(current_journal)
                if expanded and expanded != current_journal:
                    metadata["bib_journal"] = expanded
                    logger.info(f"  [{count}] [Journal Hybrid] '{current_journal}' -> '{expanded}'")

        # ── AUTHOR EXPANSION ─────────────────────────────────────────────────
        # When source had "et al." the expansion logic must be style-aware:
        #
        #   AMA: list up to 6 authors; 7+ -> first 3 + et al.
        #     - DB total <= 6  -> expand fully, DROP et al. from metadata
        #     - DB total 7+    -> expand to first 3, KEEP et al. in metadata
        #     - DB absent      -> keep source truncation (et al. remains)
        #
        #   APA: list up to 20 authors; 21+ -> first 19 ... last.
        #     - DB total <= 20 -> expand fully, DROP et al. from metadata
        #     - DB total 21+   -> expand (build_segments_apa handles ellipsis)
        #     - DB absent      -> keep source truncation (et al. remains)
        #
        # When source did NOT have et al., only expand if DB has more authors
        # than Gemini listed (guards against Gemini inventing et al.).
        source_had_etal = "et al" in raw_text.lower() or "and others" in raw_text.lower()
        if cr_it and cr_it.get("author"):
            db_authors = cr_it.get("author", [])
            if db_authors and isinstance(db_authors, list) and len(db_authors) > 0:
                db_total    = len(db_authors)
                db_families = [a.get("family", "") for a in db_authors if a.get("family")]
                db_givens   = [a.get("given",  "") for a in db_authors if a.get("family")]
                current_surnames = metadata.get("bib_surname", "").split("|") if metadata.get("bib_surname") else []

                if source_had_etal:
                    # Source was truncated — expand using DB with style-specific rules
                    full_cutoff = 6 if resolved_target == "AMA" else 20
                    truncate_to = 3 if resolved_target == "AMA" else 19
                    if db_total <= full_cutoff:
                        # Total real authors fit within style limit — always expand fully,
                        # drop et al. regardless of whether Gemini already stored the full list
                        metadata["bib_surname"] = "|".join(db_families)
                        metadata["bib_fname"]   = "|".join(db_givens)
                        logger.info(
                            f"  [{count}] [Author Expansion] DB {db_total} authors <= "
                            f"{full_cutoff} ({resolved_target} limit) — expanded fully, et al. dropped"
                        )
                    elif len(db_families) > 0:
                        # DB total exceeds limit — truncate to style-specific count, keep et al.
                        metadata["bib_surname"] = "|".join(db_families[:truncate_to]) + "|et al"
                        metadata["bib_fname"]   = "|".join(db_givens[:truncate_to])   + "|"
                        logger.info(
                            f"  [{count}] [Author Expansion] DB {db_total} authors > "
                            f"{full_cutoff} ({resolved_target} limit) — expanded to {truncate_to}, et al. retained"
                        )
                else:
                    # Source had no et al. — only expand if DB has more than Gemini listed.
                    # Guard: never replace an org author with personal authors from a
                    # type-mismatched DB record (e.g. journal article returned for a report).
                    cr_db_is_journal = (cr_it.get("type") or "").lower() in ("journal-article", "journal") \
                        or (cr_it.get("container-title") and cr_it.get("volume"))
                    if cr_db_is_journal and ref_type not in ("journal", "unknown"):
                        logger.info(
                            f"  [{count}] [Author Expansion] Skipped — DB is journal article "
                            f"but ref_type is '{ref_type}'"
                        )
                    elif len(db_authors) > len(current_surnames):
                        db_first_sn  = db_families[0].lower() if db_families else ""
                        src_first_sn = current_surnames[0].lower() if current_surnames else ""
                        if db_first_sn and src_first_sn and db_first_sn != src_first_sn:
                            logger.info(
                                f"  [{count}] [Author Expansion] Skipped — DB first author '{db_families[0]}' "
                                f"differs from source '{current_surnames[0]}'"
                            )
                        else:
                            expanded_surnames = "|".join(db_families)
                            expanded_fnames   = "|".join(db_givens)
                            if expanded_surnames:
                                metadata["bib_surname"] = expanded_surnames
                                metadata["bib_fname"]   = expanded_fnames
                                logger.info(f"  [{count}] [Author Expansion] DB provided {db_total} authors")

        # ── URL CLEANUP ──────────────────────────────────────────────────────
        # Strip "Available from:", "Retrieved from:" prefixes Gemini may include,
        # and remove trailing periods that would break the URL.
        if metadata.get("bib_url"):
            metadata["bib_url"] = _normalize_url_value(metadata["bib_url"])
        # If bib_url is still empty, try to extract URL directly from source text
        if not metadata.get("bib_url"):
            url_match = re.search(r'https?://\S+', raw_text)
            if url_match:
                metadata["bib_url"] = _normalize_url_value(url_match.group(0))
        # Always prefer URL found verbatim in source text (Gemini may mangle URL paths)
        _src_url_m = re.search(r'https?://\S+', raw_text)
        if _src_url_m:
            metadata["bib_url"] = _normalize_url_value(_src_url_m.group(0))
        metadata = _normalize_reference_metadata(metadata)

        if prefer_gemini_output and gemini_out:
            final_text = gemini_out
            # Expand abbreviated journal names with full DB name (APA/CGRN)
            expanded_journal = metadata.get("bib_journal", "")
            if pre_merge_journal and expanded_journal and pre_merge_journal != expanded_journal:
                # Apply title case to expanded journal name for APA style
                if resolved_target == "APA":
                    expanded_journal = _to_title_case(expanded_journal)
                    metadata["bib_journal"] = expanded_journal
                if pre_merge_journal in final_text:
                    final_text = final_text.replace(pre_merge_journal, expanded_journal)
                    logger.info(f"  [{count}] [Journal Expansion] '{pre_merge_journal}' -> '{expanded_journal}'")
                else:
                    clean_old = pre_merge_journal.rstrip('. ')
                    if clean_old and clean_old in final_text:
                        final_text = final_text.replace(clean_old, expanded_journal)
                        logger.info(f"  [{count}] [Journal Expansion] '{clean_old}' -> '{expanded_journal}'")
            # Fix AMA author formatting: remove commas between surname and initials
            if resolved_target == "AMA":
                final_text = _fix_ama_author_format(final_text)
                # Truncate to first 3 + et al. if 7+ authors
                final_text = _truncate_ama_authors(final_text)
            elif resolved_target == "APA":
                # Truncate to first 19 + ... + last if 21+ authors
                final_text = _truncate_apa_authors(final_text)
                # Re-capitalise first letter after ? or ! (Gemini may lowercase mid-title words)
                final_text = re.sub(
                    r'([?!])(\s+)([a-z])',
                    lambda m: m.group(1) + m.group(2) + m.group(3).upper(),
                    final_text
                )
            # Fix duplicate titles in website/ereference references
            if ref_type in ("website", "ereference") and metadata.get("bib_title"):
                final_text = _fix_duplicate_title_website(final_text, metadata["bib_title"])
            # Fix reference numbering
            final_text = _remove_reference_numbers(final_text)
            # Restore non-year parenthetical text that Gemini may have treated as a citation
            for _pm in re.finditer(r'\([^)0-9]{5,}\)', raw_text):
                _orig_p = _pm.group(0)
                if _orig_p not in final_text:
                    _pfx = re.escape(_orig_p[1:min(21, len(_orig_p) - 1)])
                    _mangled = re.search(r'\(' + _pfx + r'[^)]*\)', final_text)
                    if _mangled:
                        final_text = final_text[:_mangled.start()] + _orig_p + final_text[_mangled.end():]
            if metadata.get("bib_doi") and "doi:" not in final_text.lower() and "doi.org" not in final_text.lower():
                if resolved_target == "AMA":
                    final_text = final_text.rstrip(".") + f". doi:{metadata['bib_doi']}"
                elif resolved_target != "CGRN":
                    final_text = final_text.rstrip(".") + f". https://doi.org/{metadata['bib_doi']}"
            # Inject URL for website/ereference/report when Gemini omitted it from output
            if (ref_type in ("website", "ereference", "report") and
                    metadata.get("bib_url") and
                    metadata["bib_url"] not in final_text):
                final_text = final_text.rstrip(". ").rstrip(".") + " " + metadata["bib_url"]
            # Audit: warn if source DOI/URL is absent from final output
            _audit_doi_m = re.search(r'10\.\d{4,9}/[-._;()/:A-Za-z0-9]+', raw_text, re.IGNORECASE)
            if _audit_doi_m and _audit_doi_m.group(0).rstrip('.') not in final_text:
                logger.warning(
                    f"  [{count}] [DOI Audit] Source DOI '{_audit_doi_m.group(0).rstrip('.')}' "
                    f"not found in final output"
                )
            _audit_url_m = re.search(r'https?://\S+', raw_text)
            if _audit_url_m:
                _au = _audit_url_m.group(0).rstrip(".,;)")
                if _au not in final_text:
                    logger.warning(
                        f"  [{count}] [URL Audit] Source URL '{_au}' not found in final output"
                    )
        else:
            if resolved_target == "AMA":
                final_text = format_ama_from_metadata(metadata)
            elif resolved_target == "CGRN":
                final_text = gemini_out or raw_text
            else:
                final_text = format_apa_from_metadata(metadata)

        # Force local rebuild if Gemini stubbornly retained 'et al.' but we
        # now have the full author list AND the source itself did not have et al.
        if "et al" in final_text.lower() and metadata.get("bib_surname") and not source_had_etal:
            surnames_check = [
                s for s in metadata.get("bib_surname").split("|")
                if s.strip().lower().replace(".", "") != "et al"
            ]
            num_named = len(surnames_check)
            if num_named > 1 and num_named <= 6:
                logger.info(f"  [{count}] [Correction] Gemini retained 'et al.' despite {num_named} known authors. Rebuilding.")
                if resolved_target == "AMA":
                    final_text = format_ama_from_metadata(metadata)
                elif resolved_target == "APA":
                    final_text = format_apa_from_metadata(metadata)

        # Strip DB journal qualifiers after any forced rebuild so qualifiers
        # like "(MEDLINE, 1950-present: 2023)" are always removed from final_text
        metadata, final_text = _strip_db_journal_qualifiers(raw_text, metadata, final_text)
        metadata, final_text, val_findings, rebuilt_segs = _auto_correct_reference(
            metadata,
            final_text,
            resolved_target,
            ref_type,
            gemini_text=gemini_out,
            italic_words=task.get('italic_words', frozenset()),
        )

        # Convert straight double-quotes to curly typographic quotes (APA requirement)
        final_text = _to_curly_quotes(final_text)

        if not final_text.strip():
            error_count += 1
            entry = ConversionLogEntry(
                original=raw_text, converted="[EMPTY OUTPUT]",
                ref_type=ref_type, source_style=detected_source.value,
                target_style=target_style,
                error="Both Gemini output and metadata fallback produced empty string",
            )
            log_entries.append(entry)
            continue

        try:
            if resolved_target == "CGRN":
                _write_cgrn_runs(para, final_text, doc=doc, original_text=raw_text)
            else:
                segs = list(rebuilt_segs) if rebuilt_segs else []
                ama_metadata_usable = bool(
                    metadata and metadata.get("bib_reftype") and (
                        metadata.get("bib_title")
                        or metadata.get("bib_article")
                        or metadata.get("bib_chaptertitle")
                        or metadata.get("bib_journal")
                        or metadata.get("bib_surname")
                    )
                )
                if metadata and (metadata.get("bib_reftype") if resolved_target != "AMA" else ama_metadata_usable):
                    try:
                        if resolved_target == "AMA":
                            segs = build_segments_ama(metadata, gemini_out)
                        else:
                            segs = build_segments_apa(metadata, gemini_out, italic_words=task.get('italic_words', frozenset()))
                        segs_text = "".join(t for t, _ in segs)
                        if len(segs_text.strip()) < 10:
                            segs = []
                            logger.debug(f"  [{count}] Metadata segments too short; using Gemini text path.")
                    except Exception as _meta_err:
                        segs = []
                        logger.warning(f"  [{count}] Metadata segment build failed ({_meta_err}); falling back.")

                if not segs and final_text:
                    segs = _parse_gemini_output_to_segments(final_text)
                    logger.debug(f"  [{count}] Using Gemini text parse (fallback) for styling.")

                if segs:
                    _write_styled_runs(para, segs, doc=doc, is_conversion=(detected_source != task['target_enum']))
                else:
                    _set_paragraph_text(para, final_text, doc=doc)
        except Exception as _seg_err:
            logger.warning(f"  Segment build failed ({_seg_err}); falling back to plain text.")
            _set_paragraph_text(para, final_text, doc=doc)

        val_issues = [f.message for f in val_findings if f.severity != "info"]
        if val_issues:
            comment_msg = "Please check: " + "; ".join(val_issues)
            _add_review_comment(doc, para, comment_msg)
            logger.warning(f"  [{count}] Flagged for review — {', '.join(val_issues)}")

        converted_count += 1

        entry = ConversionLogEntry(
            original=raw_text, converted=final_text,
            ref_type=ref_type, source_style=detected_source.value,
            target_style=target_style, notes=notes,
        )
        log_entries.append(entry)
        json_records.append({
            "index": count, "ref_type": ref_type,
            "source_style": detected_source.value, "target_style": target_style,
            "original": raw_text, "converted": final_text,
            "notes": notes, "metadata": metadata,
            "validation_findings": [f.__dict__ for f in val_findings],
        })
        logger.info(f"  ✓ [{ref_type}] → {final_text[:80]}...")

    doc.save(output_docx_path)
    logger.info(f"Saved converted document: {output_docx_path}")

    summary = [
        "", "=" * 60,
        f"SUMMARY",
        f"  Total references found:  {total_count}",
        f"  Successfully converted:  {converted_count}",
        f"  Errors:                  {error_count}",
        f"  Skipped (same style):    {total_count - converted_count - error_count}",
    ]

    with open(log_file_path, "w", encoding="utf-8") as f:
        f.write("\n".join(log_header) + "\n")
        for i, entry in enumerate(log_entries, 1):
            f.write(f"[{i}]\n{entry.to_log_line()}\n\n")
        f.write("\n".join(summary) + "\n")

    logger.info(f"Log written: {log_file_path}")

    with open(json_dump_path, "w", encoding="utf-8") as f:
        json.dump(json_records, f, indent=2, ensure_ascii=False)

    logger.info(f"Metadata dump: {json_dump_path}")

    return {
        "output_docx": output_docx_path,
        "log_file":    log_file_path,
        "json_dump":   json_dump_path,
    }


# ─────────────────────────────────────────────
# CLI ENTRY POINT
# ─────────────────────────────────────────────

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        description="Convert references in a Word document between AMA, APA, and CGRN styles."
    )
    parser.add_argument("input",
                        type=str,
                        help="Path to input .docx file")
    parser.add_argument("--output-dir",
                        type=str,
                        help="Output directory (default: same as input)")
    parser.add_argument("--source-style",
                        type=str, default="Auto",
                        choices=["AMA", "APA", "CGRN", "Auto"],
                        help="Source citation style")
    parser.add_argument("--target-style",
                        type=str, default="APA",
                        choices=["AMA", "APA", "CGRN"],
                        help="Target citation style")
    parser.add_argument("--model",
                        type=str, default=DEFAULT_MODEL,
                        help="Gemini model name")
    parser.add_argument("--no-gemini-output",
                        action="store_true",
                        help="Rebuild entirely from metadata instead of Gemini formatted output")
    args = parser.parse_args()

    paths = process_conversion(
        input_docx=Path(args.input),
        output_dir=Path(args.output_dir) if args.output_dir else None,
        source_style=args.source_style,
        target_style=args.target_style,
        model_name=args.model,
        prefer_gemini_output=not args.no_gemini_output,
    )

    print("\nConversion complete:")
    for k, v in paths.items():
        print(f"  {k}: {v}")
