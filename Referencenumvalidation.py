import re
import os
import io
import zipfile
import difflib
from collections import defaultdict

from flask import Flask, request, send_file, render_template, redirect, url_for, session
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import RGBColor
from utils import track_changes
import logging

TRACK_CHANGES_ENABLED = True

app = Flask(__name__)
app.secret_key = "secret_key_for_session_encryption"
UPLOAD_DIR = "temp_reports"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# =====================================================
# Helpers & Core Logic
# =====================================================

def iter_document_paragraphs(doc):
    """
    Iterate through all paragraphs in the document body in order,
    including those inside tables.
    """
    body = doc._element.body
    for child in body:
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            table = Table(child, doc)
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p


def get_numbers(text):
    """
    Extract numbers from text like '1', '2-5', '1, 3, 5'.
    Handles ranges "1-5" -> [1, 2, 3, 4, 5].
    """
    nums = []
    # Matches: (start)-(end) OR (single)
    # Allows hyphen, en dash, em dash
    pattern = re.compile(r'(\d+)\s*[-–—]\s*(\d+)|(\d+)')
    
    for start, end, single in pattern.findall(text):
        if start and end:
            try:
                s, e = int(start), int(end)
                if s <= e:
                    nums.extend(range(s, e + 1))
            except ValueError:
                pass
        elif single:
            try:
                nums.append(int(single))
            except ValueError:
                pass
    return nums


def format_numbers(nums):
    """
    Format a list of numbers into a string like '1-3, 5'.
    Collapses ranges of 3 or more (e.g. 1,2,3 -> 1-3).
    """
    nums = sorted(set(nums))
    if not nums:
        return ""

    parts = []
    if not nums:
        return ""

    start = prev = nums[0]

    for n in nums[1:]:
        if n == prev + 1:
            prev = n
        else:
            length = prev - start + 1
            if length >= 3:
                parts.append(f"{start}-{prev}")
            elif length == 2:
                parts.append(f"{start},{prev}")
            else:
                parts.append(str(start))
            start = prev = n

    length = prev - start + 1
    if length >= 3:
        parts.append(f"{start}-{prev}")
    elif length == 2:
        parts.append(f"{start},{prev}")
    else:
        parts.append(str(start))

    return ", ".join(parts)


def _ensure_styles(doc):
    """Create cite_bib and bib_number character styles if they don't exist."""
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    try:
        doc.styles['cite_bib']
    except KeyError:
        style = doc.styles.add_style('cite_bib', WD_STYLE_TYPE.CHARACTER)
        style.font.superscript = True
    try:
        doc.styles['bib_number']
    except KeyError:
        style = doc.styles.add_style('bib_number', WD_STYLE_TYPE.CHARACTER)


BIB_NUMBER_PATTERNS = [
    re.compile(r'^\s*\[(\d+)\]\.\s+'),   # [1].
    re.compile(r'^\s*\((\d+)\)\.\s+'),   # (1).
    re.compile(r'^\s*\[(\d+)\]\s+'),     # [1]
    re.compile(r'^\s*\((\d+)\)\s+'),     # (1)
    re.compile(r'^\s*(\d+)\.\s+'),       # 1.
    re.compile(r'^\s*(\d+)\s+'),         # 1  (space only)
]


def extract_bib_number(text):
    """Extract bibliography number from text, handling all formats: 1., [1]., (1)., etc."""
    for pat in BIB_NUMBER_PATTERNS:
        m = pat.match(text)
        if m:
            return int(m.group(1)), pat
    return None, None


def get_numPr(para):
    """Detect Word auto-numbering. Returns (numId, ilvl) or (None, None)."""
    from docx.oxml.ns import qn
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return None, None
    numPr = pPr.find(qn('w:numPr'))
    if numPr is None:
        return None, None
    ilvl_el = numPr.find(qn('w:ilvl'))
    numId_el = numPr.find(qn('w:numId'))
    ilvl = int(ilvl_el.get(qn('w:val'), 0)) if ilvl_el is not None else 0
    numId = int(numId_el.get(qn('w:val'), 0)) if numId_el is not None else 0
    return numId, ilvl


def compute_list_number(doc, target_para):
    """Count preceding paragraphs with same numId+ilvl to compute displayed list number."""
    numId, ilvl = get_numPr(target_para)
    if numId is None or numId == 0:
        return None
    count = 0
    for para in doc.paragraphs:
        nid, nlvl = get_numPr(para)
        if nid == numId and nlvl == ilvl:
            count += 1
        if para._element is target_para._element:
            return count
    return None


def convert_autonumber_to_manual(para, number, doc):
    """Remove <w:numPr> and insert manual bib_number run with track changes."""
    from docx.oxml.ns import qn
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            pPr.remove(numPr)
    if number is not None:
        track_changes.add_tracked_text(para, f"{number}.", style='bib_number',
                                       author='RefBot', doc=doc)


def detect_and_tag_unstyled_citations(doc, citation_format):
    """
    Scan for unstyled citations and apply cite_bib style based on format.
    Returns {'tagged': count, 'format_used': citation_format}
    """
    tagged_count = 0
    _ensure_styles(doc)

    if citation_format == 'styled':
        return {'tagged': 0, 'format_used': 'styled'}

    if citation_format == 'superscript':
        from docx.oxml.ns import qn
        for para in iter_document_paragraphs(doc):
            for run in para.runs:
                if run.font.superscript and re.match(r'^[\d,\-–—\s]+$', run.text):
                    if not (run.style and run.style.name == 'cite_bib'):
                        run.style = doc.styles['cite_bib']
                        tagged_count += 1

    elif citation_format == 'bracket':
        for para in iter_document_paragraphs(doc):
            para_text = para.text
            if '[' not in para_text:
                continue
            if para.style and para.style.name == 'REF-N':
                continue
            for run in para.runs:
                if re.search(r'\[\d+(?:[,\-–—]\d+)*\]', run.text):
                    if not (run.style and run.style.name == 'cite_bib'):
                        new_text = re.sub(r'\[(\d+(?:[,\-–—]\d+)*)\]', r'\1', run.text)
                        run.text = new_text
                        run.style = doc.styles['cite_bib']
                        tagged_count += 1

    elif citation_format == 'paren':
        for para in iter_document_paragraphs(doc):
            para_text = para.text
            if '(' not in para_text:
                continue
            if para.style and para.style.name == 'REF-N':
                continue
            for run in para.runs:
                if re.search(r'\(\d+(?:[,\-–—]\d+)*\)', run.text) and not re.search(r'\d{4}', run.text):
                    if not (run.style and run.style.name == 'cite_bib'):
                        new_text = re.sub(r'\((\d+(?:[,\-–—]\d+)*)\)', r'\1', run.text)
                        run.text = new_text
                        run.style = doc.styles['cite_bib']
                        tagged_count += 1

    elif citation_format == 'plain':
        for para in iter_document_paragraphs(doc):
            if para.style and para.style.name in ['REF-N', 'REF-U']:
                continue
            for i, run in enumerate(para.runs):
                if i == 0:
                    continue
                if re.match(r'^\d+$', run.text) and not (run.style and run.style.name == 'cite_bib'):
                    run.style = doc.styles['cite_bib']
                    tagged_count += 1

    return {'tagged': tagged_count, 'format_used': citation_format}


def is_citation_run(run):
    """
    Determine if a run is part of a citation.
    Strictly checks for 'cite_bib' styles.
    """
    if run.style and run.style.name in ["cite_bib"]:
        return True
    return False


class ReferenceProcessor:
    def __init__(self, doc):
        self.doc = doc
        
    def get_references_in_bibliography(self):
        """
        Returns a Set of IDs found in the bibliography sections (REF-N style).
        Also returns a list of objects for reordering later.
        """
        refs_found = set()
        ref_objects = [] # list of dicts: {'id': int, 'para': p, 'run': r}

        for para in self.doc.paragraphs:
            if para.style and para.style.name == "REF-N":
                found_id = None
                bib_run = None

                # Try finding styled run
                for run in para.runs:
                    if run.style and run.style.name == "bib_number":
                        nums = get_numbers(run.text)
                        if nums:
                            found_id = nums[0]
                            bib_run = run
                            break

                # Fallback: extract from plain text using multi-format patterns
                if found_id is None:
                    found_id, pattern = extract_bib_number(para.text)

                if found_id is not None:
                    refs_found.add(found_id)
                    ref_objects.append({
                        'id': found_id,
                        'para': para,
                        'run': bib_run
                    })

        return refs_found, ref_objects

    def get_citations_in_text(self):
        """
        Scans document for citations.
        Returns:
            all_cited_ids: list of all IDs in order of appearance (with duplicates)
            appearance_order: list of unique IDs in order of first appearance
        """
        all_cited_ids = []
        appearance_order = []
        seen = set()
        
        for para in iter_document_paragraphs(self.doc):
            # 1. Process runs
            current_group = []
            
            for run in para.runs:
                if is_citation_run(run):
                    current_group.append(run)
                else:
                    if current_group:
                        # Flush group
                        text = "".join(r.text for r in current_group)
                        nums = get_numbers(text)
                        all_cited_ids.extend(nums)
                        for n in nums:
                            if n not in seen:
                                seen.add(n)
                                appearance_order.append(n)
                        current_group = []
            
            # Flush trailing group
            if current_group:
                text = "".join(r.text for r in current_group)
                nums = get_numbers(text)
                all_cited_ids.extend(nums)
                for n in nums:
                    if n not in seen:
                        seen.add(n)
                        appearance_order.append(n)
                        
        return all_cited_ids, appearance_order

    def find_duplicates(self, ref_objects):
        """
        Finds duplicate references using fuzzy matching (difflib).
        Returns a list of dicts: {'id': int, 'text': str, 'duplicate_of': int, 'score': float}
        """
        import difflib
        
        duplicates = []
        processed_refs = [] # list of (id, clean_text)
        
        # 1. Pre-process all candidates
        for obj in ref_objects:
            full_text = obj['para'].text.strip()
            # Remove leading numbering like "1. ", "[1] "
            clean_text = re.sub(r'^\[?\d+\]?[\.\s]*', '', full_text)
            processed_refs.append({'id': obj['id'], 'text': clean_text})
            
        # 2. Compare O(N^2)
        # We only check forward to avoid double reporting (A=B, B=A)
        # We assume the *earlier* ID is the "original" and later is "duplicate"
        n = len(processed_refs)
        matcher = difflib.SequenceMatcher(None, "", "")
        
        for i in range(n):
            ref_a = processed_refs[i]
            text_a = ref_a['text']
            len_a = len(text_a)
            
            if len_a == 0:
                continue
                
            matcher.set_seq1(text_a)
            
            for j in range(i + 1, n):
                ref_b = processed_refs[j]
                text_b = ref_b['text']
                len_b = len(text_b)
                
                if len_b == 0: 
                    continue
                    
                # Optimization: Length ratio check
                # If lengths differ significantly, they can't be high matches
                # If ratio > 0.85, then min_len / max_len must be roughly > 0.85
                # We use 0.6 as a conservative safety net, but 0.8 is probably safe if threshold is 0.85.
                if min(len_a, len_b) / max(len_a, len_b) < 0.6:
                    continue
                
                matcher.set_seq2(text_b)
                
                # Performance Optimization: Check cheap upper bounds first
                if matcher.real_quick_ratio() < 0.99:
                    continue
                if matcher.quick_ratio() < 0.99:
                    continue

                ratio = matcher.ratio()

                # Threshold: 0.99 (99% similar — very strict, almost identical)
                if ratio > 0.99:
                    duplicates.append({
                        'id': ref_b['id'], # The later one is the duplicate
                        'text': ref_b['text'][:100] + "...",
                        'duplicate_of': ref_a['id'],
                        'score': round(ratio * 100, 1)
                    })
                    
        return duplicates

    def resolve_duplicates(self):
        """
        Finds duplicates, remaps citations, and marks bibliography entries for deletion.
        Returns list of merge records: [{'removed_id': int, 'canonical_id': int, 'citations_updated': int}, ...]
        """
        from docx.oxml.ns import qn

        bib_refs, ref_objects = self.get_references_in_bibliography()
        duplicates = self.find_duplicates(ref_objects)

        if not duplicates:
            return []

        # Build canonical mapping, handling transitive duplicates (A→B, B→C → A→C)
        mapping = {}  # duplicate_id -> canonical_id
        for dup in duplicates:
            dup_id = dup['id']
            canonical_id = dup['duplicate_of']
            # Resolve transitively
            while canonical_id in mapping:
                canonical_id = mapping[canonical_id]
            mapping[dup_id] = canonical_id

        merge_log = []
        citations_updated_count = {dup_id: 0 for dup_id in mapping.keys()}

        # Step 1: Remap citations from duplicate IDs to canonical IDs
        for para in iter_document_paragraphs(self.doc):
            for run in para.runs:
                if is_citation_run(run):
                    nums = get_numbers(run.text)
                    new_nums = [mapping.get(n, n) for n in nums]
                    if new_nums != nums:
                        # Count how many times each duplicate ID was replaced
                        for n in nums:
                            if n in mapping:
                                citations_updated_count[n] += 1

                        new_text = format_numbers(new_nums)
                        if TRACK_CHANGES_ENABLED:
                            track_changes.delete_tracked_run(para, run)
                            run_del = run._element.getparent()
                            ins_new = track_changes.add_tracked_text(
                                para, new_text, style='cite_bib', color='008000', doc=self.doc)
                            run_del.addnext(ins_new)
                        else:
                            run.text = new_text

        # Step 2: Mark duplicate bibliography entries for deletion with track changes
        bib_refs_new, ref_objects_new = self.get_references_in_bibliography()
        for obj in ref_objects_new:
            if obj['id'] in mapping:
                # This is a duplicate — mark it for deletion
                track_changes.wrap_paragraph_content_in_del(obj['para'], author='RefBot')
                # Record the merge
                canonical_id = mapping[obj['id']]
                for dup_record in duplicates:
                    if dup_record['id'] == obj['id']:
                        merge_log.append({
                            'removed_id': obj['id'],
                            'canonical_id': canonical_id,
                            'score': dup_record['score'],
                            'citations_updated': citations_updated_count.get(obj['id'], 0)
                        })
                        break

        return merge_log

    def get_validation_stats(self):
        bib_refs, ref_objects = self.get_references_in_bibliography()
        all_cited, _ = self.get_citations_in_text()
        
        unique_cited = set(all_cited)
        
        # Missing: Cited but not in Bib
        missing = sorted(unique_cited - bib_refs)
        
        # Unused: In Bib but not Cited
        unused = sorted(bib_refs - unique_cited)
        
        # Duplicates
        duplicates = self.find_duplicates(ref_objects)
        
        # Sequence Issues
        sequence_issues = []
        seen_in_seq = []
        previous_max = 0
        
        for n in all_cited:
            if n not in seen_in_seq:
                if n < previous_max:
                     pass
                
                if n != len(seen_in_seq) + 1:
                     sequence_issues.append({
                         "position": len(seen_in_seq) + 1,
                         "current": n,
                         "expected": len(seen_in_seq) + 1
                     })
                
                seen_in_seq.append(n)
                previous_max = max(previous_max, n)
                
        return {
            "total_references": len(bib_refs),
            "total_citations": len(all_cited),
            "missing_references": missing,
            "unused_references": unused,
            "duplicate_references": duplicates,
            "sequence_issues": sequence_issues,
            "is_perfect": (not missing and not unused and not sequence_issues and not duplicates)
        }

    def _ensure_bib_numbers_styled(self):
        """Ensure all bibliography numbers have bib_number style applied."""
        from docx.enum.style import WD_STYLE_TYPE

        styles = self.doc.styles
        try:
            styles['bib_number']
        except KeyError:
            styles.add_style('bib_number', WD_STYLE_TYPE.CHARACTER)

        # Tag all bibliography entry numbers with bib_number style
        for para in self.doc.paragraphs:
            if para.style and para.style.name == 'REF-N':
                # Try to find the number run
                bib_num, pattern = extract_bib_number(para.text)
                if bib_num is not None:
                    # Find the run that contains the number
                    for run in para.runs:
                        if str(bib_num) in run.text or any(c.isdigit() for c in run.text):
                            # Check if this run starts with the number
                            if run.text.strip() and run.text[0].isdigit():
                                run.style = styles['bib_number']
                                break

    def renumber(self):
        """
        Renumber citations and reorder bibliography.
        Returns: mapping (Old -> New)
        """
        _, appearance_order = self.get_citations_in_text()

        # Ensure 'cite_bib' and 'bib_number' styles exist
        from docx.enum.style import WD_STYLE_TYPE
        styles = self.doc.styles
        try:
            styles['cite_bib']
        except KeyError:
            s = styles.add_style('cite_bib', WD_STYLE_TYPE.CHARACTER)
            s.font.superscript = True

        try:
            styles['bib_number']
        except KeyError:
            styles.add_style('bib_number', WD_STYLE_TYPE.CHARACTER)

        # Create Mapping
        mapping = {} 
        new_id = 1
        for old_id in appearance_order:
            mapping[old_id] = new_id
            new_id += 1
            
        for para in iter_document_paragraphs(self.doc):
            i = 0
            while i < len(para.runs):
                run = para.runs[i]
                
                if is_citation_run(run):
                    txt = run.text
                    nums = get_numbers(txt)
                    if nums:
                         new_nums = [mapping.get(n, n) for n in nums]
                         new_text = format_numbers(new_nums)
                         
                         is_renumbered = (nums != new_nums)
                         highlight_color = "008000" if is_renumbered else None
                         
                         style_name = run.style.name if run.style else "cite_bib"
                         
                         if TRACK_CHANGES_ENABLED:
                             # Must replace the whole run
                             track_changes.delete_tracked_run(para, run)
                             
                             run_del = run._element.getparent()
                             anchor = run_del if run_del.tag == track_changes.qn('w:del') else run._element
                             
                             ins_new = track_changes.add_tracked_text(para, new_text, style=style_name, color=highlight_color)
                             anchor.addnext(ins_new)
                         else:
                             run.text = new_text
                             if is_renumbered:
                                 run.font.color.rgb = RGBColor(0, 128, 0)
                
                i += 1
                
                i += 1

        # 2. Reorder Bibliography
        _, ref_objects = self.get_references_in_bibliography()
        
        # Sort objects into Cited and Uncited
        cited_refs = []
        uncited_refs = []
        
        for obj in ref_objects:
            if obj['id'] in mapping:
                obj['new_id'] = mapping[obj['id']]
                cited_refs.append(obj)
            else:
                uncited_refs.append(obj)
        
        if not ref_objects:
            return mapping

        # Find anchor (min index)
        body = self.doc._element.body
        
        indices = []
        for obj in ref_objects:
            try:
                idx = body.index(obj['para']._element)
                indices.append(idx)
            except ValueError:
                pass 
        
        if not indices:
            return mapping
            
        anchor = min(indices)
        
        # Remove all
        for obj in ref_objects:
             p = obj['para']._element
             if p.getparent() == body:
                 body.remove(p)
                 
        # Insert Cited (Sorted)
        cited_refs.sort(key=lambda x: x['new_id'])
        
        insert_idx = anchor
        for obj in cited_refs:
            # Update ID text
            if obj['run']:
                old_text = obj['run'].text
                new_text = str(obj['new_id'])
                
                if old_text != new_text:
                    if TRACK_CHANGES_ENABLED:
                        style_name = obj['run'].style.name if obj['run'].style else None
                        
                        track_changes.delete_tracked_run(obj['para'], obj['run'])
                        run_del = obj['run']._element.getparent()
                        anchor = run_del if run_del.tag == track_changes.qn('w:del') else obj['run']._element
                        
                        ins_new = track_changes.add_tracked_text(obj['para'], new_text, style=style_name)
                        anchor.addnext(ins_new)
                    else:
                        obj['run'].text = new_text
            
            body.insert(insert_idx, obj['para']._element)
            insert_idx += 1
            
        # Insert Uncited (Appended after cited)
        for obj in uncited_refs:
            body.insert(insert_idx, obj['para']._element)
            insert_idx += 1

        # Ensure all bibliography numbers have bib_number style applied
        self._ensure_bib_numbers_styled()

        return mapping


def process_document(file, citation_format='styled'):
    doc = Document(file)
    processor = ReferenceProcessor(doc)

    # PRE-PROCESS: Ensure styles exist
    _ensure_styles(doc)

    # PRE-PROCESS: Convert Word auto-numbered reference lists to manual
    auto_converted = 0
    for para in doc.paragraphs:
        if para.style and para.style.name == 'REF-N':
            numId, ilvl = get_numPr(para)
            if numId:
                n = compute_list_number(doc, para)
                convert_autonumber_to_manual(para, n, doc)
                auto_converted += 1

    # PRE-PROCESS: Detect and tag unstyled citations
    cite_tag_result = {}
    if citation_format != 'styled':
        cite_tag_result = detect_and_tag_unstyled_citations(doc, citation_format)

    # Check BEFORE
    before_stats = processor.get_validation_stats()
    before_stats['auto_converted'] = auto_converted
    before_stats['citations_tagged'] = cite_tag_result.get('tagged', 0)

    # DECISION:
    # 1. If Unused References exist -> ABORT renumbering.
    if before_stats["unused_references"]:
        return doc, before_stats, before_stats, {}, "Aborted: Document validation failed due to unused references.", []

    # 2. If Perfect -> No need.
    if before_stats["is_perfect"]:
        return doc, before_stats, before_stats, {}, "Validation completed.", []

    # 3. If Missing Refs -> Can't safely renumber usually
    if before_stats["missing_references"]:
        return doc, before_stats, before_stats, {}, "Aborted: Missing references detected.", []

    # PROCESS 2: Resolve duplicates with track changes
    merge_log = []
    if not before_stats["missing_references"] and not before_stats["unused_references"]:
        merge_log = processor.resolve_duplicates()
        # Re-validate after merge to get fresh stats
        before_stats = processor.get_validation_stats()

    # DO RENUMBER
    mapping = processor.renumber()

    # Check AFTER (Validate result)
    after_stats = processor.get_validation_stats()

    # Determine status message
    changes_made = False
    if mapping:
        for k, v in mapping.items():
            if k != v:
                changes_made = True
                break

    if merge_log:
        count = len(merge_log)
        status_msg = f"Renumbering completed with {count} duplicate reference{'s' if count > 1 else ''} merged."
    elif before_stats["duplicate_references"]:
        count = len(before_stats['duplicate_references'])
        prefix = "Renumbering" if changes_made else "Validation"
        status_msg = f"{prefix} completed with {count} duplicate{'s' if count > 1 else ''}."
    elif changes_made:
        status_msg = "Renumbering completed successfully."
    else:
        status_msg = "Validation completed."

    return doc, before_stats, after_stats, mapping, status_msg, merge_log


# =====================================================
# Flask Routes
# =====================================================
@app.route("/")
def upload_file():
    return render_template("upload.html")


@app.route("/process", methods=["GET", "POST"])
def process():
    if request.method == "POST":
        file = request.files.get("file")
        if not file or not file.filename.endswith(".docx"):
            return "Invalid file", 400

        doc, before, after, mapping, status_msg, merge_log = process_document(file)

        base = os.path.splitext(file.filename)[0]
        doc_path = os.path.join(UPLOAD_DIR, f"{base}_renumbered.docx")
        report_path = os.path.join(UPLOAD_DIR, f"{base}_validation.txt")

        doc.save(doc_path)

        with open(report_path, "w", encoding="utf-8") as f:
            f.write(f"STATUS: {status_msg}\n")
            f.write("VALIDATION BEFORE\n")
            f.write(str(before) + "\n\n")
            f.write("VALIDATION AFTER\n")
            f.write(str(after) + "\n\n")
            if mapping:
                f.write("RENUMBERING MAPPING (Old -> New)\n")
                for old, new in sorted(mapping.items(), key=lambda x: x[1]):
                    f.write(f"{old} -> {new}\n")

        # Create ZIP package
        zip_filename = f"{base}_results.zip"
        zip_path = os.path.join(UPLOAD_DIR, zip_filename)
        
        # Validation HTML Report (Offline)
        html_report_filename = f"{base}_results.html"
        html_report_path = os.path.join(UPLOAD_DIR, html_report_filename)
        
        # Render the template for offline use
        # Note: We pass offline_mode=True to make links relative
        html_content = render_template(
            "validation_results.html",
            filename=file.filename,
            results=after,
            before=before,
            mapping=mapping,
            status_msg=status_msg,
            report_file=os.path.basename(report_path),
            doc_file=os.path.basename(doc_path),
            zip_file=None, # No zip button in offline report
            offline_mode=True 
        )
        
        with open(html_report_path, "w", encoding="utf-8") as f:
            f.write(html_content)

        with zipfile.ZipFile(zip_path, 'w') as zf:
             # Add Doc
             zf.write(doc_path, arcname=os.path.basename(doc_path))
             # Add Text Report
             zf.write(report_path, arcname=os.path.basename(report_path))
             # Add HTML Report
             zf.write(html_report_path, arcname=os.path.basename(html_report_path))

        # Store data in session for GET request
        session['processing_result'] = {
            'filename': file.filename,
            'before': before,
            'after': after,
            'mapping': mapping,
            'status_msg': status_msg,
            'report_file': os.path.basename(report_path),
            'doc_file': os.path.basename(doc_path),
            'zip_file': zip_filename
        }
        
        return redirect(url_for('process'))

    # GET request - retrieve from session
    result = session.get('processing_result')
    if not result:
        return redirect(url_for('upload_file'))
        
    return render_template(
        "validation_results.html",
        filename=result['filename'],
        results=result['after'],
        before=result['before'],
        mapping=result['mapping'],
        status_msg=result['status_msg'],
        report_file=result['report_file'],
        doc_file=result['doc_file'],
        zip_file=result.get('zip_file')
    )


@app.route("/download/<path:filename>")
def download_file(filename):
    # Security: Ensure filename is in UPLOAD_DIR
    return send_file(os.path.join(UPLOAD_DIR, filename), as_attachment=True)


if __name__ == "__main__":
    app.run(debug=True, port=5000)
