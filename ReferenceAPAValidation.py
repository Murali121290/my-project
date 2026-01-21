import os
import re
import logging
from flask import Flask, render_template, request, send_file, redirect, url_for, flash, make_response, session
from werkzeug.utils import secure_filename
from docx import Document
import io
import zipfile
from citation_parsers import get_parser, auto_detect_style

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.secret_key = os.environ.get('SESSION_SECRET', 'dev-secret-key')

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'docx'}

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_docx(file_path):
    doc = Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        paragraphs.append(para.text)
    return paragraphs


def normalize_citation_key(author_part, year):
    author_clean = re.sub(r'[^\w\s&]', '', author_part).strip()
    author_clean = re.sub(r'\s+', ' ', author_clean)
    return f"{author_clean}|{year}"


# Ported from Referencenumvalidation.py
def find_duplicates(references, reference_details):
    """
    Finds duplicate references using fuzzy matching (difflib) on FULL TEXT.
    Args:
        references: Dictionary of reference objects.
        reference_details: Dictionary containing 'text' for each reference key.
    Returns:
        List of dicts: {'id': str, 'text': str, 'duplicate_of': str, 'score': float}
    """
    import difflib
    
    duplicates = []
    processed_refs = [] 
    for key, data in references.items():
        # Use full original text for comparison to avoid false positives on "Author (Year)"
        text = reference_details.get(key, {}).get('text', data.get('display', ''))
        processed_refs.append({'id': key, 'text': text})
        
    n = len(processed_refs)
    for i in range(n):
        ref_a = processed_refs[i]
        for j in range(i + 1, n):
            ref_b = processed_refs[j]
            
            # Simple fuzzy match on the display string
            len_a = len(ref_a['text'])
            len_b = len(ref_b['text'])
            if len_a == 0 or len_b == 0: continue
            
            if min(len_a, len_b) / max(len_a, len_b) < 0.6:
                continue
            
            ratio = difflib.SequenceMatcher(None, ref_a['text'], ref_b['text']).ratio()
            
            if ratio > 0.85:
                duplicates.append({
                    'id': ref_b['id'], 
                    'text': ref_b['text'][:100],
                    'duplicate_of': ref_a['id'],
                    'score': round(ratio * 100, 1)
                })
        

    return duplicates


def parse_single_citation(cite_text):
    cite_text = cite_text.strip()
    cite_text = re.sub(r'\[[^\]]+\]', '', cite_text).strip()
    cite_text = re.sub(r'^(see|cf\.?|e\.g\.?,?|i\.e\.?,?)\s+', '', cite_text, flags=re.IGNORECASE).strip()
    
    if re.search(r'\bp\.?\s*\d+', cite_text, re.IGNORECASE):
        return [(None, None)]
    
    if re.search(r'(January|February|March|April|May|June|July|August|September|October|November|December)', cite_text, re.IGNORECASE):
        return [(None, None)]
    
    years = re.findall(r'\b((?:19|20)\d{2})[a-z]?\b', cite_text)
    if years:
        author_part = re.sub(r',?\s*(?:19|20)\d{2}[a-z]?,?\s*', '', cite_text).strip()
        author_part = re.sub(r',\s*$', '', author_part).strip()
        author_part = re.sub(r'^\d+,?\s*', '', author_part).strip()
        
        if not author_part or len(author_part) < 2:
            return [(None, None)]
        
        results = []
        for year in years:
            results.append((author_part, year))
        return results if results else [(None, None)]
    return [(None, None)]


def find_citations_in_text(paragraphs, parser):
    """
    Find citations using the provided parser.
    """
    citations = {}
    citation_locations = {}
    
    for i, para in enumerate(paragraphs):
        # Stop if we hit the bibliography section start
        if '<ref-open>' in para:
            break
            
        # Pass the whole paragraph text to the parser
        # The parser is now responsible for finding parenthetical AND narrative citations
        found_citations = parser.parse_citation(para)
        
        for cite in found_citations:
            author = cite['author']
            year = cite['year']
            
            # Create a unique key
            citation_key = normalize_citation_key(author, year)
            
            # Construct display string
            if cite['type'] == 'parenthetical':
                display_key = f"({author}, {year})"
            else:
                display_key = f"{author} ({year})"
            
            if citation_key not in citations:
                citations[citation_key] = {
                    'display': display_key,
                    'author': author,
                    'year': year,
                    'type': cite['type'],
                    'warnings': cite.get('warnings', []),
                    'raw': cite.get('raw', '')
                }
            else:
                # Merge warnings if new ones found
                citations[citation_key]['warnings'].extend([w for w in cite.get('warnings', []) if w not in citations[citation_key]['warnings']])

            if citation_key not in citation_locations:
                citation_locations[citation_key] = []
            citation_locations[citation_key].append(i + 1)
    
    return citations, citation_locations


def find_references_in_bibliography(paragraphs, parser):
    """
    Find references strictly between <ref-open> and <ref-close> tags.
    """
    references = {}
    reference_details = {}
    abbreviation_map = {}
    
    in_references_section = False
    
    for i, para in enumerate(paragraphs):
        para_stripped = para.strip()
        
        # Check for tags
        if '<ref-open>' in para_stripped:
            in_references_section = True
            # If the tag is on its own line, continue. If inline, we might need to parse the rest?
            # Assuming tag is a section delimiter.
            continue
        
        if '<ref-close>' in para_stripped:
            in_references_section = False
            continue
            
        if in_references_section and para_stripped:
            # Parse reference using parser
            ref_data = parser.parse_reference(para_stripped)
            
            if ref_data:
                author_display = ref_data['author']
                year = ref_data['year']
                full_author = ref_data['full_author']
                abbreviations = ref_data['abbreviations']
                
                ref_key = normalize_citation_key(author_display, year)
                
                if ref_key not in references:
                    references[ref_key] = {
                        'display': f"{author_display} ({year})",
                        'author': author_display,
                        'year': year,
                        'full_author': full_author,
                        'abbreviations': abbreviations
                    }
                    reference_details[ref_key] = {
                        'line': i + 1,
                        'text': para_stripped[:150] + ('...' if len(para_stripped) > 150 else '')
                    }
                    
                    for abbr in abbreviations:
                        abbr_key = f"{abbr}|{year}"
                        abbreviation_map[abbr_key] = ref_key
    
    return references, reference_details, abbreviation_map


def extract_first_surname(author_str):
    author_str = author_str.strip()
    author_str = re.sub(r'\s*et\s+al\.?\s*', '', author_str, flags=re.IGNORECASE)
    
    if ',' in author_str:
        return author_str.split(',')[0].strip()
    
    parts = author_str.split()
    return parts[0] if parts else ''


def match_citation_to_reference(citation, references):
    cite_author = citation['author'].strip()
    cite_year = citation['year']
    cite_author_lower = cite_author.lower()
    
    for ref_key, ref_data in references.items():
        ref_year = ref_data['year']
        ref_full_author = ref_data.get('full_author', ref_data['author'])
        ref_full_lower = ref_full_author.lower()
        
        if cite_year != ref_year:
            continue
        
        cite_first = extract_first_surname(cite_author)
        ref_first = extract_first_surname(ref_full_author)
        
        if cite_first and ref_first and cite_first == ref_first:
            return ref_key
        
        if cite_author_lower in ref_full_lower or ref_full_lower.startswith(cite_author_lower):
            return ref_key
        
        # Word subset matching for basic cases
        cite_words = set(re.findall(r'\b[a-z]{2,}\b', cite_author_lower))
        cite_words -= {'et', 'al', 'and', 'the'}
        ref_words = set(re.findall(r'\b[a-z]{2,}\b', ref_full_lower))
        
        if cite_words and cite_words.issubset(ref_words):
            return ref_key
    
    return None


import difflib
from difflib import SequenceMatcher

def normalize_text_for_comparison(text):
    """Normalize text for flexible matching."""
    # Replace 'and' with '&'
    text = re.sub(r'\band\b', '&', text, flags=re.IGNORECASE)
    # Remove dots, commas, extra spaces
    text = re.sub(r'[.,]', '', text)
    text = text.strip().lower()
    # Remove leading 'the'
    text = re.sub(r'^the\s+', '', text)
    return text

def check_smart_match(cite_data, references):
    """
    Advanced matching logic for:
    1. Introduction of abbreviations: "Organization [Org]" -> Match "Organization"
    2. Et al: "Smith et al" -> Match "Smith, Jones..."
    3. Proper subsets/variations: "Smith & Jones" -> Match "Smith, ... & Jones"
    4. Narrative conjunctions: "Smith and Jones" -> Match "Smith & Jones"
    """
    cite_author = cite_data['author']
    cite_year = cite_data['year']
    
    cite_norm = normalize_text_for_comparison(cite_author)
    
    # 1. Handle Abbreviation Introduction: "Name [Abbr]"
    # Extract "Name" part
    prefix_match = re.match(r'^(.*?)\s*\[.*?\]', cite_author)
    citation_prefix_norm = None
    if prefix_match:
        citation_prefix_norm = normalize_text_for_comparison(prefix_match.group(1))

    # 2. Handle "et al"
    is_etal = 'et al' in cite_norm
    cite_first_surname = cite_norm.split()[0] if cite_norm else ""
    
    # 3. Handle list of names (e.g. "Smith & Jones")
    cite_names = [n.strip() for n in re.split(r'[&]', cite_norm)]
    
    for ref_key, ref_data in references.items():
        if cite_year and ref_data['year'] != cite_year:
            continue
            
        ref_author = ref_data.get('full_author', ref_data['author'])
        ref_norm = normalize_text_for_comparison(ref_author)
        
        # A. Direct Normalized Match (covers "Smith and Jones" vs "Smith & Jones")
        if cite_norm == ref_norm:
            return ref_key
            
        # B. Abbreviation Definition Match
        # Citation: "National Org [NO]" vs Ref: "National Org"
        if citation_prefix_norm and citation_prefix_norm == ref_norm:
            return ref_key
            
        # C. Et Al Match
        # Citation: "Smith et al" vs Ref: "Smith, Jones..."
        # Rule: First surnames match
        if is_etal:
            ref_first_surname = ref_norm.split()[0]
            if cite_first_surname == ref_first_surname:
                return ref_key
                
        # D. Word Subset Match (Robust for "Smith, Jones" vs "Smith, A., Jones, B.")
        # Only check if citation has multiple words (potential authors)
        cite_words = set(re.findall(r'\b[a-z]{2,}\b', cite_norm))
        # Remove common stopwords from citation side to avoid false positives matching "and" to "and"
        cite_words -= {'and', 'the', 'et', 'al'}
        
        if len(cite_words) > 1:
            ref_words = set(re.findall(r'\b[a-z]{2,}\b', ref_norm))
            # Check if All significant citation words are in reference
            if cite_words.issubset(ref_words):
                return ref_key

    return None

def check_spelling_mismatch(cite_author, references):
    """
    Check for spelling mismatches using difflib.
    Returns reference key if a close match is found.
    """
    cite_author_norm = normalize_text_for_comparison(cite_author)
    
    best_match = None
    max_ratio = 0.0
    
    for ref_key, ref_data in references.items():
        ref_author = ref_data['author']
        ref_norm = normalize_text_for_comparison(ref_author)
        
        # simple ratio check
        ratio = difflib.SequenceMatcher(None, cite_author_norm, ref_norm).ratio()
        
        if ratio > 0.8: # >80% similarity
            if ratio > max_ratio:
                max_ratio = ratio
                best_match = ref_key
                
    return best_match


def check_et_al_misuse(cite_data, ref_data):
    """
    Check if 'et al.' is used incorrectly in citation based on reference author count.
    
    APA 7th Edition Rules:
    - 1-2 authors: Always cite all authors
    - 3+ authors: Use 'et al.' after first author
    
    Args:
        cite_data: Citation data dict with 'author' key
        ref_data: Reference data dict with 'full_author' key
        
    Returns:
        Dict with 'has_error' (bool) and 'message' (str) if error found, None otherwise
    """
    cite_author = cite_data['author'].strip()
    ref_full_author = ref_data.get('full_author', ref_data['author'])
    
    # Check if citation uses 'et al.'
    has_et_al = 'et al' in cite_author.lower()
    
    # Count authors in reference
    # Method: Count commas and ampersands
    # "Smith, J." = 1 author (1 comma)
    # "Smith, J., & Jones, M." = 2 authors (3 commas, 1 ampersand)
    # "Smith, J., Jones, M., & Brown, K." = 3 authors (5 commas, 1 ampersand)
    
    if '&' in ref_full_author:
        # Count author segments by splitting on '&'
        parts = ref_full_author.split('&')
        # First part has N-1 authors (separated by commas)
        # Last part has 1 author
        first_part_authors = len([p for p in parts[0].split(',') if p.strip() and not re.match(r'^[A-Z]\.?$', p.strip())])
        author_count = first_part_authors + 1
    else:
        # Single author (has comma for "Last, F.")
        author_count = 1
    
    # Check for misuse
    if has_et_al and author_count <= 2:
        # Error: Using et al. with 1-2 authors
        if author_count == 1:
            correct_form = ref_data['author']
        else:
            # Extract both author surnames for 2-author case
            ref_display = ref_data['author']
            correct_form = ref_display
        
        return {
            'has_error': True,
            'message': f"Change author per reference - use '{correct_form}', not 'et al.' (reference has only {author_count} author{'s' if author_count > 1 else ''})",
            'author_count': author_count,
            'correct_form': correct_form
        }
    
    elif not has_et_al and author_count >= 3:
        # Warning: Should use et al. with 3+ authors (but this is less critical)
        first_author = extract_first_surname(ref_full_author)
        correct_form = f"{first_author} et al."
        
        return {
            'has_error': True,
            'message': f"Consider using 'et al.' - reference has {author_count} authors, APA allows '{correct_form}'",
            'author_count': author_count,
            'correct_form': correct_form,
            'severity': 'warning'  # Less severe than incorrect et al. usage
        }
    
    return None

def get_citation_matches(citations, references, abbreviation_map):
    """
    Match citations to references using Exact, Abbreviation, and Smart/Fuzzy matching.
    Returns: matched_citations (set), matched_references (set), matched_pairs (dict: cite_key -> ref_key)
    """
    matched_citations = set()
    matched_references = set()
    matched_pairs = {}
    
    for cite_key, cite_data in citations.items():
        cite_author = cite_data['author'].strip()
        cite_year = cite_data['year']
        abbr_key = f"{cite_author}|{cite_year}"
        
        matched_ref_key = None
        
        # 1. Exact Match via Key
        if cite_key in references:
            matched_ref_key = cite_key
            
        # 2. Abbreviation Match
        elif abbr_key in abbreviation_map:
            matched_ref_key = abbreviation_map[abbr_key]
            
        # 3. Smart Match / Fuzzy Match
        else:
            smart_match_key = check_smart_match(cite_data, references)
            if smart_match_key:
                matched_ref_key = smart_match_key
        
        if matched_ref_key:
            matched_citations.add(cite_key)
            matched_references.add(matched_ref_key)
            matched_pairs[cite_key] = matched_ref_key
            
    return matched_citations, matched_references, matched_pairs


def check_abbreviation_usage(matched_pairs, citations, references, citation_locations):
    """
    Validate proper usage of abbreviations (First vs Subsequent usage).
    """
    abbreviation_errors = []
    from collections import defaultdict
    ref_usage_map = defaultdict(list)
    
    # 1. Build Usage Map
    for cite_key, ref_key in matched_pairs.items():
        cite_data = citations[cite_key]
        ref_data = references.get(ref_key)
        
        if not ref_data: continue

        ref_abbrs = ref_data.get('abbreviations', [])
        if not ref_abbrs: continue

        is_intro = '[' in cite_data['author'] 
        
        # Check if citation IS the abbreviation
        is_abbr = False
        cite_author_clean = cite_data['author'].strip()
        for abbr in ref_abbrs:
             if abbr == cite_author_clean: 
                 is_abbr = True
                 break
        
        # Check if full name
        ref_full_text = ref_data.get('full_author', '').split('(')[0].strip()
        is_full = False
        if cite_author_clean.lower().startswith(ref_full_text.lower()):
            is_full = True
            
        if is_intro: 
            is_full = True 
            is_abbr = False 

        if citation_locations.get(cite_key):
             for loc in citation_locations[cite_key]:
                 ref_usage_map[ref_key].append({
                     'loc': loc,
                     'text': cite_data['display'],
                     'is_intro': is_intro,
                     'is_abbr': is_abbr,
                     'is_full': is_full,
                     'ref_abbr': ref_abbrs[0]
                 })

    # 2. Validate Usage Order
    for ref_key, usages in ref_usage_map.items():
        usages.sort(key=lambda x: x['loc'])
        
        if not usages: continue
            
        # Check FIRST usage
        first_usage = usages[0]
        if first_usage['is_abbr'] and not first_usage['is_full']:
             abbreviation_errors.append({
                'citation': first_usage['text'],
                'message': f"First confirmation of abbreviation should define it. Use 'Full Name [Abbr]' instead of '{first_usage['text']}'.",
                'locations': [first_usage['loc']],
                'ref_abbr': first_usage['ref_abbr']
            })
            
        # Check SUBSEQUENT usages
        ref_abbr = usages[0]['ref_abbr']
        for i in range(1, len(usages)):
            usage = usages[i]
            
            if usage['is_intro']: 
                 abbreviation_errors.append({
                    'citation': usage['text'],
                    'message': f"Abbreviation already introduced. Use '{ref_abbr}' instead.",
                    'locations': [usage['loc']],
                    'ref_abbr': ref_abbr
                })
            elif usage['is_full'] and not usage['is_abbr']: 
                 abbreviation_errors.append({
                    'citation': usage['text'],
                    'message': f"Abbreviation previously introduced. Consider using '{ref_abbr}' instead.",
                    'locations': [usage['loc']],
                    'ref_abbr': ref_abbr,
                    'severity': 'warning'
                })
                
    return abbreviation_errors


def validate_document(file_path, parser=None):
    if parser is None:
        parser = get_parser('apa')
        
    paragraphs = extract_text_from_docx(file_path)
    
    citations, citation_locations = find_citations_in_text(paragraphs, parser)
    references, reference_details, abbreviation_map = find_references_in_bibliography(paragraphs, parser)
    
    # 1. Match Citations
    matched_citations, matched_references, matched_pairs = get_citation_matches(citations, references, abbreviation_map)
    
    # Validation Results Containers
    missing_refs = []
    unused_refs = []
    valid_citations = []
    format_errors = []
    year_mismatches = []
    spelling_mismatches = []
    et_al_errors = []
    abbreviation_errors = []
    
    # 2. Process Matches and Identify Mismatches
    for cite_key, cite_data in citations.items():
        cite_author = cite_data['author'].strip()
        cite_year = cite_data['year']
        
        if cite_key in matched_citations:
            valid_citations.append(cite_data['display'])
            
            # Et al. Check
            ref_key = matched_pairs[cite_key]
            et_al_check = check_et_al_misuse(cite_data, references[ref_key])
            if et_al_check and et_al_check['has_error']:
                et_al_errors.append({
                    'citation': cite_data['display'],
                    'message': et_al_check['message'],
                    'correct_form': et_al_check['correct_form'],
                    'author_count': et_al_check['author_count'],
                    'locations': citation_locations.get(cite_key, []),
                    'severity': et_al_check.get('severity', 'error')
                })
        else:
            # Matches failed, check for Year Mismatch
            found_year_match = False
            for ref_key, ref_data in references.items():
                if normalize_citation_key(ref_data['author'], '0000').split('|')[0] == normalize_citation_key(cite_author, '0000').split('|')[0]:
                     year_mismatches.append({
                        'citation': cite_data['display'],
                        'cited_year': cite_year,
                        'ref_year': ref_data['year'],
                        'ref_key': ref_key,
                        'locations': citation_locations.get(cite_key, [])
                    })
                     found_year_match = True
                     break
            
            if found_year_match:
                if cite_data['warnings']:
                     format_errors.append({
                        'citation': cite_data['display'],
                        'warnings': cite_data['warnings'],
                        'locations': citation_locations.get(cite_key, [])
                    })
                continue
            
            # Check for Spelling Mismatch
            potential_match_key = check_spelling_mismatch(cite_author, references)
            if potential_match_key:
                spelling_mismatches.append({
                    'citation': cite_data['display'],
                    'cited_author': cite_author,
                    'ref_author': references[potential_match_key]['author'],
                    'ref_key': potential_match_key,
                    'locations': citation_locations.get(cite_key, [])
                })
                if cite_data['warnings']:
                     format_errors.append({
                        'citation': cite_data['display'],
                        'warnings': cite_data['warnings'],
                        'locations': citation_locations.get(cite_key, [])
                    })
                continue
                
            # Missing Reference
            missing_refs.append({
                'reference': cite_data['display'],
                'cited_at_paragraphs': citation_locations.get(cite_key, [])
            })

        if cite_data['warnings']:
             format_errors.append({
                'citation': cite_data['display'],
                'warnings': cite_data['warnings'],
                'locations': citation_locations.get(cite_key, [])
            })

    # 3. Check Abbreviation Usage
    abbreviation_errors = check_abbreviation_usage(matched_pairs, citations, references, citation_locations)
    
    # 4. Find Unused References
    for ref_key, ref_data in references.items():
        if ref_key not in matched_references:
            involved_in_mismatch = False
            for ym in year_mismatches:
                if ym['ref_key'] == ref_key: involved_in_mismatch = True
            for sm in spelling_mismatches:
                if sm['ref_key'] == ref_key: involved_in_mismatch = True
                
            if not involved_in_mismatch:
                detail = reference_details.get(ref_key, {})
                unused_refs.append({
                    'reference': ref_data['display'],
                    'line': detail.get('line', 'Unknown'),
                    'text': detail.get('text', 'N/A')
                })

    # 5. Find Duplicates
    duplicates_list = find_duplicates(references, reference_details)

    return {
        'total_citations': len(citations),
        'total_references': len(references),
        'valid_count': len(matched_citations),
        'missing_references': missing_refs,
        'unused_references': unused_refs,
        'duplicates': duplicates_list,
        'valid_citations': sorted(valid_citations),
        'format_errors': format_errors,
        'year_mismatches': year_mismatches,
        'spelling_mismatches': spelling_mismatches,
        'et_al_errors': et_al_errors,
        'abbreviation_errors': abbreviation_errors,
        'all_citations': list(citations.keys()),
        'all_references': list(references.keys()),
        'matched_citation_keys': list(matched_citations),
        'citations': citations,
        'references': references,
        'citation_locations': citation_locations,
        'reference_details': reference_details
    }

def validate_document_multi_style(file_path, citation_style=None):
    """
    Validate document with support for multiple citation styles.
    
    Args:
        file_path: Path to Word document
        citation_style: Optional style ('apa', 'vancouver', 'chicago'). 
                       If None, will auto-detect.
    
    Returns:
        Dict with validation results including detected style
    """
    paragraphs = extract_text_from_docx(file_path)
    
    # Auto-detect style if not specified
    if citation_style is None:
        sample_text = ' '.join(paragraphs[:50])  # Use first 50 paragraphs
        citation_style = auto_detect_style(sample_text)
        logger.info(f"Auto-detected citation style: {citation_style}")
    
    # Get appropriate parser
    try:
        parser = get_parser(citation_style)
        logger.info(f"Using {citation_style.upper()} parser")
    except ValueError as e:
        # Fall back to APA if style not supported
        logger.warning(f"Unsupported style '{citation_style}', falling back to APA: {e}")
        parser = get_parser('apa')
        citation_style = 'apa'
    
    # Use the detected parser with validation logic
    results = validate_document(file_path, parser=parser)
    
    # Add detected style to results
    results['citation_style'] = citation_style.upper()
    results['citation_style_name'] = {
        'apa': 'APA (American Psychological Association)',
        'vancouver': 'Vancouver',
        'chicago': 'Chicago (Author-Year)'
    }.get(citation_style, citation_style.upper())
    
    return results


def generate_report(results, filename):
    # Calculate Total Comments/Issues for Status Line
    total_issues = (
        len(results['missing_references']) + 
        len(results['unused_references']) + 
        len(results.get('format_errors', [])) + 
        len(results.get('year_mismatches', [])) + 
        len(results.get('spelling_mismatches', [])) +
        len(results.get('et_al_errors', [])) +
        len(results.get('abbreviation_errors', []))
    )
    
    # Determine Style Label
    style_label = "Name/Year"
    if "VANCOUVER" in results.get('citation_style', '').upper():
         style_label = "Numerical"

    report_lines = []
    
    # Status Header
    report_lines.append(f"STATUS: {style_label}: {total_issues} comments")
    report_lines.append("")
    
    # Main Title
    report_lines.append("=" * 60)
    report_lines.append("NAME AND YEAR VALIDATION REPORT")
    report_lines.append("=" * 60)
    
    # Previous report content follows...
    # report_lines.append("=" * 60) # User removed this dup line in example, but kept CITATION VALIDATION REPORT below?
    # Actually the user example has:
    # STATUS...
    # ===
    # NAME AND YEAR...
    # ===
    # ===
    # CITATION VALIDATION...
    
    # I will replicate this structure exactly to be safe.
    
    report_lines.append("=" * 60)
    report_lines.append("CITATION VALIDATION REPORT")
    report_lines.append("=" * 60)
    report_lines.append(f"\nDocument: {filename}")
    report_lines.append(f"Style: {results.get('citation_style_name', 'APA')}")
    report_lines.append("-" * 60)
    
    report_lines.append("\nSUMMARY:")
    report_lines.append(f"  Total in-text citations found: {results['total_citations']}")
    report_lines.append(f"  Total references in bibliography: {results['total_references']}")
    report_lines.append(f"  Valid (matched) citations: {results['valid_count']}")
    report_lines.append(f"  Missing references: {len(results['missing_references'])}")
    report_lines.append(f"  Unused references: {len(results['unused_references'])}")
    report_lines.append(f"  Format Errors: {len(results.get('format_errors', []))}")
    report_lines.append(f"  Year Mismatches: {len(results.get('year_mismatches', []))}")
    report_lines.append(f"  Spelling Mismatches: {len(results.get('spelling_mismatches', []))}")
    report_lines.append(f"  Spelling Mismatches: {len(results.get('spelling_mismatches', []))}")
    report_lines.append(f"  Et Al. Errors: {len(results.get('et_al_errors', []))}")
    report_lines.append(f"  Abbreviation Errors: {len(results.get('abbreviation_errors', []))}")
    
    if results['missing_references']:
        report_lines.append("\n" + "-" * 60)
        report_lines.append("MISSING REFERENCES (cited but not in bibliography):")
        report_lines.append("-" * 60)
        for item in results['missing_references']:
            report_lines.append(f"\n  {item['reference']}")
            report_lines.append(f"    Cited in paragraph(s): {', '.join(map(str, item['cited_at_paragraphs']))}")
    
    if results.get('year_mismatches'):
        report_lines.append("\n" + "-" * 60)
        report_lines.append("YEAR MISMATCHES (Author matches but year differs):")
        report_lines.append("-" * 60)
        for item in results['year_mismatches']:
            report_lines.append(f"\n  Citation: {item['citation']}")
            report_lines.append(f"  Reference Year: {item['ref_year']}")
            report_lines.append(f"  Cited in paragraph(s): {', '.join(map(str, item['locations']))}")

    if results.get('spelling_mismatches'):
        report_lines.append("\n" + "-" * 60)
        report_lines.append("SPELLING MISMATCHES (Author spelling differs):")
        report_lines.append("-" * 60)
        for item in results['spelling_mismatches']:
            report_lines.append(f"\n  Citation: {item['citation']}")
            report_lines.append(f"  Cited Author: {item['cited_author']}")
            report_lines.append(f"  Ref Author: {item['ref_author']}")
            report_lines.append(f"  Cited in paragraph(s): {', '.join(map(str, item['locations']))}")

    if results.get('et_al_errors'):
        report_lines.append("\n" + "-" * 60)
        report_lines.append("ET AL. ERRORS (Incorrect use of 'et al.'):") 
        report_lines.append("-" * 60)
        for item in results['et_al_errors']:
            report_lines.append(f"\n  Citation: {item['citation']}")
            report_lines.append(f"  Issue: {item['message']}")
            report_lines.append(f"  Correct Form: {item['correct_form']}")
            report_lines.append(f"  Cited in paragraph(s): {', '.join(map(str, item['locations']))}")

    if results.get('abbreviation_errors'):
        report_lines.append("\n" + "-" * 60)
        report_lines.append("ABBREVIATION ERRORS (First vs Subsequent Usage):") 
        report_lines.append("-" * 60)
        for item in results['abbreviation_errors']:
            report_lines.append(f"\n  Citation: {item['citation']}")
            report_lines.append(f"  Issue: {item['message']}")
            report_lines.append(f"  Cited in paragraph(s): {', '.join(map(str, item['locations']))}")

    if results.get('duplicates'):
        report_lines.append("\n" + "-" * 60)
        report_lines.append("DUPLICATE REFERENCES:")
        report_lines.append("-" * 60)
        for d in results['duplicates']:
            report_lines.append(f"\n  Original ID: {d['duplicate_of']}")
            report_lines.append(f"  Duplicate ID: {d['id']}")
            report_lines.append(f"  Text: {d['text']}")
            report_lines.append(f"  Similarity Score: {d['score']}%")
            
    if results.get('format_errors'):
        report_lines.append("\n" + "-" * 60)
        report_lines.append("FORMAT ERRORS (APA Style Violations):")
        report_lines.append("-" * 60)
        for item in results['format_errors']:
            report_lines.append(f"\n  Citation: {item['citation']}")
            for w in item['warnings']:
                report_lines.append(f"    - {w}")
            report_lines.append(f"    Cited in paragraph(s): {', '.join(map(str, item['locations']))}")
    
    if results['unused_references']:
        report_lines.append("\n" + "-" * 60)
        report_lines.append("UNUSED REFERENCES (in bibliography but never cited):")
        report_lines.append("-" * 60)
        for item in results['unused_references']:
            report_lines.append(f"\\n  {item['reference']}")
            report_lines.append(f"    Line: {item['line']}")
            report_lines.append(f"    Text: {item['text']}")
    
    if results['valid_citations']:
        report_lines.append("\\n" + "-" * 60)
        report_lines.append("VALID CITATIONS:")
        report_lines.append("-" * 60)
        for ref in results['valid_citations']:
            report_lines.append(f"  {ref}")
    
    report_lines.append("\\n" + "=" * 60)
    report_lines.append("END OF REPORT")
    report_lines.append("=" * 60)
    
    return "\n".join(report_lines)


def find_citation_in_runs(paragraph, citation_text, fuzzy_threshold=0.75):
    """
    Find the specific run(s) containing a citation text in a paragraph.
    Uses robust index mapping to handle citations split across multiple runs.
    
    Args:
        paragraph: python-docx Paragraph object
        citation_text: Citation text to find (e.g., "(Smith, 2020)")
        fuzzy_threshold: Minimum similarity ratio for fuzzy matching (0.0-1.0) - Unused in strict index mode but kept for signature
        
    Returns:
        List of Run objects containing the citation.
    """
    if not citation_text:
        return []

    # 1. Build full text and map run indices
    full_text = ""
    run_map = [] # List of (start_index, end_index, run_object)
    
    current_idx = 0
    for run in paragraph.runs:
        text_len = len(run.text)
        run_map.append((current_idx, current_idx + text_len, run))
        full_text += run.text
        current_idx += text_len
        
    # 2. Find citation in full text
    # Normalize for loose matching (ignore case/whitespace diffs if exact fails)
    
    # Try Exact Match first
    start_pos = full_text.find(citation_text)
    
    if start_pos == -1:
        # Try Case Insensitive
        start_pos = full_text.lower().find(citation_text.lower())
        
    if start_pos == -1:
        # Try normalizing whitespace (collapse spaces)
        # This is harder to map back to indices directly if we change length.
        # But commonly the issue is just simple splits.
        # Let's try aggressive "remove all spaces" match as a fallback?
        # No, that breaks index mapping.
        # Let's rely on the calling function to provide varying candidates (regex output).
        return []
        
    end_pos = start_pos + len(citation_text)
    
    # 3. Collect Runs Overlapping with [start_pos, end_pos]
    matched_runs = []
    
    for r_start, r_end, run_obj in run_map:
        # Check for overlap
        # Overlap if run starts before match ends AND run ends after match starts
        if r_start < end_pos and r_end > start_pos:
            matched_runs.append(run_obj)
            
    return matched_runs


def insert_comments_in_document(file_path, results, citation_locations, reference_details):
    """
    Insert Word comments for missing references, unused references, and mismatches.
    """
    doc = Document(file_path)
    comment_count = 0
    
    # 1. Missing References
    for item in results['missing_references']:
        citation_text = item['reference']
        # The parser might return display format, which might not match exact text in doc if we normalized it.
        # But we stored 'raw' in parser if we used it, but here we iterate existing list.
        # We need to rely on citation_locations which links Key -> Paragraph Index.
        # We'll need a way to link 'item' back to key or rely on display text searching.
        
        # Simpler approach: Iterate citation_locations
        pass # Logic handled below generically
    
    # helper to add comment with improved matching
    def add_comment_to_citation(cite_text, paragraphs, message):
        count = 0
        for para_num in paragraphs:
            if para_num <= len(doc.paragraphs):
                para = doc.paragraphs[para_num - 1]
                
                # Try 1: Exact index match with citation text
                runs = find_citation_in_runs(para, cite_text, fuzzy_threshold=0.75)
                
                # Try 2: Regex Based Candidates (Robust Fallback)
                if not runs:
                    import re
                    # Extract year
                    year_match = re.search(r'(\b\d{4}[a-z]?\b|n\.d\.|in press)', cite_text, re.IGNORECASE)
                    year_part = year_match.group(1) if year_match else None
                    
                    # Extract author (everything before year or parens often works)
                    author_part = None
                    if '(' in cite_text:
                        author_part = cite_text.split('(')[0].strip()
                        if not author_part: # Leading paren case: (Author, Year)
                            # Remove leading paren and extract until comma/year
                            cleaned = cite_text.replace('(', '').replace(')', '')
                            if year_part:
                                author_part = cleaned.split(year_part)[0].strip(' ,.')
                            else:
                                author_part = cleaned
                    else:
                        if year_part:
                            author_part = cite_text.split(year_part)[0].strip(' ,(')
                        else:
                            author_part = cite_text.strip()
                            
                    if author_part and year_part:
                        # Escape special chars
                        author_regex = re.escape(author_part).replace(r'\ ', r'\s+')
                        regexes = []
                        # 1. Parenthetical Pattern: (Author... Year...)
                        regexes.append(re.compile(r'\([^)]*?' + author_regex + r'.*?' + re.escape(year_part) + r'.*?\)', re.IGNORECASE))
                        # 2. Narrative Pattern: Author... (Year...)
                        regexes.append(re.compile(author_regex + r'.*?\([^)]*?' + re.escape(year_part) + r'.*?\)', re.IGNORECASE))
                        
                        candidates = []
                        for pattern in regexes:
                            candidates.extend(pattern.findall(para.text))
                            
                        for cand in candidates:
                            runs = find_citation_in_runs(para, cand)
                            if runs: break

                # Fallback 3: Just Author if all else fails
                if not runs and author_part and len(author_part) > 2 and "Unknown" not in author_part:
                     runs = find_citation_in_runs(para, author_part)
                     
                # Fallback 4: Just Year
                if not runs and year_part:
                     if year_part in para.text:
                         # Anchor to first run or try to find year specifically?
                         # Finding year in runs using find_citation_in_runs
                         runs = find_citation_in_runs(para, year_part)

                # Fallback 5: Unknown/Fail -> First run
                if not runs and "Unknown" in cite_text and para.runs:
                     runs = [para.runs[0]]
                
                if runs:
                    try:
                        doc.add_comment(
                            runs=runs,
                            text=message,
                            author="Citation Checker",
                            initials="CC"
                        )
                        count += 1
                        logger.debug(f"Successfully added comment to paragraph {para_num}")
                    except Exception as e:
                        logger.warning(f"Failed to add comment for '{cite_text}' at paragraph {para_num}: {e}")
                else:
                    logger.warning(f"Could not locate citation '{cite_text}' in paragraph {para_num}")
        
        return count

    # Add comments for Missing References
    for item in results['missing_references']:
        msg = f"⚠️ MISSING REFERENCE: '{item['reference']}' is not in the bibliography."
        comment_count += add_comment_to_citation(item['reference'], item['cited_at_paragraphs'], msg)

    # Add comments for Year Mismatches
    for item in results.get('year_mismatches', []):
        msg = f"📅 YEAR MISMATCH: You cited '{item['cited_year']}' but bibliography has '{item['ref_year']}'."
        comment_count += add_comment_to_citation(item['citation'], item['locations'], msg)

    # Add comments for Spelling Mismatches
    for item in results.get('spelling_mismatches', []):
        msg = f"🔤 SPELLING MISMATCH: Cited as '{item['cited_author']}' but bibliography has '{item['ref_author']}'."
        comment_count += add_comment_to_citation(item['citation'], item['locations'], msg)

    # Add comments for Format Errors
    for item in results.get('format_errors', []):
        warnings_str = "\\n".join(item['warnings'])
        msg = f"📝 FORMAT ERROR: {warnings_str}"
        comment_count += add_comment_to_citation(item['citation'], item['locations'], msg)
    
    # Add comments for Et Al. Errors
    for item in results.get('et_al_errors', []):
        msg = f"✏️ ET AL. ERROR: {item['message']}"
        comment_count += add_comment_to_citation(item['citation'], item['locations'], msg)
    
    # Add comments for Abbreviation Errors (NEW)
    for item in results.get('abbreviation_errors', []):
        msg = f"🔤 ABBREVIATION ERROR: {item['message']}"
        comment_count += add_comment_to_citation(item['citation'], item['locations'], msg)

    # Add comments for Unused References
    for item in results['unused_references']:


        # We need to find the reference in the bibliography section again?
        # or use reference_details which has line number.
        line_num = item['line']
        if line_num != 'Unknown' and line_num <= len(doc.paragraphs):
             para = doc.paragraphs[line_num - 1]
             if para.runs:
                try:
                    runs_to_comment = para.runs[:min(3, len(para.runs))]
                    doc.add_comment(
                        runs=runs_to_comment,
                        text=f"ℹ️ UNUSED REFERENCE: This reference is never cited in the text.",
                        author="Citation Checker",
                        initials="CC"
                    )
                    comment_count += 1
                except:
                    pass
    
    return doc, comment_count


def apply_citation_formatting(file_path, results):
    """
    Format ALL citations in the document:
    1. Shorten multi-author citations (3+ authors) to 'et al.' if not already (for valid matches).
    2. Apply 'cite_bib' character style.
    3. Highlight in Green (Valid/Matched) or Yellow (Unmatched/Mismatch).
    """
    from docx.shared import RGBColor
    from docx.enum.text import WD_COLOR_INDEX
    
    doc = Document(file_path)
    count = 0
    
    # Ensure style exists
    from docx.enum.style import WD_STYLE_TYPE
    styles = doc.styles
    try:
        styles['cite_bib']
    except KeyError:
        # Create character style if missing
        style = styles.add_style('cite_bib', WD_STYLE_TYPE.CHARACTER)
        #font = style.font
        #font.superscript = True
        # Optional: Set color or other properties if desired by default
        # font.color.rgb = RGBColor(0, 0, 0)
    
    citations = results['citations']
    references = results['references']
    citation_locations = results['citation_locations']
    matched_keys = set(results.get('matched_citation_keys', []))
    
    # Iterate through ALL detected citations
    # We sort by location availability to group work? No, just iterate dict.
    # But same citation key might have multiple locations.
    
    for cite_key, cite_data in citations.items():
        location_indices = citation_locations.get(cite_key, [])
        
        # Determine Status
        is_valid = cite_key in matched_keys
        
        # Determine Color
        highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN if is_valid else WD_COLOR_INDEX.YELLOW
        
        # Determine Shortening (Only for valid matches)
        should_shorten = False
        full_shortened_str = ""
        first_author_surname = ""
        
        if is_valid:
            target_ref_data = None
            if cite_key in references:
                target_ref_data = references[cite_key]
            else:
                match_found = match_citation_to_reference(cite_data, references)
                if not match_found:
                     match_found = check_smart_match(cite_data, references)
                if match_found:
                    target_ref_data = references[match_found]
            
            # Shortening disabled as per user request to preserve original text
            # if target_ref_data:
            #     ref_full_author = target_ref_data.get('full_author', '')
            #     comma_count = ref_full_author.count(',')
            #     if comma_count >= 3 and '&' in ref_full_author:
            #          should_shorten = True
            #     first_author_surname = extract_first_surname(ref_full_author)
            #     
            #     if cite_data['type'] == 'parenthetical':
            #          full_shortened_str = f"({first_author_surname} et al., {cite_data['year']})"
            #     else:
            #          full_shortened_str = f"{first_author_surname} et al. ({cite_data['year']})"

        # Apply to Document Paragraphs
        for para_idx in location_indices:
             if para_idx > len(doc.paragraphs): continue
             para = doc.paragraphs[para_idx - 1]
             
             citation_display = cite_data['display']
             
             # Locate text Strategy
             # 1. Try raw text
             # 2. Try display text
             # 3. Try Regex reconstruction (Robust fallback)
             
             search_candidates = []
             
             # Candidate 1: Raw text stored in parser data
             if 'raw' in cite_data and cite_data['raw']:
                 search_candidates.append(cite_data['raw'])
             
             # Candidate 2: Display text
             if citation_display and citation_display not in search_candidates:
                 search_candidates.append(citation_display)

             # Candidate 3: Regex match in this specific paragraph
             try:
                 author_part = cite_data.get('author', '').strip()
                 year_part = cite_data.get('year', '').strip()
                 
                 if author_part and year_part:
                     import re
                     # Escape special chars in author but allow for some flexibility
                     author_regex = re.escape(author_part).replace(r'\ ', r'\s+')
                     
                 if author_part and year_part:
                     import re
                     # Escape special chars in author but allow for some flexibility
                     author_regex = re.escape(author_part).replace(r'\ ', r'\s+')
                     
                     regexes = []
                     # 1. Parenthetical Pattern: (Author... Year...)
                     regexes.append(re.compile(r'\([^)]*?' + author_regex + r'.*?' + re.escape(year_part) + r'.*?\)', re.IGNORECASE))
                     # 2. Narrative Pattern: Author... (Year...)
                     regexes.append(re.compile(author_regex + r'.*?\([^)]*?' + re.escape(year_part) + r'.*?\)', re.IGNORECASE))
                     
                     for pattern in regexes:
                         matches = pattern.findall(para.text)
                         for m in matches:
                             if m not in search_candidates:
                                 search_candidates.append(m)
             except Exception as e:
                 logging.warning(f"Failed to generate regex candidate: {e}")

             
             run_group = None
             matched_candidate = None
             full_text_found = ""
             
             # Try candidates
             for candidate in search_candidates:
                 if not candidate: continue
                 # Clean candidates of trailing punctuation for search
                 clean_candidate = candidate.strip('.,; ')
                 
                 rg = find_citation_in_runs(para, clean_candidate, fuzzy_threshold=0.85)
                 if rg:
                     # Verify match
                     ft = "".join(r.text for r in rg)
                     # Check if it contains the essential parts (Author/Year) to confirm it's not a false positive
                     if clean_candidate in ft or cite_data['year'] in ft: # Loose check
                         run_group = rg
                         matched_candidate = clean_candidate
                         full_text_found = ft
                         break
             
             if run_group:
                 # Robust Replacement Logic
                 full_text = full_text_found
                 
                 # Collapse run group to single run
                 for r in run_group[1:]:
                     r.text = ""
                 
                 final_citation_text = matched_candidate
                 
                 # Apply Shortening if Valid
                 if is_valid and should_shorten and "et al" not in matched_candidate.lower():
                     has_open_paren = '(' in matched_candidate
                     has_close_paren = ')' in matched_candidate
                     
                     replacement_str = full_shortened_str
                     if not has_open_paren and not has_close_paren:
                         replacement_str = full_shortened_str.replace('(', '').replace(')', '')
                     
                     # Simple replace
                     full_text = full_text.replace(matched_candidate, replacement_str, 1)
                     final_citation_text = replacement_str
                 
                 # Split and Style
                 run = run_group[0]
                 run.text = full_text 
                 
                 try:
                     start_idx = full_text.find(final_citation_text)
                     if start_idx != -1:
                         end_idx = start_idx + len(final_citation_text)
                         
                         pre_text = full_text[:start_idx]
                         post_text = full_text[end_idx:]
                         
                         run.text = pre_text
                         
                         cite_run = para.add_run(final_citation_text)
                         cite_run.style = 'cite_bib'
                         cite_run.font.highlight_color = highlight_color
                         
                         run._element.addnext(cite_run._element)
                         
                         if post_text:
                             post_run = para.add_run(post_text)
                             if run.style and run.style.name != 'Default Paragraph Font':
                                  post_run.style = run.style
                             cite_run._element.addnext(post_run._element)
                             
                         count += 1
                     else:
                         # Fallback if text replacement made find impossible?
                         pass
                 except Exception as e:
                     logger.error(f"Error splitting runs for formatting: {e}")
                            
    doc.save(file_path)
    return count


import ReferencesStructing as RS
import shutil
import uuid
from pathlib import Path

# ... existing code ...

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file selected', 'error')
            return redirect(request.url)
        
        files = request.files.getlist('file')
        if not files or files[0].filename == '':
            flash('No file selected', 'error')
            return redirect(request.url)

        # Options
        check_validation = request.form.get('check_validation') == 'yes'
        check_structuring = request.form.get('check_structuring') == 'yes'
        citation_style = request.form.get('citation_style', 'auto')
        if citation_style == 'auto': citation_style = None

        if not check_validation and not check_structuring:
            flash('Please select at least one option (Reference Check or Structuring)', 'warning')
            return redirect(request.url)

        # Create Batch Directory
        batch_id = str(uuid.uuid4())[:8]
        batch_dir = os.path.join(app.config['UPLOAD_FOLDER'], f"batch_{batch_id}")
        os.makedirs(batch_dir, exist_ok=True)
        
        results_map = [] # To store (file_path, arcname)

        try:
            for file in files:
                if not (file.filename and allowed_file(file.filename)):
                    continue
                
                original_filename = secure_filename(file.filename)
                file_path = os.path.join(batch_dir, original_filename)
                file.save(file_path)

                # Track current working document
                current_doc_path = file_path
                current_doc_name = original_filename

                # 1. STRUCTURING PHASE
                if check_structuring:
                    try:
                        logger.info(f"Structuring references for {original_filename}")
                        # RS.process_docx_file returns dict with Path objects
                        struct_res = RS.process_docx_file(Path(file_path), Path(batch_dir))
                        
                        structured_path = str(struct_res['output_docx'])
                        log_path = str(struct_res['log_file'])
                        
                        if os.path.exists(structured_path):
                            # Add to ZIP results
                            results_map.append((structured_path, f"Structured/{current_doc_name.replace('.docx', '_Structured.docx')}"))
                            results_map.append((log_path, f"Logs/{current_doc_name.replace('.docx', '_Structuring_Log.txt')}"))
                            
                            # Update current doc for next phase (Validation)
                            current_doc_path = structured_path
                            current_doc_name = os.path.basename(structured_path)
                    except Exception as e:
                        logger.error(f"Error structuring {original_filename}: {e}")
                        # If structuring failed, continue with original for validation?
                        # Or just log error?
                        results_map.append((file_path, f"Errors/{original_filename}_Structuring_Failed.docx"))

                # 2. VALIDATION PHASE
                if check_validation:
                    try:
                        logger.info(f"Validating {current_doc_name}")
                        results = validate_document_multi_style(current_doc_path, citation_style)
                        report_text = generate_report(results, current_doc_name)
                        
                        # Save Report
                        report_filename = f"{current_doc_name}_Report.txt"
                        report_path = os.path.join(batch_dir, report_filename)
                        with open(report_path, 'w', encoding='utf-8') as f:
                            f.write(report_text)
                        
                        results_map.append((report_path, f"Reports/{report_filename}"))

                        # New Step: Auto-Formatting
                        formatted_count = apply_citation_formatting(current_doc_path, results)
                        if formatted_count > 0:
                             logger.info(f"Formatted {formatted_count} citations in {current_doc_name}")

                        # Create Annotated Doc if issues exist OR successful formatting occurred
                        has_issues = (results['missing_references'] or 
                                      results['unused_references'] or 
                                      results.get('format_errors') or 
                                      results.get('year_mismatches') or 
                                      results.get('spelling_mismatches'))
                        
                        if has_issues or formatted_count > 0:
                            doc, comment_count = insert_comments_in_document(
                                current_doc_path, 
                                results,
                                results['citation_locations'], 
                                results['reference_details']
                            )
                            if comment_count > 0 or formatted_count > 0:
                                annotated_filename = current_doc_name.replace('.docx', '_Annotated.docx')
                                annotated_path = os.path.join(batch_dir, annotated_filename)
                                doc.save(annotated_path)
                                results_map.append((annotated_path, f"Annotated/{annotated_filename}"))
                    except Exception as e:
                        logger.error(f"Error validating {current_doc_name}: {e}")

            # GENERATE ZIP
            if not results_map:
                flash('No results generated. Please check files and try again.', 'error')
                return redirect(request.url)

            memory_file = io.BytesIO()
            with zipfile.ZipFile(memory_file, 'w', zipfile.ZIP_DEFLATED) as zf:
                for src_path, arc_name in results_map:
                    if os.path.exists(src_path):
                        zf.write(src_path, arc_name)
            
            memory_file.seek(0)
            
            # Use send_file with a callback to clean up? 
            # Flask send_file doesn't support cleanup callback easily in older versions, 
            # but we can trust OS temp cleaning or do it in a finally block if we weren't returning
            # Since we are returning a stream, we can't delete immediately.
            # Best practice: schedule a cleanup or trust the unique dir is small enough until standard purge.
            # OR read bytes and delete.
            
            response = make_response(send_file(
                memory_file,
                mimetype='application/zip',
                as_attachment=True,
                download_name=f"Processed_Results_{batch_id}.zip"
            ))
            
            # Set cookie for frontend to detect download completion
            token = request.form.get('download_token')
            if token:
                response.set_cookie('download_token', token, max_age=10, path='/')
                
            return response

        except Exception as e:
            logger.error(f"Batch processing error: {e}", exc_info=True)
            flash(f'Error processing files: {str(e)}', 'error')
            return redirect(request.url)

    return render_template('index.html')

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
