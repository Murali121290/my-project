
import re
import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------
# Constants
# ---------------------------
COLOR_MAP = {
    "nohighlight": None,
    "black": "black",
    "blue": "blue",
    "turquoise": "cyan",  # Word uses 'cyan' for turquoise in XML
    "brightgreen": "green",
    "pink": "magenta",
    "red": "red",
    "yellow": "yellow",
    "white": "white",
    "darkBlue": "darkBlue",
    "teal": "darkCyan",
    "darkGreen": "darkGreen",
    "darkMagenta": "darkMagenta",
    "darkRed": "darkRed",
    "darkYellow": "darkYellow",
    "lightGray": "lightGray",
    "darkGray": "darkGray",
    # synonyms
    "cyan": "cyan",
    "green": "green",
    "magenta": "magenta",
}

def get_xml_color(name_or_index):
    """Map friendly name to Word XML highlight color value."""
    if isinstance(name_or_index, int):
        # Fallback if int passed (approximate mapping or ignore)
        return "yellow"
    return COLOR_MAP.get(name_or_index, "yellow")

# ---------------------------
# XML Helpers
# ---------------------------
def set_highlight(run, color_name):
    """
    Apply highlight to a run using OXML.
    color_name: 'yellow', 'cyan', 'magenta', 'turquoise', etc.
    Automatically converts to correct XML color value.
    """
    if not color_name:
        return
    
    # Get or create rPr
    rPr = run._element.get_or_add_rPr()
    # Check if highlight exists
    h = rPr.find(qn('w:highlight'))
    if h is None:
        h = OxmlElement('w:highlight')
        rPr.append(h)
    
    # Convert color name to XML value through get_xml_color
    xml_color = get_xml_color(color_name)
    h.set(qn('w:val'), xml_color)

def duplicate_run(run, parent_paragraph):
    """
    Create a copy of a run with the same properties/formatting
    and append it to the paragraph. Returns the new Run object.
    NOTE: usage requires careful insertion order handling.
    """
    # Create new run element
    new_r = OxmlElement('w:r')
    
    # Copy properties if they exist
    if run._element.rPr is not None:
        from copy import deepcopy
        new_r.append(deepcopy(run._element.rPr))
    
    # We don't copy text here, caller sets text
    new_run = type(run)(new_r, parent_paragraph)
    return new_run

# ---------------------------
# Core Highlight Engine
# ---------------------------
def copy_run_extras(src_run, dst_run):
    """
    Copy non-text elements (comments, footnotes, images, etc.) from src to dst.
    """
    # Elements handled by python-docx text/rPr properties that we shouldn't manually copy
    # to avoid duplication or conflict.
    SKIP_TAGS = {
        qn('w:rPr'), 
        qn('w:t'), 
        qn('w:br'), 
        qn('w:cr'), 
        qn('w:tab'), 
        qn('w:noBreakHyphen')
    }
    
    from copy import deepcopy
    for child in src_run._element:
        if child.tag not in SKIP_TAGS:
            dst_run._element.append(deepcopy(child))

# ---------------------------
# SAFETY CHECKS (COMMENTS & TRACK CHANGES)
# ---------------------------

def run_inside_comment(run):
    """
    Detect if a run is inside a Word comment range.
    """
    p = run._element.getparent()
    if p is None:
        return False

    children = list(p)
    try:
        idx = children.index(run._element)
    except ValueError:
        return False

    # Simple toggle logic for comment ranges
    # Improved to track IDs for nested/overlapping comments correctness
    open_ids = set()
    for el in children[:idx]:
        if el.tag == qn("w:commentRangeStart"):
            cid = el.get(qn("w:id"))
            open_ids.add(cid)
        elif el.tag == qn("w:commentRangeEnd"):
            cid = el.get(qn("w:id"))
            if cid in open_ids:
                open_ids.remove(cid)

    return len(open_ids) > 0


def run_inside_track_change(run):
    """
    Detect if a run is inside w:ins or w:del.
    """
    el = run._element
    while el is not None:
        if el.tag in {qn("w:ins"), qn("w:del")}:
            return True
        el = el.getparent()
    return False

# ---------------------------
# Core Highlight Engine
# ---------------------------
def highlight_paragraph(paragraph, compiled_regexes, color_name, preserve_comments=True):
    """
    Apply highlighting to a paragraph based on regex matches.
    Safe version that skips runs inside comments/track-changes to avoid corruption.
    """
    if not compiled_regexes:
        return
        
    text = paragraph.text
    if not text:
        return

    # Find all matches for all patterns
    ranges = []
    for pattern in compiled_regexes:
        for match in pattern.finditer(text):
            ranges.append((match.start(), match.end()))
            
    if not ranges:
        return
        
    # Find ignore zones based on any generic tags like <tag> ... </anytag>
    # Note: Using a heuristic where we look for `<word>` and later find another `<word>` or `</word>`
    # But for safety based on the prompt, any text wrapped between two tags, or any specific tags 
    # like <withoutspaceallcaps> we should ignore.
    ignore_zones = []
    
    # Generic regex to find anything that looks like an opening tag: <tagname>
    # or closing tag </tagname>
    # We will just collect all tags and treat them as on/off toggles or find pairwise matching.
    # Since XML tags in text might not be perfectly nested, a simple approach is:
    # 1. Find all `<[a-zA-Z0-9_]+>` (start tags)
    # 2. Find all `</[a-zA-Z0-9_]+>` or subsequent tags to close the zone.
    # However, to be highly resilient: we just find ALL tags, and assume text between 
    # an opening tag and ANY closing tag (or next tag) is an ignore zone if appropriate, 
    # but the safest is to explicitly look for `<tag>` and `</tag>` pairs.
    # Let's match `<[a-zA-Z0-9]+>` as start, and `</[a-zA-Z0-9]+>` as end.
    
    all_starts = list(re.finditer(r'<[a-zA-Z0-9_]+>', text))
    all_ends = list(re.finditer(r'<\/[a-zA-Z0-9_]+>', text))
    
    # We will match the nearest end tag that corresponds to the same tag name.
    for start_m in all_starts:
        s_pos = start_m.start()
        tag_str = start_m.group()
        tag_name = tag_str[1:-1] # strip < and >
        
        # Find closing tag
        end_pattern = f"</{tag_name}>"
        matching_end = next((e for e in all_ends if e.start() > s_pos and e.group().lower() == end_pattern.lower()), None)
        
        if matching_end:
            ignore_zones.append((s_pos, matching_end.end()))
        else:
            # If no explicit closing tag, maybe they meant the whole rest of paragraph,
            # or maybe there is just another opening tag used as a closer (like <CN> ... <CN>).
            # To be safe, we just ignore the tag itself if no closing tag is found.
            ignore_zones.append((s_pos, start_m.end()))
            
    # Filter out ranges that fall within any ignore zone
    filtered_ranges = []
    for r_s, r_e in ranges:
        is_ignored = False
        for i_s, i_e in ignore_zones:
            if r_s >= i_s and r_e <= i_e:
                is_ignored = True
                break
            elif r_s < i_e and r_e > i_s:
                # Partial overlap: since highlighting is generally at the word/phrase level,
                # we will skip the entire highlight if it overlaps with an ignore zone.
                is_ignored = True
                break
        
        if not is_ignored:
            filtered_ranges.append((r_s, r_e))
            
    ranges = filtered_ranges
    if not ranges:
        return
        
    # Merge overlaps
    ranges.sort()
    merged = []
    if ranges:
        curr_start, curr_end = ranges[0]
        for next_start, next_end in ranges[1:]:
            if next_start < curr_end:
                curr_end = max(curr_end, next_end)
            else:
                merged.append((curr_start, curr_end))
                curr_start, curr_end = next_start, next_end
        merged.append((curr_start, curr_end))
        
    # Build mask
    mask = [False] * len(text)
    for s, e in merged:
        for i in range(s, e):
            mask[i] = True
            
    # Iterate runs
    from docx.text.run import Run
    original_runs = [Run(r, paragraph) for r in paragraph._element.xpath('.//w:r')]
    p_element = paragraph._element
    global_ptr = 0
    
    for run in original_runs:
        run_text = run.text
        run_len = len(run_text)
        
        # Safety: Skip runs inside comments if requested (preserves structure)
        if preserve_comments and run_inside_comment(run):
            global_ptr += run_len
            continue
            
        # Safety: Skip track changes
        if run_inside_track_change(run):
            global_ptr += run_len
            continue

        if run_len == 0:
            continue
            
        # Identify segments
        segments = []
        current_segment_start = 0
        current_state = mask[global_ptr]
        
        for i in range(1, run_len):
            char_state = mask[global_ptr + i]
            if char_state != current_state:
                # Cut
                seg_text = run_text[current_segment_start:i]
                segments.append((seg_text, current_state, False)) # Text, Highlight, IsLast
                current_segment_start = i
                current_state = char_state
                
        # Last segment
        seg_text = run_text[current_segment_start:]
        segments.append((seg_text, current_state, True))
        
        # Optimization: Don't replace if not needed
        needs_replacement = False
        if len(segments) > 1:
            needs_replacement = True
        elif segments and segments[0][1]: # Is highlighted
            needs_replacement = True
            
        if not needs_replacement:
            global_ptr += run_len
            continue
            
        # Perform Replacement
        new_elements = []
        for txt, do_hl, is_last in segments:
            new_run = duplicate_run(run, paragraph)
            new_run.text = txt
            if do_hl:
                set_highlight(new_run, get_xml_color(color_name))
            
            # Restore extras (comments anchors, footnotes, etc.) 
            # This is crucial even if we skip "commented text", because the anchor itself resides in a run.
            if is_last:
                copy_run_extras(run, new_run)
            
            new_elements.append(new_run._element)
            
        try:
            run_parent = run._element.getparent()
            if run_parent is not None:
                run_index = run_parent.index(run._element)
                for new_el in reversed(new_elements):
                    run_parent.insert(run_index, new_el)
                run_parent.remove(run._element)
        except ValueError:
            pass
            
        global_ptr += run_len

def highlight_comments(doc, compiled_regexes, color_name):
    """
    Highlight text INSIDE comment bodies.
    """
    try:
        comments_part = doc.part.comments_part
    except Exception:
        # No comments part
        return
        
    if not comments_part:
        return

    # iterate comments
    # Namespace map
    nsmap = comments_part.element.nsmap
    
    for comment in comments_part.element.findall(".//w:comment", nsmap):
        for p in comment.findall(".//w:p", nsmap):
            # Create a proxy paragraph object
            # We can't easily instantiate a 'Paragraph' object attached to the document for a loose element.
            # But we can try hack:
            # Create a temporary paragraph in doc to get the class, then swap element?
            # Or just use the highlight_paragraph logic but adapting it?
            # highlight_paragraph needs 'paragraph.text' and 'paragraph.runs'.
            
            # Since paragraph.runs is a property that scans the element, 
            # we can instantiate a Paragraph(p, parent=None) if we import Paragraph.
            # But python-docx Paragraph requires a parent.
            
            # User code used: fake_para = Document().add_paragraph(); fake_para._element = p
            # This works if constructs a dummy document.
            
            try:
                # We need a wrapper.
                from docx.text.paragraph import Paragraph
                # A parent is needed for some ops, but maybe not for .text and .runs access if we are careful?
                # Actually, duplicate_run uses 'parent_paragraph'.
                
                # Trick: use the main doc's body as parent temporarily?
                # Or just use the 'fake_para' approach from user code.
                fake_doc = Document() # Light object?
                fake_para = fake_doc.add_paragraph()
                fake_para._element = p
                
                highlight_paragraph(fake_para, compiled_regexes, color_name, preserve_comments=False)
            except Exception as e:
                # print(f"Warning: failed to process comment: {e}")
                pass

# ---------------------------
# Regex Dictionaries (Ported from core_highlighter.py)
# ---------------------------
UNITS_PATTERN = r"\b[0-9]{1,}\sper\s(dL|µ|µL|mL|ml|L|g|mg|kg|min|h|hr|hour|day|week|month)\b"

METRIC_UNITS = [r"\b\d+[- ]?(micro|micron|liter|liters|mcg|mcm|mcl|mmHg|cc|cm|mL|mcm|mcg|gram|meter|liter|feet|inch|pound|ounce|mile|yard|g|mg|kg)\b"]

TIME_ABBR = [
    r"\b(?:a\.?m\.?|p\.?m\.?|a\.?d\.?|b\.?c\.?e?\.?|c\.?e\.?)\b"
]

PRESSURE_PATTERNS = [
    r"([pP])\s([=><-]+)\s([0-9]+(\.[0-9]+)?)",
    r"[pP]\s≤",
    r"[pP][=><]|[=><][pP]",
    r"[pP]\s[=><]",
    r"[pP]\s≤\s\?\s\d+(\.\d+)?",
    r"[pP]≥\s\?\s\d+(\.\d+)?",
    r"[pP]\s≥\s\?\s\d+(\.\d+)?",
    r"[pP]\s[-=><]\s\?\s\d+(\.\d+)?",
    r"[pP]\s[-=><]\s\?\s\d+",
]

MEDICAL_TERMS = [r"\b(DSM|COVID-19|ventilation|perfusion|V/Q|VQ|V-Q)\b"]
PERCENT_VARIATIONS = [r"(percent|per cent|percentage|%)"]
NUMBERS_1000 = [r"\b\d{4,}\b"]
NUMBERS_0_99 = [r"\b\d{1,2}\b"]
NUMBER_WORDS = [r"\b(Zero|one|two|three|four|five|six|seven|eight|nine|Ten|Eleven|Twelve|Thirteen|Fourteen|Fifteen|Sixteen|Seventeen|Eighteen|Nineteen|twenty|thirty|forty|fifty|sixty|seventy|eighty|ninety|twenty-|thirty-|forty-|fifty-|sixty-|seventy-|eighty-|ninety-|-one|-two|-three|-four|-five|-six|-seven|-eight|-nine)\b"]

DURATION_PATTERNS = [
    r"\b\d+[- ]?(year|years|month|months|week|weeks|day|days|hour|hours|minute|minutes|second|seconds)\b",
    r"\b\d+[- ]?(y|mo|wk|wks|d|h|hr|hrs|min|s|sec)\b",
    r"\b(one|two|three|four|five|six|seven|eight|nine|ten|eleven|twelve|thirteen|fourteen|fifteen|sixteen|seventeen|eighteen|nineteen|twenty|thirty|forty|fifty|sixty|seventy|eighty|ninety)[- ]?(year|years|month|months|week|weeks|day|days|hour|hours|minute|minutes|second|seconds)\b",
    r"\b(one|two|three|four|five|six|seven|eight|nine|ten|eleven|twelve|thirteen|fourteen|fifteen|sixteen|seventeen|eighteen|nineteen|twenty|thirty|forty|fifty|sixty|seventy|eighty|ninety)[- ]?(y|mo|wk|wks|d|h|hr|hrs|min|s|sec)\b"
]

CHAPTER_PATTERNS = [
    r"\bchapter\b", r"\bChapter\b", r"\bchapters\b", r"\bChapters\b",
    r"\bChap\b", r"\bsection\b", r"\bSection\b", r"\bSect\b",
    r"\b(Chap\.|Ch\.|Sec\.)"
]

COMMON_ABBR = [
    r"\b(e\.g|eg|i\.e|ie|vs|etc|et al)\b",
    r"\b(Dr|Drs|Mr|Mrs|Ms|Prof)\b",
    r"\b(M\.D|MD|M\.A|MA|M\.S|MS|Bsc|MSc)\b",
    r"\b(Blvd|St|Ste)\b"
]

LATIN_PHRASES = [r"\b(in vitro|in vivo|ex vivo|in situ|ex situ|per se|ad hoc|de novo|a priori|de facto|status quo|a posteriori|ad libitum|ad lib|supra|verbatim|cf\.|ibid|id\.)\b"]

MEDICAL_SUPERSCRIPT = [r"\b(Paco2|Pco2|Sao2|o2max|Cao2|Spo2|Pvo2|o2|Fio2|PIO2|Vo2|Pao2|Pio2|PAo2|Po2)\b"]

GREEK_PATTERNS = [
            r"\balpha\b", r"α", r"\bbeta\b", r"β", r"\bgamma\b", r"γ", r"\bdelta\b", r"δ",
            r"\bepsilon\b", r"ε", r"\bzeta\b", r"ζ", r"\btheta\b", r"θ", r"\beta\b", r"η",
            r"\biota\b", r"ι", r"\bkappa\b", r"κ", r"\blambda\b", r"λ", r"\bmu\b", r"μ",
            r"\bnu\b", r"ν", r"\bxi\b", r"ξ", r"\bomicron\b", r"ο", r"\bpi\b", r"π",
            r"\brho\b", r"ρ", r"\bsigma\b", r"σ", r"\btau\b", r"τ", r"\bupsilon\b", r"υ",
            r"\bphi\b", r"φ", r"\bchi\b", r"χ", r"\bpsi\b", r"ψ", r"\bomega\b", r"ω",
            # Uppercase
            r"\bAlpha\b", r"Α", r"\bBeta\b", r"Β", r"\bGamma\b", r"Γ", r"\bDelta\b", r"Δ",
            r"\bEpsilon\b", r"Ε", r"\bZeta\b", r"Ζ", r"\bTheta\b", r"Θ", r"\bEta\b", r"Η",
            r"\bIota\b", r"Ι", r"\bKappa\b", r"Κ", r"\bLambda\b", r"Λ", r"\bMu\b", r"Μ",
            r"\bNu\b", r"Ν", r"\bXi\b", r"Ξ", r"\bOmicron\b", r"Ο", r"\bPi\b", r"Π",
            r"\bRho\b", r"Ρ", r"\bSigma\b", r"Σ", r"\bTau\b", r"Τ", r"\bUpsilon\b", r"Υ",
            r"\bPhi\b", r"Φ", r"\bChi\b", r"Χ", r"\bPsi\b", r"Ψ", r"\bOmega\b", r"Ω",
]

NUMERIC_RANGES = [
            r"[0-9]{1,}\.[0-9]{1,}-[0-9]{1,}\.[0-9]{1,}",
            r"[0-9]{1,}-[0-9]{1,}",
            r"[0-9]{1,}\s*-\s*[0-9]{1,}",
            r"[0-9]{1,}--[0-9]{1,}",
            r"[0-9]{1,}\s–\s[0-9]{1,}",
            r"[0-9]{1,}–[0-9]{1,}",
            r"[0-9]{1,}\s—\s[0-9]{1,}",
            r"[0-9]{1,}—[0-9]{1,}",
            r"[0-9]{1,}\s+to\s+[0-9]{1,}",
]

DEGREES = [r"\b\d{1,3}-degree angle\b", r"\b\d{1,3}\sdegree angle\b", r"\b\d{1,3}\s?°\s?angle\b", r"\b\d{1,3}\s?°\b", r"\b\d+\s?°F\b", r"\b\d+\s?°C\b", r"°F", r"°C", r"\b\d+\sdegree\b"]
XRAY = [r"\b[xX]-?ray\b"]
PACO2_ETC = [r"\b(Paco2|Pco2|Sao2|Spo2|Fio2|Pao2|Po2)\b"]
TRADEMARKS = [r"\b[\w\-]+[™®©]", r"[™®©]"]
VERSUS = [r"\b(vs\.?|versus|v\.?)\b"]
SPECIAL_CHARS = [r"[§¶†‡]"]
FIGURE_TABLE = [r"\b(Figure|Table)\s*\d+"]

# New patterns for special dashes and characters
DOUBLE_DASHES = [
    r"--",        # Double hyphen
    r"––",        # Double en-dash
    r"\s-\s",     # Space-hyphen-space
    r"\s–\s",     # Space-en-dash-space
    r"\s—\s"      # Space-em-dash-space
]

# Phrases within parentheses
PHRASES_IN_PARENS = [
    r"\(\s*for\s+example[^)]*\)",
    r"\(\s*e\.g\.[^)]*\)",
    r"\(\s*that\s+is[^)]*\)",
    r"\(\s*i\.e\.[^)]*\)"
]

def compile_patterns(patterns):
    compiled = []
    for p in patterns:
        try:
            compiled.append(re.compile(p, re.IGNORECASE))
        except re.error as e:
            print(f"Warning: Invalid regex '{p}': {e}")
    return compiled

def highlight_unmatched_parentheses(doc):
    """
    Detect and highlight unmatched opening or closing parentheses.
    Scans each paragraph and highlights parentheses that don't have matching pairs.
    Uses a safer approach that only highlights entire runs containing unmatched parens.
    """
    for para in doc.paragraphs:
        if is_skip_style(para):
            continue
        
        text = para.text
        if not text or '(' not in text and ')' not in text:
            continue
        
        # Track which parens are unmatched
        open_stack = []  # Stores indices of unmatched opening parens
        unmatched = set()  # Indices of unmatched parens
        
        for i, char in enumerate(text):
            if char == '(':
                open_stack.append(i)
            elif char == ')':
                if open_stack:
                    open_stack.pop()  # Match found
                else:
                    unmatched.add(i)  # No matching open paren
        
        # Remaining unclosed opening parens
        unmatched.update(open_stack)
        
        if not unmatched:
            continue
        
        # Build mask for unmatched parens
        mask = [False] * len(text)
        for idx in unmatched:
            mask[idx] = True
        
        # Safe highlighting: iterate runs and highlight if contains unmatched parens
        from docx.text.run import Run
        original_runs = [Run(r, para) for r in para._element.xpath('.//w:r')]
        global_ptr = 0
        
        for run in original_runs:
            run_text = run.text
            run_len = len(run_text)
            
            if run_inside_comment(run) or run_inside_track_change(run):
                global_ptr += run_len
                continue
            
            if run_len == 0:
                continue
            
            # Check if this run contains any unmatched parens
            run_has_unmatched = any(mask[global_ptr + i] for i in range(run_len) if global_ptr + i < len(mask))
            
            if run_has_unmatched:
                # Highlight entire run if it contains unmatched parens
                set_highlight(run, "turquoise")
            
            global_ptr += run_len

def highlight_unmatched_quotes(doc):
    """
    Detect and highlight unmatched quotes (both straight and smart double quotes).
    Scans each paragraph and highlights quotes that don't have matching pairs.
    """
    for para in doc.paragraphs:
        if is_skip_style(para):
            continue
        
        text = para.text
        if not text:
            continue
            
        has_straight = '"' in text
        has_smart_open = '“' in text
        has_smart_close = '”' in text
        
        if not (has_straight or has_smart_open or has_smart_close):
            continue
            
        unmatched = set()
        
        # Smart quotes
        open_stack = []
        for i, char in enumerate(text):
            if char == '“':
                open_stack.append(i)
            elif char == '”':
                if open_stack:
                    open_stack.pop()
                else:
                    unmatched.add(i)
        unmatched.update(open_stack)
        
        # Straight quotes
        straight_indices = []
        for i, char in enumerate(text):
            if char == '"':
                straight_indices.append(i)
        if len(straight_indices) % 2 != 0:
            unmatched.update(straight_indices)
            
        if not unmatched:
            continue
            
        mask = [False] * len(text)
        for idx in unmatched:
            mask[idx] = True
            
        from docx.text.run import Run
        original_runs = [Run(r, para) for r in para._element.xpath('.//w:r')]
        global_ptr = 0
        
        for run in original_runs:
            run_text = run.text
            run_len = len(run_text)
            
            if run_inside_comment(run) or run_inside_track_change(run):
                global_ptr += run_len
                continue
            
            if run_len == 0:
                continue
                
            run_has_unmatched = any(mask[global_ptr + i] for i in range(run_len) if global_ptr + i < len(mask))
            
            if run_has_unmatched:
                set_highlight(run, "turquoise")
                
            global_ptr += run_len

def highlight_double_dashes(doc):
    """
    Highlight double hyphens, double en-dashes, and space-dash-space patterns.
    """
    compiled = compile_patterns(DOUBLE_DASHES)
    if not compiled:
        return
    
    for para in doc.paragraphs:
        if is_skip_style(para):
            continue
        highlight_paragraph(para, compiled, "turquoise")

def highlight_phrases_in_parentheses(doc):
    """
    Highlight phrases like 'for example', 'that is', 'e.g.', 'i.e.' within parentheses.
    """
    compiled = compile_patterns(PHRASES_IN_PARENS)
    if not compiled:
        return
    
    for para in doc.paragraphs:
        if is_skip_style(para):
            continue
        highlight_paragraph(para, compiled, "turquoise")

def highlight_symbol_fonts(doc):
    """
    Detect and highlight text that uses symbol fonts (Symbol, Wingdings, Webdings, etc.).
    Symbol fonts are mapped to regular Unicode characters but should be highlighted.
    Also highlights ® ™ © symbols that appear to be in special fonts.
    Uses safe approach without excessive run splitting.
    """
    SYMBOL_FONTS = {
        'symbol', 'wingdings', 'wingdings 2', 'wingdings 3', 'webdings',
        'symbol kitty', 'symbol smartdraw', 'symbol monospaced',
        'ms sans serif', 'arial unicode'
    }
    
    SYMBOL_CHARS = {'®', '™', '©'}  # Common symbols that often use symbol fonts
    
    def get_font_name(run):
        """Extract font name from run, checking multiple sources."""
        # Check direct run font
        if run.font.name:
            return run.font.name.lower()
        
        # Check rPr element
        rPr = run._element.rPr
        if rPr is not None:
            # Check w:rFonts element
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                # Try ascii first, then hAnsi, then cs
                font = rFonts.get(qn('w:ascii'))
                if font:
                    return font.lower()
                font = rFonts.get(qn('w:hAnsi'))
                if font:
                    return font.lower()
                font = rFonts.get(qn('w:cs'))
                if font:
                    return font.lower()
        
        return None
    
    def is_symbol_font_run(run, font_name):
        """Check if run should be highlighted as symbol font."""
        # If font is explicitly a symbol font
        if font_name and font_name in SYMBOL_FONTS:
            return True
        
        # If run contains symbol characters (common trademarked/symbol chars)
        text = run.text
        if text and any(char in SYMBOL_CHARS for char in text):
            return True
        
        return False
    
    def process_paragraph(para):
        if is_skip_style(para):
            return
        
        from docx.text.run import Run
        original_runs = [Run(r, para) for r in para._element.xpath('.//w:r')]
        
        for run in original_runs:
            if run_inside_comment(run) or run_inside_track_change(run):
                continue
            
            text = run.text
            if not text:
                continue
            
            # Get font name from multiple sources
            font_name = get_font_name(run)
            
            # Check if this run should be highlighted
            if not is_symbol_font_run(run, font_name):
                continue
            
            # Safe approach: highlight entire run instead of splitting every character
            # This prevents XML corruption from excessive element creation
            set_highlight(run, "turquoise")
    
    # Process all paragraphs
    for para in doc.paragraphs:
        process_paragraph(para)
    
    # Process table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para)



def process_docx(input_file, output_file, skip_validation=False, verbose=True):
    if verbose:
        print(f"Opening {input_file}...")
    
    doc = Document(input_file)

    def run_batch(patterns, color):
        flat = []
        if isinstance(patterns, list):
            flat = patterns
        else:
            flat = [patterns]
            
        compiled = compile_patterns(flat)
        if not compiled: return
        
        for para in doc.paragraphs:
            style_name = str(para.style.name)
            if style_name in {'REF-N', 'REF-U'}:
                 continue
            highlight_paragraph(para, compiled, color)
            
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        highlight_paragraph(para, compiled, color)

    # 0. Styles (Yellow/Green logic)
    if verbose: print("Highlighting styles (Yellow/Green)...")
    highlight_words_in_styles(doc)

    if verbose: print("Highlighting patterns...")
    
    run_batch([UNITS_PATTERN], "turquoise")
    run_batch(METRIC_UNITS, "turquoise")
    run_batch(TIME_ABBR, "turquoise")
    run_batch(PRESSURE_PATTERNS, "turquoise")
    run_batch(MEDICAL_TERMS, "turquoise")
    run_batch(PERCENT_VARIATIONS, "turquoise")
    run_batch(NUMBERS_1000, "turquoise")
    run_batch(NUMBERS_0_99, "turquoise")
    run_batch(NUMBER_WORDS, "turquoise")
    run_batch(DURATION_PATTERNS, "turquoise")
    run_batch(CHAPTER_PATTERNS, "turquoise")
    run_batch(COMMON_ABBR, "turquoise")
    run_batch(LATIN_PHRASES, "turquoise")
    run_batch(MEDICAL_SUPERSCRIPT, "turquoise")
    run_batch(GREEK_PATTERNS, "turquoise")
    run_batch(NUMERIC_RANGES, "turquoise")
    run_batch(DEGREES, "turquoise")
    run_batch(XRAY, "turquoise")
    run_batch(PACO2_ETC, "turquoise")
    run_batch(TRADEMARKS, "turquoise")
    run_batch(VERSUS, "turquoise")
    run_batch(SPECIAL_CHARS, "turquoise")
    run_batch(FIGURE_TABLE, "turquoise")

    # New functions
    if verbose: print("Highlighting comparison symbols...")
    highlight_comparison_symbols(doc)

    if verbose: print("Highlighting math symbols...")
    highlight_math_symbols(doc)

    if verbose: print("Highlighting time periods...")
    highlight_time_period_terms(doc)

    if verbose: print("Highlighting italic punctuation...")
    highlight_italic_punctuation(doc)

    if verbose: print("Highlighting multilingual chars...")
    highlight_multilingual_chars(doc)
    
    if verbose: print("Highlighting double dashes...")
    highlight_double_dashes(doc)
    
    if verbose: print("Highlighting phrases in parentheses...")
    highlight_phrases_in_parentheses(doc)
    
    if verbose: print("Checking for unmatched parentheses...")
    highlight_unmatched_parentheses(doc)
    
    if verbose: print("Checking for unmatched quotes...")
    highlight_unmatched_quotes(doc)
    
    if verbose: print("Highlighting symbol fonts...")
    highlight_symbol_fonts(doc)
    
    if not skip_validation:
        if verbose: print("Checking headings...")
        check_heading_hierarchy(doc)

        if verbose: print("Checking punctuation...")
        check_unpaired_punctuation_and_quotes(doc)

    if verbose: print(f"Saving to {output_file}...")
    doc.save(output_file)

# ---------------------------
# Extended Ported Functions
# ---------------------------

def highlight_comparison_symbols(doc):
    SYMBOLS = ["<", ">", "=", "≥", "≤", "≈"]
    SPELLINGS = ["less than", "greater than", "equal to", "approximately"]
    
    # Compile regex
    patterns = [re.escape(s) for s in SYMBOLS]
    patterns += [r"\b" + re.escape(sp) + r"\b" for sp in SPELLINGS]
    
    compiled = compile_patterns(patterns)
    
    for para in doc.paragraphs:
        if is_skip_style(para): continue
        highlight_paragraph(para, compiled, "cyan") # wdTurquoise

def highlight_math_symbols(doc):
    # original: symbols = {"-": 3, "+": 3, "×": 3, "÷": 3} (wdTurquoise)
    SYMBOLS = ["-", "+", "×", "÷"]
    patterns = [re.escape(s) for s in SYMBOLS]
    compiled = compile_patterns(patterns)
    
    for para in doc.paragraphs:
        if is_skip_style(para): continue
        highlight_paragraph(para, compiled, "cyan")

def highlight_time_period_terms(doc):
    PATTERNS = [
        r"\btimes\b", r"\bcentury\b", r"\bcenturies\b",
        r"\bdecade\b", r"\b\d+-fold\b", r"\bfold\b",
    ]
    compiled = compile_patterns(PATTERNS)
    for para in doc.paragraphs:
        if is_skip_style(para): continue
        highlight_paragraph(para, compiled, "cyan")

def highlight_italic_punctuation(doc):
    """Highlight . , : ; only when in italic."""
    PUNCT = {".", ",", ":", ";"}
    
    # We iterate runs. if run is italic, we scan its text. 
    # If we find punct, we might need to split run? 
    # Or just highlight the whole run? 
    # ORIGINAL LOGIC: r.HighlightColorIndex = wdturquoise on the CHARACTER range.
    # So we must use run splitting logic if strict.
    # BUT, reusing highlight_paragraph is hard because it doesn't know about current run's italic status.
    # We need a custom pass.
    
    for para in doc.paragraphs:
        if is_skip_style(para): continue
        
        # Snapshot runs
        from docx.text.run import Run
        original_runs = [Run(r, para) for r in para._element.xpath('.//w:r')]
        p_element = para._element

        for run in original_runs:
            text = run.text
            
            # Safety checks
            if run_inside_comment(run):
                continue
            if run_inside_track_change(run):
                continue

            segments = []
            
            # Cases for segments:
            # 1. Empty text -> preserve (is_last=True)
            # 2. Not italic -> preserve (is_last=True)
            # 3. Italic -> split by punct
            
            if not text:
                segments.append((text, None, True))
            elif not run.italic:
                segments.append((text, None, True))
            else:
                parts = re.split(r'([.,:;])', text)
                valid_parts = [p for p in parts if p]
                
                if len(valid_parts) > 1:
                    for i, p in enumerate(valid_parts):
                        color = "cyan" if p in PUNCT else None
                        is_last = (i == len(valid_parts) - 1)
                        segments.append((p, color, is_last))
                else:
                     segments.append((text, None, True))
            
            # Optimization: Check if modification needed.
            needs_replacement = False
            if len(segments) > 1:
                needs_replacement = True
            elif segments and segments[0][1]: # has color
                needs_replacement = True
            
            if not needs_replacement:
                continue

            # Build new elements
            new_elements = []
            for txt, color, is_last in segments:
                new_run = duplicate_run(run, para)
                new_run.text = txt
                if color:
                    set_highlight(new_run, get_xml_color(color))
                
                if is_last:
                    copy_run_extras(run, new_run)
                
                new_elements.append(new_run._element)
                
            try:
                run_parent = run._element.getparent()
                if run_parent is not None:
                    run_index = run_parent.index(run._element)
                    for new_el in reversed(new_elements):
                        run_parent.insert(run_index, new_el)
                    run_parent.remove(run._element)
            except ValueError:
                pass

def check_unpaired_punctuation_and_quotes(doc):
    # Collect all text
    full_text = []
    for p in doc.paragraphs:
        full_text.append(p.text)
    full_str = "\n".join(full_text)
    
    pairs = [("(", ")"), ("[", "]"), ("{", "}"), ("\"", "\""), ("'", "'"), ("\u201c", "\u201d"), ("\u2018", "\u2019")]
    messages = []
    
    for left, right in pairs:
        lc = full_str.count(left)
        rc = full_str.count(right)
        if lc != rc:
            messages.append(f"Unbalanced {left}/{right}: {lc} vs {rc}")
            
    if messages:
        # Add comment at end of doc
        try:
            para = doc.add_paragraph()
            run = para.add_run(" [Highlighter Check: " + "; ".join(messages) + "]")
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.bold = True
        except:
             pass

def check_heading_hierarchy(doc):
    prev_level = 0
    for para in doc.paragraphs:
        style_name = str(para.style.name)
        # USER PROVIDED REGEX CHANGE: Heading -> H
        m = re.search(r"H\s+([1-9])", style_name, re.IGNORECASE)
        if m:
            level = int(m.group(1))
            if level > prev_level + 1 and prev_level != 0:
                # Add inline warning/comment
                try:
                    run = para.add_run(f" [HIERARCHY ERROR: Jumped {prev_level}->{level}]")
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    run.bold = True
                except:
                    pass
            prev_level = level

def highlight_multilingual_chars(doc):
    # Unicode blocks
    UNICODE_BLOCKS = [
        ("Chinese",       19968, 40959, "cyan"),
        ("Greek",         0x370, 0x3FF, "cyan"),
        ("Cyrillic",      0x400, 0x4FF, "magenta"),
        ("Hebrew",        0x590, 0x5FF, "green"),
        ("Arabic",        0x600, 0x6FF, "blue"),
        ("Arabic",        0x750, 0x77F, "blue"),
        ("Devanagari",    0x900, 0x97F, "red"),
        ("Japanese",      0x3040, 0x309F, "darkMagenta"), # Violet approx
        ("Japanese",      0x30A0, 0x30FF, "darkMagenta"),
        ("Korean",        0xAC00, 0xD7AF, "darkMagenta"),
        ("Thai",          0x0E00, 0x0E7F, "darkRed"),
        ("Currency",      0x20A0, 0x20CF, "darkBlue"),
    ]
    
    # Strategy: Build regex for ranges or iterate chars?
    # Building regex for unicode ranges is efficient.
    # range regex: [\u4e00-\u9fff]
    
    patterns_by_color = {}
    for name, start, end, color in UNICODE_BLOCKS:
        # Construct char class
        # Python re supports \uXXXX
        c_pattern = f"[{chr(start)}-{chr(end)}]"
        if color not in patterns_by_color:
            patterns_by_color[color] = []
        patterns_by_color[color].append(c_pattern)
        
    for color, pat_list in patterns_by_color.items():
        # combine into one regex: ([range1]|[range2])
        combined = "|".join(pat_list)
        compiled = re.compile(combined)
        for para in doc.paragraphs:
             highlight_paragraph(para, [compiled], color)

def highlight_words_in_styles(doc):
    target_styles = {"T1", "CT", "H1", "H2", "H2A", "H3", "H3A", "NBX1-TTL"}
    # Original logic:
    # 1. Highlight entire paragraph in "turquoise" (cyan)
    # 2. Highlight words <= 4 chars in "brightgreen" (green)
    
    wdBrightGreen = "green"
    wdTurquoise = "cyan"
    
    # Pre-compile word finder
    # Words <= 4 chars. Regex: \b\w{1,4}\b ?
    # Original used split() and then .Find matching whole word.
    # Regex equivalent: \b\S{1,4}\b (approx) or \b\w{1,4}\b
    short_word_re = re.compile(r"\b\w{1,4}\b")
    
    for para in doc.paragraphs:
        if para.style.name not in target_styles:
            continue
            
        # 1. Apply Turfuoise to whole para
        # We can just set highlight on all existing runs
        from docx.text.run import Run
        for run in [Run(r, para) for r in para._element.xpath('.//w:r')]:
             set_highlight(run, wdTurquoise)
             
        # 2. Highlight short words Green
        # This requires run splitting of the ALREADY highlighted runs.
        # highlight_paragraph handles this! It creates new runs and preserves/sets highlight.
        highlight_paragraph(para, [short_word_re], wdBrightGreen)

def is_skip_style(para):
    return para.style.name in {'REF-N', 'REF-U'}

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Highlight DOCX using python-docx (Linux compatible)")
    parser.add_argument("input", help="Input file path")
    parser.add_argument("output", help="Output file path")
    args = parser.parse_args()
    
    if not os.path.exists(args.input):
        print(f"Error: {args.input} not found.")
        exit(1)
        
    # Updated process call
    process_docx(args.input, args.output)

