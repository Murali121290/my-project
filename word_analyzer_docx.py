import os
import re
import datetime
import subprocess
from pathlib import Path
from collections import defaultdict
from typing import List, Tuple, Dict, Any, Set, Optional
import jinja2
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

import importlib.util as _importlib_util
HAS_PDFPLUMBER: bool = _importlib_util.find_spec("pdfplumber") is not None

def _normalize_for_match(text: str) -> str:
    """Strip all non-alphanumeric chars and lowercase for robust matching."""
    return re.sub(r'\W+', '', text.lower())

# ------------------------------
# 1. HTML Templates & Helpers
# (Templates moved to templates/ directory and loaded via Jinja2)
# ------------------------------

def escape_html(s: str) -> str:
    if not isinstance(s, str): return str(s)
    return (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")
            .replace("\n", "<br>"))


# ------------------------------
# 2. Citation Analyzer Class (Same Logic)
# ------------------------------

class CitationAnalyzer:
    def __init__(self):
        self.supported_types = ["Figure", "Table", "Box", "Exhibit", "Appendix", "Case Study"]
        self.regex_patterns = self._setup_regex_patterns()

    def _setup_regex_patterns(self) -> Dict[str, re.Pattern]:
        patterns = {}
        patterns['single'] = re.compile(
            r'(?:\(|\b)(Figures?|Figs?\.?|Tables?|Tabs?\.?|Box(?:es)?|BX|Exhibits?|Appendix|Appendices|Case\s+Stud(?:y|ies))\.?\s*([0-9]+(?:[.\-][0-9]+)*)([A-Za-z]?)(?:\)|\b)',
            re.IGNORECASE
        )
        patterns['range'] = re.compile(
            r'(?:\(|\b)(Figures?|Figs?\.?|Tables?|Tabs?\.?|Boxes?|Exhibits?|Appendices?|Case\s+Studies?)\.?\s+'
            r'([0-9]+(?:\.[0-9]+)*)([A-Za-z]?)'
            r'(?:\s+(?:to|through)\s+|\s*[\u2013\u2014]\s*|[ \t]*-[ \t]*)'
            r'([0-9]+(?:\.[0-9]+)*)([A-Za-z]?)(?:\)|\b)',
            re.IGNORECASE
        )
        patterns['and'] = re.compile(
            r'(?:\(|\b)(Figures?|Figs?\.?|Tables?|Tabs?\.?|Boxes?|Exhibits?|Appendices?|Case\s+Studies?)\.?\s+([0-9]+(?:[\.\-][0-9]+)*)([A-Za-z]?)\s+(?:and|&)\s*([0-9]+(?:[\.\-][0-9]+)*)([A-Za-z]?)(?:\)|\b)',
            re.IGNORECASE
        )
        patterns['list'] = re.compile(
            r'(?:\(|\b)(Figures?|Figs?\.?|Tables?|Tabs?\.?|Boxes?|Exhibits?|Appendices?|Case\s+Studies?)\.?\s+'
            r'([0-9]+(?:[\.\-][0-9]+)*[A-Za-z]?(?:\s*,\s*[0-9]+(?:[\.\-][0-9]+)*[A-Za-z]?)*'
            r'(?:\s*,?\s+(?:and|&)\s+[0-9]+(?:[\.\-][0-9]+)*[A-Za-z]?)?)(?:\)|\b)',
            re.IGNORECASE
        )
        patterns['tag'] = re.compile(
            r'<\s*(FIG|TAB|BX|CS|EX|APP|Fig(?:ure)?|Tab(?:le)?|Box|Exhibit|App(?:endix)?|Case\s+Study)\s*(\d+[\.\-]\d+)([A-Za-z]?)\s*>',
            re.IGNORECASE
        )
        return patterns

    def normalize_for_regex(self, text: str) -> str:
        text = re.sub(r'[\u00a0\u2000-\u200b\u202f\u205f]', ' ', text)
        text = re.sub(r'(\d)\s+([.\-])\s*(\d)', r'\1\2\3', text)
        text = re.sub(r'(\d)([.\-])\s+(\d)', r'\1\2\3', text)
        return text

    def normalize_type(self, label: str) -> str:
        if not label:
            return "Figure"
        lbl = label.lower()
        if lbl.startswith('fig'):
            return "Figure"
        if lbl.startswith('tab'):
            return "Table"
        if lbl.startswith('box') or lbl.startswith('bx'):
            return "Box"
        if lbl.startswith('ex') or lbl.startswith('exhibit'):
            return "Exhibit"
        if lbl.startswith('app'):
            return "Appendix"
        if lbl.startswith('case') or lbl.startswith('cs'):
            return "Case Study"
        return "Figure"

    def normalize_fig_number(self, fig_ref: str) -> str:
        if not fig_ref:
            return ""
        fig_ref = fig_ref.strip()
        fig_ref = fig_ref.replace('--', '-').replace('\u2013', '-').replace('\u2014', '-')
        for ch in ['[', ']', '°']:
            fig_ref = fig_ref.replace(ch, '')
        m = re.search(r'([0-9]+(?:[.\-][0-9]+)*)([A-Za-z]?)', fig_ref)
        if m:
            base = m.group(1).replace('-', '.')
            suffix = m.group(2)
            if base.endswith('.'):
                base = base[:-1]
            return base + suffix
        return fig_ref

    def is_caption_paragraph(self, text: str, style_name: str = "") -> bool:
        if style_name:
            s_low = style_name.strip().lower()
            if s_low in ['fig-leg', 'fgc', 't1', 'tt', 'figurelegend', 'tablecaption', 'cs-ttl','nbx1-num','nbx1-ttl','nbx2-num','nbx2-ttl', 'exhibitcaption']:
                return True

        t_norm_orig = self.normalize_for_regex(text.strip())
        has_title_tag = bool(re.search(r'<(TITLE|CAPTION|BX_TTL|BX-TTL|CS-TTL)>', t_norm_orig, re.IGNORECASE))
        has_type_tag  = bool(re.search(r'^<(FIG|TAB|BX|CS|EX|APP)>', t_norm_orig.strip(), re.IGNORECASE))

        t_norm = re.sub(r'^(?:<[^>]*>\s*)+', '', t_norm_orig)
        if not t_norm:
            return False
        if len(t_norm.splitlines()) > 7:
            return False

        match = re.match(r'(?i)^(figure|fig\.|table|tab\.|box|exhibit|appendix|case\s+study)\s+([0-9]+(?:[.\-][0-9]+)*[a-zA-Z]?)(.*)', t_norm)
        if match:
            if has_title_tag or has_type_tag:
                return True
            remainder = match.group(3).strip()
            if not re.search(r'[A-Za-z0-9]', remainder):
                return False
            # If the remainder starts with a separator (. : - — –), the
            # lowercase check is skipped: "Table 4.2. pH dependence…" is a
            # valid caption even though "pH" starts with a lowercase letter.
            has_separator = bool(re.match(r'^[.\:\-\u2013\u2014]\s', remainder))
            if not has_separator:
                first_word_char = re.sub(r'^[\W_]+', '', remainder)
                if first_word_char and first_word_char[0].islower():
                    return False
            return True

        return False

    def analyze_document_citations(self, document_content: List[Tuple[str, int, bool]]) -> Dict[str, Any]:
        dict_types = {t: {
            "Caption": {}, "Citation": {}, "CaptionPage": {}, "CitationPage": {},
            "CaptionCount": {}, "CaptionExtraPages": {},
            "DerivedCaption": {}, "DerivedCaptionPage": {}
        } for t in self.supported_types}

        for text, page_no, is_caption in document_content:
            txt = self.normalize_for_regex(text)

            # Process FIG/TABLE tags BEFORE stripping leading tags, so <FIG2.3> at
            # the start of a paragraph is captured as a citation and not removed.
            for m in self.regex_patterns['tag'].finditer(txt):
                label = self.normalize_type(m.group(1))
                main_no = m.group(2)
                suffix = m.group(3) or ""
                item_id = f"{label} {self.normalize_fig_number(main_no + suffix)}"
                self._store(dict_types, label, item_id, page_no, False)

            txt = re.sub(r'^(?:<[^>]*>\s*)+', '', txt)

            label_boundary = 60
            if is_caption:
                match_boundary = re.search(r'[\.\:\-\u2013\u2014]\s', txt)
                if match_boundary and match_boundary.start() < 60:
                    label_boundary = match_boundary.end() + 5

            for m in self.regex_patterns['range'].finditer(txt):
                is_match_caption = is_caption and m.start() <= label_boundary
                label = self.normalize_type(m.group(1))
                start_num = self.normalize_fig_number(m.group(2))
                end_num = self.normalize_fig_number(m.group(4))
                try:
                    sp = start_num.split('.')
                    ep = end_num.split('.')
                    if int(sp[0]) == int(ep[0]) and len(sp) > 1 and len(ep) > 1:
                        start_minor = int(sp[1])
                        end_minor = int(ep[1])
                        for n in range(start_minor, end_minor + 1):
                            item_id = f"{label} {sp[0]}.{n}"
                            self._store(dict_types, label, item_id, page_no, is_match_caption)
                    else:
                        has_word_sep = bool(re.search(r'\s+(?:to|through)\s+', m.group(0), re.IGNORECASE))
                        if not has_word_sep and len(sp) == 1 and len(ep) == 1:
                            # Hyphen/dash + pure integers → hyphenated chapter.item ID (e.g. Fig. 1-4 → Figure 1.4)
                            combined = self.normalize_fig_number(f"{m.group(2)}-{m.group(4)}")
                            self._store(dict_types, label, f"{label} {combined}", page_no, is_match_caption)
                        else:
                            self._store(dict_types, label, f"{label} {start_num}", page_no, is_match_caption)
                            self._store(dict_types, label, f"{label} {end_num}", page_no, is_match_caption)
                except Exception:
                    self._store(dict_types, label, f"{label} {start_num}", page_no, is_match_caption)
                    self._store(dict_types, label, f"{label} {end_num}", page_no, is_match_caption)

            for m in self.regex_patterns['and'].finditer(txt):
                is_match_caption = is_caption and m.start() <= label_boundary
                label = self.normalize_type(m.group(1))
                first_num = self.normalize_fig_number(m.group(2))
                second_num = self.normalize_fig_number(m.group(4))
                self._store(dict_types, label, f"{label} {first_num}", page_no, is_match_caption)
                self._store(dict_types, label, f"{label} {second_num}", page_no, is_match_caption)

            for m in self.regex_patterns.get('list', re.compile('')).finditer(txt):
                is_match_caption = is_caption and m.start() <= label_boundary
                label = self.normalize_type(m.group(1))
                items_str = m.group(2)
                parts = re.split(r',\s*and\s+|\s+and\s+|,\s*&\s+|\s+&\s+|,\s*', items_str)
                for part in parts:
                    part = part.strip()
                    if not part: continue
                    range_m = re.match(r'^([0-9]+(?:[\.\-][0-9]+)*[A-Za-z]?)(?:\s+(?:to|through)\s+|\s*[\u2013\u2014]\s*|\s*-\s*)([0-9]+(?:[\.\-][0-9]+)*[A-Za-z]?)$', part)
                    if range_m:
                        start_num = self.normalize_fig_number(range_m.group(1))
                        end_num = self.normalize_fig_number(range_m.group(2))
                        try:
                            sp = start_num.split('.')
                            ep = end_num.split('.')
                            if int(sp[0]) == int(ep[0]) and len(sp) > 1 and len(ep) > 1:
                                start_minor = int(sp[1])
                                end_minor = int(ep[1])
                                for n in range(start_minor, end_minor + 1):
                                    item_id = f"{label} {sp[0]}.{n}"
                                    self._store(dict_types, label, item_id, page_no, is_match_caption)
                            else:
                                has_word_sep = bool(re.search(r'\s+(?:to|through)\s+', part, re.IGNORECASE))
                                if not has_word_sep and len(sp) == 1 and len(ep) == 1:
                                    # Hyphen/dash + pure integers → hyphenated chapter.item ID
                                    combined = self.normalize_fig_number(f"{range_m.group(1)}-{range_m.group(2)}")
                                    self._store(dict_types, label, f"{label} {combined}", page_no, is_match_caption)
                                else:
                                    self._store(dict_types, label, f"{label} {start_num}", page_no, is_match_caption)
                                    self._store(dict_types, label, f"{label} {end_num}", page_no, is_match_caption)
                        except Exception:
                            self._store(dict_types, label, f"{label} {start_num}", page_no, is_match_caption)
                            self._store(dict_types, label, f"{label} {end_num}", page_no, is_match_caption)
                    else:
                        norm_num = self.normalize_fig_number(part)
                        if norm_num.count('.') > 1:
                            continue
                        item_id = f"{label} {norm_num}"
                        self._store(dict_types, label, item_id, page_no, is_match_caption)

            for m in self.regex_patterns['single'].finditer(txt):
                is_match_caption = is_caption and m.start() <= label_boundary
                label = self.normalize_type(m.group(1))
                main_no = m.group(2)
                suffix = m.group(3) or ""
                norm_num = self.normalize_fig_number(main_no + suffix)
                # Skip if greedy match swallowed a range (e.g. "3.33-3.36" → "3.33.3.36")
                if norm_num.count('.') > 1:
                    continue
                item_id = f"{label} {norm_num}"
                self._store(dict_types, label, item_id, page_no, is_match_caption)

            if is_caption:
                base_m = re.match(
                    r'(?i)(figures?|figs?\.?|tables?|tabs?\.?|box(?:es)?|exhibits?|appendix|case\s+stud(?:y|ies))'
                    r'\.?\s*([0-9]+(?:[.\-][0-9]+)*)',
                    txt
                )
                if base_m:
                    base_label = self.normalize_type(base_m.group(1))
                    base_num   = self.normalize_fig_number(base_m.group(2))
                    remainder  = txt[label_boundary:]
                    for pm in _PANEL_LETTER_RE.finditer(remainder):
                        panel_letter = (pm.group(1) or pm.group(2)).upper()
                        sub_id = f"{base_label} {base_num}{panel_letter}"
                        tdict = dict_types.get(base_label)
                        if tdict is not None and sub_id not in tdict['DerivedCaption']:
                            tdict['DerivedCaption'][sub_id] = True
                            tdict['DerivedCaptionPage'][sub_id] = page_no

        return dict_types

    def _store(self, dict_types, label, item_id, page_no, is_caption):
        tdict = dict_types.get(label)
        if tdict is None:
            return
        if is_caption:
            if item_id not in tdict['Caption']:
                tdict['Caption'][item_id] = True
                tdict['CaptionPage'][item_id] = page_no
                tdict['CaptionCount'][item_id] = 1
                tdict['CaptionExtraPages'][item_id] = []
            else:
                tdict['CaptionCount'][item_id] = tdict['CaptionCount'].get(item_id, 1) + 1
                tdict['CaptionExtraPages'].setdefault(item_id, []).append(page_no)
        else:
            if item_id not in tdict['Citation']:
                tdict['Citation'][item_id] = True
                tdict['CitationPage'][item_id] = page_no


# ------------------------------
# Box Tag Linker
# ------------------------------
_NBX_TYPE_DEF_RE  = re.compile(
    r'<BX_TYPE>\s*Box\s+(\d+[\.\-]\d+)[^<]*<BX_TTL>(.*)', re.IGNORECASE
)
_NBX_TTL_SAME_RE  = re.compile(
    r'<BX(\d+[\.\-]\d+)>\s*<NBX-TTL>(.*)', re.IGNORECASE
)
_NBX_TTL_NEXT_RE  = re.compile(r'^<NBX-TTL>(.*)', re.IGNORECASE)
_BX_TAG_ONLY_RE   = re.compile(r'^<BX(\d+[\.\-]\d+)>$', re.IGNORECASE)
_BX_TAG_RE        = re.compile(r'<BX(\d+[\.\-]\d+)>',   re.IGNORECASE)
_BX_PLAIN_DEF_RE  = re.compile(
    r'^Box\s+(\d+[\.\-]\d+)\u2003\s*(\S.*)', re.IGNORECASE
)
_BX_PLAIN_SPC_RE  = re.compile(
    r'^Box\s+(\d+[\.\-]\d+)\s+(\S.*)', re.IGNORECASE
)
_BX_CAPTION_WORD_LIMIT = 15
_BX_TEXT_RE       = re.compile(r'\bBox(?:es)?\s+(\d+[\.\-]\d+)\b', re.IGNORECASE)
_BX_TITLE_NEXT_RE = re.compile(r'^<TITLE>Box\s+\d+[\.\-]\d+\s+(.*)', re.IGNORECASE)
_BX_CLOSE_RE      = re.compile(r'^</BX>', re.IGNORECASE)
_BOX_OPEN_RE      = re.compile(r'^<BOX>\s*(?:BOX\s*)?$', re.IGNORECASE)
_BOX_WORD_RE      = re.compile(r'^BOX\s*$', re.IGNORECASE)
_BOX_NUM_ONLY_RE  = re.compile(r'^(\d+[\.\-]\d+)\s*$')

_PANEL_LETTER_RE = re.compile(
    r'(?:(?:^|[.]\s+)([A-F])\.\s+[A-Z\d])'
    r'|'
    r'(?:\(([A-F])\)\s+[A-Z\d])'
)


class BoxTagLinker:
    def __init__(self, chapter_number: Optional[str] = None):
        raw = str(chapter_number).strip() if chapter_number else ""
        _m = re.search(r'\d+', raw)
        self.chapter_number: Optional[str] = _m.group() if _m else (raw or None)
        self.citations:     Dict[str, List[int]] = {}
        self.definitions:   Dict[str, Dict[str, Any]] = {}
        self.cross_chapter: Dict[str, List[int]] = {}
        self.errors:        List[str] = []

    def _norm(self, raw: str) -> str:
        return raw.replace('.', '-')

    def _is_cross_chapter(self, norm_id: str) -> bool:
        if not self.chapter_number:
            return False
        first_part = norm_id.split('-')[0]
        return first_part != self.chapter_number

    def scan(self, paragraphs: List[Any]) -> None:
        texts: list[tuple[str, int]] = [(str(item[0]).strip(), int(item[1])) for item in paragraphs]

        # Pre-process: merge multi-line <BOX>/BOX/1-1/Title into "Box 1-1 Title"
        merged: list[tuple[str, int]] = []
        i = 0
        while i < len(texts):
            s, page_no = texts[i]
            if _BOX_OPEN_RE.match(s):
                j = i + 1
                # Skip a standalone "BOX" word line if present
                if j < len(texts) and _BOX_WORD_RE.match(texts[j][0]):
                    j += 1
                # Expect the box number on the next line
                if j < len(texts):
                    nm = _BOX_NUM_ONLY_RE.match(texts[j][0])
                    if nm:
                        num = nm.group(1)
                        j += 1
                        # Expect the title on the line after the number
                        if j < len(texts) and texts[j][0].strip():
                            title: str = texts[j][0].strip()
                            merged.append((f"Box {num} {title}", page_no))
                            i = j + 1
                            continue
            merged.append((s, page_no))
            i += 1
        texts = merged

        skip_next = False
        for line_no, (s, page_no) in enumerate(texts):
            if skip_next:
                skip_next = False
                continue
            if _BX_CLOSE_RE.match(s):
                continue
            m = _NBX_TYPE_DEF_RE.search(s)
            if m:
                self._store_def(self._norm(m.group(1)), m.group(2).strip(), line_no, page_no)
                continue
            m = _BX_TAG_ONLY_RE.match(s)
            if m:
                next_s = texts[line_no + 1][0] if line_no + 1 < len(texts) else ""
                nm = _NBX_TTL_NEXT_RE.match(next_s)
                if nm:
                    self._store_def(self._norm(m.group(1)), nm.group(1).strip(), line_no, page_no)
                    skip_next = True
                    continue
                nm = _BX_TITLE_NEXT_RE.match(next_s)
                if nm:
                    self._store_def(self._norm(m.group(1)), nm.group(1).strip(), line_no, page_no)
                    skip_next = True
                    continue
                self._store_citation(self._norm(m.group(1)), line_no)
                continue
            m = _NBX_TTL_SAME_RE.search(s)
            if m:
                self._store_def(self._norm(m.group(1)), m.group(2).strip(), line_no, page_no)
                continue
            m = _BX_PLAIN_DEF_RE.match(s)
            if m:
                self._store_def(self._norm(m.group(1)), m.group(2).strip(), line_no, page_no)
                continue
            m = _BX_PLAIN_SPC_RE.match(s)
            if m:
                nid  = self._norm(m.group(1))
                rest = m.group(2).strip()
                if len(rest.split()) <= _BX_CAPTION_WORD_LIMIT:
                    self._store_def(nid, rest, line_no, page_no)
                else:
                    self._store_citation(nid, line_no)
                continue
            for m in _BX_TAG_RE.finditer(s):
                self._store_citation(self._norm(m.group(1)), line_no)
            for m in _BX_TEXT_RE.finditer(s):
                self._store_citation(self._norm(m.group(1)), line_no)

    def _store_def(self, nid: str, caption: str, line_no: int, page_no: int) -> None:
        if nid in self.definitions:
            prev = self.definitions[nid]['line']
            self.errors.append(f"Duplicate box definition: Box {nid} (lines {prev} and {line_no})")
        else:
            self.definitions[nid] = {"caption": caption, "line": line_no, "page": page_no}

    def _store_citation(self, nid: str, line_no: int) -> None:
        if self._is_cross_chapter(nid):
            if nid not in self.cross_chapter:
                self.cross_chapter[nid] = []
            self.cross_chapter[nid].append(line_no)
        else:
            if nid not in self.citations:
                self.citations[nid] = []
            self.citations[nid].append(line_no)

    def validate(self):
        for nid in self.citations:
            if nid not in self.definitions:
                self.errors.append(f"Missing caption for Box {nid}")
        # Orphan box check suppressed — not displayed in report

    def results(self):
        def sort_key(x):
            return [int(p) for p in x.replace('-', '.').split('.') if p.isdigit()]
        all_ids = sorted(set(self.citations) | set(self.definitions), key=sort_key)
        rows = []
        for nid in all_ids:
            cit  = nid in self.citations
            defn = self.definitions.get(nid)
            status = "Matched" if (cit and defn) else ("Missing Caption" if cit else "Orphan Box")
            rows.append({
                "box_id": f"Box {nid}", "citation_found": cit,
                "caption_found": bool(defn),
                "caption_text": defn["caption"] if defn else "", "status": status,
            })
        for nid in sorted(self.cross_chapter, key=sort_key):
            rows.append({
                "box_id": f"Box {nid}", "citation_found": True,
                "caption_found": False, "caption_text": "", "status": "Cross-Chapter Ref",
            })
        return rows

    def build_html(self):
        rows = self.results()
        if not rows and not self.errors:
            return ""
        
        # Sort rows by citation page if available, else 999999
        def _safe_page(r):
            page_str = str(r.get("citation_page", ""))
            m = re.search(r'\d+', page_str)
            if m: return int(m.group())
            return 999999
        rows.sort(key=_safe_page)

        icon = {
            "Matched": "✅ Matched", "Missing Caption": "⚠️ Missing Caption",
            "Orphan Box": "⚠️ Orphan Box", "Cross-Chapter Ref": "ℹ️ Cross-Chapter Ref",
        }
        lines = [
            '<h3>Box Citation ↔ Caption Mapping</h3>',
            '<table><thead><tr><th>Box ID</th><th>Cited in Text</th>'
            '<th>Caption Found</th><th>Status</th><th>Caption Title</th></tr></thead><tbody>',
        ]
        for r in rows:
            if r["status"] == "Matched":
                continue
            lines.append(
                f'<tr><td>{r["box_id"]}</td>'
                f'<td>{"Yes" if r["citation_found"] else "No"}</td>'
                f'<td>{"Yes" if r["caption_found"] else "No"}</td>'
                f'<td>{icon.get(r["status"], r["status"])}</td>'
                f'<td>{r["caption_text"]}</td></tr>'
            )
        lines.append('</tbody></table>')
        return "\n".join(lines)


def build_element_mapping_html(
    dict_types: Dict[str, Any],
    type_key: str,
    chapter_number: str = "",
) -> str:
    data: Dict[str, Any] = dict_types.get(type_key) or {}
    if not data:
        return ""

    captions:     Dict[str, Any] = {**data.get("Caption", {}), **data.get("DerivedCaption", {})}
    citations:    Dict[str, Any] = data.get("Citation",    {})
    cap_pages:    Dict[str, Any] = {**data.get("CaptionPage", {}), **data.get("DerivedCaptionPage", {})}
    cit_pages:    Dict[str, Any] = data.get("CitationPage",{})
    cap_counts:   Dict[str, Any] = data.get("CaptionCount", {})
    cap_extra_pg: Dict[str, Any] = data.get("CaptionExtraPages", {})

    _cm = re.search(r'\d+', chapter_number) if chapter_number else None
    ch_digit = _cm.group() if _cm else ""

    def _id_prefix(label: str) -> str:
        m = re.search(r'(\d+)[.\-]\d+', label)
        return m.group(1) if m else ""

    def _norm(label: str) -> str:
        return re.sub(r'[.\-]', '-', label.strip().lower())

    all_ids: List[str] = sorted(
        set(captions.keys()) | set(citations.keys()),
        key=lambda x: [int(d) for d in re.findall(r'\d+', str(x))]
    )

    # Detects panel-letter IDs: any ID ending with a letter directly after a digit
    # e.g. "Figure 10.1D", "Figure 3.2A"
    _panel_label_re = re.compile(r'.*\d[A-Za-z]$')

    icon = {
        "Matched":             "✅ Matched",
        "Duplicate Caption":   "⚠️ Duplicate Caption",
        "Missing Caption":     "⚠️ Missing Caption",
        "Missing Part Label":  "⚠️ Missing Part Label in Caption",
        "Orphan":              "⚠️ Missing citation",
        "Orphan Part Label":   "⚠️ Missing citation for part label",
        "Cross-Chapter Ref":   "ℹ️ Cross-Chapter Ref",
    }

    rows: List[Dict[str, Any]] = []
    cross_ids: List[str] = []
    # Plain-number pattern: "Figure 6", "Figure 18B" — no chapter dot/dash separator
    _plain_num_re = re.compile(
        r'(?i)^(?:figures?|figs?\.?|tables?|tabs?\.?|box(?:es)?|exhibits?|appendix|case\s+stud(?:y|ies))\s+\d+[A-Za-z]?$'
    )

    for label in all_ids:
        prefix = _id_prefix(label)
        if ch_digit and prefix and prefix != ch_digit:
            cross_ids.append(label)
            continue
        # Bare-number refs (e.g. "Figure 6", "Figure 18B") have no chapter prefix;
        # when a chapter number is known treat them as cross-chapter references.
        if ch_digit and not prefix and _plain_num_re.match(label.strip()):
            cross_ids.append(label)
            continue
        norm = _norm(label)
        cap_found = any(_norm(k) == norm for k in captions)
        cit_found = any(_norm(k) == norm for k in citations)

        if not cap_found:
            base_m = re.match(r'^(.*\d+)[A-Za-z]$', label.strip())
            if base_m:
                base_norm = _norm(base_m.group(1).strip())
                cap_found = any(_norm(k) == base_norm for k in captions)

        if not cit_found:
            cit_found = any(
                re.match(r'^' + re.escape(norm) + r'[a-z]$', _norm(k))
                for k in citations
            )

        # Fix 4: panel DerivedCaption entry (e.g., "Figure 3.3A") with no individual
        # citation — if the BASE figure ("Figure 3.3") is cited, the base citation
        # covers all its panels; treat as matched.
        if not cit_found:
            base_m2 = re.match(r'^(.*\d+)[A-Za-z]$', label.strip())
            if base_m2:
                base_norm2 = _norm(base_m2.group(1).strip())
                cit_found = any(_norm(k) == base_norm2 for k in citations)

        is_panel = bool(_panel_label_re.match(label.strip()))
        is_dup   = cap_counts.get(label, 1) > 1
        status = (
            "Duplicate Caption"  if cap_found and is_dup    else
            "Matched"            if cap_found and cit_found else
            "Missing Part Label" if cit_found and is_panel  else
            "Missing Caption"    if cit_found               else
            "Orphan Part Label"  if is_panel                else
            "Orphan"
        )

        # Build caption page display — show extra pages when duplicated
        cap_pg_raw = cap_pages.get(label, "")
        if is_dup:
            extras = cap_extra_pg.get(label, [])
            cap_pg_display = f"{cap_pg_raw} (also p.{', '.join(map(str, extras))})" if extras else str(cap_pg_raw)
        else:
            cap_pg_display = cap_pg_raw

        rows.append({
            "label":     label,
            "cit_found": cit_found,
            "cap_found": cap_found,
            "status":    status,
            "cit_page":  cit_pages.get(label, ""),
            "cap_page":  cap_pg_display,
        })

    for label in cross_ids:
        rows.append({
            "label":     label,
            "cit_found": label in citations,
            "cap_found": label in captions,
            "status":    "Cross-Chapter Ref",
            "cit_page":  cit_pages.get(label, ""),
            "cap_page":  cap_pages.get(label, ""),
        })

    if not rows:
        return ""

    def _safe_page(r):
        cit_p = str(r.get("cit_page", ""))
        cap_p = str(r.get("cap_page", ""))
        # Try finding the first digit sequence in the page fields
        cit_m = re.search(r'\d+', cit_p)
        if cit_m: return int(cit_m.group())
        cap_m = re.search(r'\d+', cap_p)
        if cap_m: return int(cap_m.group())
        return 999999

    # Sort rows by available Page Number
    rows.sort(key=_safe_page)

    lines = [
        f'<h3>{type_key} Citation \u2194 Caption Mapping</h3>',
        '<table><thead><tr>'
        f'<th>{type_key} ID</th>'
        '<th>Cited in Text</th><th>Caption Found</th>'
        '<th>Status</th><th>Citation Page</th><th>Caption Page</th>'
        '</tr></thead><tbody>',
    ]
    _HIDE_STATUSES = {"Matched", "Duplicate Caption"}
    issue_rows = [r for r in rows if r["status"] not in _HIDE_STATUSES]

    if not issue_rows:
        lines.append("<tr><td colspan='6' style='text-align:center;color:#27ae60;padding:12px;'>"
                     "&#10003; No issues found — all elements matched.</td></tr>")
    else:
        for r in issue_rows:
            lines.append(
                f'<tr>'
                f'<td>{r["label"]}</td>'
                f'<td>{"Yes" if r["cit_found"] else "No"}</td>'
                f'<td>{"Yes" if r["cap_found"] else "No"}</td>'
                f'<td>{icon.get(r["status"], r["status"])}</td>'
                f'<td>{r["cit_page"]}</td>'
                f'<td>{r["cap_page"]}</td>'
                f'</tr>'
            )
    lines.append('</tbody></table>')
    return "\n".join(lines)


def build_detailed_summary_table(
    dict_types: dict,
    figure_count: int,
    table_count: int,
    footnote_count: int,
    endnote_count: int,
    fmt_content: str,
    spec_content: str,
    comment_content: str,
    ref_count: int = 0,
    ref_style: str = "—",
    unnumbered_counts: dict = None,
    chapter_number: str = "",
    box_linker: Optional["BoxTagLinker"] = None,
) -> str:

    def count_items(section_html: str, token: str) -> int:
        return section_html.lower().count(token.lower())

    def extract_num(item: str) -> str:
        parts = item.strip().split()
        return parts[-1] if parts else item

    def format_num_list(items: List[str]) -> str:
        def _num_key(s: str):
            return [int(x) for x in re.findall(r'\d+', s)]
        nums = sorted(set(extract_num(i) for i in items), key=_num_key)
        return ", ".join(nums)

    def _chap_digit(chapter_number: str) -> str:
        m = re.search(r'\d+', chapter_number)
        return m.group() if m else ""

    def _item_chap(item: str) -> str:
        m = re.search(r'(\d+)', item.strip())
        return m.group(1) if m else ""

    def normalize_ref(ref: str) -> str:
        return ref.replace("-", ".").strip().lower()

    cn = _chap_digit(chapter_number) if chapter_number else ""

    def _is_cross_chapter(item: str) -> bool:
        """Return True for cross-chapter refs or plain-number refs (e.g. 'Figure 6')."""
        if not cn:
            return False
        norm_item = normalize_ref(item)
        # Plain number refs (no dot/dash, e.g. "Figure 8") → always treat as cross-chapter
        is_plain = bool(re.match(
            r'(?:figure|table|box|exhibit|appendix|case\s+study)\s+\d+[a-z]?$',
            norm_item
        ))
        if is_plain:
            return True
        return _item_chap(item) != cn

    # Panel-label regex: ID ending with a letter after a digit e.g. "Figure 10.1D"
    _panel_item_re = re.compile(r'.*\d[A-Za-z]$')

    def _format_action(verb: str, kind: str, items: List[str],
                       cross_chapter_items: Optional[List[str]] = None) -> str:
        """
        Build a human-readable action sentence.

        items               — current-chapter items that need attention
        cross_chapter_items — refs/captions belonging to other chapters (shown separately)
        """
        cross_chapter_items = cross_chapter_items or []
        parts: List[str] = []

        # Split current-chapter items into panel-label vs normal
        panel_items  = [i for i in items if _panel_item_re.match(str(i).strip())]
        normal_items = [i for i in items if not _panel_item_re.match(str(i).strip())]

        if normal_items:
            parts.append(
                f"Missing {len(normal_items)} {kind}(s): {verb} {format_num_list(normal_items)}."
            )
        if panel_items:
            parts.append(
                f"Missing {len(panel_items)} part label(s): {verb} {format_num_list(panel_items)}."
            )

        # Cross-chapter items grouped by chapter number
        if cross_chapter_items:
            from collections import defaultdict
            ch_groups: dict = defaultdict(list)
            for item in cross_chapter_items:
                num = extract_num(item)
                ch = num.split('.')[0] if '.' in num else (num.split('-')[0] if '-' in num else num)
                ch_groups[ch].append(num)
            group_strs = []
            for ch in sorted(ch_groups, key=lambda x: int(x) if x.isdigit() else 0):
                nums = ", ".join(sorted(
                    ch_groups[ch],
                    key=lambda x: [int(d) for d in re.findall(r'\d+', x)]
                ))
                group_strs.append(f"Chapter {ch}: {nums}")
            parts.append("Other chapters: " + " | ".join(group_strs) + ".")

        return "<br>".join(parts) if parts else ""

    def build_action_text(miss_cap_items: List[str],
                          miss_cit_items: List[str], chapter_number: str = "") -> str:
        cap_icon = "<i class='fas fa-times-circle' style='color:#e74c3c;'></i> "
        cit_icon = "<i class='fas fa-exclamation-triangle' style='color:#f39c12;'></i> "
        ch = _chap_digit(chapter_number)

        def _split(items: List[str]):
            """Split into (same_chapter_items, other_chapter_items)."""
            if not ch:
                return items, []
            same  = []
            other = []
            for i in items:
                norm = normalize_ref(i)
                # Plain single-number refs (e.g. "Figure 8") → other chapter
                is_plain = bool(re.match(
                    r'(?:figure|table|box|exhibit|appendix|case\s+study)\s+\d+[a-z]?$',
                    norm
                ))
                if is_plain or _item_chap(i) != ch:
                    other.append(i)
                else:
                    same.append(i)
            return same, other

        same_cap, other_cap = _split(miss_cap_items)
        same_cit, other_cit = _split(miss_cit_items)

        parts = []
        cap_text = _format_action(
            "Provide captions for", "caption", same_cap,
            cross_chapter_items=other_cap
        )
        cit_text = _format_action(
            "Insert citations for", "citation", same_cit,
            cross_chapter_items=other_cit
        )
        if cap_text:
            parts.append(cap_icon + cap_text)
        if cit_text:
            parts.append(cit_icon + cit_text)
        return "<br>".join(parts) if parts else "No action required"

    def build_progress_row(title: str, cap_cnt: int, cit_cnt: int, miss_cap: int, miss_cit: int,
                           action_text: str = "No action required",
                           miss_cap_items: Optional[List[str]] = None,
                           miss_cit_items: Optional[List[str]] = None,
                           chapter_number: str = "",
                           tab_target: str = "citations") -> str:
        miss_cap_items = miss_cap_items or []
        miss_cit_items = miss_cit_items or []
        total = max(cap_cnt, cit_cnt)
        complete_pct = round(((total - miss_cap - miss_cit) / total * 100), 1) if total else 0
        html = (
            f"<tr class='summary-table-row'>\n"
            f"  <td style='cursor:pointer;' onclick=\"showTabFromRow('{tab_target}', this.closest('tr'))\"><strong>{title}</strong></td>\n"
            f"  <td>{total}</td>\n"
            f"  <td>\n"
            f"    <div style='display:flex;align-items:center;gap:10px;'>\n"
            f"      <div style='width:100px;height:20px;background:#f0f0f0;border-radius:10px;overflow:hidden;'>\n"
            f"        <div style='width:0%;height:100%;background:linear-gradient(90deg,#27ae60,#2ecc71);transition:width 1s ease-in-out;' data-w='{complete_pct}'></div>\n"
            f"      </div>\n"
            f"      <a href='javascript:void(0);' onclick=\"showTab('{tab_target}');\" style='font-size:12px;color:#27ae60;text-decoration:none;'>{complete_pct}% Complete</a>\n"
            f"    </div>\n"
            f"  </td>\n"
            f"  <td><i class='fas fa-check-circle' style='color:#27ae60;'></i> {cit_cnt} citation(s)</td>\n"
            f"  <td><i class='fas fa-check-circle' style='color:#27ae60;'></i> {cap_cnt} caption(s)</td>\n"
            f"  <td>{action_text}</td>\n"
            f"</tr>\n"
        )
        return html

    fmt_count = count_items(fmt_content, "<tr><td>")
    spec_count = count_items(spec_content, "<tr><td>")
    comment_count_val = count_items(comment_content, "<tr><td>")

    global_stats = {
        "fmt_issues": fmt_count,
        "spec_count": spec_count,
        "total_citations": 0,
        "total_captions": 0,
        "missing_citations": 0,
        "missing_captions": 0,
        "fig_missing_cap": 0,
        "fig_missing_cit": 0,
        "tab_missing_cap": 0,
        "tab_missing_cit": 0,
        "box_missing_cap": 0,
        "box_missing_cit": 0,
        "fig_total_cit": 0,
        "fig_total_cap": 0,
        "tab_total_cit": 0,
        "tab_total_cap": 0,
        "box_total_cit": 0,
        "box_total_cap": 0
    }

    fig_cap = fig_cit = fig_miss_cap = fig_miss_cit = 0
    tab_cap = tab_cit = tab_miss_cap = tab_miss_cit = 0
    fig_miss_cap_items: List[str] = []
    fig_miss_cit_items: List[str] = []
    tab_miss_cap_items: List[str] = []
    tab_miss_cit_items: List[str] = []

    for type_key in dict_types.keys():
        if type_key == "Figure":
            fig_cap = len(dict_types[type_key]["Caption"])
            fig_cit = len(dict_types[type_key]["Citation"])
            cap_norms = set(
                [normalize_ref(x) for x in dict_types[type_key]["Caption"]] +
                [normalize_ref(x) for x in dict_types[type_key].get("DerivedCaption", {})]
            )
            cit_norms = set(normalize_ref(x) for x in dict_types[type_key]["Citation"])
            for k in dict_types[type_key]["Citation"]:
                if _is_cross_chapter(k):
                    continue
                norm = normalize_ref(k)
                if norm in cap_norms:
                    continue
                # Fix 2: panel-letter citation (e.g. "Figure 3.2A") — check base figure caption
                base_m = re.match(r'^(.*\d+)[A-Za-z]$', k.strip())
                if base_m and normalize_ref(base_m.group(1).strip()) in cap_norms:
                    continue
                fig_miss_cap += 1
                fig_miss_cap_items.append(str(k))
            for k in dict_types[type_key]["Caption"]:
                if _is_cross_chapter(k):
                    continue
                norm = normalize_ref(k)
                if norm in cit_norms:
                    continue
                # Fix 3: base figure orphan — check if any panel variant is cited
                if any(cn.startswith(norm) and len(cn) == len(norm) + 1 and cn[-1].isalpha()
                       for cn in cit_norms):
                    continue
                fig_miss_cit += 1
                fig_miss_cit_items.append(str(k))
        elif type_key == "Table":
            tab_cap = len(dict_types[type_key]["Caption"])
            tab_cit = len(dict_types[type_key]["Citation"])
            tab_cap_norms = set(
                [normalize_ref(x) for x in dict_types[type_key]["Caption"]] +
                [normalize_ref(x) for x in dict_types[type_key].get("DerivedCaption", {})]
            )
            tab_cit_norms = set(normalize_ref(x) for x in dict_types[type_key]["Citation"])
            for k in dict_types[type_key]["Citation"]:
                if _is_cross_chapter(k):
                    continue
                norm = normalize_ref(k)
                if norm in tab_cap_norms:
                    continue
                base_m = re.match(r'^(.*\d+)[A-Za-z]$', k.strip())
                if base_m and normalize_ref(base_m.group(1).strip()) in tab_cap_norms:
                    continue
                tab_miss_cap += 1
                tab_miss_cap_items.append(str(k))
            for k in dict_types[type_key]["Caption"]:
                if _is_cross_chapter(k):
                    continue
                norm = normalize_ref(k)
                if norm in tab_cit_norms:
                    continue
                if any(cn.startswith(norm) and len(cn) == len(norm) + 1 and cn[-1].isalpha()
                       for cn in tab_cit_norms):
                    continue
                tab_miss_cit += 1
                tab_miss_cit_items.append(str(k))

    html = """
    <div class='header'>
        <table style='margin-bottom:20px;width:100%;border-collapse:collapse;'>
        <thead>
          <tr>
            <th>Element Type</th>
            <th>Totals</th>
            <th>Status Overview</th>
            <th>Citations Status</th>
            <th>Captions Status</th>
            <th>Recommended Actions</th>
          </tr>
        </thead><tbody>
    """

    html += build_progress_row("Figures", fig_cap, fig_cit, fig_miss_cap, fig_miss_cit,
                               build_action_text(fig_miss_cap_items, fig_miss_cit_items, chapter_number),
                               fig_miss_cap_items, fig_miss_cit_items, chapter_number)
    html += build_progress_row("Tables", tab_cap, tab_cit, tab_miss_cap, tab_miss_cit,
                               build_action_text(tab_miss_cap_items, tab_miss_cit_items, chapter_number),
                               tab_miss_cap_items, tab_miss_cit_items, chapter_number)

    global_stats["missing_captions"] += fig_miss_cap + tab_miss_cap
    global_stats["missing_citations"] += fig_miss_cit + tab_miss_cit
    global_stats["total_captions"] += fig_cap + tab_cap
    global_stats["total_citations"] += fig_cit + tab_cit
    global_stats["fig_missing_cap"] = fig_miss_cap
    global_stats["fig_missing_cit"] = fig_miss_cit
    global_stats["tab_missing_cap"] = tab_miss_cap
    global_stats["tab_missing_cit"] = tab_miss_cit
    global_stats["fig_total_cit"] = fig_cit
    global_stats["fig_total_cap"] = fig_cap
    global_stats["tab_total_cit"] = tab_cit
    global_stats["tab_total_cap"] = tab_cap

    if box_linker is not None:
        bx_cit_cnt = len(box_linker.citations)
        bx_cap_cnt = len(box_linker.definitions)
        bx_miss_cap = [nid for nid in box_linker.citations if nid not in box_linker.definitions]
        bx_miss_cit = [nid for nid in box_linker.definitions if nid not in box_linker.citations]
        global_stats["missing_captions"] += len(bx_miss_cap)
        global_stats["missing_citations"] += len(bx_miss_cit)
        global_stats["total_captions"] += bx_cap_cnt
        global_stats["total_citations"] += bx_cit_cnt
        global_stats["box_missing_cap"] = len(bx_miss_cap)
        global_stats["box_missing_cit"] = len(bx_miss_cit)
        global_stats["box_total_cit"] = bx_cit_cnt
        global_stats["box_total_cap"] = bx_cap_cnt
        html += build_progress_row(
            "Boxes", bx_cap_cnt, bx_cit_cnt,
            len(bx_miss_cap), len(bx_miss_cit),
            build_action_text(bx_miss_cap, bx_miss_cit, chapter_number),
            bx_miss_cap, bx_miss_cit, chapter_number
        )

    other_types: List[str] = [str(k) for k in dict_types.keys() if k not in ("Figure", "Table")]
    for type_key in other_types:
        if box_linker is not None and type_key == "Box":
            continue
        o_cap = len(dict_types[type_key]["Caption"])
        o_cit = len(dict_types[type_key]["Citation"])
        o_miss_cap = 0
        o_miss_cit = 0
        o_miss_cap_items: List[str] = []
        o_miss_cit_items: List[str] = []
        for k in dict_types[type_key]["Citation"]:
            norm = normalize_ref(k)
            if not any(normalize_ref(x) == norm for x in dict_types[type_key]["Caption"]):
                o_miss_cap += 1
                o_miss_cap_items.append(str(k))
        for k in dict_types[type_key]["Caption"]:
            norm = normalize_ref(k)
            if not any(normalize_ref(x) == norm for x in dict_types[type_key]["Citation"]):
                o_miss_cit += 1
                o_miss_cit_items.append(str(k))

        global_stats["missing_captions"] += o_miss_cap
        global_stats["missing_citations"] += o_miss_cit
        global_stats["total_captions"] += o_cap
        global_stats["total_citations"] += o_cit

        if o_cap > 0 or o_cit > 0:
            html += build_progress_row(str(type_key) + "s", o_cap, o_cit, o_miss_cap, o_miss_cit,
                                       build_action_text(o_miss_cap_items, o_miss_cit_items, chapter_number),
                                       o_miss_cap_items, o_miss_cit_items, chapter_number)

    html += f"""
    <tr class='summary-table-row'><td style='cursor:pointer;' onclick="showTabFromRow('special-chars', this.closest('tr'))"><strong>Non-Std characters/typesetting instructions</strong></td><td>{spec_count}</td>
        <td colspan='3'><a href='javascript:void(0);' onclick="showTab('special-chars');"
        style='color:#667eea;text-decoration:none;'>Review non-standard characters</a></td>
        <td>Review non-standard characters and typesetting instructions</td></tr>

    <tr class='summary-table-row'><td style='cursor:pointer;' onclick="showTabFromRow('formatting', this.closest('tr'))"><strong>Formatting Issues</strong></td><td>{fmt_count}</td>
        <td colspan='3'><a href='javascript:void(0);' onclick="showTab('formatting');"
        style='color:#f39c12;text-decoration:none;'>View formatting issues</a></td>
        <td>Review formatting anomalies</td></tr>

    <tr class='summary-table-row'><td style='cursor:pointer;' onclick="showTabFromRow('comments', this.closest('tr'))"><strong>Comments</strong></td><td>{comment_count_val}</td>
        <td colspan='3'><a href='javascript:void(0);' onclick="showTab('comments');"
        style='color:#3498db;text-decoration:none;'>Review Comments and highlighted text</a></td>
        <td>No action required</td></tr>

    <tr class='summary-table-row'><td><strong>Notes</strong></td><td>{footnote_count + endnote_count}</td>
        <td colspan='3'>{footnote_count} Footnotes, {endnote_count} Endnotes</td>
        <td>No action required</td></tr>
    """

    if figure_count > 0:
        html += f"""
        <tr class='summary-table-row'><td style='cursor:pointer;' onclick="showTabFromRow('media', this.closest('tr'))"><strong>Images</strong></td><td>{figure_count}</td>
        <td colspan='3'><a href='javascript:void(0);' onclick="showTab('media');"
        style='color:#27ae60;text-decoration:none;'><i class='fas fa-check-circle'></i> {figure_count} image(s) detected</a></td>
        <td>No action required</td></tr>
        """
    else:
        html += """
        <tr class='summary-table-row'><td style='cursor:pointer;' onclick="showTabFromRow('media', this.closest('tr'))"><strong>Images</strong></td><td>0</td>
        <td colspan='3'><span style='color:#e67e22;'><i class='fas fa-exclamation-triangle'></i> No images detected</span></td>
        <td>No action required</td></tr>
        """

    if ref_count > 0:
        style_badge = ""
        if ref_style and ref_style != "—":
            badge_color = "#2980b9" if "AMA" in ref_style else "#8e44ad" if "APA" in ref_style else "#e67e22"
            style_badge = (
                f" &nbsp;<span style='display:inline-block;background:{badge_color};"
                f"color:white;font-size:0.75rem;font-weight:600;padding:1px 8px;"
                f"border-radius:20px;vertical-align:middle;'>{ref_style}</span>"
            )
        html += f"""
        <tr class='summary-table-row'><td style='cursor:pointer;' onclick="showTabFromRow('media', this.closest('tr'))"><strong>References</strong></td><td>{ref_count}</td>
        <td colspan='3'><span style='color:#27ae60;'><i class='fas fa-check-circle'></i> {ref_count} reference(s) detected</span>{style_badge}</td>
        <td>No action required</td></tr>
        """

    if unnumbered_counts:
        u_figs         = unnumbered_counts.get("unnumbered_images",    0)
        u_tabs         = unnumbered_counts.get("unnumbered_tables",    0)
        u_boxes        = unnumbered_counts.get("unnumbered_boxes",     0)
        u_callouts     = unnumbered_counts.get("callouts",             0)
        u_eq_omml      = unnumbered_counts.get("equations_omml",       0)
        u_eq_mt        = unnumbered_counts.get("equations_mathtype",   0)
        u_placeholders = unnumbered_counts.get("image_placeholders",   0)
        u_total = u_figs + u_tabs + u_boxes + u_callouts + u_eq_omml + u_eq_mt + u_placeholders
        if u_total > 0:
            detail = ", ".join(filter(None, [
                f"{u_figs} fig(s)"                       if u_figs          else "",
                f"{u_callouts} callout(s)"               if u_callouts      else "",
                f"{u_eq_omml} OMML eq(s)"                if u_eq_omml       else "",
                f"{u_eq_mt} MathType eq(s)"              if u_eq_mt         else "",
                f"{u_placeholders} image placeholder(s)" if u_placeholders  else "",
            ]))
            html += f"""
            <tr class='summary-table-row'><td style='cursor:pointer;' onclick="showTabFromRow('unnumbered', this.closest('tr'))"><strong>Unnumbered Elements</strong></td><td>{u_total}</td>
            <td colspan='3'><span style='color:#e67e22;'><i class='fas fa-exclamation-triangle'></i> {detail}</span></td>
            <td>Add numbers/captions where needed</td></tr>
            """

    html += "</tbody></table>"

    html += f"""
<style>
  .summary-row-active td {{ background:#f0f4ff !important; border-left:4px solid #667eea; }}
  .summary-table-row td:first-child:hover {{ background:#eef2ff; }}
</style>
<script>
(function(){{
  document.addEventListener('DOMContentLoaded', function(){{
    document.querySelectorAll('[data-w]').forEach(function(bar){{
      var w = bar.getAttribute('data-w');
      setTimeout(function(){{ bar.style.width = w + '%'; }}, 150);
    }});
  }});
}})();
</script>"""

    html += "</div>"

    return html, global_stats

def build_comments_html(comments: List[Tuple[str, str, str]]):
    if not comments:
        return "<p>No comments found.</p>"
    
    def _safe_page(val):
        m = re.search(r'\d+', str(val))
        return int(m.group()) if m else 999999

    # Sort comments by Page Number (third element)
    comments = sorted(comments, key=lambda x: _safe_page(x[2]))

    html = "<table><thead><tr><th>#</th><th>Author</th><th>Comment</th><th>Page</th></tr></thead><tbody>"
    for i, (author, text, page) in enumerate(comments, start=1):
        html += f"<tr><td>{i}</td><td>{escape_html(author)}</td><td>{escape_html(text)}</td><td>{page}</td></tr>"
    html += "</tbody></table>"
    return html


def build_unnumbered_tab_html(unnumbered_counts: dict) -> str:
    if not unnumbered_counts:
        return "<p>No unnumbered elements data available.</p>"

    def _with_pages(base, pages):
        if pages:
            return base + f"{', '.join(map(str, pages))}"
        return base

    callout_notes = "Multiple cross-references, including “see Figure X,” “see table above,” and page citations, appear on pages: "
    callout_pages = unnumbered_counts.get("callout_pages", [])
    if callout_pages:
        callout_notes += f"{', '.join(map(str, callout_pages))}"

    rows = [
        ("Figures",   unnumbered_counts.get("unnumbered_images",  0),
         _with_pages("Images without caption appear on pages: ",
                     unnumbered_counts.get("unnumbered_figure_pages", []))),
        ("Callouts",  unnumbered_counts.get("callouts",           0), callout_notes),
        ("OMML Equations",     unnumbered_counts.get("equations_omml",     0),
         _with_pages("Inline OMML math appear on pages: ",
                     unnumbered_counts.get("equations_omml_pages", []))),
        ("MathType Equations", unnumbered_counts.get("equations_mathtype", 0),
         _with_pages("MathType Equations appear on pages: ", unnumbered_counts.get("equations_mathtype_pages", []))),
    ]
    total = sum(r[1] for r in rows)
    html = "<table><thead><tr><th>Element Type</th><th>Count</th><th>Notes</th></tr></thead><tbody>"
    if total == 0:
        html += "<tr><td colspan='3'>No unnumbered elements found.</td></tr>"
    else:
        for label, count, notes in rows:
            if count > 0:
                html += f"<tr><td>{label}</td><td>{count}</td><td>{notes}</td></tr>"
    html += "</tbody></table>"
    return html


def build_export_highlight_html(paragraphs_full):
    highlights = []
    for t, p, is_cap, is_high in paragraphs_full:
        if is_high in ("Full", "Partial") or is_high is True:
            highlights.append((t, p))

    if not highlights:
        return "<p>No highlighted Text found.</p>"

    def _safe_page(val):
        m = re.search(r'\d+', str(val))
        return int(m.group()) if m else 999999

    # Sort highlights by Page Number (second element)
    highlights.sort(key=lambda x: _safe_page(x[1]))

    html = "<table><thead><tr><th>Highlighted Text</th><th>Page</th></tr></thead><tbody>\n"
    for t, p in highlights:
        html += f"<tr><td>{escape_html(t)}</td><td>{p}</td></tr>\n"
    html += "</tbody></table>\n"
    return html


# ------------------------------
# 3. New docx-based Implementations
# ------------------------------

def get_xml_comments(doc, text_page_map: Optional[List[str]] = None,
                     paras: Optional[List[Any]] = None):
    """Parses word/comments.xml to extract comments with PDF page numbers."""
    # Build text→page lookup from remapped paras (most accurate)
    text_page: dict[str, int] = {}
    if paras:
        for entry in paras:
            t, pg = entry[0], entry[1]
            if t and t not in text_page:
                text_page[t] = pg

    # Walk body to build comment_id → anchor paragraph text mapping
    comment_anchor: dict[str, str] = {}
    _w_crs = qn('w:commentRangeStart')
    _w_t   = qn('w:t')
    try:
        for p_el in doc.element.body.iter(qn('w:p')):
            para_text = ''.join(n.text or '' for n in p_el.iter(_w_t)).strip()
            for crs in p_el.iter(_w_crs):
                cid = crs.get(qn('w:id'), '')
                if cid and cid not in comment_anchor:
                    comment_anchor[cid] = para_text
    except Exception:
        pass

    def _page_for_anchor(anchor_text: str) -> str:
        if anchor_text and text_page:
            pg = text_page.get(anchor_text)
            if pg:
                return str(pg)
        if anchor_text and text_page_map:
            pg = _page_from_map(text_page_map, anchor_text, 0)
            if pg:
                return str(pg)
        return "N/A"

    comments = []
    try:
        for part in doc.part.package.parts:
            if part.partname.endswith('comments.xml'):
                comments_xml = part.blob
                root = etree.fromstring(comments_xml)
                namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                for comment in root.findall('.//w:comment', namespaces):
                    author  = comment.get(qn('w:author'), 'Unknown')
                    cid     = comment.get(qn('w:id'), '')
                    text_nodes = comment.findall('.//w:t', namespaces)
                    text    = "".join([t.text for t in text_nodes if t.text])
                    anchor  = comment_anchor.get(cid, '')
                    page    = _page_for_anchor(anchor)
                    comments.append((author, text, page))
    except Exception:
        pass
    return comments

def get_xml_note_count(doc, note_type='footnotes'):
    """Counts note references in the main document body."""
    count = 0
    try:
        if note_type == 'footnotes':
            tag = qn('w:footnoteReference')
        else:
            tag = qn('w:endnoteReference')
        elements = doc.element.findall(f'.//{tag}')
        count = len(elements)
    except Exception:
        pass
    return count

def _find_libreoffice() -> Optional[str]:
    """Find the LibreOffice executable path."""
    if os.name == 'nt':
        paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
        ]
        for path in paths:
            if os.path.exists(path):
                return path
    else:
        import shutil
        soffice = shutil.which("soffice")
        if soffice: return soffice
        libreoffice = shutil.which("libreoffice")
        if libreoffice: return libreoffice
    return None

def _pdf_page_map(doc_path: str, docx_paragraphs: List[str]) -> Dict[int, int]:
    """
    Converts to PDF and uses pdfplumber to map docx paragraph indices to PDF page numbers.
    Returns a dict mapping docx_paragraph_index to page_number (1-indexed).
    """
    if not HAS_PDFPLUMBER:
        return {}

    lo_path = _find_libreoffice()
    if not lo_path:
        return {}

    import tempfile
    import pdfplumber

    page_map = {}
    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            subprocess.run([
                lo_path, "--headless", "--convert-to", "pdf",
                "--outdir", tmpdir, doc_path
            ], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=60)
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired):
            return {}

        pdf_filename = os.path.splitext(os.path.basename(doc_path))[0] + ".pdf"
        pdf_path = os.path.join(tmpdir, pdf_filename)

        if not os.path.exists(pdf_path):
            return {}

        try:
            pdf_pages_text = []
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    pdf_pages_text.append(text if text else "")
        except Exception:
            return {}

        current_pdf_page_idx = 0
        num_pdf_pages = len(pdf_pages_text)
        normalized_pdf_pages = [_normalize_for_match(t) for t in pdf_pages_text]

        for p_idx, text in enumerate(docx_paragraphs):
            norm_text = _normalize_for_match(text)
            if not norm_text:
                page_map[p_idx] = current_pdf_page_idx + 1
                continue

            search_str = norm_text[:30]
            found = False
            for page_idx in range(current_pdf_page_idx, num_pdf_pages):
                if search_str and search_str in normalized_pdf_pages[page_idx]:
                    current_pdf_page_idx = page_idx
                    page_map[p_idx] = current_pdf_page_idx + 1
                    found = True
                    break

            if not found:
                page_map[p_idx] = current_pdf_page_idx + 1

    return page_map

def extract_with_docx(doc_path: str, doc: Optional[Any] = None):
    """
    Robust extraction using python-docx + lxml.
    Returns: paragraphs, comments, img_count, footnotes, endnotes
    """
    if doc is None:
        if not os.path.exists(doc_path):
            raise FileNotFoundError(f"{doc_path} not found")
        doc = Document(doc_path)

    def _is_sup(r):
        if r.font.superscript:
            return True
        if r.style:
            if r.style.font.superscript:
                return True
            sname = (r.style.name or "").lower()
            if "superscript" in sname or "footnote reference" in sname or "endnote reference" in sname:
                return True
            # "citation" alone flags numbered reference-superscript styles, but
            # FigureCitation/TableCitation (applied by docx_pipeline step8) mark
            # genuine in-text Figure/Table references and must be kept.
            if "citation" in sname and "figure" not in sname and "table" not in sname:
                return True
        return False

    analyzer = CitationAnalyzer()
    paragraphs = []

    has_lrpb = len(doc.element.findall('.//' + qn('w:lastRenderedPageBreak'))) > 0

    if has_lrpb:
        current_page = 1
        for p in doc.paragraphs:
            lrpb_count = len(p._element.findall('.//' + qn('w:lastRenderedPageBreak')))
            if lrpb_count:
                current_page += lrpb_count

            text = "".join(run.text for run in p.runs if not _is_sup(run)).strip()
            if not text:
                continue

            try:
                s_name = p.style.name
            except:
                s_name = ""
            is_caption = analyzer.is_caption_paragraph(text, style_name=s_name)

            has_highlighted_text = False
            has_unhighlighted_text = False
            for run in p.runs:
                if _is_sup(run):
                    continue
                if run.text.strip():
                    if run.font.highlight_color:
                        has_highlighted_text = True
                    else:
                        has_unhighlighted_text = True

            if has_highlighted_text:
                is_highlighted = "Partial" if has_unhighlighted_text else "Full"
            else:
                is_highlighted = False

            paragraphs.append((text, current_page, is_caption, is_highlighted))
    else:
        raw_paragraphs = []
        for p in doc.paragraphs:
            text = "".join(run.text for run in p.runs if not _is_sup(run)).strip()
            if not text:
                continue

            try:
                s_name = p.style.name
            except:
                s_name = ""
            is_caption = analyzer.is_caption_paragraph(text, style_name=s_name)

            has_highlighted_text = False
            has_unhighlighted_text = False
            for run in p.runs:
                if _is_sup(run):
                    continue
                if run.text.strip():
                    if run.font.highlight_color:
                        has_highlighted_text = True
                    else:
                        has_unhighlighted_text = True

            if has_highlighted_text:
                is_highlighted = "Partial" if has_unhighlighted_text else "Full"
            else:
                is_highlighted = False

            raw_paragraphs.append((text, is_caption, is_highlighted))

        texts_for_mapping = [p[0] for p in raw_paragraphs]
        pdf_pages = _pdf_page_map(doc_path, texts_for_mapping)

        for p_idx, (text, is_caption, is_highlighted) in enumerate(raw_paragraphs):
            page_num = pdf_pages.get(p_idx, 1)
            paragraphs.append((text, page_num, is_caption, is_highlighted))

    # Scan first row of each table for embedded captions
    for tbl_idx, table in enumerate(doc.tables):
        if not table.rows:
            continue
        first_row = table.rows[0]
        found_caption = False
        for cell in first_row.cells:
            for cp in cell.paragraphs:
                cell_text = "".join(run.text for run in cp.runs if not _is_sup(run)).strip()
                if not cell_text:
                    continue
                try:
                    s_name = cp.style.name or ""
                except Exception:
                    s_name = ""
                clean_text = re.sub(r'^(?:<[^>]*>\s*)+', '', cell_text)
                if analyzer.is_caption_paragraph(clean_text, style_name=str(s_name)):
                    page_est = paragraphs[-1][1] if paragraphs else 1
                    paragraphs.append((cell_text, page_est, True, False))
                    found_caption = True
                    break
            if found_caption:
                break

    comments = get_xml_comments(doc)

    img_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            img_count += 1

    # Include SmartArt and Charts as images
    for drawing in doc.element.findall('.//' + qn('w:drawing')):
        xml_str = etree.tostring(drawing, encoding='unicode')
        if 'drawingml/2006/diagram' in xml_str or 'drawingml/2006/chart' in xml_str:
            img_count += 1

    footnotes = get_xml_note_count(doc, 'footnotes')
    endnotes = get_xml_note_count(doc, 'endnotes')

    return paragraphs, comments, img_count, footnotes, endnotes

def remove_tags_keep_formatting_docx(doc_path: str, doc: Optional[Any] = None):
    """Removes <tags> using regex on run text, preserving other formatting and LRPB elements.

    CRITICAL: Do NOT use run.text = ... which rebuilds the XML and destroys LRPB elements.
    Instead, directly modify <w:t> elements to preserve sibling elements like <w:lastRenderedPageBreak/>.
    """
    should_save = False
    if doc is None:
        if not os.path.exists(doc_path):
            return
        doc = Document(doc_path)
        should_save = True

    tag_cleaner = re.compile(r'<[^>]+>')
    _w_t = qn('w:t')
    modified = False

    def _clean_run_text_preserve_xml(run):
        """Modify <w:t> elements directly without using run.text = which destroys LRPB."""
        nonlocal modified
        r = run._r  # Get the XML element (<w:r>)
        for t_elem in r.findall(_w_t):
            old_text = t_elem.text or ''
            if '<' in old_text and '>' in old_text:
                new_text = tag_cleaner.sub('', old_text)
                if new_text != old_text:
                    t_elem.text = new_text
                    modified = True

    for p in doc.paragraphs:
        for run in p.runs:
            _clean_run_text_preserve_xml(run)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        _clean_run_text_preserve_xml(run)

    if modified and should_save:
        doc.save(doc_path)

    return doc_path


def extract_unnumbered_image_markups(doc_path: str, doc: Optional[Any] = None, text_page_map: Optional[List[str]] = None) -> List[Dict[str, Any]]:
    """
    Scans paragraph text for image placeholder markup and returns a list of dicts.
    Must be called BEFORE remove_tags_keep_formatting_docx() so tags are still present.
    """
    if not os.path.exists(doc_path):
        return []
    if doc is None:
        doc = Document(doc_path)
    if text_page_map is None:
        _, text_page_map = build_text_page_map(doc_path)

    results = []
    for para_idx, p in enumerate(doc.paragraphs):
        text = p.text
        matches = _UNIMG_MARKUP_RE.findall(text)
        if matches:
            page = _page_from_map(text_page_map, text, para_idx // 40 + 1)
            for m in matches:
                results.append({"markup": m.strip(), "page": page})
    return results


def build_unnumbered_image_markups_html(items: List[Dict[str, Any]]) -> str:
    """Renders a list from extract_unnumbered_image_markups() as an HTML table."""
    if not items:
        return "<p>No unnumbered image placeholders found.</p>"

    def _safe_page(val):
        m = re.search(r'\d+', str(val))
        return int(m.group()) if m else 999999

    # Sort items by Page Number
    items = sorted(items, key=lambda x: _safe_page(x.get("page", "")))

    rows = []
    for i, item in enumerate(items, 1):
        markup_escaped = item["markup"].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        rows.append(
            f'<tr><td style="text-align:center">{i}</td>'
            f'<td><code>{markup_escaped}</code></td>'
            f'<td style="text-align:center">{item["page"]}</td></tr>'
        )

    rows_html = "\n".join(rows)
    return (
        '<table border="1" cellpadding="6" cellspacing="0" style="border-collapse:collapse;width:100%">'
        "<thead><tr>"
        '<th style="text-align:center">#</th>'
        "<th>Markup Text</th>"
        '<th style="text-align:center">Page</th>'
        "</tr></thead>"
        f"<tbody>{rows_html}</tbody>"
        "</table>"
    )


# ------------------------------
# PDF Conversion & Page Lookup
# ------------------------------

def convert_to_pdf(docx_path: str) -> str:
    """Convert a DOCX file to PDF. Returns the PDF path on success, or "" on failure."""
    import shutil
    pdf_path = str(Path(docx_path).with_suffix(".pdf"))

    if os.path.exists(pdf_path):
        return pdf_path

    lo_cmd = shutil.which("libreoffice") or shutil.which("soffice")
    if not lo_cmd and os.name == "nt" and os.path.exists(r"C:\Program Files\LibreOffice\program\soffice.exe"):
        lo_cmd = r"C:\Program Files\LibreOffice\program\soffice.exe"

    if not lo_cmd:
        print("DEBUG: LibreOffice not found in PATH. Conversion will fail.")
        return ""

    try:
        out_dir = str(Path(docx_path).parent)
        result = subprocess.run(
            [lo_cmd, "--headless", "--convert-to", "pdf", "--outdir", out_dir, os.path.abspath(docx_path)],
            timeout=60, capture_output=True
        )
        if result.returncode == 0 and os.path.exists(pdf_path):
            return pdf_path
    except Exception:
        pass

    return ""


def build_text_page_map(docx_path: str) -> Tuple[int, List[str]]:
    """
    Convert DOCX to PDF then build a list of normalized text for each page.
    Returns (total_pages, list_of_page_texts).
    """
    if not HAS_PDFPLUMBER:
        return 0, []
    pdf_path = convert_to_pdf(docx_path)
    if not pdf_path:
        return 0, []

    page_texts: List[str] = []
    total_pages = 0
    try:
        import pdfplumber as _pdfplumber
        with _pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)
            for page in pdf.pages:
                raw = page.extract_text() or ""
                page_texts.append(_normalize_for_match(raw))
    except Exception:
        pass
    return total_pages, page_texts


def _page_from_map(text_page_map: Optional[List[str]], text: str, fallback: int) -> int:
    """Look up real page number by searching normalized text; fall back to estimated value."""
    if text_page_map:
        search_str = _normalize_for_match(text)[:60]
        if search_str:
            for page_idx, p_text in enumerate(text_page_map):
                if search_str in p_text:
                    return page_idx + 1
    return fallback


def generate_formatting_html(doc_path: str, used_word: bool = False,
                             text_page_map: Optional[List[str]] = None,
                             doc: Optional[Any] = None,
                             paras: Optional[List[Any]] = None) -> str:
    """Scans for Strikethrough, Hidden, Page Breaks, Line Breaks, Text Boxes, Section Breaks."""
    if doc is None:
        doc = Document(doc_path)
    rows: list[tuple[str, str, str]] = []

    WNS      = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    TXBX_TAG = f"{{{WNS}}}txbxContent"
    _w_t     = qn('w:t')
    _w_br    = qn('w:br')
    _w_type  = qn('w:type')

    # Build text→page lookup from paras (same approach as count_unnumbered_elements)
    text_page: dict[str, int] = {}
    if paras:
        for entry in paras:
            t, pg = entry[0], entry[1]
            if t and t not in text_page:
                text_page[t] = pg

    has_lrpb = len(doc.element.findall('.//' + qn('w:lastRenderedPageBreak'))) > 0
    _w_lrpb  = qn('w:lastRenderedPageBreak')

    current_page = 1
    body_children = list(doc.element.body)

    for child in body_children:
        local = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if local in ['p', 'tbl']:
            if has_lrpb:
                lrpb_count = len(child.findall('.//' + _w_lrpb))
                if lrpb_count:
                    current_page += lrpb_count

        if local == 'p':
            if not has_lrpb:
                raw_text = ''.join(r.text or '' for r in child.iter(_w_t)).strip()
                if raw_text in text_page:
                    current_page = text_page[raw_text]
                elif not text_page and text_page_map:
                    # fallback to text_page_map when paras not provided
                    current_page = _page_from_map(text_page_map, raw_text, current_page)

            raw_text = ''.join(r.text or '' for r in child.iter(_w_t)).strip()
            page = str(current_page)
            para_preview = escape_html(raw_text[:60]) if raw_text else "Empty paragraph"

            # Walk runs for font flags and break elements
            for run_el in child.findall(f"{{{WNS}}}r"):
                # Font flags
                rpr = run_el.find(f"{{{WNS}}}rPr")
                if rpr is not None:
                    if rpr.find(f"{{{WNS}}}strike") is not None or rpr.find(f"{{{WNS}}}dstrike") is not None:
                        rows.append((page, "Strikethrough", para_preview))
                    if rpr.find(f"{{{WNS}}}vanish") is not None:
                        rows.append((page, "Hidden Text", para_preview))

                # Break elements
                for br in run_el.findall(_w_br):
                    br_type = br.get(_w_type, "")
                    if br_type == "page":
                        rows.append((page, "Page Break", para_preview))
                    elif br_type == "column":
                        rows.append((page, "Column Break", para_preview))
                    else:
                        rows.append((page, "Line Break", para_preview))

            # Text boxes
            for txbx in child.iter(TXBX_TAG):
                inner_text = " ".join(n.text for n in txbx.iter(_w_t) if n.text)
                rows.append((page, "Text Box", escape_html(inner_text[:60]) or "(empty text box)"))

    # Section breaks — find paragraph-level <w:sectPr> in body to get real page numbers
    stype_map = {
        "nextPage": "Next Page", "continuous": "Continuous",
        "evenPage": "Even Page", "oddPage": "Odd Page", "nextColumn": "Next Column"
    }
    _w_ppr    = f"{{{WNS}}}pPr"
    _w_sectpr = f"{{{WNS}}}sectPr"
    _w_type   = f"{{{WNS}}}type"
    sect_page = 1  # re-track pages for this second pass
    for child in body_children:
        local = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if local != 'p':
            continue
        if has_lrpb:
            lrpb_count = len(child.findall('.//' + _w_lrpb))
            if lrpb_count:
                sect_page += lrpb_count
        else:
            raw_text_s = ''.join(n.text or '' for n in child.iter(_w_t)).strip()
            if raw_text_s in text_page:
                sect_page = text_page[raw_text_s]
        ppr = child.find(_w_ppr)
        if ppr is None:
            continue
        sect_pr = ppr.find(_w_sectpr)
        if sect_pr is None:
            continue
        # Update current_page from this paragraph's text
        raw_text = ''.join(n.text or '' for n in child.iter(_w_t)).strip()
        sect_type_el = sect_pr.find(_w_type)
        raw_val = sect_type_el.get(f"{{{WNS}}}val", "nextPage") if sect_type_el is not None else "nextPage"
        stype = stype_map.get(raw_val, raw_val)
        rows.append((str(sect_page), f"Section Break ({stype})", raw_text[:60] or "(no text)"))

    def _safe_page(val):
        m = re.search(r'\d+', str(val))
        return int(m.group()) if m else 999999

    # Sort rows by Page Number (first element)
    rows.sort(key=lambda x: _safe_page(x[0]))

    html = "<table><thead><tr><th>Category</th><th>Details</th><th>Page</th></tr></thead><tbody>"
    if rows:
        for page, category, details in rows:
            html += f"<tr><td>{category}</td><td>{details}</td><td>{page}</td></tr>"
    else:
        html += "<tr><td colspan='3'>There is no formatting changes.</td></tr>"
    html += "</tbody></table>"
    return html

def generate_multilingual_html(doc_path: str,
                               text_page_map: Optional[List[str]] = None,
                               doc: Optional[Any] = None) -> str:
    """Scans for multilingual chars and keywords using python-docx without modifying the file."""
    if doc is None:
        doc = Document(doc_path)
    keyword_map: dict     = defaultdict(list)   # keyword  → [(page, text), ...]
    multilingual_map: dict = defaultdict(list)  # language → [(page, text), ...]

    keywords = [
        "Refer", "Insert", "Pick-up", "pickup", "See",
        "COMP", "AU", "AQ", "SPU", "Compositor",
        "Ph", "Photo", "video", "images"
    ]
    keyword_pattern = re.compile(r'\b(' + '|'.join(re.escape(k) for k in keywords) + r')\b\s+(\S+)', re.IGNORECASE)

    # Words that, when following "See"/"Refer", indicate a real cross-reference instruction.
    # Any other follower (pronoun, article, preposition) means it's regular prose.
    _SEE_VALID_FOLLOWERS = {
        'figure', 'fig', 'table', 'tab', 'box', 'plate',
        'above', 'below', 'chapter', 'section', 'appendix',
        'note', 'footnote', 'sidebar', 'insert', 'panel', 'also', 'p', 'pp'
    }
    # Words that, when following "images"/"video"/"Photo"/"Ph", mean regular English prose.
    _IMAGES_STOP = {
        'of', 'is', 'are', 'was', 'were', 'and', 'or', 'to', 'in',
        'on', 'at', 'by', 'for', 'with', 'from', 'the', 'a', 'an',
        'that', 'which', 'take', 'taken', 'show', 'shows', 'showing', 'shown'
    }
    _VIDEO_STOP = {
        'game', 'games', 'cassette', 'cassettes', 'clip', 'clips',
        'player', 'recording', 'screen', 'monitor', 'conference', 'conferencing'
    }

    analyzer = CitationAnalyzer()

    multilingual_ranges = [
        ("Chinese",      0x4E00, 0x9FFF),
        ("Greek",        0x0370, 0x03FF),
        ("Cyrillic",     0x0400, 0x04FF),
        ("Hebrew",       0x0590, 0x05FF),
        ("Arabic",       0x0600, 0x06FF),
        ("Devanagari",   0x0900, 0x097F),
        ("Japanese",     0x3040, 0x309F),
        ("Korean",       0xAC00, 0xD7AF),
        ("Thai",         0x0E00, 0x0E7F),
    ]

    for i, p in enumerate(doc.paragraphs):
        text = p.text
        if not text: continue
        page = _page_from_map(text_page_map, text, i // 40 + 1)

        try:
            s_name = p.style.name
        except Exception:
            s_name = ""
        if analyzer.is_caption_paragraph(text, style_name=s_name):
            continue  # skip: keyword matches inside captions are not instructions

        for match in keyword_pattern.finditer(text):
            kw = match.group(1).lower()
            follower = match.group(2).strip('.,;:()\'"').lower()
            # "See"/"Refer": only flag when followed by a reference term
            if kw in ('see', 'refer') and follower not in _SEE_VALID_FOLLOWERS:
                continue
            # "images"/"photo"/"ph": skip when followed by common prose words
            if kw in ('images', 'photo', 'ph') and follower in _IMAGES_STOP:
                continue
            # "video": skip common compound-noun followers
            if kw == 'video' and follower in _VIDEO_STOP:
                continue

            key = f"Keyword: {match.group(0).strip()}"
            # Extract only the sentence containing the match
            sent_start = 0
            for i in range(match.start() - 1, -1, -1):
                if text[i] in '.!?':
                    sent_start = i + 1
                    break
            sent_end = len(text)
            for i in range(match.end(), len(text)):
                if text[i] in '.!?':
                    sent_end = i + 1
                    break
            sentence = text[sent_start:sent_end].strip()
            keyword_map[key].append((page, sentence or text.strip()))

        # Collect unique multilingual chars per language for this paragraph
        ml_chars: dict = {}
        for char in text:
            code = ord(char)
            for lang, low, high in multilingual_ranges:
                if low <= code <= high:
                    ml_chars.setdefault(lang, [])
                    if char not in ml_chars[lang]:
                        ml_chars[lang].append(char)
                    break
        for lang, chars in ml_chars.items():
            multilingual_map[lang].append((page, ' '.join(chars)))

    def _snippet(text: str, keyword: str, window: int = 120) -> str:
        """Return a short context window around the keyword match."""
        m = re.search(r'\b' + re.escape(keyword) + r'\b', text, re.IGNORECASE)
        if not m:
            return (text[:220] + '…') if len(text) > 220 else text
        s = max(0, m.start() - window)
        e = min(len(text), m.end() + window)
        return ('…' if s > 0 else '') + text[s:e].strip() + ('…' if e < len(text) else '')

    def _make_table(title, data, empty_msg, line_col_label="Text"):
        t = f"<h3>{title}</h3>"
        # TYPE and text columns non-sortable; PAGE column gets default sort via dt-page-col class
        t += (f"<table><thead><tr>"
              f"<th data-orderable='false'>Type</th>"
              f"<th data-orderable='false'>{line_col_label}</th>"
              f"<th class='dt-page-col'>Page</th>"
              f"</tr></thead><tbody>")
        if data:
            all_entries = []
            for key, entries in data.items():
                raw_keyword = key.replace("Keyword: ", "")
                is_keyword  = key.startswith("Keyword:")
                for pg, para_text in entries:
                    if is_keyword:
                        snippet = _snippet(para_text, raw_keyword)
                        kw_escaped = re.escape(raw_keyword)
                        trailing_b = r'\b' if raw_keyword and (raw_keyword[-1].isalnum() or raw_keyword[-1] == '_') else ''
                        display_html = re.sub(
                            r'(\b' + kw_escaped + trailing_b + r')',
                            r'<mark>\1</mark>',
                            escape_html(snippet),
                            flags=re.IGNORECASE
                        )
                    else:
                        # Multilingual: para_text is already just the characters found
                        display_html = escape_html(para_text)
                    all_entries.append((key, display_html, pg))

            def _safe_page(val):
                m = re.search(r'\d+', str(val))
                return int(m.group()) if m else 999999

            all_entries.sort(key=lambda x: _safe_page(x[2]))

            for key, display_html, pg in all_entries:
                t += (
                    f"<tr>"
                    f"<td>{escape_html(key)}</td>"
                    f"<td>{display_html}</td>"
                    f"<td>{pg}</td>"
                    f"</tr>"
                )
        else:
            t += f"<tr><td colspan='3'>{empty_msg}</td></tr>"
        t += "</tbody></table>"
        return t

    html  = _make_table("Non-Standard Characters", multilingual_map, "No non-standard characters found", line_col_label="Characters")
    html += _make_table("Keywords", keyword_map, "No keywords found", line_col_label="Line")
    return html

# ------------------------------
# 4. Document Metadata & Extended Counts
# ------------------------------

_REFS_HEADING_RE = re.compile(
    r'^\s*('
    r'references?'
    r'|bibliographys?'
    r'|notes\s+and\s+bibliography'
    r'|works\s+citeds?'
    r'|cited\s+works?'
    r'|literature\s+cited'
    r'|reference\s+list'
    r'|selected\s+bibliography'
    r'|further\s+reading'
    r')\s*[:\.]?\s*$',
    re.IGNORECASE
)
_MARKUP_TAG_RE   = re.compile(r'<[^>]+>')
_FIGURE_LEGENDS_RE = re.compile(r'^\s*(figure\s+legends?|list\s+of\s+(figures?|tables?|illustrations?))\s*:?\s*$', re.IGNORECASE)
_AUTHOR_STYLE_RE = re.compile(r'author|by.?line|^a[0-9]$|^au\b', re.IGNORECASE)
_TITLE_STYLE_RE  = re.compile(r'heading\s*1|chapter\s*title|^ct\b|^title$', re.IGNORECASE)
_CALLOUT_RE      = re.compile(r'\b(see\s+(figure|fig\.?|table|tab\.?|box)\s+(above|below|following|on\s+page|[0-9]+(?:[\.\-][0-9]+)*[A-Za-z]?))\b', re.IGNORECASE)
_PAGE_REF_RE     = re.compile(r'\(p\.?\s*\d+\)', re.IGNORECASE)
_UNIMG_MARKUP_RE = re.compile(
    r'<\s*(?:[a-zA-Z0-9]+_)*'
    r'(?:unfig(?:ure)?'
    r'|csimage'
    r'|coimage'
    r'|insert\s+(?:photo|unf(?:ig(?:ure)?)?'
    r'|fig(?:ure)?|here)'
    r'|icon\s+here'
    r'|unf\b)'
    r'[^>]*>?',
    re.IGNORECASE
)


def extract_chapter_metadata(doc_path: str, doc: Optional[Any] = None):
    """Returns (chapter_number, chapter_title, authors) extracted from the chapter opener.

    Handles documents that begin with a section intro (SECTION I, section title,
    section overview, TOC) before the actual chapter opener paragraph.  The function
    scans up to 120 paragraphs, skipping section-level content, until it finds an
    explicit "Chapter N" marker and its associated title.
    """
    if not os.path.exists(doc_path):
        return "", ""
    if doc is None:
        doc = Document(doc_path)
    chapter_number = ""
    chapter_title = ""
    authors = []
    found_title = False

    _TAG_PREFIX_RE   = re.compile(r'^<([^>]+)>(.*)', re.DOTALL)
    _CHAP_NUM_RE     = re.compile(r'^(?:(?:chapter|ch\.?)\s*)?\d+\s*$', re.IGNORECASE)
    # Matches "Chapter 1" or "Chapter 1:" or "Chapter 1 —" at start, then title text
    _CHAP_INLINE_RE  = re.compile(
        r'^((?:chapter|ch\.?)\s*\d+)\s*[:\-\u2013\u2014]?\s+(.+)$', re.IGNORECASE)
    # Matches a bare "Chapter N" line (no title on same line)
    _CHAP_BARE_RE    = re.compile(r'^(?:chapter|ch\.?)\s*(\d+)\s*$', re.IGNORECASE)
    _TITLE_TAGS      = {'ct', 'chapter-title', 'chaptertitle', 'chap-title', 'chaptitle'}
    _AUTHOR_TAGS     = {'cau', 'au', 'author', 'byline', 'by-line', 'contrib'}
    _CHAP_NUM_TAGS   = {'cn', 'chapternum', 'chnum', 'cn1'}
    # Section-level noise to skip (all-caps section headers, bracketed instructions, quotes)
    _SKIP_RE         = re.compile(
        r'^(SECTION\s+[IVX\d]+|section\s+[IVX\d]+|\[.+\]|"[^"]+"\s*[\u2013\u2014\-])',
        re.IGNORECASE)

    paragraphs = doc.paragraphs[:60]

    # --- Pass 1: look for explicit Chapter N marker (tag-based or inline) ---
    for idx, p in enumerate(paragraphs):
        raw = p.text.strip()
        if not raw:
            continue
        try:
            s_name = p.style.name.strip()
        except:
            s_name = ""

        tag_match = _TAG_PREFIX_RE.match(raw)
        tag_name  = tag_match.group(1).lower().strip() if tag_match else ""
        text      = _MARKUP_TAG_RE.sub('', raw).strip()
        if not text:
            continue

        # Tag-based explicit markers — highest priority
        if tag_name in _TITLE_TAGS:
            chapter_title = text
            found_title = True
            continue
        if tag_name in _AUTHOR_TAGS:
            authors.append(text)
            continue
        if tag_name in _CHAP_NUM_TAGS:
            chapter_number = text
            continue
        if tag_name and (found_title or authors):
            break
        if tag_name:
            continue

        # Style-based
        if re.match(r'^cn\b', s_name, re.IGNORECASE) and not chapter_number:
            chapter_number = text
            continue
        if _TITLE_STYLE_RE.search(s_name):
            if _CHAP_BARE_RE.match(text) and not chapter_number:
                # e.g. "Chapter 1" styled as 'Title' — it's the chapter number, not the title
                chapter_number = text
            else:
                chapter_title = text
                found_title = True
            continue
        if _AUTHOR_STYLE_RE.search(s_name):
            authors.append(text)
            continue

        # Inline "Chapter N: Title" on a single line — stop scanning once found
        inline_m = _CHAP_INLINE_RE.match(text)
        if inline_m:
            chapter_number = inline_m.group(1).strip()
            chapter_title  = re.sub(r'^[:\-\u2013\u2014]\s*', '', inline_m.group(2)).strip()
            found_title = True
            break

        # Bare "Chapter N" line — look ahead for title on next non-empty paragraph
        bare_m = _CHAP_BARE_RE.match(text)
        if bare_m:
            chapter_number = text
            for look_idx, look_p in enumerate(paragraphs[idx + 1:idx + 5]):
                look_text = _MARKUP_TAG_RE.sub('', look_p.text.strip()).strip()
                if (look_text
                        and not _CHAP_BARE_RE.match(look_text)
                        and len(look_text.splitlines()) <= 3):
                    if not found_title:
                        chapter_title = look_text
                        found_title = True
                        
                        # Once title is found, grab the very next non-empty paragraph as a fallback author
                        # if it doesn't look like a standard heading (e.g. "Introduction" or "Contents")
                        try:
                            next_p_text = _MARKUP_TAG_RE.sub('', paragraphs[idx + 1 + look_idx + 1].text.strip()).strip()
                            if next_p_text and len(next_p_text.split()) < 15 and not re.match(r'^(introduction|contents|abstract|summary|objective)', next_p_text, re.IGNORECASE):
                                # Check if we already found authors via tags/styles, if not, use this
                                if not authors:
                                    authors.append(next_p_text)
                        except IndexError:
                            pass
                    break
            # Do not break here if we haven't found authors yet, let it continue scanning just in case 
            # there are explicitly tagged authors further down.
            if found_title and authors:
                break
            continue

        if chapter_number and not found_title:
            words = text.split()
            if len(words) <= 15:
                chapter_title = text
                found_title = True
            
            # If we found the title but no authors, grab the next line as a fallback author
            try:
                next_p_text = _MARKUP_TAG_RE.sub('', paragraphs[idx + 1].text.strip()).strip()
                if next_p_text and len(next_p_text.split()) < 15 and not re.match(r'^(introduction|contents|abstract|summary|objective)', next_p_text, re.IGNORECASE):
                    if not authors:
                        authors.append(next_p_text)
            except IndexError:
                pass
            
            if authors:
                break

    # --- Pass 2 fallback: if no explicit Chapter marker found, use first short
    #     non-section non-noise paragraphs (original behaviour for simple files) ---
    if not chapter_number and not found_title:
        for p in paragraphs[:10]:
            raw = p.text.strip()
            if not raw:
                continue
            text = _MARKUP_TAG_RE.sub('', raw).strip()
            if not text or _SKIP_RE.match(text):
                continue
            try:
                s_name = p.style.name.strip()
            except:
                s_name = ""
            if _TITLE_STYLE_RE.search(s_name):
                chapter_title = text
                found_title = True
                break
            words = text.split()
            if _CHAP_NUM_RE.match(text):
                chapter_number = text
            elif not found_title and len(words) <= 15:
                chapter_title = text
                found_title = True

    # Strip any leading separator from title that crept in
    if chapter_title:
        chapter_title = re.sub(r'^[:\-\u2013\u2014\s]+', '', chapter_title).strip()

    return chapter_number, chapter_title, "; ".join(authors) if authors else ""


def count_references_and_body_wc(doc_path: str, doc: Optional[Any] = None):
    """Returns (ref_count, body_wc, total_wc, ref_style).
    ref_style is one of: 'AMA (Numbered)', 'APA (Name–Year)', 'Mixed', or '—'.
    """
    if not os.path.exists(doc_path):
        return 0, 0, 0, "—"
    if doc is None:
        doc = Document(doc_path)

    table_para_texts: Set[int] = set()
    total_wc = 0
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for tp in cell.paragraphs:
                    table_para_texts.add(id(tp))
                    total_wc += len(tp.text.split())

    # Count words in text boxes (w:txbxContent) which doc.paragraphs misses
    _txbx_para_ids: Set[int] = set()
    for txbx in doc.element.findall('.//' + qn('w:txbxContent')):
        for p_elem in txbx.findall('.//' + qn('w:p')):
            _txbx_para_ids.add(id(p_elem))
            text = ''.join(t.text or '' for t in p_elem.findall('.//' + qn('w:t'))).strip()
            if text:
                total_wc += len(text.split())

    body_wc = 0
    ref_count = 0
    numbered_count = 0
    year_count = 0
    in_refs = False

    _CAPTION_START_RE = re.compile(r'^\s*(figure|fig\.?|table|tab\.?|box|exhibit|appendix)\s', re.IGNORECASE)

    for p in doc.paragraphs:
        raw_text = p.text.strip()
        if not raw_text:
            continue

        total_wc += len(raw_text.split())

        text = _MARKUP_TAG_RE.sub('', raw_text).strip()
        if not text:
            continue
        try:
            s_name = p.style.name.strip()
        except:
            s_name = ""

        if _REFS_HEADING_RE.match(text) and ("heading" in s_name.lower() or len(text.split()) <= 3):
            in_refs = True
            continue

        if in_refs:
            if _FIGURE_LEGENDS_RE.match(text):
                in_refs = True
                continue
            if "heading" in s_name.lower() and not re.match(r'^\d', text):
                in_refs = True
                continue
            if _CAPTION_START_RE.match(text):
                in_refs = True
                continue
            if re.match(r'^\s*(source|note[s]?|adapted\s+from|information\s+(based|from)|'
                        r'abbreviation|\*not\s+a\s+U\.S\.|IM,|IV,|PO,)', text, re.IGNORECASE):
                continue
            is_numbered = bool(re.match(r'^\[?\d+\]?[\.\)\t\s]', text))
            # Word auto-numbered list: the "1." prefix lives in XML (w:numPr), not in p.text
            if not is_numbered:
                try:
                    pPr = p._element.find(qn('w:pPr'))
                    if pPr is not None and pPr.find(qn('w:numPr')) is not None:
                        is_numbered = True
                except Exception:
                    pass
            has_year    = bool(re.search(r'(?<!\d)(?:19|20)\d{2}(?!\d)', text))
            # APA entries start with an author/org name (capital-letter word) + contain a year.
            # Continuation lines of numbered refs lack a leading digit and an author pattern,
            # so they are skipped — preventing year_count from being inflated by them.
            is_apa_start = (not is_numbered
                            and has_year
                            and bool(re.match(r'^[A-Z][a-zA-Z\-]+[\s,]', text)))
            if not (is_numbered or is_apa_start):
                continue
            ref_count += 1
            if is_numbered:
                numbered_count += 1
            else:
                year_count += 1
        else:
            body_wc += len(raw_text.split())

    # Determine style
    if ref_count == 0:
        ref_style = "—"
    elif numbered_count >= ref_count * 0.7:
        ref_style = "Numeric style"
    elif year_count >= ref_count * 0.7:
        ref_style = "Author-date style"
    else:
        ref_style = "Mixed"

    return ref_count, body_wc, total_wc, ref_style


def count_total_words_via_txt(doc_path: str) -> int:
    """Convert DOCX to plain text via LibreOffice and count words.
    Strips markup tags before counting to match Word's word count. Returns 0 if unavailable."""
    lo_path = _find_libreoffice()
    if not lo_path or not doc_path or not os.path.exists(doc_path):
        return 0
    import tempfile
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            subprocess.run(
                [lo_path, '--headless', '--convert-to', 'txt:Text', '--outdir', tmpdir, doc_path],
                capture_output=True, timeout=60
            )
            txt_name = os.path.splitext(os.path.basename(doc_path))[0] + '.txt'
            txt_path = os.path.join(tmpdir, txt_name)
            if not os.path.exists(txt_path):
                return 0
            with open(txt_path, encoding='utf-8', errors='replace') as f:
                text = f.read()
            text = _MARKUP_TAG_RE.sub('', text)
            return len(text.split())
    except Exception:
        return 0


def count_equations(doc: Any) -> Dict[str, int]:
    """Count equations in a python-docx Document."""
    _M_NS        = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    _W_NS        = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    _O_NS        = 'urn:schemas-microsoft-com:office:office'
    _omath_tag   = f'{{{_M_NS}}}oMath'
    _obj_tag     = f'{{{_W_NS}}}object'
    _ole_tag     = f'{{{_O_NS}}}OLEObject'
    _progid_attr = f'{{{_O_NS}}}ProgID'

    root = doc.element
    omml_count = len(root.findall(f'.//{_omath_tag}'))

    mathtype_count = 0
    for obj in root.findall(f'.//{_obj_tag}'):
        for ole in obj.findall(f'.//{_ole_tag}'):
            prog_id = ole.get(_progid_attr, '')
            if 'Equation' in prog_id or 'MathType' in prog_id:
                mathtype_count += 1

    return {"omml": omml_count, "mathtype": mathtype_count}


def count_unnumbered_elements(doc_path: str, dtypes: dict, doc: Optional[Any] = None, paras: Optional[list] = None):
    """Returns a dict with counts of unnumbered/uncaptioned elements, including page lists."""
    if not os.path.exists(doc_path):
        return {}
    if doc is None:
        doc = Document(doc_path)

    img_count = sum(1 for rel in doc.part.rels.values() if "image" in rel.reltype)

    # Include SmartArt and Charts as images
    for drawing in doc.element.findall('.//' + qn('w:drawing')):
        xml_str = etree.tostring(drawing, encoding='unicode')
        if 'drawingml/2006/diagram' in xml_str or 'drawingml/2006/chart' in xml_str:
            img_count += 1

    numbered_figs = len(dtypes.get("Figure", {}).get("Caption", {}))
    unnumbered_images = max(0, img_count - numbered_figs)

    numbered_tabs = len(dtypes.get("Table", {}).get("Caption", {}))
    unnumbered_tables = max(0, len(doc.tables) - numbered_tabs)

    numbered_boxes = len(dtypes.get("Box", {}).get("Caption", {}))
    box_style_re = re.compile(r'nbx|box|sidebar', re.IGNORECASE)
    box_para_count = sum(1 for p in doc.paragraphs
                         if p.text.strip() and box_style_re.search(getattr(p.style, 'name', '') or ''))
    unnumbered_boxes = max(0, box_para_count - numbered_boxes)

    callout_count = 0
    for p in doc.paragraphs:
        t = p.text
        if _CALLOUT_RE.search(t) or _PAGE_REF_RE.search(t):
            callout_count += 1

    eq = count_equations(doc)

    # --- Page tracking for unnumbered elements ---
    unnumbered_figure_pages = []
    unnumbered_table_pages  = []
    unnumbered_box_pages    = []
    omml_pages              = []

    if paras:
        # Build text→page lookup from paragraph list
        text_page = {}
        for entry in paras:
            t, pg = entry[0], entry[1]
            if t and t not in text_page:
                text_page[t] = pg

        # Caption text sets for quick lookup
        fig_caption_texts = set(dtypes.get("Figure", {}).get("Caption", {}).keys())
        tab_caption_ids   = set(dtypes.get("Table",  {}).get("Caption", {}).keys())

        _w_drawing = qn('w:drawing')
        _w_t       = qn('w:t')
        _w_p       = qn('w:p')
        _w_tbl     = qn('w:tbl')
        _m_omath   = qn('m:oMath')

        current_page = 1
        tbl_index    = 0
        _fig_cap_re  = re.compile(r'(?i)^(figure|fig\.?)\s+\d+')
        _tab_cap_re  = re.compile(r'(?i)^(?:<[A-Za-z]+>\s*)?(?:table|tab\.?)\s+\d+[\.\-]\d+')
        body_children = list(doc.element.body)

        for i, child in enumerate(body_children):
            local = child.tag.split('}')[-1] if '}' in child.tag else child.tag

            if local == 'p':
                raw_text = ''.join(r.text or '' for r in child.iter(_w_t)).strip()
                if raw_text in text_page:
                    current_page = text_page[raw_text]

                # Unnumbered inline image: paragraph contains <w:drawing> and
                # neither the prev nor next sibling paragraph is a Figure caption.
                # (Image paragraphs have no text themselves; the caption is adjacent.)
                if child.find('.//' + _w_drawing) is not None:
                    has_nearby_caption = False
                    for delta in (-1, 1):
                        idx = i + delta
                        if 0 <= idx < len(body_children):
                            sib = body_children[idx]
                            sib_local = sib.tag.split('}')[-1] if '}' in sib.tag else sib.tag
                            if sib_local == 'p':
                                sib_text = ''.join(r.text or '' for r in sib.iter(_w_t)).strip()
                                if _fig_cap_re.match(sib_text):
                                    has_nearby_caption = True
                                    break
                    if not has_nearby_caption and current_page not in unnumbered_figure_pages:
                        unnumbered_figure_pages.append(current_page)

                # OMML equations in this paragraph
                if child.find('.//' + _m_omath) is not None:
                    if current_page not in omml_pages:
                        omml_pages.append(current_page)

            elif local == 'tbl':
                if tbl_index < len(doc.tables):
                    tbl = doc.tables[tbl_index]
                    has_caption = False

                    # Check adjacent body paragraphs first (covers <TAB>Table N.N format)
                    for delta in (-2, -1, 1, 2):
                        idx = i + delta
                        if 0 <= idx < len(body_children):
                            sib = body_children[idx]
                            sib_local = sib.tag.split('}')[-1] if '}' in sib.tag else sib.tag
                            if sib_local == 'p':
                                sib_text = ''.join(r.text or '' for r in sib.iter(_w_t)).strip()
                                if _tab_cap_re.match(sib_text) or any(cap_id in sib_text or sib_text in cap_id for cap_id in tab_caption_ids):
                                    has_caption = True
                                    break

                    # Fallback: check inside table cells
                    if not has_caption:
                        for row in tbl.rows:
                            if has_caption:
                                break
                            for cell in row.cells:
                                if has_caption:
                                    break
                                for cp in cell.paragraphs:
                                    ct = cp.text.strip()
                                    if any(cap_id in ct or ct in cap_id for cap_id in tab_caption_ids):
                                        has_caption = True
                                        break

                    if not has_caption and current_page not in unnumbered_table_pages:
                        unnumbered_table_pages.append(current_page)
                tbl_index += 1

        # Box-style pages: iterate doc.paragraphs paired with paras for page info
        for p, para_entry in zip(doc.paragraphs, paras):
            if not p.text.strip():
                continue
            sname = getattr(p.style, 'name', '') or ''
            if box_style_re.search(sname):
                pg = para_entry[1]
                is_cap = para_entry[2]
                if not is_cap and pg not in unnumbered_box_pages:
                    unnumbered_box_pages.append(pg)

    return {
        "unnumbered_images":       unnumbered_images,
        "unnumbered_figure_pages": sorted(unnumbered_figure_pages),
        "unnumbered_tables":       unnumbered_tables,
        "unnumbered_table_pages":  sorted(unnumbered_table_pages),
        "unnumbered_boxes":        unnumbered_boxes,
        "unnumbered_box_pages":    sorted(unnumbered_box_pages),
        "callouts":                callout_count,
        "equations_omml":          eq["omml"],
        "equations_omml_pages":    sorted(omml_pages),
        "equations_mathtype":      eq["mathtype"],
    }


def build_combined_dashboard_html(chapters_data: list, css: str, js: str, logo_b64: str) -> str:
    """Builds a single Combined_Dashboard.html from a list of per-chapter data dicts."""
    template_dir = os.path.join(os.getcwd(), 'templates')

    if not os.path.exists(template_dir):
        template_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'templates')

    env = jinja2.Environment(loader=jinja2.FileSystemLoader(template_dir))

    try:
        template = env.get_template('word_analyzer_dashboard.html')
        css_content = env.get_template('word_analyzer_styles.css').render()
        js_content = env.get_template('word_analyzer_scripts.js').render()
    except jinja2.exceptions.TemplateNotFound as e:
        raise FileNotFoundError(f"Template not found: {e}. Ensure templates/ contains the HTML/CSS/JS dashboard files.")

    return template.render(
        chapters_data=chapters_data,
        css_content=css_content,
        js_content=js_content,
        logo_b64=logo_b64
    )



def build_new_combined_dashboard_html(chapters_data: list, logo_b64: str, book_title: str = "") -> str:
    """
    Renders the new dark-sidebar combined dashboard (combined_dashboard.html).
    Used for 2+ chapter runs. Individual chapter dashboards still use build_combined_dashboard_html().
    """
    template_dir = os.path.join(os.getcwd(), 'templates')
    if not os.path.exists(template_dir):
        template_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'templates')
    env = jinja2.Environment(loader=jinja2.FileSystemLoader(template_dir))
    try:
        template = env.get_template('combined_dashboard.html')
    except jinja2.exceptions.TemplateNotFound as e:
        raise FileNotFoundError(f"Template not found: {e}")
    return template.render(chapters_data=chapters_data, logo_b64=logo_b64, book_title=book_title)


# ------------------------------
# 5. Exports
# ------------------------------
__all__ = [
    "CitationAnalyzer",
    "BoxTagLinker",
    "build_element_mapping_html",
    "extract_with_docx",
    "generate_formatting_html",
    "generate_multilingual_html",
    "build_comments_html",
    "build_detailed_summary_table",
    "build_combined_dashboard_html",
    "build_export_highlight_html",
    "remove_tags_keep_formatting_docx",
    "extract_chapter_metadata",
    "count_references_and_body_wc",
    "count_total_words_via_txt",
    "count_unnumbered_elements",
    "extract_unnumbered_image_markups",
    "build_unnumbered_image_markups_html",
    "convert_to_pdf",
    "build_text_page_map",
    "_page_from_map",
    "HAS_PDFPLUMBER",
]


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python word_analyzer_docx.py <path_to_docx>")
        sys.exit(1)

    file_path = sys.argv[1]
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        sys.exit(1)

    print(f"Analyzing {file_path}...")
    try:
        doc = Document(file_path)
        paras, comments, imgs, footnotes, endnotes = extract_with_docx(file_path, doc=doc)
        print(f"Extraction Success:")
        print(f" - Paragraphs: {len(paras)}")
        print(f" - Comments: {len(comments)}")
        print(f" - Images: {imgs}")
        print(f" - Footnotes: {footnotes}")
        print(f" - Endnotes: {endnotes}")

        print("\nChecking Formatting...")
        fmt_html = generate_formatting_html(file_path, doc=doc)
        print("Formatting HTML generated (length: {} chars)".format(len(fmt_html)))

        print("\nChecking Multilingual...")
        multi_html = generate_multilingual_html(file_path, doc=doc)
        print("Multilingual HTML generated (length: {} chars)".format(len(multi_html)))

    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()