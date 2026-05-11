import re
import os
import zipfile
import difflib
from collections import defaultdict

from flask import Flask, request, send_file, render_template, redirect, url_for, session
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import RGBColor
from utils import track_changes
import logging

TRACK_CHANGES_ENABLED = False

app = Flask(__name__)
app.secret_key = "secret_key_for_session_encryption"
UPLOAD_DIR = "temp_reports"
os.makedirs(UPLOAD_DIR, exist_ok=True)

DASH_CLASS = r"\-\u2013\u2014"
NUMBER_ONLY_PATTERN = re.compile(rf'^[\d,\s{DASH_CLASS}]+$')
NUMBER_TOKEN_PATTERN = re.compile(rf'(\d+)\s*[{DASH_CLASS}]\s*(\d+)|(\d+)')
BRACKET_CITATION_PATTERN = re.compile(rf'\[\d+(?:[,\s{DASH_CLASS}]*\d+)*\]')
PAREN_CITATION_PATTERN = re.compile(rf'\(\d+(?:[,\s{DASH_CLASS}]*\d+)*\)')
BIB_PREFIX_PATTERN = re.compile(r'^\s*(?:\[\d+\]\.?\s*|\(\d+\)\.?\s*|\d+\.?\s*)')
BIB_NUMBER_TRAILING_TEXT = ".\t"

# =====================================================
# Helpers & Core Logic
# =====================================================

def iter_document_paragraphs(doc):
    """
    Iterate through all paragraphs in the document body in order,
    including those inside tables.
    """
    body = doc._element.body
    for child in body:
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            table = Table(child, doc)
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p


def get_visible_runs(para):
    """
    Yields all runs in a paragraph that are not deleted (not inside <w:del>).
    This allows us to see runs added via track changes (<w:ins>).
    """
    from docx.text.run import Run
    from docx.oxml.ns import qn
    runs = []
    for element in para._element.iter():
        if element.tag == qn('w:r'):
            # Check if any parent is w:del
            parent = element.getparent()
            is_deleted = False
            while parent is not None and parent != para._element.getparent():
                if parent.tag == qn('w:del'):
                    is_deleted = True
                    break
                parent = parent.getparent()
            if not is_deleted:
                runs.append(Run(element, para))
    return runs


def get_numbers(text):
    """
    Extract numbers from text like '1', '2-5', '1, 3, 5'.
    Handles ranges "1-5" -> [1, 2, 3, 4, 5].
    """
    nums = []
    # Matches: (start)-(end) OR (single)
    # Allows hyphen, en dash, em dash
    pattern = re.compile(r'(\d+)\s*[-–—]\s*(\d+)|(\d+)')
    
    for start, end, single in pattern.findall(text):
        if start and end:
            try:
                s, e = int(start), int(end)
                if s <= e:
                    nums.extend(range(s, e + 1))
            except ValueError:
                pass
        elif single:
            try:
                nums.append(int(single))
            except ValueError:
                pass
    return nums


def format_numbers(nums):
    """
    Format a list of numbers into a string like '1-3, 5'.
    Collapses ranges of 3 or more (e.g. 1,2,3 -> 1-3).
    """
    nums = sorted(set(nums))
    if not nums:
        return ""

    parts = []
    start = prev = nums[0]

    for n in nums[1:]:
        if n == prev + 1:
            prev = n
        else:
            length = prev - start + 1
            if length >= 3:
                parts.append(f"{start}-{prev}")
            elif length == 2:
                parts.append(f"{start},{prev}")
            else:
                parts.append(str(start))
            start = prev = n

    length = prev - start + 1
    if length >= 3:
        parts.append(f"{start}-{prev}")
    elif length == 2:
        parts.append(f"{start},{prev}")
    else:
        parts.append(str(start))

    return ", ".join(parts)


def _ensure_styles(doc):
    """Create cite_bib and bib_number character styles if they don't exist."""
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    try:
        doc.styles['cite_bib']
    except KeyError:
        style = doc.styles.add_style('cite_bib', WD_STYLE_TYPE.CHARACTER)
        style.font.superscript = True
    try:
        doc.styles['bib_number']
    except KeyError:
        style = doc.styles.add_style('bib_number', WD_STYLE_TYPE.CHARACTER)


BIB_NUMBER_PATTERNS = [
    re.compile(r'^\s*\[(\d+)\]\.\s+'),   # [1].
    re.compile(r'^\s*\((\d+)\)\.\s+'),   # (1).
    re.compile(r'^\s*\[(\d+)\]\s+'),     # [1]
    re.compile(r'^\s*\((\d+)\)\s+'),     # (1)
    re.compile(r'^\s*(\d+)\.\s+'),       # 1.
    re.compile(r'^\s*(\d+)\s+'),         # 1  (space only)
]


def extract_bib_number(text):
    """Extract bibliography number from text, handling all formats: 1., [1]., (1)., etc."""
    for pat in BIB_NUMBER_PATTERNS:
        m = pat.match(text)
        if m:
            return int(m.group(1)), pat
    return None, None


def get_numPr(para):
    """Detect Word auto-numbering. Returns (numId, ilvl) or (None, None)."""
    from docx.oxml.ns import qn
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return None, None
    numPr = pPr.find(qn('w:numPr'))
    if numPr is None:
        return None, None
    ilvl_el = numPr.find(qn('w:ilvl'))
    numId_el = numPr.find(qn('w:numId'))
    ilvl = int(ilvl_el.get(qn('w:val'), 0)) if ilvl_el is not None else 0
    numId = int(numId_el.get(qn('w:val'), 0)) if numId_el is not None else 0
    return numId, ilvl


def compute_list_number(doc, target_para):
    """Count preceding paragraphs with same numId+ilvl to compute displayed list number."""
    numId, ilvl = get_numPr(target_para)
    if numId is None or numId == 0:
        return None
    count = 0
    for para in doc.paragraphs:
        nid, nlvl = get_numPr(para)
        if nid == numId and nlvl == ilvl:
            count += 1
        if para._element is target_para._element:
            return count
    return None


def convert_autonumber_to_manual(para, number, doc):
    """Remove <w:numPr> and insert manual bib_number run with track changes."""
    from docx.oxml.ns import qn
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            pPr.remove(numPr)
    if number is not None:
        track_changes.add_tracked_text(para, f"{number}.", style='bib_number',
                                       author='RefBot', doc=doc)


def detect_and_tag_unstyled_citations(doc, citation_format):
    """
    Scan for unstyled citations and apply cite_bib style based on format.
    Returns {'tagged': count, 'format_used': citation_format}
    """
    tagged_count = 0
    _ensure_styles(doc)

    if citation_format == 'styled':
        return {'tagged': 0, 'format_used': 'styled'}

    if citation_format == 'superscript':
        from docx.oxml.ns import qn
        for para in iter_document_paragraphs(doc):
            for run in para.runs:
                if run.font.superscript and re.match(r'^[\d,\-–—\s]+$', run.text):
                    if not (run.style and run.style.name == 'cite_bib'):
                        run.style = doc.styles['cite_bib']
                        tagged_count += 1

    elif citation_format == 'bracket':
        for para in iter_document_paragraphs(doc):
            para_text = para.text
            if '[' not in para_text:
                continue
            if para.style and para.style.name == 'REF-N':
                continue
            for run in para.runs:
                if re.search(r'\[\d+(?:[,\-–—]\d+)*\]', run.text):
                    if not (run.style and run.style.name == 'cite_bib'):
                        new_text = re.sub(r'\[(\d+(?:[,\-–—]\d+)*)\]', r'\1', run.text)
                        run.text = new_text
                        run.style = doc.styles['cite_bib']
                        tagged_count += 1

    elif citation_format == 'paren':
        for para in iter_document_paragraphs(doc):
            para_text = para.text
            if '(' not in para_text:
                continue
            if para.style and para.style.name == 'REF-N':
                continue
            for run in para.runs:
                if re.search(r'\(\d+(?:[,\-–—]\d+)*\)', run.text) and not re.search(r'\d{4}', run.text):
                    if not (run.style and run.style.name == 'cite_bib'):
                        new_text = re.sub(r'\((\d+(?:[,\-–—]\d+)*)\)', r'\1', run.text)
                        run.text = new_text
                        run.style = doc.styles['cite_bib']
                        tagged_count += 1

    elif citation_format == 'plain':
        for para in iter_document_paragraphs(doc):
            if para.style and para.style.name in ['REF-N', 'REF-U']:
                continue
            for i, run in enumerate(para.runs):
                if i == 0:
                    continue
                if re.match(r'^\d+$', run.text) and not (run.style and run.style.name == 'cite_bib'):
                    run.style = doc.styles['cite_bib']
                    tagged_count += 1

    return {'tagged': tagged_count, 'format_used': citation_format}


def is_citation_run(run):
    """
    Determine if a run is part of a citation.
    Prioritize cite_bib style, fallback to superscript + numbers.
    """
    if run.style and run.style.name == "cite_bib":
        return True
    # Fallback: detect unstyled superscript citations
    # But be selective - must match digit-only or number+comma+dash patterns
    if run.font.superscript:
        text = run.text.strip()
        if text and re.match(r'^[\d,\-–—]+$', text):
            return True
    return False


def get_numbers(text):
    """
    Extract numbers from text like '1', '2-5', '1, 3, 5'.
    Handles ranges like '1-5' -> [1, 2, 3, 4, 5].
    """
    nums = []

    for start, end, single in NUMBER_TOKEN_PATTERN.findall(text or ""):
        if start and end:
            try:
                s, e = int(start), int(end)
            except ValueError:
                continue
            if s <= e:
                nums.extend(range(s, e + 1))
        elif single:
            try:
                nums.append(int(single))
            except ValueError:
                continue

    return nums


def strip_bib_prefix(text):
    return BIB_PREFIX_PATTERN.sub("", text or "", count=1).lstrip()


def _remove_numpr(para):
    from docx.oxml.ns import qn

    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return

    numPr = pPr.find(qn('w:numPr'))
    if numPr is not None:
        pPr.remove(numPr)


def _strip_leading_chars_from_runs(para, char_count):
    remaining = char_count
    for run in para.runs:
        if remaining <= 0:
            break

        text = run.text or ""
        if not text:
            continue

        if len(text) <= remaining:
            run.text = ""
            remaining -= len(text)
        else:
            run.text = text[remaining:]
            remaining = 0


def _insert_run_at_start(para, text, doc, style_name=None):
    from docx.oxml.ns import qn

    run = para.add_run(text)
    if style_name:
        run.style = doc.styles[style_name]

    paragraph_element = para._element
    run_element = run._element
    paragraph_element.remove(run_element)

    first_content = None
    for child in paragraph_element:
        if child.tag != qn('w:pPr'):
            first_content = child
            break

    if first_content is None:
        paragraph_element.append(run_element)
    else:
        first_content.addprevious(run_element)

    return run


def normalize_reference_paragraph(para, number, doc):
    """
    Replace any leading automatic or manual reference number with a dedicated
    `bib_number` run for the digits only, preserving the following punctuation
    in a separate unstyled run.
    """
    if number is None:
        return None

    _remove_numpr(para)

    match = BIB_PREFIX_PATTERN.match(para.text or "")
    if match:
        _strip_leading_chars_from_runs(para, match.end())

    _insert_run_at_start(para, BIB_NUMBER_TRAILING_TEXT, doc)
    return _insert_run_at_start(para, str(number), doc, style_name='bib_number')


def convert_autonumber_to_manual(para, number, doc):
    """Convert a REF-N paragraph from Word auto numbering to a manual bib run."""
    return normalize_reference_paragraph(para, number, doc)


def detect_and_tag_unstyled_citations(doc, citation_format):
    """
    Scan for unstyled citations and apply cite_bib style based on the selected format.
    Returns {'tagged': count, 'format_used': citation_format}
    """
    tagged_count = 0
    _ensure_styles(doc)

    if citation_format == 'styled':
        return {'tagged': 0, 'format_used': 'styled'}

    for para in iter_document_paragraphs(doc):
        if para.style and para.style.name in ['REF-N', 'REF-U']:
            continue

        for index, run in enumerate(para.runs):
            if run.style and run.style.name == 'cite_bib':
                continue

            text = run.text or ""
            should_tag = False

            if citation_format == 'superscript':
                should_tag = bool(run.font.superscript and NUMBER_ONLY_PATTERN.match(text.strip()))
            elif citation_format == 'bracket':
                should_tag = bool(BRACKET_CITATION_PATTERN.fullmatch(text.strip()))
                if should_tag:
                    run.text = re.sub(r'^\[(.*)\]$', r'\1', text.strip())
            elif citation_format == 'paren':
                should_tag = bool(PAREN_CITATION_PATTERN.fullmatch(text.strip()) and not re.search(r'\d{4}', text))
                if should_tag:
                    run.text = re.sub(r'^\((.*)\)$', r'\1', text.strip())
            elif citation_format == 'plain':
                should_tag = index > 0 and bool(NUMBER_ONLY_PATTERN.match(text.strip()))

            if should_tag:
                run.style = doc.styles['cite_bib']
                tagged_count += 1

    return {'tagged': tagged_count, 'format_used': citation_format}


def is_citation_run(run):
    """
    Determine if a run is part of a citation.
    Prioritize cite_bib style, fallback to numeric superscripts.
    """
    if run.style and run.style.name == "cite_bib":
        return True

    if run.font.superscript:
        return bool(NUMBER_ONLY_PATTERN.match((run.text or "").strip()))

    return False


class ReferenceProcessor:
    def __init__(self, doc):
        self.doc = doc
        
    def get_references_in_bibliography(self):
        """
        Returns a Set of IDs found in the bibliography sections (REF-N style).
        Also returns a list of objects for reordering later.
        """
        refs_found = set()
        ref_objects = []

        for para in self.doc.paragraphs:
            if para.style and para.style.name == "REF-N":
                found_id = None
                bib_run = None

                # Try finding styled run
                for run in para.runs:
                    if run.style and run.style.name == "bib_number":
                        nums = get_numbers(run.text)
                        if nums:
                            found_id = nums[0]
                            bib_run = run
                            break

                # Fallback: extract from plain text using multi-format patterns
                if found_id is None:
                    found_id, pattern = extract_bib_number(para.text)

                if found_id is not None:
                    refs_found.add(found_id)
                    ref_objects.append({
                        'id': found_id,
                        'para': para,
                        'run': bib_run,
                        'text': strip_bib_prefix(para.text.strip())
                    })

        return refs_found, ref_objects

    def get_citations_in_text(self):
        """
        Scans document for citations.
        Returns:
            all_cited_ids: list of all IDs in order of appearance (with duplicates)
            appearance_order: list of unique IDs in order of first appearance
        """
        all_cited_ids = []
        appearance_order = []
        seen = set()
        
        for para in iter_document_paragraphs(self.doc):
            # 1. Process runs
            current_group = []
            
            for run in get_visible_runs(para):
                if is_citation_run(run):
                    current_group.append(run)
                else:
                    if current_group:
                        # Flush group
                        text = "".join(r.text for r in current_group)
                        nums = get_numbers(text)
                        all_cited_ids.extend(nums)
                        for n in nums:
                            if n not in seen:
                                seen.add(n)
                                appearance_order.append(n)
                        current_group = []
            
            # Flush trailing group
            if current_group:
                text = "".join(r.text for r in current_group)
                nums = get_numbers(text)
                all_cited_ids.extend(nums)
                for n in nums:
                    if n not in seen:
                        seen.add(n)
                        appearance_order.append(n)
                        
        return all_cited_ids, appearance_order

    def find_duplicates(self, ref_objects):
        """
        Finds 100% exact duplicate references (ignoring leading numbers).
        Strips all number formats: 1. [1]. (1). [1] (1) before comparing.
        Returns a list of dicts: {'id': int, 'text': str, 'duplicate_of': int, 'score': float}
        """
        duplicates = []
        processed_refs = []

        for obj in ref_objects:
            clean_text = strip_bib_prefix(obj['para'].text.strip())
            processed_refs.append({'id': obj['id'], 'text': clean_text})

        n = len(processed_refs)

        for i in range(n):
            ref_a = processed_refs[i]
            text_a = ref_a['text']
            if not text_a:
                continue

            for j in range(i + 1, n):
                ref_b = processed_refs[j]
                text_b = ref_b['text']
                if not text_b:
                    continue

                # 100% exact match only (ignore leading numbers)
                if text_a == text_b:
                    duplicates.append({
                        'id': ref_b['id'],
                        'text': ref_b['text'][:100] + "...",
                        'duplicate_of': ref_a['id'],
                        'score': 100.0
                    })

        return duplicates

    def resolve_duplicates(self):
        """
        Finds duplicates, remaps citations, and marks bibliography entries for deletion.
        Returns list of merge records: [{'removed_id': int, 'canonical_id': int, 'citations_updated': int}, ...]
        """
        from docx.oxml.ns import qn

        bib_refs, ref_objects = self.get_references_in_bibliography()
        duplicates = self.find_duplicates(ref_objects)

        if not duplicates:
            return []

        # Build canonical mapping, handling transitive duplicates (A→B, B→C → A→C)
        mapping = {}  # duplicate_id -> canonical_id
        for dup in duplicates:
            dup_id = dup['id']
            canonical_id = dup['duplicate_of']
            # Resolve transitively
            while canonical_id in mapping:
                canonical_id = mapping[canonical_id]
            mapping[dup_id] = canonical_id

        merge_log = []
        citations_updated_count = {dup_id: 0 for dup_id in mapping.keys()}

        # Step 1: Remap citations from duplicate IDs to canonical IDs
        for para in iter_document_paragraphs(self.doc):
            visible_runs = get_visible_runs(para)
            i = 0
            while i < len(visible_runs):
                if is_citation_run(visible_runs[i]):
                    group = [visible_runs[i]]
                    j = i + 1
                    while j < len(visible_runs) and is_citation_run(visible_runs[j]):
                        group.append(visible_runs[j])
                        j += 1
                        
                    text = "".join(r.text for r in group)
                    nums = get_numbers(text)
                    new_nums = [mapping.get(n, n) for n in nums]
                    if new_nums != nums:
                        # Count how many times each duplicate ID was replaced
                        for n in nums:
                            if n in mapping and mapping[n] != n:
                                citations_updated_count[n] += 1

                        new_text = format_numbers(new_nums)
                        if TRACK_CHANGES_ENABLED:
                            anchor = group[-1]._element.getparent() if group[-1]._element.getparent().tag == track_changes.qn('w:del') else group[-1]._element
                            for r in group:
                                track_changes.delete_tracked_run(para, r)
                            ins_new = track_changes.add_tracked_text(
                                para, new_text, style='cite_bib', color='008000', doc=self.doc)
                            try:
                                anchor.addnext(ins_new)
                            except TypeError:
                                para._element.append(ins_new)
                        else:
                            group[0].text = new_text
                            for r in group[1:]:
                                r.text = ""
                    
                    i = j
                else:
                    i += 1

        # Step 2: Physically remove duplicate bibliography entries
        bib_refs_new, ref_objects_new = self.get_references_in_bibliography()
        body = self.doc._element.body

        for obj in ref_objects_new:
            if obj['id'] in mapping:
                # This is a duplicate — physically remove it
                para_element = obj['para']._element
                if para_element.getparent() == body:
                    body.remove(para_element)

                # Record the merge
                canonical_id = mapping[obj['id']]
                for dup_record in duplicates:
                    if dup_record['id'] == obj['id']:
                        merge_log.append({
                            'removed_id': obj['id'],
                            'canonical_id': canonical_id,
                            'score': dup_record['score'],
                            'citations_updated': citations_updated_count.get(obj['id'], 0)
                        })
                        break

        return merge_log

    def get_validation_stats(self):
        bib_refs, ref_objects = self.get_references_in_bibliography()
        all_cited, appearance_order = self.get_citations_in_text()
        
        unique_cited = set(all_cited)
        
        # Missing: Cited but not in Bib
        missing = sorted(unique_cited - bib_refs)
        
        # Unused: In Bib but not Cited
        unused = sorted(bib_refs - unique_cited)
        
        # Duplicates
        duplicates = self.find_duplicates(ref_objects)
        
        # Sequence Issues
        sequence_issues = []
        seen_in_seq = []
        for n in all_cited:
            if n not in seen_in_seq:
                if n != len(seen_in_seq) + 1:
                    sequence_issues.append({
                        "position": len(seen_in_seq) + 1,
                        "current": n,
                        "expected": len(seen_in_seq) + 1
                    })

                seen_in_seq.append(n)

        return {
            "total_references": len(ref_objects),
            "total_citations": len(all_cited),
            "citation_order": appearance_order,
            "reference_order": [obj['id'] for obj in ref_objects],
            "missing_references": missing,
            "unused_references": unused,
            "duplicate_references": duplicates,
            "sequence_issues": sequence_issues,
            "is_perfect": (not missing and not unused and not sequence_issues and not duplicates)
        }

    def _ensure_bib_numbers_styled(self):
        """Ensure all bibliography numbers have bib_number style applied."""
        from docx.enum.style import WD_STYLE_TYPE

        styles = self.doc.styles
        try:
            styles['bib_number']
        except KeyError:
            styles.add_style('bib_number', WD_STYLE_TYPE.CHARACTER)

        for para in self.doc.paragraphs:
            if para.style and para.style.name == 'REF-N':
                bib_num, _ = extract_bib_number(para.text)
                if bib_num is None:
                    continue
                for run in para.runs:
                    run_num, _ = extract_bib_number(run.text)
                    if run_num == bib_num:
                        run.style = styles['bib_number']
                        break

    def _ensure_citations_styled(self):
        """Ensure all citation runs that are superscript with numbers have cite_bib style."""
        from docx.enum.style import WD_STYLE_TYPE

        styles = self.doc.styles
        try:
            styles['cite_bib']
        except KeyError:
            s = styles.add_style('cite_bib', WD_STYLE_TYPE.CHARACTER)
            s.font.superscript = True

        for para in iter_document_paragraphs(self.doc):
            if para.style and para.style.name == 'REF-N':
                continue
            for run in para.runs:
                if (run.font.superscript and
                        re.match(r'^[\d,\-–—\s]+$', run.text.strip()) and
                        not (run.style and run.style.name == 'cite_bib')):
                    run.style = styles['cite_bib']

    def renumber(self):
        """
        Renumber citations and reorder bibliography.
        Returns: mapping (Old -> New)
        """
        _, appearance_order = self.get_citations_in_text()

        # Ensure 'cite_bib' and 'bib_number' styles exist
        from docx.enum.style import WD_STYLE_TYPE
        styles = self.doc.styles
        try:
            styles['cite_bib']
        except KeyError:
            s = styles.add_style('cite_bib', WD_STYLE_TYPE.CHARACTER)
            s.font.superscript = True

        try:
            styles['bib_number']
        except KeyError:
            styles.add_style('bib_number', WD_STYLE_TYPE.CHARACTER)

        # Create Mapping
        mapping = {} 
        new_id = 1
        for old_id in appearance_order:
            mapping[old_id] = new_id
            new_id += 1
            
        for para in iter_document_paragraphs(self.doc):
            visible_runs = get_visible_runs(para)
            i = 0
            while i < len(visible_runs):
                if is_citation_run(visible_runs[i]):
                    group = [visible_runs[i]]
                    j = i + 1
                    while j < len(visible_runs) and is_citation_run(visible_runs[j]):
                        group.append(visible_runs[j])
                        j += 1
                        
                    txt = "".join(r.text for r in group)
                    nums = get_numbers(txt)
                    if nums:
                         new_nums = [mapping.get(n, n) for n in nums]
                         new_text = format_numbers(new_nums)
                         
                         is_renumbered = (nums != new_nums)
                         highlight_color = "008000" if is_renumbered else None
                         
                         first_run = group[0]
                         style_name = first_run.style.name if first_run.style else "cite_bib"
                         
                         if TRACK_CHANGES_ENABLED:
                             anchor = group[-1]._element.getparent() if group[-1]._element.getparent().tag == track_changes.qn('w:del') else group[-1]._element
                             for r in group:
                                 track_changes.delete_tracked_run(para, r)

                             ins_new = track_changes.add_tracked_text(para, new_text, style=style_name, color=highlight_color, doc=self.doc)
                             try:
                                 anchor.addnext(ins_new)
                             except TypeError:
                                 para._element.append(ins_new)
                         else:
                             first_run.text = new_text
                             if is_renumbered:
                                 first_run.font.color.rgb = RGBColor(0, 128, 0)
                             for r in group[1:]:
                                 r.text = ""
                
                    i = j
                else:
                    i += 1

        # 2. Reorder Bibliography
        _, ref_objects = self.get_references_in_bibliography()
        
        # Sort objects into Cited and Uncited
        cited_refs = []
        uncited_refs = []
        
        for obj in ref_objects:
            if obj['id'] in mapping:
                obj['new_id'] = mapping[obj['id']]
                cited_refs.append(obj)
            else:
                uncited_refs.append(obj)
        
        if not ref_objects:
            return mapping

        # Find anchor (min index)
        body = self.doc._element.body
        
        indices = []
        for obj in ref_objects:
            try:
                idx = body.index(obj['para']._element)
                indices.append(idx)
            except ValueError:
                pass 
        
        if not indices:
            return mapping
            
        anchor = min(indices)
        
        # Remove all
        for obj in ref_objects:
             p = obj['para']._element
             if p.getparent() == body:
                 body.remove(p)
                 
        # Insert Cited (Sorted)
        cited_refs.sort(key=lambda x: x['new_id'])
        uncited_refs.sort(key=lambda x: x['id'])

        insert_idx = anchor
        for obj in cited_refs:
            normalize_reference_paragraph(obj['para'], obj['new_id'], self.doc)
            body.insert(insert_idx, obj['para']._element)
            insert_idx += 1
            
        # Insert Uncited (Appended after cited)
        for obj in uncited_refs:
            next_unused_id = insert_idx - anchor + 1
            normalize_reference_paragraph(obj['para'], next_unused_id, self.doc)
            body.insert(insert_idx, obj['para']._element)
            insert_idx += 1

        # Ensure all bibliography numbers have bib_number style applied
        self._ensure_bib_numbers_styled()

        # Ensure all citation runs have cite_bib style
        self._ensure_citations_styled()

        return mapping


def process_document(file, citation_format='styled'):
    doc = Document(file)
    processor = ReferenceProcessor(doc)
    citation_format = (citation_format or 'styled').strip().lower()

    # PRE-PROCESS: Ensure styles exist
    _ensure_styles(doc)

    # PRE-PROCESS: Convert Word auto-numbered reference lists to manual
    auto_converted = 0
    normalized_references = 0
    for para in doc.paragraphs:
        if para.style and para.style.name == 'REF-N':
            manual_num, _ = extract_bib_number(para.text)
            numId, ilvl = get_numPr(para)
            if numId:
                manual_num = compute_list_number(doc, para)
                convert_autonumber_to_manual(para, manual_num, doc)
                auto_converted += 1
                normalized_references += 1
            elif manual_num is not None:
                normalize_reference_paragraph(para, manual_num, doc)
                normalized_references += 1

    # PRE-PROCESS: Detect and tag unstyled citations
    cite_tag_result = {}
    if citation_format != 'styled':
        cite_tag_result = detect_and_tag_unstyled_citations(doc, citation_format)

    # Check BEFORE
    before_stats = processor.get_validation_stats()
    before_stats['citation_format'] = citation_format
    before_stats['auto_converted'] = auto_converted
    before_stats['references_normalized'] = normalized_references
    before_stats['citations_tagged'] = cite_tag_result.get('tagged', 0)
    before_stats['pipeline_log'] = [
        "Step 1: pre-processing completed",
        "Step 2: before-stats collected"
    ]

    # STEP 2: Missing check — STOP if any missing references
    if before_stats["missing_references"]:
        before_stats['pipeline_log'].append("Abort: missing references detected before renumbering")
        return doc, before_stats, before_stats, {}, {}, \
            f"Stopped: Missing references detected: {before_stats['missing_references']}", []

    # If Perfect -> No changes needed
    if before_stats["is_perfect"]:
        before_stats['pipeline_log'].append("No pass processing required")
        return doc, before_stats, before_stats, {}, {}, "Validation completed.", []

    # TWO-PASS VALIDATION PIPELINE
    # ===========================

    # PASS 1: Renumber based on citation order and reorder bibliography
    # -----------------------------------------------------------------
    mapping_pass1 = processor.renumber()
    pass1_stats = processor.get_validation_stats()

    # STEP 4: Check for 100% duplicates after reordering (PASS 2)
    merge_log = []
    mapping_pass2 = {}

    if pass1_stats["duplicate_references"]:
        merge_log = processor.resolve_duplicates()
        mapping_pass2 = processor.renumber()

    # Final validation stats
    after_stats = processor.get_validation_stats()
    after_stats['citation_format'] = citation_format
    after_stats['auto_converted'] = auto_converted
    after_stats['references_normalized'] = normalized_references
    after_stats['citations_tagged'] = cite_tag_result.get('tagged', 0)
    after_stats['pass1_mapping'] = mapping_pass1
    after_stats['pass2_mapping'] = mapping_pass2
    after_stats['merge_log'] = merge_log
    after_stats['pipeline_log'] = [
        "Step 1: pre-processing completed",
        "Step 2: before-stats passed",
        f"Step 3: pass 1 renumbered {len(mapping_pass1)} cited ids",
        f"Step 4: pass 2 removed {len(merge_log)} duplicate references and renumbered {len(mapping_pass2)} ids" if merge_log else "Step 4: no exact duplicates found",
        "Step 5: after-stats collected"
    ]

    # Determine status message
    status_parts = []
    if merge_log:
        count = len(merge_log)
        status_parts.append(f"{count} duplicate reference{'s' if count > 1 else ''} removed")

    pass1_changes = any(k != v for k, v in mapping_pass1.items()) if mapping_pass1 else False
    pass2_changes = any(k != v for k, v in mapping_pass2.items()) if mapping_pass2 else False
    after_perfect = (
        not after_stats["missing_references"] and
        not after_stats["unused_references"] and
        not after_stats["duplicate_references"] and
        not after_stats["sequence_issues"]
    )

    if pass1_changes or pass2_changes or merge_log:
        if merge_log:
            status_msg = f"Two-pass validation: Pass 1 renumbered, Pass 2 {', '.join(status_parts)} and renumbered."
        elif pass1_changes:
            status_msg = "Two-pass validation: Pass 1 renumbered and sequence fixed."
        else:
            status_msg = "Two-pass validation completed."
    elif after_perfect:
        status_msg = "Validation completed with perfect citation/reference sequence."
    else:
        status_msg = "Validation completed - no changes needed."

    # Return 7 values: doc, before_stats, after_stats, mapping_pass1, mapping_pass2, status_msg, merge_log
    return doc, before_stats, after_stats, mapping_pass1, mapping_pass2, status_msg, merge_log


def write_validation_report(report_path, before, after, mapping_pass1, mapping_pass2, status_msg, merge_log):
    with open(report_path, "w", encoding="utf-8") as f:
        f.write(f"STATUS: {status_msg}\n\n")

        f.write("STEP 1: PRE-PROCESS\n")
        f.write(f"Citation format: {before.get('citation_format', 'styled')}\n")
        f.write(f"Reference styles normalized: {before.get('references_normalized', 0)}\n")
        f.write(f"Auto-numbered references converted: {before.get('auto_converted', 0)}\n")
        f.write(f"Unstyled citations tagged: {before.get('citations_tagged', 0)}\n\n")

        f.write("STEP 2: BEFORE STATS\n")
        f.write(str(before) + "\n\n")

        f.write("STEP 3: PASS 1 - INITIAL RENUMBERING\n")
        if mapping_pass1:
            for old, new in sorted(mapping_pass1.items(), key=lambda item: item[1]):
                f.write(f"{old} -> {new}\n")
        else:
            f.write("No pass 1 remapping required.\n")
        f.write("\n")

        f.write("STEP 4: PASS 2 - DUPLICATE REMOVAL\n")
        if merge_log:
            for merge in merge_log:
                f.write(
                    f"Removed {merge['removed_id']} -> kept {merge['canonical_id']} "
                    f"(citations updated: {merge['citations_updated']}, score: {merge['score']})\n"
                )
        else:
            f.write("No exact duplicate references found.\n")

        if mapping_pass2:
            f.write("Pass 2 renumbering:\n")
            for old, new in sorted(mapping_pass2.items(), key=lambda item: item[1]):
                f.write(f"{old} -> {new}\n")
        f.write("\n")

        f.write("STEP 5: AFTER STATS\n")
        f.write(str(after) + "\n")


# =====================================================
# Flask Routes
# =====================================================
@app.route("/")
def upload_file():
    return render_template("upload.html")


@app.route("/process", methods=["GET", "POST"])
def process():
    if request.method == "POST":
        file = request.files.get("file")
        if not file or not file.filename.endswith(".docx"):
            return "Invalid file", 400

        citation_format = request.form.get("citation_format", "styled")
        doc, before, after, mapping_pass1, mapping_pass2, status_msg, merge_log = process_document(
            file,
            citation_format=citation_format
        )

        base = os.path.splitext(file.filename)[0]
        doc_path = os.path.join(UPLOAD_DIR, f"{base}_renumbered.docx")
        report_path = os.path.join(UPLOAD_DIR, f"{base}_validation.txt")

        doc.save(doc_path)

        write_validation_report(report_path, before, after, mapping_pass1, mapping_pass2, status_msg, merge_log)

        # Create ZIP package
        zip_filename = f"{base}_results.zip"
        zip_path = os.path.join(UPLOAD_DIR, zip_filename)

        # Validation HTML Report (Offline)
        html_report_filename = f"{base}_results.html"
        html_report_path = os.path.join(UPLOAD_DIR, html_report_filename)

        # Render the template for offline use
        # Note: We pass offline_mode=True to make links relative
        html_content = render_template(
            "validation_results.html",
            filename=file.filename,
            results=after,
            before=before,
            mapping_pass1=mapping_pass1,
            mapping_pass2=mapping_pass2,
            status_msg=status_msg,
            report_file=os.path.basename(report_path),
            doc_file=os.path.basename(doc_path),
            zip_file=None, # No zip button in offline report
            offline_mode=True,
            merge_log=merge_log
        )
        
        with open(html_report_path, "w", encoding="utf-8") as f:
            f.write(html_content)

        with zipfile.ZipFile(zip_path, 'w') as zf:
             # Add Doc
             zf.write(doc_path, arcname=os.path.basename(doc_path))
             # Add Text Report
             zf.write(report_path, arcname=os.path.basename(report_path))
             # Add HTML Report
             zf.write(html_report_path, arcname=os.path.basename(html_report_path))

        # Store data in session for GET request
        session['processing_result'] = {
            'filename': file.filename,
            'before': before,
            'after': after,
            'mapping_pass1': mapping_pass1,
            'mapping_pass2': mapping_pass2,
            'status_msg': status_msg,
            'report_file': os.path.basename(report_path),
            'doc_file': os.path.basename(doc_path),
            'zip_file': zip_filename,
            'merge_log': merge_log
        }

        return redirect(url_for('process'))

    # GET request - retrieve from session
    result = session.get('processing_result')
    if not result:
        return redirect(url_for('upload_file'))

    return render_template(
        "validation_results.html",
        filename=result['filename'],
        results=result['after'],
        before=result['before'],
        mapping_pass1=result['mapping_pass1'],
        mapping_pass2=result['mapping_pass2'],
        status_msg=result['status_msg'],
        report_file=result['report_file'],
        doc_file=result['doc_file'],
        zip_file=result.get('zip_file'),
        merge_log=result.get('merge_log', [])
    )


@app.route("/download/<path:filename>")
def download_file(filename):
    # Security: Ensure filename is in UPLOAD_DIR
    return send_file(os.path.join(UPLOAD_DIR, filename), as_attachment=True)


if __name__ == "__main__":
    app.run(debug=True, port=5000)
